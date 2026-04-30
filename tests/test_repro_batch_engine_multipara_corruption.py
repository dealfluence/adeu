import io

import docx

from adeu import ModifyText, RedlineEngine


def test_batch_engine_corrupts_multipara_insertion():
    """
    Reproduces a bug where replacing a space with multiple paragraphs
    (e.g., 'Cell Text' -> 'Cell\n\nNew\n\nText') causes the engine to
    swallow the trailing text and corrupt the OOXML structure, generating
    out-of-order text fragments.
    """
    d = docx.Document()
    d.add_paragraph("Cell Text")
    stream = io.BytesIO()
    d.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA")
    engine.process_batch([ModifyText(target_text="Cell Text", new_text="Cell\n\nNew\n\nText")])

    # We must accept all revisions here to properly inspect text structure
    # because docx.Paragraph.text natively ignores <w:ins> tag contents
    engine.accept_all_revisions()
    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    d2 = docx.Document(out_stream)

    texts = [p.text for p in d2.paragraphs]
    assert texts == ["Cell", "New", "Text"], f"BUG: Text mapping corruption! Got {texts}"
