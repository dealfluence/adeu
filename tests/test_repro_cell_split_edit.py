# FILE: tests/test_repro_dk1_cell_split.py
from io import BytesIO

from docx import Document

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_dk1_cell_split_empty_cell_placement():
    """
    Tests that a cell-spanning edit properly targets an empty cell,
    even if the LLM omits the trailing space in the target string.
    """
    # 1. Setup: 1x2 table, Cell 1 is empty.
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Site Organization"

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Apply Edit
    engine = RedlineEngine(stream, author="Test AI")
    edit = ModifyText(
        type="modify",
        target_text="Site Organization |",  # Notice the missing trailing space
        new_text="Site Organization | 10%",
    )

    stats = engine.process_batch([edit])

    # Verify the engine didn't skip it
    assert stats["edits_applied"] == 1
    assert stats["edits_skipped"] == 0

    out_stream = engine.save_to_stream()
    out_doc = Document(out_stream)

    # 3. Verify XML directly (The <w:ins> should be in Cell 1)
    cell_0_xml = out_doc.tables[0].cell(0, 0)._tc.xml
    cell_1_xml = out_doc.tables[0].cell(0, 1)._tc.xml

    assert "10%" not in cell_0_xml, "The text '10%' was wrongly inserted into the first cell!"
    assert "<w:ins" in cell_1_xml, "The insertion tag is missing from the second cell!"
    assert "10%" in cell_1_xml, "The text '10%' is missing from the second cell!"

    # 4. Verify Final Accepted State
    out_stream.seek(0)
    final_engine = RedlineEngine(out_stream)
    final_engine.accept_all_revisions()
    final_doc = Document(final_engine.save_to_stream())

    assert final_doc.tables[0].cell(0, 0).text == "Site Organization"
    assert final_doc.tables[0].cell(0, 1).text == "10%"
