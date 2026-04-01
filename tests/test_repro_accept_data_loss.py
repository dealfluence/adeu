import io

import structlog
from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import AcceptChange, ModifyText
from adeu.redline.engine import RedlineEngine

logger = structlog.get_logger(__name__)


def test_repro_accept_resolves_paired_modification():
    """
    Scenario:
    1. Modifying text creates a Deletion (ID 1) and paired Insertion (ID 2).
    2. User targets the Deletion ID to accept.
    3. The paired Insertion MUST also be resolved (accepted) automatically to mimic Word's atomic handling.
    """
    doc = Document()
    doc.add_paragraph("Old Text")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Apply Edit
    engine = RedlineEngine(stream, author="Me")
    edit = ModifyText(target_text="Old Text", new_text="New Text")
    engine.apply_edits([edit])

    stream_edited = engine.save_to_stream()
    text_mid = extract_text_from_stream(stream_edited)

    assert "{--Old--}" in text_mid
    assert "{++New++}" in text_mid

    # 2. Accept the Deletion
    import re

    ids = re.findall(r"\[Chg:(\d+)\]", text_mid)
    del_id = ids[0] if ids else "1"

    engine2 = RedlineEngine(stream_edited, author="Reviewer")
    action = AcceptChange(target_id=f"Chg:{del_id}")
    engine2.apply_review_actions([action])

    stream_final = engine2.save_to_stream()
    text_final = extract_text_from_stream(stream_final)

    # Check:
    # 1. Old Text should be GONE (Accepted Deletion)
    assert "Old" not in text_final, "Old Text should be removed"

    # 2. New Text should be PRESENT as normal text (Accepted paired Insertion)
    assert "{++New++}" not in text_final, "New Text tracking wrapper should be removed by paired resolution"
    assert "New Text" in text_final, "New Text content should remain in the document"


def test_id_collision_prevention():
    """
    Ensure RedlineEngine respects existing IDs in the document.
    """
    doc = Document()
    p = doc.add_paragraph("Start")
    # Manually inject an ID=5 insertion
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "5")
    ins.set(qn("w:author"), "Existing")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Existing"
    r.append(t)
    ins.append(r)
    p._element.append(ins)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Create Engine
    engine = RedlineEngine(stream)

    # The next ID should be > 5 (i.e., 6)
    next_id = engine._get_next_id()
    assert int(next_id) > 5, f"Engine should pick ID > 5, got {next_id}"
