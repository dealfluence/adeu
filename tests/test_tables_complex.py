import io

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine


def test_interleaved_tables_and_text():
    """
    Verifies that the extractor respects document order:
    Paragraph -> Table -> Paragraph.
    Previously, all tables were extracted at the end of the section.
    """
    doc = Document()
    doc.add_paragraph("Section 1")

    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "TableContent"

    doc.add_paragraph("Section 2")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # 1. Content check
    assert "Section 1" in text
    assert "TableContent" in text
    assert "Section 2" in text

    # 2. Order check
    p1 = text.find("Section 1")
    tbl = text.find("TableContent")
    p2 = text.find("Section 2")

    assert p1 < tbl < p2, f"Table content out of order! Indicies: P1={p1}, Tbl={tbl}, P2={p2}"


def test_nested_tables_extraction_and_editing():
    """
    Verifies recursive extraction logic.
    Structure: Table -> Cell -> Table -> Cell -> Text
    """
    doc = Document()
    outer_table = doc.add_table(rows=1, cols=1)
    outer_cell = outer_table.cell(0, 0)

    # Add nested table inside the cell
    nested_table = outer_cell.add_table(rows=1, cols=1)
    nested_table.cell(0, 0).text = "InnerSecret"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Verify Ingest finds it
    text = extract_text_from_stream(stream)
    assert "InnerSecret" in text, "Nested table content failed to extract"

    # 2. Verify Mapping/Editing can reach it
    edit = DocumentEdit(target_text="InnerSecret", new_text="OuterSecret")
    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1
    assert skipped == 0

    res_stream = engine.save_to_stream()
    res_text = extract_text_from_stream(res_stream)

    # Expect clean insertion logic
    assert "{--InnerSecret--}{++OuterSecret++}" in res_text


def test_merged_cells_no_duplication():
    """
    Verifies that merged cells are extracted exactly once.
    python-docx iterates a 2-col merged row as [Cell A, Cell A].
    We must deduplicate to avoid "Text | Text".
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)

    c1.merge(c2)
    c1.text = "MergedUnique"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # Should appear exactly once. If deduplication fails, count will be 2.
    count = text.count("MergedUnique")
    assert count == 1, f"Merged cell content appeared {count} times (expected 1)"

    # Verify Edit targets correctly (coordinates shouldn't be confused by skip)
    edit = DocumentEdit(target_text="MergedUnique", new_text="ChangedUnique")
    engine = RedlineEngine(stream)
    applied, _ = engine.apply_edits([edit])
    assert applied == 1


def test_empty_row_alignment():
    """
    Verifies that Ingest and Mapper stay synchronized even with empty rows.
    If Ingest skips empty rows but Mapper counts them (or vice versa),
    subsequent edits will drift and target the wrong text.
    """
    doc = Document()
    table = doc.add_table(rows=3, cols=1)

    table.cell(0, 0).text = "RowA"
    table.cell(1, 0).text = ""  # Empty Row
    table.cell(2, 0).text = "RowB"  # Target

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # If alignment is broken, "RowB" index will be calculated wrong
    edit = DocumentEdit(target_text="RowB", new_text="RowC")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1, "Edit failed - likely due to mapping index drift caused by empty row"

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream)

    # RowA should still exist
    assert "RowA" in text
    # RowB should be modified
    assert "{--RowB--}{++RowC++}" in text
