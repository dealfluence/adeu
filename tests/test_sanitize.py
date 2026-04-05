"""Tests for adeu sanitize — DOCX metadata scrubber."""

import io
import os
import tempfile

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.sanitize.core import SanitizeError, SanitizeMode, SanitizeResult, sanitize_docx
from adeu.sanitize import transforms


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_doc_with_track_changes() -> io.BytesIO:
    """Create a DOCX with track changes (insertion + deletion)."""
    doc = Document()
    p = doc.add_paragraph()

    # Normal text
    p.add_run("The ")

    # Deletion
    d = OxmlElement("w:del")
    d.set(qn("w:id"), "1")
    d.set(qn("w:author"), "Opposing Counsel")
    d.set(qn("w:date"), "2025-01-15T10:00:00Z")
    rd = OxmlElement("w:r")
    rt = OxmlElement("w:delText")
    rt.set(qn("xml:space"), "preserve")
    rt.text = "Vendor"
    rd.append(rt)
    d.append(rd)
    p._element.append(d)

    # Insertion
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "2")
    ins.set(qn("w:author"), "Opposing Counsel")
    ins.set(qn("w:date"), "2025-01-15T10:00:00Z")
    ri = OxmlElement("w:r")
    ti = OxmlElement("w:t")
    ti.set(qn("xml:space"), "preserve")
    ti.text = "Supplier"
    ri.append(ti)
    ins.append(ri)
    p._element.append(ins)

    p.add_run(" shall provide services.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _make_doc_with_rsids() -> io.BytesIO:
    """Create a DOCX with rsid attributes on paragraphs and runs."""
    doc = Document()
    p = doc.add_paragraph("Hello World")
    p._element.set(qn("w:rsidR"), "00A21F3B")
    p._element.set(qn("w:rsidRDefault"), "004C12DE")
    p._element.set(qn("w:rsidP"), "00B33E21")

    for run in p.runs:
        run._element.set(qn("w:rsidR"), "00A21F3B")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _make_doc_with_comments(resolved: bool = False) -> io.BytesIO:
    """Create a DOCX with a comment."""
    doc = Document()
    p = doc.add_paragraph()

    # Comment range start
    crs = OxmlElement("w:commentRangeStart")
    crs.set(qn("w:id"), "0")
    p._element.append(crs)

    p.add_run("Target text for comment")

    # Comment range end
    cre = OxmlElement("w:commentRangeEnd")
    cre.set(qn("w:id"), "0")
    p._element.append(cre)

    # Comment reference in a run
    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), "0")
    ref_run.append(ref)
    p._element.append(ref_run)

    # Create comments.xml part manually via the engine approach
    from adeu.redline.comments import CommentsManager
    cm = CommentsManager(doc)
    cm.add_comment(
        comment_id="0",
        author="Internal Reviewer",
        text="Check this with the client",
    )

    # If resolved, mark it
    if resolved and cm.extended_part:
        for child in cm.extended_part.element:
            child.set(qn("w15:done"), "1")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _save_to_tmp(stream: io.BytesIO, suffix=".docx") -> str:
    """Save a BytesIO stream to a temp file, return the path."""
    fd, path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, "wb") as f:
        f.write(stream.getvalue())
    return path


# ---------------------------------------------------------------------------
# Transform unit tests
# ---------------------------------------------------------------------------


class TestStripRsid:
    def test_removes_rsid_attributes(self):
        stream = _make_doc_with_rsids()
        doc = Document(stream)

        # Verify rsids exist
        found = False
        for el in doc.element.iter():
            if qn("w:rsidR") in el.attrib:
                found = True
                break
        assert found, "Expected rsid attributes in test doc"

        lines = transforms.strip_rsid(doc)
        assert len(lines) > 0
        assert "rsid" in lines[0].lower()

        # Verify rsids are gone
        for el in doc.element.iter():
            for attr in transforms.RSID_ATTRS:
                assert attr not in el.attrib


class TestStripParaIds:
    def test_removes_para_ids(self):
        doc = Document()
        p = doc.add_paragraph("Test")
        p._element.set(f"{{{transforms.W14_NS}}}paraId", "3F2A91BC")
        p._element.set(f"{{{transforms.W14_NS}}}textId", "77D61234")

        lines = transforms.strip_para_ids(doc)
        assert len(lines) > 0

        for el in doc.element.iter():
            for attr in transforms.W14_ATTRS:
                assert attr not in el.attrib


class TestCountTrackedChanges:
    def test_counts_insertions_and_deletions(self):
        stream = _make_doc_with_track_changes()
        doc = Document(stream)
        ins, dels = transforms.count_tracked_changes(doc)
        assert ins == 1
        assert dels == 1


class TestAcceptAllTrackedChanges:
    def test_accepts_changes(self):
        stream = _make_doc_with_track_changes()
        doc = Document(stream)

        lines = transforms.accept_all_tracked_changes(doc)
        assert len(lines) > 0

        # Verify no track changes remain
        ins, dels = transforms.count_tracked_changes(doc)
        assert ins == 0
        assert dels == 0

        # Verify text content: "Vendor" deleted, "Supplier" kept
        text = doc.paragraphs[0].text
        assert "Supplier" in text
        assert "Vendor" not in text


class TestStripHiddenText:
    def test_removes_vanish_runs(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Hidden text")
        # Add w:vanish to the run's rPr
        rpr = run._element.get_or_add_rPr()
        vanish = OxmlElement("w:vanish")
        rpr.append(vanish)

        lines = transforms.strip_hidden_text(doc)
        assert len(lines) > 0
        assert "hidden" in lines[0].lower()


# ---------------------------------------------------------------------------
# Orchestrator integration tests
# ---------------------------------------------------------------------------


class TestSanitizeFull:
    def test_full_sanitize_clean_doc(self):
        """Full sanitize on a doc with no track changes works."""
        doc = Document()
        doc.add_paragraph("Clean document.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        input_path = _save_to_tmp(stream)

        try:
            result = sanitize_docx(input_path)
            assert result.status == "clean"
            assert os.path.exists(result.output_path)

            # Verify output opens
            out_doc = Document(result.output_path)
            assert "Clean document" in out_doc.paragraphs[0].text
        finally:
            os.unlink(input_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)

    def test_full_sanitize_blocks_on_unresolved_changes(self):
        """Full sanitize refuses if unresolved track changes exist."""
        stream = _make_doc_with_track_changes()
        input_path = _save_to_tmp(stream)

        try:
            with pytest.raises(SanitizeError) as exc_info:
                sanitize_docx(input_path)
            assert "unresolved" in str(exc_info.value).lower()
        finally:
            os.unlink(input_path)

    def test_full_sanitize_with_accept_all(self):
        """Full sanitize with --accept-all accepts changes and succeeds."""
        stream = _make_doc_with_track_changes()
        input_path = _save_to_tmp(stream)

        try:
            result = sanitize_docx(input_path, accept_all=True)
            assert result.status == "clean"
            assert result.tracked_changes_accepted > 0

            out_doc = Document(result.output_path)
            text = out_doc.paragraphs[0].text
            assert "Supplier" in text
        finally:
            os.unlink(input_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)

    def test_full_sanitize_strips_rsids(self):
        """Full sanitize removes rsid attributes."""
        stream = _make_doc_with_rsids()
        input_path = _save_to_tmp(stream)

        try:
            result = sanitize_docx(input_path)
            out_doc = Document(result.output_path)

            for el in out_doc.element.iter():
                for attr in transforms.RSID_ATTRS:
                    assert attr not in el.attrib, f"rsid attribute {attr} should be stripped"
        finally:
            os.unlink(input_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)

    def test_report_text_generated(self):
        """Sanitize produces a non-empty report."""
        stream = _make_doc_with_rsids()
        input_path = _save_to_tmp(stream)

        try:
            result = sanitize_docx(input_path)
            assert result.report_text
            assert "Sanitize Report" in result.report_text
        finally:
            os.unlink(input_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)


class TestSanitizeKeepMarkup:
    def test_keeps_track_changes(self):
        """--keep-markup preserves track changes."""
        stream = _make_doc_with_track_changes()
        input_path = _save_to_tmp(stream)

        try:
            result = sanitize_docx(input_path, keep_markup=True)
            assert result.status == "clean"
            assert result.tracked_changes_found > 0

            out_doc = Document(result.output_path)
            ins_count = len(out_doc.element.findall(f".//{qn('w:ins')}"))
            del_count = len(out_doc.element.findall(f".//{qn('w:del')}"))
            assert ins_count > 0 or del_count > 0, "Track changes should be preserved"
        finally:
            os.unlink(input_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)

    def test_strips_rsids_but_keeps_markup(self):
        """--keep-markup strips metadata but keeps changes."""
        stream = _make_doc_with_track_changes()
        input_path = _save_to_tmp(stream)

        try:
            result = sanitize_docx(input_path, keep_markup=True)
            out_doc = Document(result.output_path)

            # rsids should be gone
            for el in out_doc.element.iter():
                for attr in transforms.RSID_ATTRS:
                    assert attr not in el.attrib

            # but track changes should remain
            ins_els = out_doc.element.findall(f".//{qn('w:ins')}")
            assert len(ins_els) > 0
        finally:
            os.unlink(input_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)

    def test_warns_when_no_markup_found(self):
        """--keep-markup warns if document has no track changes or comments."""
        doc = Document()
        doc.add_paragraph("No changes here.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        input_path = _save_to_tmp(stream)

        try:
            result = sanitize_docx(input_path, keep_markup=True)
            assert result.status == "clean_with_warnings"
            assert any("no tracked changes" in w.lower() for w in result.warnings)
        finally:
            os.unlink(input_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)

    def test_replaces_author_names(self):
        """--keep-markup with --author replaces author names on markup."""
        stream = _make_doc_with_track_changes()
        input_path = _save_to_tmp(stream)

        try:
            result = sanitize_docx(input_path, keep_markup=True, author="Firm A")
            out_doc = Document(result.output_path)

            for ins in out_doc.element.findall(f".//{qn('w:ins')}"):
                assert ins.get(qn("w:author")) == "Firm A"
            for d in out_doc.element.findall(f".//{qn('w:del')}"):
                assert d.get(qn("w:author")) == "Firm A"
        finally:
            os.unlink(input_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)


class TestSanitizeBaseline:
    def test_baseline_recomputes_delta(self):
        """--baseline produces track changes showing the diff."""
        # Create baseline
        baseline_doc = Document()
        baseline_doc.add_paragraph("The Vendor shall provide services.")
        baseline_stream = io.BytesIO()
        baseline_doc.save(baseline_stream)
        baseline_stream.seek(0)
        baseline_path = _save_to_tmp(baseline_stream)

        # Create working version (different text)
        working_doc = Document()
        working_doc.add_paragraph("The Supplier shall provide services.")
        working_stream = io.BytesIO()
        working_doc.save(working_stream)
        working_stream.seek(0)
        working_path = _save_to_tmp(working_stream)

        try:
            result = sanitize_docx(
                working_path,
                baseline_path=baseline_path,
                author="Test Firm",
            )
            assert result.status in ("clean", "clean_with_warnings")
            assert result.tracked_changes_found > 0

            out_doc = Document(result.output_path)
            # Should have track changes
            all_changes = (
                out_doc.element.findall(f".//{qn('w:ins')}")
                + out_doc.element.findall(f".//{qn('w:del')}")
            )
            assert len(all_changes) > 0, "Baseline mode should produce track changes"
        finally:
            os.unlink(baseline_path)
            os.unlink(working_path)
            if os.path.exists(result.output_path):
                os.unlink(result.output_path)


class TestSanitizeExitCodes:
    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            sanitize_docx("/nonexistent/file.docx")

    def test_baseline_not_found(self):
        doc = Document()
        doc.add_paragraph("Test")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        input_path = _save_to_tmp(stream)

        try:
            with pytest.raises(FileNotFoundError):
                sanitize_docx(input_path, baseline_path="/nonexistent/baseline.docx")
        finally:
            os.unlink(input_path)
