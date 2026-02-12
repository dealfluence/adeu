# FILE: tests/test_repro_markdown_boundary.py

import io

from docx import Document

from adeu.markup import apply_edits_to_markdown
from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine


def test_repro_ui_markdown_boundary_leak_no_space():
    """
    Reproduction of the UI bug where the target text consumes the
    closing bold marker of the PRECEDING entity when there is NO SPACE.

    Original Text: "**Header.**Body"
    Target: "Body"

    If the matcher is greedy or fuzzy regex allows preceding markers to be consumed,
    it might grab the closing '**' of Header as the start of Body match.
    """
    text = "**Header.**Body"

    edit = DocumentEdit(target_text="Body", new_text="NewBody")

    result = apply_edits_to_markdown(text, [edit])

    # Expected: "**Header.**{--Body--}{++NewBody++}"
    # Buggy: "**Header.{--**Body--}{++NewBody++}"

    assert "**Header.**" in result, f"The header's closing bold marker was lost! Result: {result}"
    assert "{--**" not in result, f"The deletion block incorrectly starts with '**'! Result: {result}"


def test_repro_engine_skipped_edit_on_boundary_fixed_assertion():
    """
    Reproduction of the Engine skipping edits when the target immediately
    follows a bold run.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run 1: Bold Header
    r1 = p.add_run("3.2 Payment Terms.")
    r1.bold = True

    # Run 2: Space + Content (Normal)
    # We explicitly verify the engine can handle this common boundary
    r2 = p.add_run(" Standard payment terms are Net 90.")
    r2.bold = False

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Define Edit targeting the plain text
    edit = DocumentEdit(
        target_text="Standard payment terms are Net 90.",
        new_text="Standard payment terms are Net 30.",
    )

    stream.seek(0)
    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    # If the bug is present in the engine matching logic, this will fail
    assert applied == 1, f"Engine skipped the edit! Skipped count: {skipped}"

    # Verify content changed using XML inspection
    res_stream = engine.save_to_stream()
    doc_res = Document(res_stream)

    # Find the paragraph
    p = doc_res.paragraphs[0]
    xml = p._element.xml

    # We expect an insertion containing "Net 30"
    # It might be split like <w:t>Standard payment terms are Net 30.</w:t>
    # or <w:t>Net 30</w:t>

    assert "Net 30" in xml or "Net 30" in "".join(t.text for t in p._element.xpath(".//w:t")), (
        f"New text 'Net 30' not found in paragraph XML: {xml}"
    )


def test_repro_engine_skip_with_formatting_noise():
    """
    Reproduction of Engine skipping when target matches text that has internal
    formatting markers in the DOCX (e.g. **Net 90**).

    Scenario:
    Doc: "...are **Net 90 Days**..."
    Target: "...are Net 90 Days..."
    New: "...are Net 30 Days..."

    The engine must fuzzy match 'Net 90 Days' against '**Net 90 Days**'.
    """
    doc = Document()
    p = doc.add_paragraph("Terms are ")

    # Bold part
    r = p.add_run("Net 90 Days")
    r.bold = True

    p.add_run(".")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Target provided by LLM is usually plain text
    edit = DocumentEdit(target_text="Terms are Net 90 Days.", new_text="Terms are Net 30 Days.")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1, f"Engine failed to match plain text target against bolded doc text. Skipped: {skipped}"
