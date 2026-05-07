import io
import zipfile

from docx import Document

from adeu.ingest import _extract_text_from_doc
from adeu.models import ModifyText
from adeu.outline import extract_outline
from adeu.pagination import paginate
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_paragraph_prefix

W_NS_DECL = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _wrap_docx(document_xml: str, styles_xml: str) -> io.BytesIO:
    """Bundles a document.xml and styles.xml into a minimal valid DOCX zip."""
    rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rId1"'
        b' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"'
        b' Target="word/document.xml"/>'
        b"</Relationships>"
    )
    doc_rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rId1"'
        b' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles"'
        b' Target="styles.xml"/>'
        b"</Relationships>"
    )
    ct = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        b'<Default Extension="rels"'
        b' ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        b'<Override PartName="/word/document.xml"'
        b' ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        b'<Override PartName="/word/styles.xml"'
        b' ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        b"</Types>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("[Content_Types].xml", ct)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("word/document.xml", document_xml.encode("utf-8"))
        z.writestr("word/styles.xml", styles_xml.encode("utf-8"))
    buf.seek(0)
    return buf


class TestQaIssuesV2:
    def test_issue_10_accept_all_changes_removes_comments(self):
        """Issue 10: accept_all_changes doesn't remove comments."""
        doc = Document()
        doc.add_paragraph("Text with comment.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        # Add a comment
        engine.process_batch([ModifyText(target_text="comment", new_text="comment", comment="QA Comment")])

        assert len(engine.comments_manager.extract_comments_data()) == 1

        engine.accept_all_revisions()

        assert len(engine.comments_manager.extract_comments_data()) == 0, (
            "Comments should be removed by accept_all_revisions"
        )

    def test_issue_1_outline_custom_heading_styles(self):
        """Issue 1: Outline detector misses long custom-derived heading styles."""
        # This style name was reported as failing
        style_name = "StyleHeading2NotItalicBefore0ptAfter0ptLinespa"

        styles_xml = (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f"<w:styles {W_NS_DECL}>"
            f'<w:style w:type="paragraph" w:styleId="{style_name}">'
            f'<w:name w:val="{style_name}"/>'
            # No explicit outlineLvl, so it relies on the name heuristic
            f"</w:style>"
            f"</w:styles>"
        )
        document_xml = (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f"<w:document {W_NS_DECL}><w:body>"
            f'<w:p><w:pPr><w:pStyle w:val="{style_name}"/></w:pPr>'
            f"<w:r><w:t>Sub Heading</w:t></w:r></w:p>"
            f"<w:sectPr/></w:body></w:document>"
        )
        doc = Document(_wrap_docx(document_xml, styles_xml))
        p = doc.paragraphs[0]

        prefix = get_paragraph_prefix(p)
        assert prefix == "## ", f"Custom heading style '{style_name}' was not recognized as H2. Got prefix: '{prefix}'"

    def test_issue_4_heading_2_with_leading_break(self):
        """Issue 4: Some Heading 2s render as ## \\n<text>."""
        doc = Document()
        p = doc.add_paragraph(style="Heading 2")
        r1 = p.add_run()
        r1.add_break()
        p.add_run("Heading Text")

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        text = _extract_text_from_doc(Document(stream))

        # We expect "## Heading Text" (with or without a space)
        # But NOT "## \nHeading Text"
        assert "## \n" not in text, f"Heading 2 rendered with a newline after prefix: {repr(text)}"

    def test_issue_8_heading_level_change_leaves_empty_paragraph(self):
        """Issue 8: Changing ## X -> # Y leaves behind an empty paragraph after accept_all_changes."""
        doc = Document()
        doc.add_paragraph("Heading 2 Content", style="Heading 2")
        doc.add_paragraph("Body text.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        # Change ## X to # Y
        # In projected text, H2 is "## Heading 2 Content"
        edit = ModifyText(target_text="## Heading 2 Content", new_text="# Heading 1 Content")
        engine.process_batch([edit])

        engine.accept_all_revisions()

        # Check all paragraphs, including empty ones
        paragraphs = [p.text for p in engine.doc.paragraphs]

        # The bug is that the old paragraph remains as an empty one.
        # So we expect ["", "Heading 1 Content", "Body text."] if it fails.
        assert len(paragraphs) == 2, f"Expected 2 paragraphs, but got {len(paragraphs)}: {paragraphs}"
        assert paragraphs == ["Heading 1 Content", "Body text."]

    def test_issue_9_reject_multi_paragraph_leaves_break(self):
        """Issue 9: Rejecting multi-paragraph insertion leaves behind the paragraph break."""
        from adeu.models import RejectChange

        doc = Document()
        doc.add_paragraph("Paragraph 1.")
        doc.add_paragraph("Paragraph 2.")
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        # Insert a multi-paragraph text
        edit = ModifyText(target_text="Paragraph 1.", new_text="Paragraph 1.\n\nNew Para.")
        engine.process_batch([edit])

        # REBUILD MAP to see the new insertions
        engine.mapper._build_map()

        # Get the change ID from mapper spans
        chg_id = None
        for span in engine.mapper.spans:
            if span.ins_id:
                chg_id = f"Chg:{span.ins_id}"
                break

        assert chg_id is not None, "No insertion found to reject"

        # Reject it
        engine.process_batch([RejectChange(target_id=chg_id)])

        # Verify document content
        paragraphs = [p.text for p in engine.doc.paragraphs]
        # Should be ["Paragraph 1.", "Paragraph 2."]
        # If it fails, it might be ["Paragraph 1.", "", "Paragraph 2."]
        assert len(paragraphs) == 2, f"Extra paragraph left behind after rejection: {paragraphs}"
        assert paragraphs == ["Paragraph 1.", "Paragraph 2."]

    def test_issue_2_outline_has_table_aggregation(self):
        """Issue 2: (has table) annotation aggregates upward to ancestor headings."""
        doc = Document()
        doc.add_heading("H1", level=1)
        doc.add_heading("H2", level=2)
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "Table here"
        doc.add_heading("H2 No Table", level=2)
        doc.add_paragraph("Just text")

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        doc_obj = Document(stream)
        text = _extract_text_from_doc(doc_obj)
        body = text  # Simplified
        pagination_result = paginate(body)

        nodes = extract_outline(doc_obj, body, pagination_result.body_pages, pagination_result.body_page_offsets)

        # Expected:
        # H1: has_table=True (because it owns H2 which has a table)
        # H2: has_table=True
        # H2 No Table: has_table=False

        # Wait, the report says "aggregates upward to ancestor headings; misleading for navigation."
        # If H2 has a table, and H1 is its parent, H1 usually "owns" everything until the next H1.
        # But if H2 is there, H1 should only own until H2?

        # Let's check the docstring of extract_outline:
        # "Heading ownership: a heading owns the document range from its position up to
        # (but not including) the next heading of equal or higher level."

        # So H1 owns H2. Thus H1 will have has_table=True if H2 has a table.
        # This is what the report says is "misleading".

        # If H1 owns until the next H1, then it includes H2.
        # But if we want it to NOT aggregate, it should own only until the next heading of ANY level?
        # No, that's not how outlines usually work.

        # However, if H2 *also* shows (has table), then it's redundant on H1.

        node_map = {n.text: n for n in nodes}
        assert node_map["H2"].has_table is True
        assert node_map["H2 No Table"].has_table is False

        # The issue is probably that H1 also shows it.
        assert node_map["H1"].has_table is False, "H1 should not inherit has_table from its children headings"
