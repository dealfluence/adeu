import io

from docx import Document

from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_visible_runs


def test_markdown_headers_leak_into_docx():
    doc = Document()
    p1 = doc.add_paragraph("Section 1.")
    p1.style = "Heading 1"
    p2 = doc.add_paragraph("Section 2.")
    p2.style = "Heading 1"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    target_text = "# Section 1.\n\n# Section 2."
    new_text = "# Section 1.\n\n# New Section.\n\n# Section 2."

    edit = DocumentEdit(target_text=target_text, new_text=new_text)

    engine = RedlineEngine(stream)

    # Unit Test the Helper directly
    clean, style = engine._parse_markdown_style("# New Section.")
    print(f"DEBUG: Parse '# New Section.' -> '{clean}', '{style}'")
    assert clean == "New Section."
    assert style == "Heading 1"

    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    print("\n--- Result XML ---")
    print(doc_result.element.xml)
    print("------------------\n")

    assert len(doc_result.paragraphs) == 3
    p2_text = "".join(r.text for r in get_visible_runs(doc_result.paragraphs[1]))
    assert "New Section." in p2_text
