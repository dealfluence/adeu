# FILE: tests/test_repro_dirty_docs.py

import io

from docx import Document
from docx.oxml import OxmlElement

from adeu.ingest import extract_text_from_stream


def test_extract_handles_tabs_and_breaks():
    """
    REPRO: Documents using tabs for spacing (e.g. 'Word<tab>Word')
    should extract as 'Word Word', not 'WordWord'.
    """
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run()

    # Manually inject w:t, w:tab, w:t to simulate "Word<tab>Word"
    # python-docx run.text doesn't support adding tabs easily, so we manipulate XML
    t1 = OxmlElement("w:t")
    t1.text = "Word"
    run._element.append(t1)

    tab = OxmlElement("w:tab")
    run._element.append(tab)

    t2 = OxmlElement("w:t")
    t2.text = "One"
    run._element.append(t2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # EXPECTATION: "Word One", not "WordOne"
    assert "Word One" in text
    assert "WordOne" not in text


def test_heuristic_header_detection():
    """
    REPRO: 'Normal' style paragraphs that are BOLD and ALL-CAPS should
    be detected as headers (##) to give the LLM structural context.

    UPDATED: Now includes inline markdown **markers** for bold text.
    """
    doc = Document()
    p = doc.add_paragraph("LIABILITY CAP")
    p.style = doc.styles["Normal"]

    # Make it Bold
    p.runs[0].bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # EXPECTATION: "## **LIABILITY CAP**"
    # The header heuristic adds ##. The inline processor adds **...**.
    assert "## **LIABILITY CAP**" in text
