import io
import pytest
from docx import Document
from adeu.redline.engine import RedlineEngine, BatchValidationError
from adeu.models import ModifyText

def _build_simple_doc(text: str = "Hello world.") -> io.BytesIO:
    doc = Document()
    doc.add_paragraph(text)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

def test_process_batch_returns_detailed_edit_reports():
    """
    Verifies that process_batch returns a detailed edits list
    containing contextual previews, status, and engine versioning.
    """
    stream = _build_simple_doc("The quick brown fox jumps over the lazy dog.")
    engine = RedlineEngine(stream, author="Reviewer AI")

    stats = engine.process_batch([
        ModifyText(target_text="quick brown fox", new_text="fast red fox")
    ])

    # Check baseline presence
    assert "edits" in stats
    assert len(stats["edits"]) == 1

    edit_report = stats["edits"][0]
    assert edit_report["status"] == "applied"
    assert edit_report["target_text"] == "quick brown fox"
    assert edit_report["new_text"] == "fast red fox"

    # Contextual check (should include surrounding words)
    assert "{--quick brown fox--}{++fast red fox++}" in edit_report["critic_markup"]
    assert "The " in edit_report["critic_markup"]
    assert " jumps over" in edit_report["critic_markup"]

    assert "The fast red fox jumps over" in edit_report["clean_text"]
    assert "engine" in stats and stats["engine"] == "python"
    assert "version" in stats

def test_punctuation_anchor_triggers_warning():
    """
    Verifies that target_text containing tokenization-splitting punctuation
    like underscores or hyphens generates a warning.
    """
    stream = _build_simple_doc("Refer to sample_term_name in Section 4.")
    engine = RedlineEngine(stream, author="Reviewer AI")

    stats = engine.process_batch([
        ModifyText(target_text="sample_term_name", new_text="validated_term_name")
    ])

    edit_report = stats["edits"][0]
    assert edit_report["warning"] is not None
    assert "punctuation" in edit_report["warning"].lower()
    assert "sample_term_name" in edit_report["warning"]

def test_dry_run_does_not_mutate_and_reports_safely():
    """
    Verifies that dry_run=True does not mutate the document
    and reports validation errors gracefully without raising exceptions.
    """
    stream = _build_simple_doc("Baseline text.")
    engine = RedlineEngine(stream, author="Reviewer AI")

    # 1. Valid Dry Run
    stats = engine.process_batch([
        ModifyText(target_text="Baseline", new_text="Modified Preview")
    ], dry_run=True)

    assert stats["edits_applied"] == 1
    assert stats["edits"][0]["status"] == "applied"
    assert "Modified Preview" in stats["edits"][0]["clean_text"]

    # Verify original document was NOT mutated
    out_doc = Document(engine.save_to_stream())
    para_text = out_doc.paragraphs[0].text
    assert "Modified Preview" not in para_text
    assert "Baseline text" in para_text

    # 2. Invalid Dry Run (should report without throwing)
    stats_invalid = engine.process_batch([
        ModifyText(target_text="NON_EXISTENT", new_text="fail")
    ], dry_run=True)

    assert stats_invalid["edits_skipped"] == 1
    assert stats_invalid["edits"][0]["status"] == "failed"
    assert stats_invalid["edits"][0]["error"] is not None
    assert "not found" in stats_invalid["edits"][0]["error"].lower()