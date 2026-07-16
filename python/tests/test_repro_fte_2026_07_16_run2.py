r"""
Regression tests for the second 2026-07-16 first-time-user QA run
(mailroom: room/qa-runs/adeu-fte-2026-07-16/report.md).

Findings covered:

- F1 (correctness-risk): `extract` output — stdout and `-o` alike — begins
  with a decorative `> **File Path:** ...` header line. Feeding that output
  back through the tool's own core workflow (extract → edit → diff) makes
  `diff` report the header as a genuine inserted-text edit; applying that
  batch would physically insert the header line into the document.

  Pinned contract: an *unedited* extract → diff round trip reports zero
  changes.

- F2 (confusing): error usage-lines are inconsistently scoped. An invalid
  value for a known flag (`--mode bogus`) prints `adeu extract`'s usage,
  but an unrecognized flag on the very same command prints the top-level
  `adeu` usage — pointing a self-correcting agent at the wrong option
  surface.

  Pinned contract: both argparse error paths for a subcommand invocation
  show that subcommand's usage line, exit code 2.

Findings reviewed but deliberately not tested here:

- F3 (graceful non-UTF-8 diff) and F8 (`[x]`/`[ok]` ASCII icons) are the
  terminal-encoding strategy from f6bf92c working as designed — already
  pinned by tests/test_cli_encoding.py and tests/test_repro_fte_2026_07_16.py.
- F9 ("planted" pre-existing report): resolved as genuine, not planted — it
  was the same day's *earlier* QA run against the 1.19.1 release state
  (build 1aea2a8: no accept-all subcommand, strict-UTF-8 diff read present),
  whose headline crash was independently reproduced and fixed in f6bf92c.
"""

import json
import sys
from unittest.mock import patch

import pytest


def _run_cli(argv) -> int:
    from adeu.cli import main

    with patch.object(sys, "argv", ["adeu", *argv]):
        try:
            main()
        except SystemExit as e:
            return e.code if e.code is not None else 0
    return 0


@pytest.fixture
def roundtrip_docx(tmp_path):
    from docx import Document

    path = tmp_path / "roundtrip.docx"
    doc = Document()
    doc.add_paragraph("First paragraph of the round trip test.")
    doc.add_paragraph("Second paragraph stays unchanged.")
    doc.save(str(path))
    return path


def test_f1_extract_to_diff_round_trip_reports_zero_changes(roundtrip_docx, tmp_path, capsys):
    """F1: extract → (no edits) → diff must be a no-op, not a spurious edit."""
    extracted = tmp_path / "extracted.txt"

    exit_code = _run_cli(["extract", str(roundtrip_docx), "--clean-view", "-o", str(extracted)])
    assert exit_code == 0
    capsys.readouterr()  # discard extract's status output

    exit_code = _run_cli(["diff", str(roundtrip_docx), str(extracted), "--json"])
    assert exit_code == 0

    edits = json.loads(capsys.readouterr().out.strip())
    assert edits == [], f"unedited round trip must produce no edits, got: {edits}"


def test_f1_round_trip_edit_produces_only_the_real_change(roundtrip_docx, tmp_path, capsys):
    """F1 (workflow variant): after one real edit, diff must report exactly
    that edit — no header artifact alongside it."""
    extracted = tmp_path / "extracted.txt"

    exit_code = _run_cli(["extract", str(roundtrip_docx), "--clean-view", "-o", str(extracted)])
    assert exit_code == 0
    capsys.readouterr()

    content = extracted.read_text(encoding="utf-8")
    extracted.write_text(content.replace("stays unchanged", "was modified"), encoding="utf-8")

    exit_code = _run_cli(["diff", str(roundtrip_docx), str(extracted), "--json"])
    assert exit_code == 0

    edits = json.loads(capsys.readouterr().out.strip())
    assert len(edits) == 1, f"expected exactly the one real edit, got: {edits}"
    assert "File Path" not in json.dumps(edits)
    assert edits[0]["type"] == "modify"
    assert "was modified" in edits[0]["new_text"]


def test_f2_unrecognized_flag_shows_subcommand_usage(capsys):
    """F2: an unrecognized flag on `adeu extract ...` must show the usage of
    `adeu extract`, not of top-level `adeu`."""
    from adeu.cli import main

    test_args = ["adeu", "extract", "whatever.docx", "--nonexistent-flag"]
    with patch.object(sys, "argv", test_args):
        with pytest.raises(SystemExit) as exc_info:
            main()
        assert exc_info.value.code == 2

    err = capsys.readouterr().err
    assert "unrecognized arguments: --nonexistent-flag" in err
    assert "usage: adeu extract" in err


def test_f2_invalid_choice_shows_subcommand_usage(capsys):
    """F2 guard: the already-correct half — invalid choice for a known flag
    shows the subcommand usage — must stay correct."""
    from adeu.cli import main

    test_args = ["adeu", "extract", "whatever.docx", "--mode", "bogus"]
    with patch.object(sys, "argv", test_args):
        with pytest.raises(SystemExit) as exc_info:
            main()
        assert exc_info.value.code == 2

    err = capsys.readouterr().err
    assert "invalid choice: 'bogus'" in err
    assert "usage: adeu extract" in err
