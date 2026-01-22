import io
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from adeu.redline.engine import RedlineEngine
from adeu.models import DocumentEdit, ReviewAction

def test_threading_creates_extended_part():
    """
    Verifies that adding comments to a clean doc creates commentsExtended.xml,
    which is required for visible threading in modern Word.
    """
    doc = Document()
    doc.add_paragraph("Content")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    # 1. Add Root Comment
    engine = RedlineEngine(stream)
    # Note: Must change text so the edit is not skipped as no-op
    engine.apply_edits([DocumentEdit(target_text="Content", new_text="Content Modified", comment="Root")])
    stream_1 = engine.save_to_stream()
    
    # Check if Extended part exists immediately
    doc_1 = Document(stream_1)
    extended_rel = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
    has_extended = any(rel.reltype == extended_rel for rel in doc_1.part.rels.values())
    assert has_extended, "commentsExtended.xml should be created with the first comment"
    
    # 2. Add Reply
    engine2 = RedlineEngine(stream_1)
    # Get root ID
    root_id = list(engine2.comments_manager.extract_comments_data().keys())[0]
    
    engine2.apply_review_actions([ReviewAction(action="REPLY", target_id=f"Com:{root_id}", text="Reply")])
    stream_2 = engine2.save_to_stream()
    
    # 3. Inspect XML for Threading
    doc_2 = Document(stream_2)
    extended_part = None
    for rel in doc_2.part.rels.values():
        if rel.reltype == extended_rel:
            extended_part = rel.target_part
            break
            
    xml = extended_part.blob.decode("utf-8")
    print(xml)
    
    # Should contain paraIdParent linking
    assert "w15:paraIdParent" in xml