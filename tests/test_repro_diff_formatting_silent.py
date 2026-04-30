def test_diff_engine_ignores_formatting_changes(tmp_path):
    """
    Reproduces a bug where `diff_docx_files` with compare_clean=True
    fails to detect formatting-only changes (e.g., Bold to Italic).
    """
    import docx

    d1_path = tmp_path / "format_1_pytest.docx"
    d2_path = tmp_path / "format_2_pytest.docx"

    d1 = docx.Document()
    p1 = d1.add_paragraph()
    r1 = p1.add_run("Silent Change")
    r1.bold = True
    d1.save(str(d1_path))

    d2 = docx.Document()
    p2 = d2.add_paragraph()
    r2 = p2.add_run("Silent Change")
    r2.italic = True
    d2.save(str(d2_path))

    import subprocess

    diff_output = subprocess.check_output(["uv", "run", "adeu", "diff", str(d1_path), str(d2_path)]).decode("utf-8")

    assert "[~]" in diff_output or "Found 1 changes:" in diff_output, (
        "Diff engine completely missed the bold -> italic formatting change."
    )
