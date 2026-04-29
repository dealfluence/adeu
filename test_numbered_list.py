import io

from docx import Document
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_visible_runs

doc = Document()
doc.add_paragraph("Some text.")
stream = io.BytesIO()
doc.save(stream)
stream.seek(0)

edit = ModifyText(target_text="Some text.", new_text="Some text.\n\n1. Numbered Item")
engine = RedlineEngine(stream)
engine.apply_edits([edit])
doc_result = Document(engine.save_to_stream())

p_new = doc_result.paragraphs[1]
visible = "".join(r.text for r in get_visible_runs(p_new))
print("Visible:", repr(visible))

has_numPr = False
if p_new._element.pPr is not None:
    if p_new._element.pPr.find(qn("w:numPr")) is not None:
        has_numPr = True

print("Has List Formatting:", has_numPr)
