import io

from docx import Document

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_visible_runs


def test_paragraph_merge_on_newline_deletion():
    """
    Test Case: When replacing text that spans a paragraph boundary (`\\n\\n`)
    with text that does not have a paragraph boundary, the engine should
    merge the two paragraphs into one.
    """
    doc = Document()
    doc.add_paragraph("Paragraph 1 end.")
    doc.add_paragraph("Paragraph 2 start.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(target_text="1 end.\n\nParagraph 2", new_text="1 end. Paragraph 2")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])
    assert applied == 1, "Edit should be applied"

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    # Initially there were 2 paragraphs. They should be merged into 1.
    assert len(doc_result.paragraphs) == 1, "Paragraphs were not merged"

    visible_text = "".join(r.text for r in get_visible_runs(doc_result.paragraphs[0]))
    assert "Paragraph 1 end. Paragraph 2 start." in visible_text
