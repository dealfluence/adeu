# FILE: tests/test_document_settings_warnings.py
"""
Tests for the privacy-flag warning emitted in the Structural Appendix when
word/settings.xml contains removePersonalInformation or removeDateAndTime.

These flags cause Microsoft Word to silently strip attribution from tracked
changes and comments on next save. The appendix surfaces this at read time so
the agent can decide how to proceed.
"""

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml import parse_xml

from adeu.domain import build_structural_appendix, extract_document_settings_warnings

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
SETTINGS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"


def _install_settings_part(doc, settings_xml_bytes: bytes) -> None:
    """
    Replaces (or installs) word/settings.xml on the given document with the
    provided XML payload. Self-contained so each test sees exactly the XML it
    declares — no dependence on whatever python-docx ships by default.
    """
    pkg = doc.part.package

    # Remove any existing settings part + relationship so we don't end up with
    # a stale Part still referenced from the main document.
    partname = PackURI("/word/settings.xml")

    # Sever existing relationship from the main document part.
    rels_to_drop = [
        rId
        for rId, rel in doc.part.rels.items()
        if not rel.is_external and getattr(rel.target_part, "partname", None) == partname
    ]
    for rId in rels_to_drop:
        del doc.part.rels[rId]

    # Drop from package parts list.
    if hasattr(pkg, "_parts") and isinstance(pkg._parts, list):
        pkg._parts[:] = [p for p in pkg._parts if p.partname != partname]
    elif hasattr(pkg, "parts") and isinstance(pkg.parts, list):
        pkg.parts[:] = [p for p in pkg.parts if p.partname != partname]

    # Install the new part.
    new_part = XmlPart(partname, SETTINGS_CT, parse_xml(settings_xml_bytes), pkg)
    pkg.parts.append(new_part)
    doc.part.relate_to(new_part, RT.SETTINGS)


def _make_settings_xml(inner: str) -> bytes:
    return (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:settings xmlns:w="{W_NS}">{inner}</w:settings>'
    ).encode("utf-8")


def _make_doc_with_settings(inner_settings_xml: str):
    doc = Document()
    doc.add_paragraph("Body content.")
    _install_settings_part(doc, _make_settings_xml(inner_settings_xml))
    return doc


# ---------------------------------------------------------------------------
# Test 1: Both flags enabled produces both warnings.
# ---------------------------------------------------------------------------
def test_both_privacy_flags_produce_both_warnings():
    doc = _make_doc_with_settings('<w:removePersonalInformation w:val="true"/><w:removeDateAndTime w:val="true"/>')

    warnings = extract_document_settings_warnings(doc)
    assert len(warnings) == 2

    appendix = build_structural_appendix(doc, "Body content.")

    assert "## Document Settings" in appendix
    assert "[Warning]" in appendix
    assert "`removePersonalInformation`" in appendix
    assert "`removeDateAndTime`" in appendix
    # Both should be described as enabled.
    assert appendix.count("is enabled in word/settings.xml") == 2


# ---------------------------------------------------------------------------
# Test 2: No w:val attribute → defaults to enabled (OOXML rule).
# ---------------------------------------------------------------------------
def test_no_val_attribute_defaults_to_enabled():
    doc = _make_doc_with_settings("<w:removePersonalInformation/>")

    warnings = extract_document_settings_warnings(doc)
    assert len(warnings) == 1
    assert "`removePersonalInformation`" in warnings[0]
    assert "is enabled" in warnings[0]

    appendix = build_structural_appendix(doc, "Body content.")
    assert "## Document Settings" in appendix
    assert "`removePersonalInformation`" in appendix


# ---------------------------------------------------------------------------
# Test 3: Explicitly-disabled values suppress the warning.
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("disabled_val", ["0", "false", "off", "FALSE", "False", "OFF"])
def test_explicitly_disabled_values_suppress_warning(disabled_val):
    doc = _make_doc_with_settings(
        f'<w:removePersonalInformation w:val="{disabled_val}"/><w:removeDateAndTime w:val="{disabled_val}"/>'
    )

    warnings = extract_document_settings_warnings(doc)
    assert warnings == [], f"Disabled value '{disabled_val}' should yield no warnings"

    appendix = build_structural_appendix(doc, "Body content.")
    assert "## Document Settings" not in appendix
    assert "removePersonalInformation" not in appendix
    assert "removeDateAndTime" not in appendix


# ---------------------------------------------------------------------------
# Test 4: Settings.xml present but no privacy flags → no section.
# ---------------------------------------------------------------------------
def test_no_privacy_flags_no_section():
    doc = _make_doc_with_settings('<w:zoom w:percent="100"/>')

    warnings = extract_document_settings_warnings(doc)
    assert warnings == []

    appendix = build_structural_appendix(doc, "Body content.")
    assert "## Document Settings" not in appendix


# ---------------------------------------------------------------------------
# Test 5: Section ordering — Document Settings appears before Defined Terms.
# ---------------------------------------------------------------------------
def test_document_settings_appears_before_defined_terms():
    doc = Document()
    # Defined term to populate ## Defined Terms.
    doc.add_paragraph('"Agreement" means this contract.')
    doc.add_paragraph("The Agreement is binding.")
    _install_settings_part(
        doc,
        _make_settings_xml('<w:removePersonalInformation w:val="true"/>'),
    )

    base_text = "\n".join(p.text for p in doc.paragraphs)
    appendix = build_structural_appendix(doc, base_text)

    assert "## Document Settings" in appendix
    assert "## Defined Terms" in appendix

    settings_idx = appendix.index("## Document Settings")
    defs_idx = appendix.index("## Defined Terms")
    assert settings_idx < defs_idx, "## Document Settings must appear before ## Defined Terms in the appendix"


# ---------------------------------------------------------------------------
# Bonus: settings.xml entirely absent → no warnings, no section.
# ---------------------------------------------------------------------------
def test_no_settings_part_at_all():
    """
    Defensive: an in-memory document we never gave a settings part to should
    still extract cleanly (empty list, no exception).
    """
    doc = Document()
    doc.add_paragraph("Body content.")

    # Strip any settings part python-docx may have implicitly created.
    pkg = doc.part.package
    settings_partname = PackURI("/word/settings.xml")
    if hasattr(pkg, "_parts") and isinstance(pkg._parts, list):
        pkg._parts[:] = [p for p in pkg._parts if p.partname != settings_partname]

    # Whether or not settings.xml was present, the call must return [] without raising.
    warnings = extract_document_settings_warnings(doc)
    assert warnings == []
