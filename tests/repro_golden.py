import io
import os

from docx import Document

from adeu.models import DocumentEdit, ReviewAction
from adeu.redline.engine import RedlineEngine


def generate_golden_replica(output_path: str):
    """
    Generates a file with the exact structure of golden.docx using Adeu engine.
    Structure:
    1. Base paragraph
    2. Root comment
    3. Reply to root
    4. Second Reply (Triad)
    """
    # Use initial.docx as base to avoid huge style XMLs from default template
    base_path = "tests/fixtures/initial.docx"
    if os.path.exists(base_path):
        doc = Document(base_path)
        # Clear existing paragraphs to start fresh
        doc._body.clear_content()
    else:
        doc = Document()

    doc.add_paragraph("Original placeholder text")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Add Root Comment
    engine = RedlineEngine(stream, author="Mikko Korpela")
    # Force edit to attach comment
    edit = DocumentEdit(
        target_text="Original placeholder text",
        new_text="Start of comment thread",
        comment="Start of comment thread",  # Comment body same as text in golden example
    )
    engine.apply_edits([edit])

    # Get ID
    comments = engine.comments_manager.extract_comments_data()
    root_id = list(comments.keys())[0]

    # 2. Add First Reply
    # Action: REPLY
    action1 = ReviewAction(action="REPLY", target_id=f"Com:{root_id}", text="Second comment")
    engine.apply_review_actions([action1])

    # 3. Add Second Reply
    # In golden.docx, it seems linear or threaded?
    # Usually replying to the root creates a thread.
    action2 = ReviewAction(action="REPLY", target_id=f"Com:{root_id}", text="Third comment in the thread")
    engine.apply_review_actions([action2])

    # Save
    with open(output_path, "wb") as f:
        f.write(engine.save_to_stream().getvalue())

    print(f"Generated replica at: {output_path}")


if __name__ == "__main__":
    generate_golden_replica("tests/fixtures/replica_golden.docx")
