import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_multiline_insert_does_not_create_nested_paragraphs():
    """
    Validates Issue Fix: Paragraph merge during wholesale replacement.
    Simulates anchoring a multiline insert onto a run that is already wrapped
    inside a w:ins tag to ensure we don't accidentally create <w:p><w:p> nested structures.
    """
    # 1. Setup Document
    doc = Document()
    doc.add_paragraph()
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Inject a mocked "Word-style" inline insertion
    engine = RedlineEngine(stream, author="TestAuthor")

    ins_tag = OxmlElement("w:ins")
    ins_tag.set(qn("w:id"), "99")

    r_tag = OxmlElement("w:r")
    t_tag = OxmlElement("w:t")
    t_tag.text = "10. Force Majeure"

    r_tag.append(t_tag)
    ins_tag.append(r_tag)

    p_element = engine.doc.paragraphs[0]._element
    p_element.append(ins_tag)

    # 3. Act: Apply a multiline text edit targeting the tracked text
    # This will trigger track_delete_run (splitting the ins) and track_insert
    edit = ModifyText(
        target_text="10. Force Majeure",
        new_text="10. Force Majeure\n\n11. Entire Agreement",
    )

    applied, _ = engine.apply_edits([edit])
    assert applied == 1, "Edit should apply successfully"

    # 4. Assert: Validate XML structure
    doc_xml = engine.doc.element.xml

    # Word strictly forbids a w:p element being a direct child of another w:p.
    # The new paragraph (11. Entire Agreement) must be a sibling to the original paragraph.
    nested_p_check = engine.doc.element.xpath("//w:p//w:p")
    assert len(nested_p_check) == 0, "FATAL: Nested <w:p> tags detected. Word will merge these!"

    # Ensure the new paragraph was successfully inserted
    assert "11. Entire Agreement" in doc_xml


def test_val_obs_new_7_paragraph_break_tracking():
    """
    VAL-OBS-NEW-7: When a multi-line string is inserted, the paragraph break
    itself must be tracked inside the <w:pPr><w:rPr> of the newly created paragraph.
    """
    doc = Document()
    doc.add_paragraph("First paragraph")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="TestAuthor")
    edit = ModifyText(target_text="paragraph", new_text="paragraph\n\nSecond paragraph")

    engine.apply_edits([edit])

    # Find the newly created paragraph
    p_elements = engine.doc.element.xpath("//w:p")
    assert len(p_elements) == 2, "Should have exactly 2 paragraphs"

    # Assert the break is tracked: <w:pPr><w:rPr><w:ins/></w:rPr></w:pPr>
    new_p = p_elements[1]
    ins_marker = new_p.xpath("./w:pPr/w:rPr/w:ins")
    assert len(ins_marker) > 0, "Paragraph break must be tracked with an <w:ins> inside <w:pPr>"
