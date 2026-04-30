import io

import docx

from adeu import ModifyText, RedlineEngine


def test_batch_engine_formatting_removal_fails():
    """
    Reproduces a bug where the Batch Engine fails to remove formatting.
    If the target_text has formatting (like **text**), but the new_text
    is plain ('text'), the inserted run inherits the old formatting
    instead of dropping it.
    """
    d = docx.Document()
    p = d.add_paragraph("Body ")
    r = p.add_run("text")
    r.bold = True
    p.add_run(" here.")
    stream = io.BytesIO()
    d.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA")
    engine.apply_edits([ModifyText(target_text="Body **text** here.", new_text="Body text here.")])

    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    d2 = docx.Document(out_stream)

    # We look for the inserted run (w:ins) in the XML.
    xml = d2.paragraphs[0]._element.xml
    assert "<w:b/>" in xml.split("<w:ins")[1], "Bug fixed? The bold tag was properly stripped from the insertion."
