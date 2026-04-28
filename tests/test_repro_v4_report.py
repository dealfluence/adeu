import io

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine
from adeu.utils.docx import get_visible_runs


def test_val_crit_7_header_acceptance_and_namespace():
    """
    VAL-CRIT-7 & VAL-OBS-1B:
    1. Header edits must successfully accept.
    2. Header xml parts must have w16du declared at the root.
    """
    doc = Document()
    doc.add_paragraph("Body text.")

    # Add a header
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "CONFIDENTIAL DRAFT"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Apply Redline to Header
    engine = RedlineEngine(stream)
    engine.apply_edits([ModifyText(target_text="CONFIDENTIAL DRAFT", new_text="FINAL EXECUTED")])

    stream_edited = engine.save_to_stream()
    doc_edited = Document(stream_edited)

    # Verify edit landed
    header_xml_edited = doc_edited.sections[0].header._element.xml
    assert "w:ins" in header_xml_edited

    # VAL-OBS-1B check: The namespace should be at the root, NOT injected per-element as ns0
    assert 'xmlns:w16du="http://schemas.microsoft.com/office/word/2023/wordml/word16du"' in header_xml_edited
    assert "xmlns:ns0" not in header_xml_edited

    # 2. Accept Changes
    engine2 = RedlineEngine(stream_edited)
    engine2.accept_all_revisions()

    stream_accepted = engine2.save_to_stream()
    doc_accepted = Document(stream_accepted)

    # VAL-CRIT-7 check: The header should have no tracking markup left
    header_xml_accepted = doc_accepted.sections[0].header._element.xml
    assert "w:ins" not in header_xml_accepted, "Header acceptance failed: w:ins still present"
    assert "w:del" not in header_xml_accepted, "Header acceptance failed: w:del still present"
    assert "FINAL EXECUTED" in doc_accepted.sections[0].header.paragraphs[0].text


def test_val_obs_8_bullet_markdown_leak():
    """
    VAL-OBS-8:
    When inserting text into a list-styled paragraph, markdown bullets (*, -, 1.)
    should be stripped from the OOXML text because the <w:numPr> provides the bullet natively.
    """
    doc = Document()
    p = doc.add_paragraph("First Item")

    # Manually inject <w:numPr> to simulate a list item without relying on local Word templates
    pPr = OxmlElement("w:pPr")
    numPr = OxmlElement("w:numPr")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(numId)
    pPr.append(numPr)
    p._element.insert(0, pPr)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    edit = ModifyText(target_text="First Item", new_text="First Item\n* Second Item")
    engine.apply_edits([edit])

    stream_edited = engine.save_to_stream()
    doc_edited = Document(stream_edited)

    # The second paragraph should have inherited the <w:numPr>
    p2 = doc_edited.paragraphs[1]
    assert p2._element.find(f".//{qn('w:numPr')}") is not None

    # The literal text should NOT contain the asterisk
    visible_text = "".join(r.text for r in get_visible_runs(p2))
    assert visible_text == "Second Item", f"Markdown bullet leaked into OOXML text: {visible_text}"


def test_val_obs_13_nested_redline_fragmentation():
    """
    VAL-OBS-13:
    A Modification targeting text strictly inside an active <w:ins> by a DIFFERENT author
    should trigger a Strict Refusal via BatchValidationError to prevent nested tag fragmentation.
    """
    doc = Document()
    doc.add_paragraph("Original baseline.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Author A makes an insertion
    engine_a = RedlineEngine(stream, author="Author A")
    engine_a.apply_edits([ModifyText(target_text="Original baseline.", new_text="Original baseline. Inserted by A.")])
    stream_a = engine_a.save_to_stream()

    # 2. Author B tries to modify Author A's pending insertion
    engine_b = RedlineEngine(stream_a, author="Author B")
    edit = ModifyText(target_text="Inserted by A", new_text="Modified by B")

    # We expect the validator to catch this and raise
    with pytest.raises(BatchValidationError) as excinfo:
        engine_b.process_batch([edit])

    assert "targets an active insertion from another author" in str(excinfo.value)


def test_val_obs_4_appendix_over_rejection():
    """
    VAL-OBS-4:
    The validator shouldn't reject a body edit just because the target_text
    happens to coincidentally appear in the structural appendix.
    """
    doc = Document()
    p = doc.add_paragraph("1. Definitions")

    # Inject a bookmark to force it to show up in the appendix
    bstart = OxmlElement("w:bookmarkStart")
    bstart.set(qn("w:name"), "_Ref123")
    bstart.set(qn("w:id"), "0")
    p._element.insert(0, bstart)
    bend = OxmlElement("w:bookmarkEnd")
    bend.set(qn("w:id"), "0")
    p._element.append(bend)

    doc.add_paragraph("Some body text referencing it.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    # The appendix text currently contains: `Anchored to: "1. Definitions"`
    # So "1. Definitions" exists both before and after the appendix boundary.
    edit = ModifyText(target_text="1. Definitions", new_text="Article 1. Definitions")

    # This should NOT return errors (over-rejection)
    errors = engine.validate_edits([edit])
    assert not errors, f"Validator over-rejected body edit due to appendix collision: {errors}"
