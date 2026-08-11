import asyncio
from unittest.mock import patch

import pytest
from docx import Document
from fastmcp.exceptions import ToolError

from adeu.cli import _main_impl
from adeu.mcp_components.tools.document import apply_text_revision


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


@pytest.fixture
def sample_docx(tmp_path) -> str:
    """Creates a basic DOCX file for testing."""
    doc = Document()
    doc.add_paragraph("This is the original paragraph one of the document.")
    doc.add_paragraph("This is paragraph two, containing more text for testing purposes.")
    doc.add_paragraph("And paragraph three concludes the baseline document content.")
    path = tmp_path / "sample.docx"
    doc.save(path)
    return str(path)


def test_apply_text_revision_produces_tracked_changes(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = str(tmp_path / "output.docx")
    revised = (
        "This is the revised paragraph one of the document.\n\n"
        "This is paragraph two, containing more text for testing purposes.\n\n"
        "And paragraph three concludes the baseline document content."
    )

    result = asyncio.run(
        apply_text_revision(
            file_path=sample_docx,
            revised_text=revised,
            ctx=ctx,
            output_path=out_path,
            author="TestAuthor",
        )
    )

    assert out_path in result or "Saved" in result or "Batch complete" in result or "Applied" in result
    doc = Document(out_path)
    xml = doc.element.xml
    assert "w:del" in xml or "w:ins" in xml
    assert "revised" in xml


def test_apply_text_revision_refuses_major_deletion_without_flag(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = str(tmp_path / "output.docx")

    # Major character deletion (>30% of characters deleted)
    short_revised = "This is short."

    with pytest.raises(ToolError) as exc_info:
        asyncio.run(
            apply_text_revision(
                file_path=sample_docx,
                revised_text=short_revised,
                ctx=ctx,
                output_path=out_path,
                allow_major_deletions=False,
            )
        )
    assert "major" in str(exc_info.value).lower() or "deletion" in str(exc_info.value).lower()

    # Allowed with flag
    result = asyncio.run(
        apply_text_revision(
            file_path=sample_docx,
            revised_text=short_revised,
            ctx=ctx,
            output_path=out_path,
            allow_major_deletions=True,
        )
    )
    assert out_path in result or "Saved" in result or "Applied" in result


def test_apply_text_revision_refuses_criticmarkup_input(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = str(tmp_path / "output.docx")
    critic_revised = "This is original text {++with inserted text++} and {--deleted text--}."

    with pytest.raises(ToolError) as exc_info:
        asyncio.run(
            apply_text_revision(
                file_path=sample_docx,
                revised_text=critic_revised,
                ctx=ctx,
                output_path=out_path,
            )
        )
    assert "criticmarkup" in str(exc_info.value).lower()


def test_apply_text_revision_verification_failure_writes_unverified_sibling(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = tmp_path / "output.docx"
    unverified_sibling = tmp_path / "output.unverified.docx"

    revised = (
        "This is the revised paragraph one of the document.\n\n"
        "This is paragraph two, containing more text for testing purposes.\n\n"
        "And paragraph three concludes the baseline document content."
    )

    # Force verification failure by patching text extraction on output verification
    from adeu import text_revision

    def bad_extract(doc):
        return "Mismatch text that does not match revised"

    with patch.object(text_revision, "_extract_clean_text_from_doc", side_effect=bad_extract):
        with pytest.raises(ToolError) as exc_info:
            asyncio.run(
                apply_text_revision(
                    file_path=sample_docx,
                    revised_text=revised,
                    ctx=ctx,
                    output_path=str(out_path),
                )
            )

    assert not out_path.exists()
    assert unverified_sibling.exists()
    assert "verification" in str(exc_info.value).lower() or "unverified" in str(exc_info.value).lower()


def test_cli_text_apply_still_behaves_identically(sample_docx, tmp_path):
    revised_file = tmp_path / "revised.txt"
    revised_text = (
        "This is the revised paragraph one of the document.\n\n"
        "This is paragraph two, containing more text for testing purposes.\n\n"
        "And paragraph three concludes the baseline document content."
    )
    revised_file.write_text(revised_text, encoding="utf-8")

    out_docx = tmp_path / "cli_out.docx"

    import sys

    test_args = ["adeu", "apply", sample_docx, str(revised_file), "-o", str(out_docx)]
    with patch.object(sys, "argv", test_args):
        try:
            _main_impl()
        except SystemExit as e:
            assert e.code in (0, None)

    assert out_docx.exists()
    doc = Document(str(out_docx))
    xml = doc.element.xml
    assert "revised" in xml
