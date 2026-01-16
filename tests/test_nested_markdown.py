# FILE: tests/test_nested_markdown.py

from io import BytesIO

from docx import Document

from adeu.redline.engine import RedlineEngine


def _parse_and_check(engine, text, expected_tokens):
    """
    Helper to run the engine's internal parser and check output structure.
    expected_tokens: list of (text, props_dict)
    """
    tokens = engine._parse_inline_markdown(text)

    # Debug print if failure
    if tokens != expected_tokens:
        print(f"\nFailed: '{text}'")
        print("Expected:", expected_tokens)
        print("Actual:  ", tokens)

    assert tokens == expected_tokens


def test_recursive_nested_styles():
    doc = Document()
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    engine = RedlineEngine(stream)

    # Case 1: Simple Bold inside Italic
    # "Outer _Inner **Deep** Inner_ Outer"
    text = "A _B **C** B_ A"
    expected = [
        ("A ", {}),
        ("B ", {"italic": True}),
        ("C", {"italic": True, "bold": True}),
        (" B", {"italic": True}),
        (" A", {}),
    ]
    _parse_and_check(engine, text, expected)


def test_complex_llm_failure_case():
    """
    Reproduces the user reported failure:
    _Either Party... **(i)** ..._
    """
    doc = Document()
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    engine = RedlineEngine(stream)

    text = "_Start **Bold** End_"
    expected = [
        ("Start ", {"italic": True}),
        ("Bold", {"italic": True, "bold": True}),
        (" End", {"italic": True}),
    ]
    _parse_and_check(engine, text, expected)


def test_sequential_tags():
    doc = Document()
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    engine = RedlineEngine(stream)

    text = "**Bold**_Italic_"
    expected = [("Bold", {"bold": True}), ("Italic", {"italic": True})]
    _parse_and_check(engine, text, expected)
