import io
import pytest
from docx import Document

@pytest.fixture
def simple_docx_stream():
    """Returns a BytesIO stream containing a simple DOCX."""
    doc = Document()
    doc.add_heading("Contract Agreement", 0)
    doc.add_paragraph("This is a simple contract.")
    doc.add_paragraph("The party of the first part shall be known as the Seller.")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream