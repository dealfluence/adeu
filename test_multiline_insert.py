import io
import logging

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine

logging.basicConfig(level=logging.DEBUG)

doc = Document()
doc.add_paragraph()
stream = io.BytesIO()
doc.save(stream)
stream.seek(0)

engine = RedlineEngine(stream, author="TestAuthor")

ins_tag = OxmlElement("w:ins")
ins_tag.set(qn("w:id"), "99")

r_tag = OxmlElement("w:r")
t_tag = OxmlElement("w:t")
t_tag.text = "10. Force Majeure"

r_tag.append(t_tag)
ins_tag.append(r_tag)

p_element = engine.doc.paragraphs[0]._element
p_element.append(ins_tag)

edit = ModifyText(
    target_text="10. Force Majeure",
    new_text="10. Force Majeure\n\n11. Entire Agreement",
)

applied, _ = engine.apply_edits([edit])
print(engine.doc.element.xml)
