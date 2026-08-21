"""CC-5 — `set_field` (A4).

Resolution and fill semantics for the `set_field` change type. The XML
assertions read the SAVED package, not the in-memory tree: the whole point of
a fill is what Word opens, and the two have diverged before.
"""

import pytest

from adeu.fields import FieldResolutionError, collect_fields, resolve_field
from tests.cc_fixture import cc_fixture_bytes


@pytest.fixture(scope="module")
def entries():
    import io

    from docx import Document

    from adeu.ingest import _extract_text_from_doc

    doc = Document(io.BytesIO(cc_fixture_bytes()))
    text = _extract_text_from_doc(doc, clean_view=False, include_appendix=False)
    if isinstance(text, tuple):
        text = text[0]
    return collect_fields(doc, text, None)


class TestA42Resolution:
    """A4.2 — field resolution order and ambiguity."""

    def test_resolves_by_cc_ordinal(self, entries):
        hits = resolve_field(entries, "CC:2")
        assert [e.ordinal for e in hits] == [2]

    def test_resolves_by_exact_tag(self, entries):
        tagged = next(e for e in entries if e.tag)
        hits = resolve_field(entries, tagged.tag)
        assert tagged.ordinal in [e.ordinal for e in hits]

    def test_resolves_by_exact_alias(self, entries):
        tags = {e.tag for e in entries if e.tag}
        aliased = next(e for e in entries if e.alias and e.alias not in tags)
        hits = resolve_field(entries, aliased.alias)
        assert [e.ordinal for e in hits] == [aliased.ordinal]

    def test_a_tag_beats_an_alias_of_the_same_string(self, entries):
        """Resolution order is tag before alias, so a string that is one
        control's tag and another's alias resolves to the tagged one."""
        from dataclasses import replace

        tagged = replace(entries[0], ordinal=201, tag="shared_name", alias=None)
        aliased = replace(entries[0], ordinal=202, tag=None, alias="shared_name")
        hits = resolve_field([aliased, tagged], "shared_name")
        assert [e.ordinal for e in hits] == [201]

    def test_ordinal_wins_over_a_tag_that_looks_like_one(self, entries):
        """A document may legally tag a control `CC:2`. The published id wins.

        Otherwise the addressing scheme this engine advertises could be
        shadowed by the document it is addressing.
        """
        from dataclasses import replace

        decoy = replace(entries[-1], tag="CC:2")
        hits = resolve_field(list(entries) + [decoy], "CC:2")
        assert [e.ordinal for e in hits] == [2]

    def test_matching_is_case_sensitive(self, entries):
        tagged = next(e for e in entries if e.tag and e.tag.lower() != e.tag.upper())
        with pytest.raises(FieldResolutionError):
            resolve_field(entries, tagged.tag.upper() if tagged.tag.islower() else tagged.tag.lower())

    def test_unresolvable_field_teaches_the_alternatives(self, entries):
        with pytest.raises(FieldResolutionError) as exc:
            resolve_field(entries, "nonexistent")
        msg = str(exc.value)
        assert "nonexistent" in msg
        assert "mode='fields'" in msg
        assert any(e.tag and e.tag in msg for e in entries)

    def test_unknown_ordinal_names_the_id(self, entries):
        with pytest.raises(FieldResolutionError) as exc:
            resolve_field(entries, "CC:9999")
        assert "CC:9999" in str(exc.value)

    def test_empty_field_is_a_clean_error_not_a_crash(self, entries):
        """Clients drop primitive `required[]` entries, so this arrives empty."""
        with pytest.raises(FieldResolutionError) as exc:
            resolve_field(entries, "")
        assert "requires 'field'" in str(exc.value)


class TestA42Ambiguity:
    """A4.2 — a tag shared by several controls, the repeating-section reality."""

    @pytest.fixture
    def dupes(self, entries):
        from dataclasses import replace

        a = replace(entries[0], ordinal=101, tag="item_name")
        b = replace(entries[0], ordinal=102, tag="item_name")
        return list(entries) + [a, b]

    def test_strict_rejects_listing_the_candidates(self, dupes):
        with pytest.raises(FieldResolutionError) as exc:
            resolve_field(dupes, "item_name")
        msg = str(exc.value)
        assert "CC:101" in msg and "CC:102" in msg
        assert "match_mode" in msg

    def test_first_takes_document_order(self, dupes):
        assert [e.ordinal for e in resolve_field(dupes, "item_name", "first")] == [101]

    def test_all_fans_out(self, dupes):
        assert [e.ordinal for e in resolve_field(dupes, "item_name", "all")] == [101, 102]


# ---------------------------------------------------------------------------
# A4.1 — the end-to-end fill
# ---------------------------------------------------------------------------

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _saved_sdt(raw_bytes, ordinal):
    """The `w:sdt` element with the given CC ordinal, read from SAVED bytes."""
    import io
    import zipfile

    from lxml import etree

    with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
        tree = etree.fromstring(z.read("word/document.xml"))
    sdts = tree.iter(f"{W}sdt")
    return list(sdts)[ordinal - 1]


def _fill(field, value, **kw):
    """Run one `set_field` through the real batch pipeline; return saved bytes."""
    import io

    from adeu.models import SetField
    from adeu.redline.engine import RedlineEngine

    engine = RedlineEngine(io.BytesIO(cc_fixture_bytes()), author="Test Author")
    result = engine.process_batch([SetField(field=field, value=value, **kw)])
    return engine.save_to_stream().getvalue(), result


class TestA41FillEmptyTextField:
    """A4.1 — fill an empty text field by tag, checked in the saved package."""

    @pytest.fixture(scope="class")
    @classmethod
    def filled(cls):
        return _fill("client_name", "Acme Legal Services Ltd.")

    def test_the_edit_applies(self, filled):
        _raw, result = filled
        assert result["edits_applied"] == 1, result.get("skipped_details")

    def test_showing_placeholder_is_gone(self, filled):
        raw, _ = filled
        sdt = _saved_sdt(raw, 2)
        assert sdt.find(f".//{W}showingPlcHdr") is None

    def test_no_run_keeps_the_placeholder_style(self, filled):
        """CC-6(a): Word's own fill carries no rStyle PlaceholderText at all."""
        raw, _ = filled
        sdt = _saved_sdt(raw, 2)
        styles = [s.get(f"{W}val") for s in sdt.iter(f"{W}rStyle")]
        assert "PlaceholderText" not in styles

    def test_the_ghost_run_left_without_a_deletion(self, filled):
        """CONFIRMED CC-6(a): filling an empty control makes ONE revision.

        A `w:del` here would strike through prompt text the author never
        wrote, which is worse than wrong - it is libellous to the document.
        """
        raw, _ = filled
        sdt = _saved_sdt(raw, 2)
        assert sdt.find(f".//{W}del") is None

    def test_the_value_lands_inside_a_tracked_insertion(self, filled):
        raw, _ = filled
        sdt = _saved_sdt(raw, 2)
        ins = sdt.findall(f".//{W}ins")
        assert ins, "no tracked insertion in the filled control"
        text = "".join(t.text or "" for i in ins for t in i.iter(f"{W}t"))
        assert text == "Acme Legal Services Ltd."

    def test_the_insertion_is_attributed_to_the_acting_author(self, filled):
        raw, _ = filled
        sdt = _saved_sdt(raw, 2)
        authors = {i.get(f"{W}author") for i in sdt.iter(f"{W}ins")}
        assert authors == {"Test Author"}

    def test_the_report_names_the_field(self, filled):
        """spec-fields-ledger §6 — audit-trail symmetry with heading_path."""
        _raw, result = filled
        rep = result["edits"][0]
        assert rep["field"] == 'CC:2 "Client Name" (tag: client_name)'

    def test_raw_view_shows_the_insertion_inside_the_anchor_pair(self, filled):
        import io

        raw, _ = filled
        from adeu.ingest import extract_text_from_stream

        text = extract_text_from_stream(io.BytesIO(raw), clean_view=False, include_appendix=False)
        if isinstance(text, tuple):
            text = text[0]
        assert "{#cc:2}{++Acme Legal Services Ltd.++}" in text

    def test_clean_view_shows_the_value_as_settled_text(self, filled):
        import io

        raw, _ = filled
        from adeu.ingest import extract_text_from_stream

        text = extract_text_from_stream(io.BytesIO(raw), clean_view=True, include_appendix=False)
        if isinstance(text, tuple):
            text = text[0]
        assert "{#cc:2}Acme Legal Services Ltd.{#/cc:2}" in text


# ---------------------------------------------------------------------------
# Surfaces — the fill has to be reachable, not merely implemented
# ---------------------------------------------------------------------------


class TestSetFieldSurfaces:
    @staticmethod
    def _flat_props():
        from pydantic import TypeAdapter

        from adeu.models import FlatSchemaDocumentChange

        return TypeAdapter(list[FlatSchemaDocumentChange]).json_schema()["items"]["properties"]

    def test_field_and_value_publish_exactly_like_the_other_optional_strings(self):
        """Asserted as PARITY with `target_text`, not against a literal shape.

        Every optional property on the flat schema is a nullable string
        (`anyOf: [string, null]`), which is the convention this surface
        already publishes. Pinning a literal here would either freeze a
        second convention for the two new props or fail the moment the
        existing ones changed - so the assertion is that they are the same.
        """
        props = self._flat_props()
        shape = lambda d: {k: v for k, v in d.items() if k not in ("description", "title")}  # noqa: E731
        for name in ("field", "value"):
            assert shape(props[name]) == shape(props["target_text"]), name

    def test_neither_new_property_introduces_a_variant_union(self):
        """AI_CONTEXT §7a: the flat schema stays ONE object (no oneOf)."""
        from pydantic import TypeAdapter

        from adeu.models import FlatSchemaDocumentChange

        schema = TypeAdapter(list[FlatSchemaDocumentChange]).json_schema()
        assert "$defs" not in schema
        assert "oneOf" not in schema["items"]
        for name in ("field", "value"):
            assert "oneOf" not in schema["items"]["properties"][name]

    def test_set_field_is_in_the_published_type_enum(self):
        from pydantic import TypeAdapter

        from adeu.models import FlatSchemaDocumentChange

        schema = TypeAdapter(list[FlatSchemaDocumentChange]).json_schema()
        assert "set_field" in schema["items"]["properties"]["type"]["enum"]

    def test_a_missing_type_discriminator_is_inferred_from_field_plus_value(self):
        """Clients drop primitive `required[]` entries; `field` belongs to no
        other variant, so the pair is unambiguous."""
        from pydantic import TypeAdapter

        from adeu.models import BatchChanges

        changes = TypeAdapter(BatchChanges).validate_python([{"field": "client_name", "value": "X"}])
        assert changes[0].type == "set_field"

    def test_the_cli_strict_schema_accepts_set_field(self):
        from pydantic import TypeAdapter

        from adeu.models import StrictBatchChanges

        changes = TypeAdapter(StrictBatchChanges).validate_python(
            [{"type": "set_field", "field": "client_name", "value": "X"}]
        )
        assert changes[0].value == "X"

    def test_the_cli_strict_schema_still_demands_an_explicit_type(self):
        """Surface-specific requiredness: inference is an MCP tolerance."""
        import pydantic
        from pydantic import TypeAdapter

        from adeu.models import StrictBatchChanges

        with pytest.raises(pydantic.ValidationError):
            TypeAdapter(StrictBatchChanges).validate_python([{"field": "a", "value": "b"}])

    def test_a_value_that_fabricates_an_anchor_is_refused(self):
        """CC-1e, reached through set_field's `value` rather than `new_text`.

        A hard batch failure rather than a skip: writing `{#cc:3}` into a
        document would fabricate a control that does not exist, and the
        transactional contract says such a batch touches nothing.
        """
        from adeu.redline.engine import BatchValidationError

        with pytest.raises(BatchValidationError) as exc:
            _fill("client_name", "Acme {#cc:3} Ltd.")
        assert "anchor" in str(exc.value).lower()


class TestSetFieldThroughTheCli:
    """The whole point of a skeleton: reachable from a real command."""

    def test_apply_fills_the_field_from_a_changes_file(self, tmp_path):
        import json
        import zipfile

        from lxml import etree

        from tests.utils import run_cli

        docx = tmp_path / "cc.docx"
        docx.write_bytes(cc_fixture_bytes())
        changes = tmp_path / "changes.json"
        changes.write_text(
            json.dumps([{"type": "set_field", "field": "client_name", "value": "Acme Ltd."}]),
            encoding="utf-8",
        )
        out = tmp_path / "out.docx"

        res = run_cli("apply", str(docx), str(changes), "-o", str(out))
        assert res.returncode == 0, res.stderr
        assert out.exists()

        with zipfile.ZipFile(out) as z:
            tree = etree.fromstring(z.read("word/document.xml"))
        inserted = ["".join(t.text or "" for t in ins.iter(f"{W}t")) for ins in tree.iter(f"{W}ins")]
        assert "Acme Ltd." in inserted

    def test_an_unresolvable_field_fails_the_command_with_the_available_list(self, tmp_path):
        import json

        from tests.utils import run_cli

        docx = tmp_path / "cc.docx"
        docx.write_bytes(cc_fixture_bytes())
        changes = tmp_path / "changes.json"
        changes.write_text(
            json.dumps([{"type": "set_field", "field": "nope", "value": "x"}]),
            encoding="utf-8",
        )
        res = run_cli("apply", str(docx), str(changes), "-o", str(tmp_path / "o.docx"))
        assert res.returncode != 0
        assert "nope" in (res.stdout + res.stderr)


# ---------------------------------------------------------------------------
# A4.11 / A4.7 — the refusals
# ---------------------------------------------------------------------------


def _expect_refusal(field, value, **kw):
    """Run a set_field expected to be rejected; return the combined message."""
    from adeu.redline.engine import BatchValidationError

    try:
        _raw, result = _fill(field, value, **kw)
    except BatchValidationError as exc:
        return str(exc)
    assert result["edits_applied"] == 0, f"expected a refusal, got {result['edits_applied']} applied"
    return " ".join(result.get("skipped_details") or []) + " " + str((result["edits"][0] or {}).get("error") or "")


class TestA411NonValueClasses:
    """A4.11 — classes that hold no single value refuse the whole operation.

    Writing to these is not merely unsupported, it is destructive: a group's
    "content" is other controls, and replacing it with a string would delete
    every field inside it.
    """

    def test_a_group_is_refused_and_named(self):
        msg = _expect_refusal("std_terms", "anything")
        assert "not a value-bearing field" in msg
        assert "group" in msg

    def test_a_repeating_section_is_refused_and_named(self):
        msg = _expect_refusal("deliverables", "x")
        assert "not a value-bearing field" in msg
        assert "repeating" in msg

    def test_the_group_refusal_points_at_the_nested_fields(self):
        """A refusal that does not say what to do instead just costs a turn."""
        msg = _expect_refusal("std_terms", "anything")
        assert "nested" in msg.lower() or "inside" in msg.lower()


class TestA47StructureRules:
    """A4.7 — a plain-text control cannot hold structure."""

    def test_paragraphs_are_refused_in_a_plain_text_control(self):
        msg = _expect_refusal("counterparty", "Line1\n\nLine2")
        assert "paragraph" in msg.lower()
        assert "CC:3" in msg or "counterparty" in msg

    def test_a_line_break_is_refused_without_multiline(self):
        msg = _expect_refusal("counterparty", "Line1\nLine2")
        assert "multiline" in msg.lower().replace(" ", "") or "line break" in msg.lower()

    def test_a_richtext_control_accepts_paragraphs(self):
        """The same value that a w:text control must refuse."""
        raw, result = _fill("indemnity", "Line1\n\nLine2")
        assert result["edits_applied"] == 1, result.get("skipped_details")
        sdt = _saved_sdt(raw, 1)
        inserted = "".join(t.text or "" for ins in sdt.iter(f"{W}ins") for t in ins.iter(f"{W}t"))
        assert "Line1" in inserted and "Line2" in inserted


class TestA46Checkbox:
    """A4.6 — the toggle, which has no anchor pair to edit."""

    @pytest.fixture(scope="class")
    @classmethod
    def unchecked(cls):
        return _fill("confidential", "false")

    def test_the_edit_applies(self, unchecked):
        _raw, result = unchecked
        assert result["edits_applied"] == 1, result.get("skipped_details")

    def test_the_state_attribute_flips(self, unchecked):
        raw, _ = unchecked
        sdt = _saved_sdt(raw, 6)
        W14 = "{http://schemas.microsoft.com/office/word/2010/wordml}"
        checked = sdt.find(f".//{W14}checked")
        assert checked is not None and checked.get(f"{W14}val") == "0"

    def test_the_attribute_flip_carries_no_revision_of_its_own(self, unchecked):
        """URL_RETARGET precedent: one act, one redline."""
        raw, _ = unchecked
        sdt = _saved_sdt(raw, 6)
        assert len(sdt.findall(f".//{W}ins")) == 1
        assert len(sdt.findall(f".//{W}del")) == 1

    def test_the_insertion_precedes_the_deletion(self, unchecked):
        """CC-6(b): Word's order, and visible - the projection reads document
        order, so the reverse renders the toggle backwards."""
        raw, _ = unchecked
        sdt = _saved_sdt(raw, 6)
        tags = [el.tag for el in sdt.iter() if el.tag in (f"{W}ins", f"{W}del")]
        assert tags == [f"{W}ins", f"{W}del"]

    def test_the_glyph_swaps_to_the_controls_own_unchecked_character(self, unchecked):
        raw, _ = unchecked
        sdt = _saved_sdt(raw, 6)
        ins_text = "".join(t.text or "" for i in sdt.iter(f"{W}ins") for t in i.iter(f"{W}t"))
        assert ins_text == "\u2610"

    def test_a_value_that_names_neither_state_is_refused(self):
        msg = _expect_refusal("confidential", "maybe")
        assert "checkbox" in msg and "true/false" in msg

    @pytest.mark.parametrize("value", ["true", "x", "[x]", "1", "yes", "checked"])
    def test_truthy_spellings_all_check_the_box(self, value):
        raw, result = _fill("confidential", value)
        assert result["edits_applied"] == 1
        W14 = "{http://schemas.microsoft.com/office/word/2010/wordml}"
        assert _saved_sdt(raw, 6).find(f".//{W14}checked").get(f"{W14}val") == "1"

    def test_raw_view_shows_the_pending_toggle(self, unchecked):
        import io

        raw, _ = unchecked
        from adeu.ingest import extract_text_from_stream

        text = extract_text_from_stream(io.BytesIO(raw), clean_view=False, include_appendix=False)
        if isinstance(text, tuple):
            text = text[0]
        line = next(ln for ln in text.split("\n") if "Confidentiality" in ln)
        assert "{++ ++}" in line and "{--x--}" in line

    def test_clean_view_shows_exactly_one_checkbox(self, unchecked):
        """The deleted half must not leave a second, permanently empty box."""
        import io

        raw, _ = unchecked
        from adeu.ingest import extract_text_from_stream

        text = extract_text_from_stream(io.BytesIO(raw), clean_view=True, include_appendix=False)
        if isinstance(text, tuple):
            text = text[0]
        line = next(ln for ln in text.split("\n") if "Confidentiality" in ln)
        assert line.count("[") == 1 and line.count("]") == 1
        assert line.endswith("[ ]")

    def test_the_two_projections_agree_on_the_toggled_document(self, unchecked):
        """Mapper and ingest are twins; CC-12 was them disagreeing."""
        import io

        from docx import Document

        from adeu.ingest import extract_text_from_stream
        from adeu.redline.mapper import DocumentMapper

        raw, _ = unchecked
        for clean in (False, True):
            text = extract_text_from_stream(io.BytesIO(raw), clean_view=clean, include_appendix=False)
            if isinstance(text, tuple):
                text = text[0]
            mapper = DocumentMapper(Document(io.BytesIO(raw)), clean_view=clean)
            mapper._build_map()
            line_i = next(ln for ln in text.split("\n") if "Confidentiality" in ln)
            line_m = next(ln for ln in mapper.full_text.split("\n") if "Confidentiality" in ln)
            assert line_i == line_m, f"clean_view={clean}"


class TestA43Dropdown:
    """A4.3 — G10 option validation."""

    def test_a_listed_display_text_replaces_the_current_selection(self):
        raw, result = _fill("governing_law", "British Columbia")
        assert result["edits_applied"] == 1, result.get("skipped_details")
        sdt = _saved_sdt(raw, 4)
        ins = "".join(t.text or "" for i in sdt.iter(f"{W}ins") for t in i.iter(f"{W}t"))
        dele = "".join(t.text or "" for d in sdt.iter(f"{W}del") for t in d.iter(f"{W}delText"))
        assert ins == "British Columbia" and dele == "Ontario"

    def test_the_last_value_attribute_follows_the_selection(self):
        raw, _ = _fill("governing_law", "British Columbia")
        sdt = _saved_sdt(raw, 4)
        node = sdt.find(f".//{W}dropDownList")
        assert node is not None and node.get(f"{W}lastValue") == "British Columbia"

    def test_a_machine_value_resolves_to_its_display_text(self):
        """The document must read like every other row, not like a database."""
        raw, result = _fill("governing_law", "BC")
        assert result["edits_applied"] == 1, result.get("skipped_details")
        sdt = _saved_sdt(raw, 4)
        ins = "".join(t.text or "" for i in sdt.iter(f"{W}ins") for t in i.iter(f"{W}t"))
        assert ins == "British Columbia"

    def test_an_unlisted_option_is_refused_with_the_list(self):
        msg = _expect_refusal("governing_law", "Manitoba")
        assert "Ontario" in msg and "British Columbia" in msg and "Federal" in msg


class TestA45Date:
    """A4.5 — G12 format handling and the silent fullDate sync."""

    def test_a_canonical_date_is_written_and_tracked(self):
        raw, result = _fill("effective_date", "2026-03-01")
        assert result["edits_applied"] == 1, result.get("skipped_details")
        sdt = _saved_sdt(raw, 5)
        ins = "".join(t.text or "" for i in sdt.iter(f"{W}ins") for t in i.iter(f"{W}t"))
        assert ins == "2026-03-01"

    def test_full_date_syncs_with_no_revision_of_its_own(self):
        raw, _ = _fill("effective_date", "2026-03-01")
        sdt = _saved_sdt(raw, 5)
        date_el = sdt.find(f".//{W}date")
        assert date_el is not None
        assert date_el.get(f"{W}fullDate") == "2026-03-01T00:00:00Z"
        assert len(sdt.findall(f".//{W}ins")) == 1

    def test_a_non_canonical_date_is_refused_naming_the_format(self):
        msg = _expect_refusal("effective_date", "01.03.2026")
        assert "YYYY-MM-DD" in msg

    def test_an_impossible_date_is_refused(self):
        """A regex-only check would accept 2026-02-30 and write it."""
        msg = _expect_refusal("effective_date", "2026-02-30")
        assert "date" in msg.lower()


class TestDateFormatRendering:
    """spec §5 — the `yyyy/MM/dd/d/M` subset, and honesty beyond it."""

    @pytest.mark.parametrize(
        "fmt,expected",
        [
            (None, "2026-03-01"),
            ("yyyy-MM-dd", "2026-03-01"),
            ("MM/dd/yyyy", "03/01/2026"),
            ("d/M/yyyy", "1/3/2026"),
        ],
    )
    def test_supported_formats_render(self, fmt, expected):
        from adeu.utils.field_write import render_date

        assert render_date((2026, 3, 1), fmt) == (expected, False)

    @pytest.mark.parametrize("fmt", ["dddd, MMMM d", "EEE dd", "MMM yyyy"])
    def test_name_bearing_formats_fall_back_and_say_so(self, fmt):
        """`dddd` is the day NAME. A substring test sees the supported `dd`
        inside it and rendered `dddd, MMMM d` as `0101, 0303 1`."""
        from adeu.utils.field_write import render_date

        text, unsupported = render_date((2026, 3, 1), fmt)
        assert unsupported is True
        assert text == "2026-03-01"


def _fill_variant(field, value, **fixture_kw):
    """As `_fill`, against a fixture variant."""
    import io

    from adeu.models import SetField
    from adeu.redline.engine import RedlineEngine

    engine = RedlineEngine(io.BytesIO(cc_fixture_bytes(**fixture_kw)), author="Test Author")
    result = engine.process_batch([SetField(field=field, value=value)])
    return engine.save_to_stream().getvalue(), result


class TestA48BoundControls:
    """A4.8 — the dual-write, and the disclosure when it cannot happen."""

    STORE = "<root><matter>M-2026-001</matter></root>"

    def test_a_resolving_binding_updates_the_store_silently(self):
        """CC-6(e): the store WINS ON OPEN. A content-only write to a bound
        control is destroyed the next time anyone opens the document."""
        import io
        import zipfile

        raw, result = _fill_variant("matter_number", "M-2026-002", custom_xml=self.STORE)
        assert result["edits_applied"] == 1, result.get("skipped_details")
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            store = z.read("customXml/item1.xml").decode("utf-8")
        assert "M-2026-002" in store
        assert "M-2026-001" not in store

    def test_the_content_change_is_still_tracked(self):
        raw, _ = _fill_variant("matter_number", "M-2026-002", custom_xml=self.STORE)
        sdt = _saved_sdt(raw, 10)
        ins = "".join(t.text or "" for i in sdt.iter(f"{W}ins") for t in i.iter(f"{W}t"))
        assert ins == "M-2026-002"

    def test_the_report_discloses_the_store_write(self):
        _raw, result = _fill_variant("matter_number", "M-2026-002", custom_xml=self.STORE)
        note = result["edits"][0].get("warning") or ""
        assert "bound store" in note and "/root[1]/matter[1]" in note

    def test_a_dangling_binding_applies_content_only_with_a_warning(self):
        """Dangling bindings exist in the wild - sanitize's scrub is one
        producer - so refusing would be worse than disclosing."""
        _raw, result = _fill("matter_number", "M-2026-002")
        assert result["edits_applied"] == 1, result.get("skipped_details")
        note = result["edits"][0].get("warning") or ""
        assert "WARNING" in note and "could not be resolved" in note

    def test_the_dangling_warning_says_what_will_happen_later(self):
        _raw, result = _fill("matter_number", "M-2026-002")
        note = result["edits"][0].get("warning") or ""
        assert "overwrite" in note.lower()


_TEMPORARY_BODY = (
    '<w:p><w:r><w:t xml:space="preserve">Prepared by </w:t></w:r>'
    "<w:sdt><w:sdtPr>"
    '<w:alias w:val="Preparer"/><w:tag w:val="preparer"/><w:id w:val="900"/>'
    "<w:temporary/><w:showingPlcHdr/><w:text/>"
    "</w:sdtPr><w:sdtContent>"
    '<w:r><w:rPr><w:rStyle w:val="PlaceholderText"/></w:rPr>'
    "<w:t>Click here to enter a name.</w:t></w:r>"
    "</w:sdtContent></w:sdt>"
    '<w:r><w:t xml:space="preserve">.</w:t></w:r></w:p>'
)


class TestA49TemporaryUnwrap:
    """A4.9 — Word dissolves a temporary control on any content edit."""

    def test_the_sdt_wrapper_is_gone_from_the_saved_xml(self):
        import io
        import zipfile

        from lxml import etree

        raw, result = _fill_variant("preparer", "Jane Roe", body_xml=_TEMPORARY_BODY)
        assert result["edits_applied"] == 1, result.get("skipped_details")
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            tree = etree.fromstring(z.read("word/document.xml"))
        assert not list(tree.iter(f"{W}sdt")), "the temporary control survived the edit"

    def test_the_inserted_text_stands_as_a_tracked_insertion(self):
        """The revision outlives the wrapper (CC-6(c)), so the value is still
        reviewable even though the control is gone."""
        import io
        import zipfile

        from lxml import etree

        raw, _ = _fill_variant("preparer", "Jane Roe", body_xml=_TEMPORARY_BODY)
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            tree = etree.fromstring(z.read("word/document.xml"))
        ins = "".join(t.text or "" for i in tree.iter(f"{W}ins") for t in i.iter(f"{W}t"))
        assert ins == "Jane Roe"

    def test_the_ledger_no_longer_lists_it(self):
        import io

        from docx import Document

        from adeu.fields import collect_fields
        from adeu.ingest import _extract_text_from_doc

        raw, _ = _fill_variant("preparer", "Jane Roe", body_xml=_TEMPORARY_BODY)
        doc = Document(io.BytesIO(raw))
        text = _extract_text_from_doc(doc, clean_view=False, include_appendix=False)
        if isinstance(text, tuple):
            text = text[0]
        assert collect_fields(doc, text, None) == []

    def test_the_report_discloses_the_unwrap(self):
        """The control vanishing is a surprise unless the report says so."""
        _raw, result = _fill_variant("preparer", "Jane Roe", body_xml=_TEMPORARY_BODY)
        note = result["edits"][0].get("warning") or ""
        assert "temporary" in note and "unwrapped" in note


class TestA410TextFirstParity:
    """A4.10 — typing between the anchors must equal calling set_field.

    Two routes to one outcome. They agree because they run the same fill
    code, not because someone compared them once.
    """

    @staticmethod
    def _text_first(tmp_path):
        import io

        from adeu.ingest import extract_text_from_stream
        from adeu.text_revision import apply_text_revision_core

        src = tmp_path / "cc.docx"
        src.write_bytes(cc_fixture_bytes())
        clean = extract_text_from_stream(io.BytesIO(src.read_bytes()), clean_view=True, include_appendix=False)
        if isinstance(clean, tuple):
            clean = clean[0]
        assert "{#cc:2}{#/cc:2}" in clean, "the empty pair is the edit surface A4.10 describes"
        revised = clean.replace("{#cc:2}{#/cc:2}", "{#cc:2}Acme Legal Services Ltd.{#/cc:2}")
        _res, out = apply_text_revision_core(src, revised, output_path=tmp_path / "out.docx", author="Test Author")
        return out.read_bytes()

    def test_the_value_lands_inside_the_control(self, tmp_path):
        """Without fill semantics this inserts BESIDE the field, leaving the
        control empty and the document reading `Acme Ltd.{#cc:2}{#/cc:2}`."""
        sdt = _saved_sdt(self._text_first(tmp_path), 2)
        ins = "".join(t.text or "" for i in sdt.iter(f"{W}ins") for t in i.iter(f"{W}t"))
        assert ins == "Acme Legal Services Ltd."

    def test_the_placeholder_state_is_cleared_exactly_as_set_field_clears_it(self, tmp_path):
        sdt = _saved_sdt(self._text_first(tmp_path), 2)
        assert sdt.find(f".//{W}showingPlcHdr") is None
        styles = [s.get(f"{W}val") for s in sdt.iter(f"{W}rStyle")]
        assert "PlaceholderText" not in styles

    def test_no_deletion_is_emitted_for_the_ghost_run(self, tmp_path):
        sdt = _saved_sdt(self._text_first(tmp_path), 2)
        assert sdt.find(f".//{W}del") is None

    def test_both_routes_agree_on_the_control_state(self, tmp_path):
        """The assertion A4.10 is actually about: same field, same result."""
        text_first = _saved_sdt(self._text_first(tmp_path), 2)
        explicit = _saved_sdt(_fill("client_name", "Acme Legal Services Ltd.")[0], 2)

        def state(sdt):
            return (
                sdt.find(f".//{W}showingPlcHdr") is None,
                sdt.find(f".//{W}del") is None,
                "".join(t.text or "" for i in sdt.iter(f"{W}ins") for t in i.iter(f"{W}t")),
                [s.get(f"{W}val") for s in sdt.iter(f"{W}rStyle") if s.get(f"{W}val") == "PlaceholderText"],
            )

        assert state(text_first) == state(explicit)
