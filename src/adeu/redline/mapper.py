from copy import deepcopy
from dataclasses import dataclass
from typing import List, Optional, Set, Tuple

import structlog
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from adeu.redline.comments import CommentsManager
from adeu.utils.docx import (
    DocxEvent,
    get_paragraph_prefix,
    get_run_style_markers,
    get_run_text,
    iter_document_parts,
    iter_paragraph_content,
)

logger = structlog.get_logger(__name__)


@dataclass
class TextSpan:
    start: int
    end: int
    text: str
    run: Optional[Run]
    paragraph: Paragraph
    ins_id: Optional[str] = None
    del_id: Optional[str] = None


class DocumentMapper:
    def __init__(self, doc: DocumentObject):
        self.doc = doc
        self.comments_mgr = CommentsManager(doc)
        self.comments_map = self.comments_mgr.extract_comments_data()
        self.full_text = ""
        self.spans: List[TextSpan] = []
        self._build_map()

    def _build_map(self):
        current_offset = 0
        self.spans = []
        self.full_text = ""

        for part in iter_document_parts(self.doc):
            # 1. Loose Paragraphs
            for p in part.paragraphs:
                prefix = get_paragraph_prefix(p)
                if prefix:
                    self._add_virtual_text(prefix, current_offset, p)
                    current_offset += len(prefix)

                current_offset = self._map_paragraph_content(p, current_offset)

                self._add_virtual_text("\n\n", current_offset, p)
                current_offset += 2

            # 2. Tables
            for table in part.tables:
                for row in table.rows:
                    row_has_content = False
                    for cell in row.cells:
                        cell_paras = list(cell.paragraphs)
                        # Check visible content (approximate)
                        cell_non_empty = any(p.runs for p in cell_paras)

                        if cell_non_empty:
                            if row_has_content:
                                self._add_virtual_text(" | ", current_offset, cell_paras[0])
                                current_offset += 3

                            row_has_content = True

                            for j, p in enumerate(cell_paras):
                                if j > 0:
                                    self._add_virtual_text("\n", current_offset, p)
                                    current_offset += 1

                                prefix = get_paragraph_prefix(p)
                                if prefix:
                                    self._add_virtual_text(prefix, current_offset, p)
                                    current_offset += len(prefix)

                                current_offset = self._map_paragraph_content(p, current_offset)

                    if row_has_content:
                        self._add_virtual_text("\n\n", current_offset, row.cells[0].paragraphs[0])
                        current_offset += 2

        if self.spans and self.spans[-1].text == "\n\n":
            self.spans.pop()
            self.full_text = self.full_text[:-2]

    def _map_paragraph_content(self, paragraph: Paragraph, start_offset: int) -> int:
        """
        Maps Runs to Spans, handling Flattened CriticMarkup generation.
        Matches logic in ingest.py _build_paragraph_text.
        """
        current = start_offset

        active_ids: set[str] = set()
        active_ins_event: Optional[DocxEvent] = None
        active_del_event: Optional[DocxEvent] = None

        # We need to buffer Run items because we might wrap them in virtual text
        # But we need to yield them one by one for the mapper.

        pending_spans: List[Tuple[str, str, Optional[Run], Optional[str], Optional[str]]] = []

        for item in iter_paragraph_content(paragraph):
            if isinstance(item, Run):
                # 1. Prefix
                prefix, suffix = get_run_style_markers(item)

                run_items: List[Tuple[str, str, Optional[Run]]] = []
                if prefix:
                    run_items.append(("virtual", prefix, None))

                text = get_run_text(item)
                if text:
                    run_items.append(("real", text, item))

                if suffix:
                    run_items.append(("virtual", suffix, None))

                # Attach context
                current_ins_id = active_ins_event.id if active_ins_event else None
                current_del_id = active_del_event.id if active_del_event else None

                for k, t, r in run_items:
                    # Tuple structure: (kind, text, run, ins_id, del_id)
                    pending_spans.append((k, t, r, current_ins_id, current_del_id))

            elif isinstance(item, DocxEvent):
                # Flush
                if pending_spans:
                    current = self._flush_spans(
                        pending_spans,
                        active_ids,
                        current,
                        paragraph,
                        active_ins_event,
                        active_del_event,
                    )
                    pending_spans = []

                # Update State
                if item.type == "start":
                    active_ids.add(item.id)
                elif item.type == "end":
                    if item.id in active_ids:
                        active_ids.remove(item.id)
                elif item.type == "ins_start":
                    active_ins_event = item
                elif item.type == "ins_end":
                    active_ins_event = None
                elif item.type == "del_start":
                    active_del_event = item
                elif item.type == "del_end":
                    active_del_event = None

        # Final Flush
        if pending_spans:
            current = self._flush_spans(
                pending_spans,
                active_ids,
                current,
                paragraph,
                active_ins_event,
                active_del_event,
            )

        return current

    def _flush_spans(
        self,
        items: List[Tuple],
        active_ids: Set[str],
        current_offset: int,
        paragraph: Paragraph,
        ins_event: Optional[DocxEvent],
        del_event: Optional[DocxEvent],
    ) -> int:
        # Determine if we are inside any block that requires wrapping
        # items[0] = (kind, text, run, ins_id, del_id)
        _, _, _, ins_id, del_id = items[0]

        is_wrapped = bool(active_ids) or bool(ins_id) or bool(del_id)

        # 1. Virtual Start
        if is_wrapped:
            if del_id:
                marker = "{--"
            elif ins_id:
                marker = "{++"
            else:
                marker = "{=="

            self._add_virtual_text(marker, current_offset, paragraph)
            current_offset += len(marker)

        # 2. Content
        for kind, text, run_obj, i_id, d_id in items:
            if kind == "virtual":
                self._add_virtual_text(text, current_offset, paragraph)
                current_offset += len(text)
            else:
                span = TextSpan(
                    start=current_offset,
                    end=current_offset + len(text),
                    text=text,
                    run=run_obj,
                    paragraph=paragraph,
                    ins_id=i_id,
                    del_id=d_id,
                )
                self.spans.append(span)
                self.full_text += text
                current_offset += len(text)

        # 3. Virtual End + Meta
        if is_wrapped:
            # Close with matching marker
            if del_id:
                marker = "--}"
            elif ins_id:
                marker = "++}"
            else:
                marker = "==}"

            self._add_virtual_text(marker, current_offset, paragraph)
            current_offset += len(marker)

            meta_content = self._build_meta_block(active_ids, ins_event, del_event)
            meta_block = f"{{>>{meta_content}<<}}"
            self._add_virtual_text(meta_block, current_offset, paragraph)
            current_offset += len(meta_block)

        return current_offset

    def _build_meta_block(
        self,
        active_ids: Set[str],
        ins_event: Optional[DocxEvent],
        del_event: Optional[DocxEvent],
    ) -> str:
        # Match logic in ingest.py _format_segment
        lines = []

        # 1. Ins/Del Metadata
        # (Since we flush on any event change, there is at most one active ins/del context in items)
        if ins_event:
            auth = ins_event.author or "Unknown"
            lines.append(f"[Chg:{ins_event.id}] {auth}")

        if del_event:
            auth = del_event.author or "Unknown"
            lines.append(f"[Chg:{del_event.id}] {auth}")

        # 2. Comment Metadata
        sorted_ids = sorted(list(active_ids))
        for cid in sorted_ids:
            if cid not in self.comments_map:
                continue
            data = self.comments_map[cid]
            author = data["author"]
            body = data["text"]
            date_str = data["date"]
            resolved = data["resolved"]

            header_parts = [f"[Com:{cid}] {author}"]
            if date_str:
                short_date = date_str.split("T")[0]
                header_parts.append(f"@ {short_date}")
            if resolved:
                header_parts.append("(RESOLVED)")
            header = " ".join(header_parts)
            lines.append(f"{header}: {body}")

        return "\n".join(lines)

    def _add_virtual_text(self, text: str, offset: int, context_paragraph: Paragraph):
        span = TextSpan(
            start=offset,
            end=offset + len(text),
            text=text,
            run=None,  # Virtual
            paragraph=context_paragraph,
        )
        self.spans.append(span)
        self.full_text += text

    def _replace_smart_quotes(self, text: str) -> str:
        return text.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")

    def find_match_index(self, target_text: str) -> int:
        start_idx = self.full_text.find(target_text)
        if start_idx == -1:
            norm_full = self._replace_smart_quotes(self.full_text)
            norm_target = self._replace_smart_quotes(target_text)
            start_idx = norm_full.find(norm_target)
        return start_idx

    def find_target_runs(self, target_text: str) -> List[Run]:
        start_idx = self.find_match_index(target_text)
        if start_idx == -1:
            return []
        return self._resolve_runs_at_range(start_idx, start_idx + len(target_text))

    def find_target_runs_by_index(self, start_index: int, length: int) -> List[Run]:
        end_index = start_index + length
        return self._resolve_runs_at_range(start_index, end_index)

    def _resolve_runs_at_range(self, start_idx: int, end_idx: int) -> List[Run]:
        affected_spans = [s for s in self.spans if s.end > start_idx and s.start < end_idx]
        if not affected_spans:
            return []

        working_runs = [s.run for s in affected_spans if s.run is not None]
        if not working_runs:
            return []

        dom_modified = False

        # 1. Start Split
        first_real_span = next((s for s in affected_spans if s.run is not None), None)
        start_split_adjustment = 0

        if first_real_span:
            local_start = start_idx - first_real_span.start
            if local_start > 0:
                idx_in_working = 0
                _, right_run = self._split_run_at_index(working_runs[idx_in_working], local_start)
                working_runs[idx_in_working] = right_run
                dom_modified = True
                start_split_adjustment = local_start

        # 2. End Split
        last_real_span = next((s for s in reversed(affected_spans) if s.run is not None), None)

        if last_real_span:
            is_same_run = first_real_span is last_real_span
            run_to_split = working_runs[-1]
            overlap_end = min(last_real_span.end, end_idx)
            local_end = overlap_end - last_real_span.start

            if is_same_run and start_split_adjustment > 0:
                local_end -= start_split_adjustment

            if 0 < local_end < len(run_to_split.text):
                left_run, _ = self._split_run_at_index(run_to_split, local_end)
                working_runs[-1] = left_run
                dom_modified = True

        if dom_modified:
            self._build_map()

        return working_runs

    def get_insertion_anchor(self, index: int) -> Optional[Run]:
        preceding = [s for s in self.spans if s.end == index]
        if preceding:
            if preceding[-1].run:
                return preceding[-1].run
        containing = [s for s in self.spans if s.start < index < s.end]
        if containing:
            span = containing[0]
            if span.run is None:
                pass
            else:
                offset = index - span.start
                left, _ = self._split_run_at_index(span.run, offset)
                return left

        if index == 0 and self.spans:
            for s in self.spans:
                if s.run:
                    return s.run
            return None

        preceding_gap = [s for s in self.spans if s.end < index]
        if preceding_gap:
            for s in reversed(preceding_gap):
                if s.run:
                    return s.run
        return None

    def _split_run_at_index(self, run: Run, split_index: int) -> Tuple[Run, Run]:
        text = run.text
        left_text = text[:split_index]
        right_text = text[split_index:]

        run.text = left_text
        new_r_element = deepcopy(run._element)
        t_list = new_r_element.findall(qn("w:t"))
        for t in t_list:
            new_r_element.remove(t)

        new_t = OxmlElement("w:t")
        new_t.text = right_text
        if right_text.strip() != right_text:
            new_t.set(qn("xml:space"), "preserve")
        new_r_element.append(new_t)
        run._element.addnext(new_r_element)
        new_run = Run(new_r_element, run._parent)
        return run, new_run

    def get_context_at_range(self, start_idx: int, end_idx: int) -> Optional[TextSpan]:
        """
        Returns the first real TextSpan in the range to check context.
        Useful for detecting if we are editing inside an Insertion.
        """
        real_spans = [s for s in self.spans if s.run and s.end > start_idx and s.start < end_idx]
        if real_spans:
            return real_spans[0]
        return None
