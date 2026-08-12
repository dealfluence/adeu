import io
import json

import pytest
from docx import Document
from docx.oxml import parse_xml

from adeu.cli import handle_apply
from adeu.mcp_components.tools.document import process_document_batch
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine

W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _create_clean_docx_stream() -> io.BytesIO:
    doc = Document()
    doc.add_paragraph("This is a simple document with baseline text for testing author impersonation warnings.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _create_docx_with_pending_revision(author: str) -> io.BytesIO:
    stream = _create_clean_docx_stream()
    engine = RedlineEngine(stream, author=author)
    edit = ModifyText(
        target_text="simple document",
        new_text="modified document",
    )
    engine.process_batch([edit])
    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    return out_stream


def _create_docx_with_pending_comment(author: str) -> io.BytesIO:
    stream = _create_clean_docx_stream()
    engine = RedlineEngine(stream, author=author)
    edit = ModifyText(
        target_text="simple document",
        new_text="simple document",
        comment="Review note by author",
    )
    engine.process_batch([edit])
    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    return out_stream


def _create_docx_with_tracked_move(author: str, tag_name: str = "w:moveTo") -> io.BytesIO:
    doc = Document()
    p = doc.add_paragraph("This is a simple document with baseline text for testing author impersonation warnings.")
    move_xml = (
        f'<{tag_name} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:id="1" w:author="{author}" w:date="2026-08-12T00:00:00Z">'
        f"<w:r><w:t>moved text</w:t></w:r>"
        f"</{tag_name}>"
    )
    p._element.append(parse_xml(move_xml))
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _create_docx_with_table_revision(author: str, marker: str = "tblPrChange") -> io.BytesIO:
    """Builds a docx whose only pending revision is a table-level marker.

    `marker` is either "tblPrChange" (injected into `w:tblPr`) or "cellIns"
    (injected into the cell's `w:tcPr`). Neither is a run- or paragraph-level
    revision element, so they are only found by scanning for `w:author` itself.
    """
    doc = Document()
    doc.add_paragraph("This is a simple document with baseline text for testing author impersonation warnings.")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "cell text"

    if marker == "tblPrChange":
        table._tbl.tblPr.append(
            parse_xml(
                f'<w:tblPrChange {W_NS} w:id="10" w:author="{author}" w:date="2026-08-12T00:00:00Z">'
                f"<w:tblPr/></w:tblPrChange>"
            )
        )
    else:
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:{marker} {W_NS} w:id="11" w:author="{author}" w:date="2026-08-12T00:00:00Z"/>')
        )

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_warning_when_acting_author_impersonates_a_pending_author():
    # Document has pending revision by "Alice"
    doc_stream = _create_docx_with_pending_revision("Alice")

    # Acting author is also "Alice"
    engine = RedlineEngine(doc_stream, author="Alice")
    edit = ModifyText(target_text="baseline text", new_text="updated text")
    stats = engine.process_batch([edit])

    assert stats.get("author_impersonation_warning") is not None
    assert "Alice" in stats["author_impersonation_warning"]
    assert "matches an author with pending revisions" in stats["author_impersonation_warning"]

    # Document has pending comment by "Charlie"
    comment_doc_stream = _create_docx_with_pending_comment("Charlie")
    engine_comment = RedlineEngine(comment_doc_stream, author="Charlie")
    stats_comment = engine_comment.process_batch([edit])

    assert stats_comment.get("author_impersonation_warning") is not None
    assert "Charlie" in stats_comment["author_impersonation_warning"]


def test_warning_when_acting_author_matches_tracked_move_author():
    # Test w:moveTo
    moveto_stream = _create_docx_with_tracked_move("Dave", tag_name="w:moveTo")
    engine_moveto = RedlineEngine(moveto_stream, author="Dave")
    edit = ModifyText(target_text="baseline text", new_text="updated text")
    stats_moveto = engine_moveto.process_batch([edit])

    assert stats_moveto.get("author_impersonation_warning") is not None
    assert "Dave" in stats_moveto["author_impersonation_warning"]
    assert "matches an author with pending revisions" in stats_moveto["author_impersonation_warning"]

    # Test w:moveFrom
    movefrom_stream = _create_docx_with_tracked_move("Eve", tag_name="w:moveFrom")
    engine_movefrom = RedlineEngine(movefrom_stream, author="Eve")
    stats_movefrom = engine_movefrom.process_batch([edit])

    assert stats_movefrom.get("author_impersonation_warning") is not None
    assert "Eve" in stats_movefrom["author_impersonation_warning"]
    assert "matches an author with pending revisions" in stats_movefrom["author_impersonation_warning"]


@pytest.mark.parametrize("marker", ["tblPrChange", "cellIns"])
def test_warning_when_acting_author_matches_table_revision_author(marker):
    doc_stream = _create_docx_with_table_revision("Frank", marker=marker)
    engine = RedlineEngine(doc_stream, author="Frank")
    edit = ModifyText(target_text="baseline text", new_text="updated text")
    stats = engine.process_batch([edit])

    assert stats.get("author_impersonation_warning") is not None, f"{marker} author was not detected"
    assert "Frank" in stats["author_impersonation_warning"]
    assert "matches an author with pending revisions" in stats["author_impersonation_warning"]

    # A distinct acting author on the same document stays silent.
    other_stream = _create_docx_with_table_revision("Frank", marker=marker)
    other_engine = RedlineEngine(other_stream, author="Grace")
    assert other_engine.process_batch([edit]).get("author_impersonation_warning") is None


def test_no_warning_for_a_distinct_author():
    # Document has pending revision by "Alice"
    doc_stream = _create_docx_with_pending_revision("Alice")

    # Acting author is "Bob" (distinct)
    engine = RedlineEngine(doc_stream, author="Bob")
    edit = ModifyText(target_text="baseline text", new_text="updated text")
    stats = engine.process_batch([edit])

    assert stats.get("author_impersonation_warning") is None


def test_no_warning_on_a_clean_document():
    # Clean document without any pending revisions
    doc_stream = _create_clean_docx_stream()

    # Acting author is "Alice"
    engine = RedlineEngine(doc_stream, author="Alice")
    edit = ModifyText(target_text="baseline text", new_text="updated text")
    stats = engine.process_batch([edit])

    assert stats.get("author_impersonation_warning") is None


def test_cli_surfaces_the_warning(tmp_path, capsys):
    # Prepare a document with pending revision by "Alice"
    doc_stream = _create_docx_with_pending_revision("Alice")
    orig_path = tmp_path / "orig.docx"
    orig_path.write_bytes(doc_stream.getvalue())

    changes_path = tmp_path / "changes.json"
    changes_path.write_text(
        json.dumps([{"type": "modify", "target_text": "baseline text", "new_text": "updated text"}]),
        encoding="utf-8",
    )

    out_path = tmp_path / "out.docx"

    # 1. Test CLI human mode (stdout/stderr output)
    class ArgsHuman:
        original = orig_path
        changes = changes_path
        output = out_path
        author = "Alice"
        json = False
        live = False
        partial = False
        report = "standard"
        allow_major_deletions = False

    handle_apply(ArgsHuman())
    captured = capsys.readouterr()
    assert "Alice" in captured.err
    assert "pending revisions" in captured.err

    # 2. Test CLI JSON mode
    out_path_json = tmp_path / "out_json.docx"

    class ArgsJSON:
        original = orig_path
        changes = changes_path
        output = out_path_json
        author = "Alice"
        json = True
        live = False
        partial = False
        report = "standard"
        allow_major_deletions = False

    handle_apply(ArgsJSON())
    captured_json = capsys.readouterr()
    stdout_data = json.loads(captured_json.out)
    assert stdout_data.get("author_impersonation_warning") is not None
    assert "Alice" in stdout_data["author_impersonation_warning"]
    assert "pending revisions" in stdout_data["author_impersonation_warning"]


@pytest.mark.anyio
async def test_mcp_batch_response_surfaces_the_warning(tmp_path):
    doc_path = tmp_path / "pending.docx"
    doc_path.write_bytes(_create_docx_with_pending_revision("Alice").getvalue())
    out_path = tmp_path / "out.docx"

    warnings: list[str] = []

    class FakeContext:
        async def info(self, *a, **kw):
            pass

        async def warning(self, *a, **kw):
            warnings.append(str(a[0]) if a else "")

        async def debug(self, *a, **kw):
            pass

        async def error(self, *a, **kw):
            pass

    result = await process_document_batch(
        reasoning="Testing the impersonation warning in the MCP response",
        original_docx_path=str(doc_path),
        output_path=str(out_path),
        author_name="Alice",
        ctx=FakeContext(),  # type: ignore
        changes=[{"type": "modify", "target_text": "baseline text", "new_text": "updated text"}],
    )

    assert "Batch complete" in result
    assert "*Warning:*" in result
    assert "Alice" in result
    assert "matches an author with pending revisions" in result
    assert any("matches an author with pending revisions" in w for w in warnings)
