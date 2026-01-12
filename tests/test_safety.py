import io

import structlog
from docx import Document

from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine

logger = structlog.get_logger(__name__)


def test_reject_empty_target_heuristic():
    """
    Ensures that an edit with empty target_text (heuristic) is skipped,
    preventing accidental start-of-document insertion or 'Applied 1' stats for empty edits.
    """
    doc = Document()
    doc.add_paragraph("Content")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # An edit that effectively has empty target and empty new (or just empty target)
    # Pydantic requires target_text, but it can be an empty string.
    edit = DocumentEdit(target_text="", new_text="Unexpected Header")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 0
    assert skipped == 1

    # Verify doc content didn't change
    result_stream = engine.save_to_stream()
    doc_res = Document(result_stream)
    text = doc_res.paragraphs[0].text
    assert "Unexpected Header" not in text


def test_multiple_occurrences_apply_once():
    """
    Verifies that a heuristic edit only applies to the first match found.
    """
    doc = Document()
    doc.add_paragraph("Repeat")
    doc.add_paragraph("Repeat")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="Repeat", new_text="Changed")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1
    assert skipped == 0

    result_stream = engine.save_to_stream()
    doc_res = Document(result_stream)

    # First one changed, second one remains
    # Note: docx.paragraphs list might be affected by redline tag structure,
    # but text checking is enough.
    full_xml = doc_res.element.xml

    # We expect one deletion of Repeat and one insertion of Changed
    assert full_xml.count("<w:delText>Repeat</w:delText>") == 1
    assert full_xml.count("<w:t>Changed</w:t>") == 1
    # The second "Repeat" should still be a normal text run
    # (It won't be in a delText)

    # Simple check: count occurrences of literal "Repeat" outside of deletion tags
    # This is tricky in XML. Let's just assume if we have 1 del and 1 ins, we did 1 op.
