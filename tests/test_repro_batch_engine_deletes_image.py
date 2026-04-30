import io

import docx

from adeu import ModifyText, RedlineEngine


def test_batch_engine_deletes_special_content():
    """
    Reproduces a bug where modifying text in a paragraph that contains
    special content (like images/drawings) deletes the drawing element,
    violating Safety Constraint #2.
    """
    # 1. Setup a doc with a paragraph containing text and an image.
    d = docx.Document()
    p = d.add_paragraph("Normal paragraph")
    r = p.add_run(" ")
    # Add a drawing element (use a dummy file or we can just inject XML directly to avoid file dependencies)
    from docx.oxml import parse_xml

    drawing_xml = parse_xml('<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    r._r.append(drawing_xml)

    stream = io.BytesIO()
    d.save(stream)
    stream.seek(0)

    # Verify the drawing is there before
    d_before = docx.Document(stream)
    assert (
        len(
            d_before.paragraphs[0]._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
        )
        == 1
    )
    stream.seek(0)

    # 2. Modify the text preceding the drawing.
    engine = RedlineEngine(stream, author="QA")
    engine.apply_edits([ModifyText(target_text="Normal paragraph ", new_text="Modified ")])

    # 3. Verify the drawing was destroyed.
    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    d_after = docx.Document(out_stream)

    drawings_after = d_after.paragraphs[0]._element.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    )
    assert len(drawings_after) == 1, "BUG: The drawing element was destroyed by the text modification!"
