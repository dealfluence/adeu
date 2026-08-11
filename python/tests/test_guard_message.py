import io

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


def _create_doc_with_foreign_ins(author: str = "Supplier's Counsel", ins_id: str = "201") -> io.BytesIO:
    doc = Document()
    p = doc.add_paragraph("The party shall provide ")
    p_el = p._element

    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), ins_id)
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), "2026-06-30T08:00:00Z")

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "written notice"
    r.append(t)
    ins.append(r)
    p_el.append(ins)

    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = " within 30 days."
    r2.append(t2)
    p_el.append(r2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _create_doc_with_many_foreign_ins() -> io.BytesIO:
    """A paragraph whose target text is covered by 6 insertions from 3 foreign authors."""
    doc = Document()
    p = doc.add_paragraph()
    p_el = p._element

    authors = ["Supplier's Counsel", "Regulatory Reviewer", "Buyer's Counsel"]
    words = ["alpha ", "beta ", "gamma ", "delta ", "epsilon ", "zeta"]

    for idx, word in enumerate(words):
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), str(201 + idx))
        ins.set(qn("w:author"), authors[idx % len(authors)])
        ins.set(qn("w:date"), "2026-06-30T08:00:00Z")

        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = word
        r.append(t)
        ins.append(r)
        p_el.append(ins)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_guard_names_the_accept_action():
    stream = _create_doc_with_foreign_ins(author="Supplier's Counsel", ins_id="201")
    engine = RedlineEngine(stream, author="Reviewer AI")
    edit = ModifyText(target_text="written notice", new_text="email notification", match_mode="all")

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch([edit])

    msg = exc_info.value.errors[0]
    assert '{"type": "accept", "target_id": "Chg:201"}' in msg


def test_guard_names_the_narrowing_alternative():
    stream = _create_doc_with_foreign_ins(author="Supplier's Counsel", ins_id="201")
    engine = RedlineEngine(stream, author="Reviewer AI")
    edit = ModifyText(target_text="written notice", new_text="email notification", match_mode="all")

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch([edit])

    msg = exc_info.value.errors[0]
    assert 'or use match_mode="strict" or "first"' in msg
    assert "scope your edit outside of it" in msg


def test_guard_still_names_author_and_ids():
    """author name and change IDs preserved by the actionable rewrite (holds pre- and post-change)."""
    stream = _create_doc_with_foreign_ins(author="Supplier's Counsel", ins_id="201")
    engine = RedlineEngine(stream, author="Reviewer AI")
    edit = ModifyText(target_text="written notice", new_text="email notification", match_mode="all")

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch([edit])

    msg = exc_info.value.errors[0]
    assert "Supplier's Counsel" in msg
    assert "(e.g. Chg:201)" in msg


def test_guard_recommends_scoping_outside_for_strict_straddle():
    stream = _create_doc_with_foreign_ins(author="Supplier's Counsel", ins_id="201")
    engine = RedlineEngine(stream, author="Reviewer AI")
    edit = ModifyText(target_text="provide written notice", new_text="provide email notification", match_mode="strict")

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch([edit])

    msg = exc_info.value.errors[0]
    assert '{"type": "accept", "target_id": "Chg:201"}' in msg
    assert "scope your edit outside of it" in msg
    assert 'use match_mode="strict"' not in msg
    assert "use match_mode" not in msg


def test_guard_recommends_scoping_outside_for_first_straddle():
    stream = _create_doc_with_foreign_ins(author="Supplier's Counsel", ins_id="201")
    engine = RedlineEngine(stream, author="Reviewer AI")
    edit = ModifyText(target_text="provide written notice", new_text="provide email notification", match_mode="first")

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch([edit])

    msg = exc_info.value.errors[0]
    assert '{"type": "accept", "target_id": "Chg:201"}' in msg
    assert "scope your edit outside of it" in msg
    assert "use match_mode" not in msg


def test_guard_message_token_budget():
    stream = _create_doc_with_foreign_ins(author="Supplier's Counsel", ins_id="201")
    engine = RedlineEngine(stream, author="Reviewer AI")
    edit_all = ModifyText(target_text="written notice", new_text="email notification", match_mode="all")

    with pytest.raises(BatchValidationError) as exc_info_all:
        engine.process_batch([edit_all])

    msg_all = exc_info_all.value.errors[0]
    approx_tokens_all = len(msg_all) // 4
    assert approx_tokens_all <= 70

    stream_strict = _create_doc_with_foreign_ins(author="Supplier's Counsel", ins_id="201")
    engine_strict = RedlineEngine(stream_strict, author="Reviewer AI")
    edit_strict = ModifyText(
        target_text="provide written notice", new_text="provide email notification", match_mode="strict"
    )

    with pytest.raises(BatchValidationError) as exc_info_strict:
        engine_strict.process_batch([edit_strict])

    msg_strict = exc_info_strict.value.errors[0]
    approx_tokens_strict = len(msg_strict) // 4
    assert approx_tokens_strict <= 70

    stream_many = _create_doc_with_many_foreign_ins()
    engine_many = RedlineEngine(stream_many, author="Reviewer AI")
    edit_many = ModifyText(target_text="alpha beta gamma delta epsilon zeta", new_text="one two", match_mode="all")

    with pytest.raises(BatchValidationError) as exc_info_many:
        engine_many.process_batch([edit_many])

    msg_many = exc_info_many.value.errors[0]
    approx_tokens_many = len(msg_many) // 4
    assert approx_tokens_many <= 70, f"{approx_tokens_many} tokens: {msg_many}"
    # The bounded hint must still name one author, one usable ID and the omitted count.
    assert "(+2 more)" in msg_many


def test_strict_edit_inside_foreign_insertion_still_allowed():
    stream = _create_doc_with_foreign_ins(author="Supplier's Counsel", ins_id="201")
    engine = RedlineEngine(stream, author="Reviewer AI")
    edit = ModifyText(target_text="written notice", new_text="email notification", match_mode="strict")

    res = engine.process_batch([edit])
    assert res["edits_applied"] == 1
