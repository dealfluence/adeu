# FILE: tests/test_styling_and_lists.py

import io

from docx import Document
from docx.oxml.ns import qn

from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_visible_runs


def _create_initial_docx() -> io.BytesIO:
    doc = Document()

    # 1. Heading
    doc.add_heading("Project Scope", level=1)

    # 2. List Items (Simulated with text for simplicity in setup)
    p1 = doc.add_paragraph("Phase 1: Initial")
    # Setting a list style if available would be ideal, but for unit tests
    # relying on text structure and generic style inheritance is sufficient.
    # In real world, style="List Paragraph" would be present.
    try:
        p1.style = "List Paragraph"
    except KeyError:
        pass  # Fallback if style doesn't exist in default template

    p2 = doc.add_paragraph("Phase 3: Final")
    try:
        p2.style = "List Paragraph"
    except KeyError:
        pass

    # 3. Inline Styling Paragraphs
    doc.add_paragraph("The quick brown fox.")
    doc.add_paragraph("Jump over the dog.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _get_run_property(run, prop_name):
    """
    Helper to check rPr.b or rPr.i
    """
    # runs from get_visible_runs wrap the element.
    # run._element gives access to the underlying xml element.
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return False

    prop_el = rPr.find(qn(f"w:{prop_name}"))
    if prop_el is None:
        return False
    # If tag exists (e.g. <w:b/>), it means true unless val="0" or "false"
    val = prop_el.get(qn("w:val"))
    if val in ("0", "false"):
        return False
    return True


def test_list_injection_with_nested_styling():
    """
    Test Case:
    - Insert "Phase 2" between "Phase 1" and "Phase 3".
    - "Phase 2" has **_Development_**.
    """
    stream = _create_initial_docx()
    engine = RedlineEngine(stream)

    # Target text must match what ingest produces.
    # adeu.ingest adds \n\n between paragraphs usually.
    # We target "Phase 1: Initial" and rely on the engine to find it.

    edit = DocumentEdit(
        target_text="Phase 1: Initial",
        new_text="Phase 1: Initial\nPhase 2: **_Development_** starts here.",
        comment="Inserting list item",
    )

    applied, skipped = engine.apply_edits([edit])
    assert applied == 1
    assert skipped == 0

    # Verification
    res_stream = engine.save_to_stream()
    doc = Document(res_stream)

    # Check Paragraph Order: Heading -> Phase 1 -> Phase 2 -> Phase 3
    assert "Phase 1: Initial" in "".join(r.text for r in get_visible_runs(doc.paragraphs[1]))

    p_inserted = doc.paragraphs[2]
    # Visible text should be clean (no markdown markers)
    visible_text = "".join(r.text for r in get_visible_runs(p_inserted))
    assert "Phase 2: Development starts here." in visible_text
    assert "**" not in visible_text

    # Check Runs for Formatting
    runs = get_visible_runs(p_inserted)

    found_formatted = False
    for r in runs:
        if "Development" in r.text:
            is_bold = _get_run_property(r, "b")
            is_italic = _get_run_property(r, "i")
            if is_bold and is_italic:
                found_formatted = True
                break

    assert found_formatted, "Did not find a run with text 'Development' that is both Bold and Italic"


def test_inline_italic_modification():
    """
    Test Case: Change "brown" to "_red_".
    """
    stream = _create_initial_docx()
    engine = RedlineEngine(stream)

    edit = DocumentEdit(target_text="brown", new_text="_red_", comment="Color change")

    applied, skipped = engine.apply_edits([edit])
    assert applied == 1

    res_stream = engine.save_to_stream()
    doc = Document(res_stream)

    # Find paragraph
    target_p = None
    for p in doc.paragraphs:
        # docx.Paragraph.text doesn't show deletions, but shows insertions (usually)
        # We check full visible text
        p_text = "".join(r.text for r in get_visible_runs(p))
        if "red" in p_text:
            target_p = p
            break

    assert target_p is not None
    assert "red" in "".join(r.text for r in get_visible_runs(target_p))

    runs = get_visible_runs(target_p)

    found_italic = False
    for r in runs:
        if "red" in r.text:
            if _get_run_property(r, "i"):
                found_italic = True

    assert found_italic, "New word 'red' should be italicized"


def test_inline_bold_modification():
    """
    Test Case: Change "dog" to "**cat**".
    """
    stream = _create_initial_docx()
    engine = RedlineEngine(stream)

    edit = DocumentEdit(target_text="dog", new_text="**cat**", comment="Animal change")

    applied, skipped = engine.apply_edits([edit])
    assert applied == 1

    res_stream = engine.save_to_stream()
    doc = Document(res_stream)

    target_p = None
    for p in doc.paragraphs:
        p_text = "".join(r.text for r in get_visible_runs(p))
        if "cat" in p_text:
            target_p = p
            break

    assert target_p is not None

    runs = get_visible_runs(target_p)

    found_bold = False
    for r in runs:
        if "cat" in r.text:
            if _get_run_property(r, "b"):
                found_bold = True

    assert found_bold, "New word 'cat' should be bold"
