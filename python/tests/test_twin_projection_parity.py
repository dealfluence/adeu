"""
Virtual Text contract regression gate: the ingest reader
(_extract_text_from_doc) and the editor-side mapper (DocumentMapper) must
produce BYTE-IDENTICAL body text — agents read the ingest projection, and
edits resolve against the mapper projection, so any drift makes strict
target_text matches fail on text the agent just read.

Fixed 2026-07-24 (mapper aligned to the reader, which is canonical). The
mechanisms below each reproduce a divergence found on real documents
(BIGDOC/VVBIG goldens, 196 diff hunks):

  1. Marker elision across boundary whitespace: bold "Request for " +
     bold "Bids" must merge to "**Request for Bids**" (mapper emitted
     "**Request for** **Bids**").
  2. Elision safety: a whitespace-only same-style run must not destroy the
     closing marker (mapper emitted "**March 2012 " with no closer).
  3. Redline state transitions must not split wrapper groups: adjacent
     w:del elements (one per Chg id) coalesce into ONE {--...--} block.
  4. Empty parts contribute nothing — not even part separators (mapper
     emitted 4 extra leading newlines on BIGDOC).
  5. A styled run whose only child is a drawing/reference (empty text)
     contributes no dangling markers (mapper emitted "(docx-image:1)****").

The parity assertion is extraction(include_appendix=False) ==
mapper.full_text — the mapper never projects the appendix. (Note that
split_structural_appendix(raw_with_appendix).body keeps the appendix's
"\\n\\n---" ruler, so it is NOT the right-hand side of this contract.)
"""

from io import BytesIO

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from adeu.ingest import _extract_text_from_doc
from adeu.redline.mapper import DocumentMapper
from adeu.utils.docx import strip_bom_from_docx_bytes


def _roundtrip(doc) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return strip_bom_from_docx_bytes(buf.getvalue())


def assert_twins_identical(doc_bytes: bytes) -> list[tuple[bool, str]]:
    """Asserts reader == mapper for raw and clean views; returns
    [(clean_view, reader_text), ...] for additional content assertions."""
    results = []
    for clean in (False, True):
        reader = _extract_text_from_doc(Document(BytesIO(doc_bytes)), clean_view=clean, include_appendix=False)
        mapper = DocumentMapper(Document(BytesIO(doc_bytes)), clean_view=clean)
        assert reader == mapper.full_text, (
            f"twin drift in {'clean' if clean else 'raw'} view:\nreader: {reader!r}\nmapper: {mapper.full_text!r}"
        )
        results.append((clean, reader))
    return results


def _add_bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def test_adjacent_bold_runs_with_boundary_whitespace():
    doc = Document()
    p = doc.add_paragraph()
    _add_bold_run(p, "Request for ")
    _add_bold_run(p, "Bids")
    data = _roundtrip(doc)
    for _clean, reader in assert_twins_identical(data):
        assert "**Request for Bids**" in reader
        assert "** **" not in reader


def test_whitespace_only_same_style_run_keeps_marker_balance():
    doc = Document()
    p = doc.add_paragraph()
    _add_bold_run(p, "March 2012")
    _add_bold_run(p, " ")
    doc.add_paragraph("after")
    data = _roundtrip(doc)
    for _clean, reader in assert_twins_identical(data):
        assert "**March 2012**" in reader


def test_multiline_and_leading_whitespace_styled_runs():
    doc = Document()
    p = doc.add_paragraph()
    _add_bold_run(p, "Alpha")
    _add_bold_run(p, " Beta")  # leading whitespace before same-style run
    p2 = doc.add_paragraph()
    r = p2.add_run("Line1\nLine2")
    r.bold = True
    data = _roundtrip(doc)
    for _clean, reader in assert_twins_identical(data):
        assert "**Alpha Beta**" in reader


def test_adjacent_del_elements_coalesce_into_one_block():
    """A replacement stored as several sibling w:del elements (distinct ids,
    mixed run styling) must project as ONE {--...--} block in both twins."""
    w = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    p_xml = (
        f"<w:p {w}>"
        '<w:del w:id="55" w:author="A" w:date="2020-01-01T00:00:00Z">'
        '<w:r><w:delText xml:space="preserve">plain </w:delText></w:r>'
        "</w:del>"
        '<w:del w:id="56" w:author="A" w:date="2020-01-01T00:00:00Z">'
        '<w:r><w:rPr><w:i/></w:rPr><w:delText xml:space="preserve">italic part</w:delText></w:r>'
        "</w:del>"
        '<w:ins w:id="57" w:author="A" w:date="2020-01-01T00:00:00Z">'
        "<w:r><w:t>replacement</w:t></w:r>"
        "</w:ins>"
        "</w:p>"
    )
    doc = Document()
    anchor = doc.add_paragraph("anchor")
    new_p = parse_xml(p_xml)
    anchor._p.addnext(new_p)
    data = _roundtrip(doc)
    for clean, reader in assert_twins_identical(data):
        if not clean:
            assert reader.count("{--") == 1, reader
            assert "{--plain _italic part_--}" in reader
            assert "{++replacement++}" in reader


# --------------------------------------------------------------------------
# Deleted paragraph mark, nothing surviving inside (QA round 3, finding 2.4).
#
# The reader drops such a paragraph from the clean view (ingest._extract_blocks)
# because accepting the mark deletion merges the paragraph away. The mapper must
# drop it identically: it previously counted the paragraph as a block and ran
# 2 chars ahead per occurrence, and since caller-pinned _match_start_index
# offsets are bound to clean_mapper (and validate_edits skips pinned edits),
# every later edit resolved mid-word and applied with no error raised.
#
# The run-level w:del case above does NOT cover this — the mark lives in
# w:pPr/w:rPr, which is what makes the container itself disappear.
# --------------------------------------------------------------------------

_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _deleted_mark_para_xml(*, style: str | None = None, text: str = "removed entirely") -> str:
    """A paragraph whose pilcrow is a tracked deletion and whose only content
    is tracked-deleted, so the clean view has nothing left to render."""
    pstyle = f'<w:pStyle w:val="{style}"/>' if style else ""
    return (
        f"<w:p {_W}>"
        f"<w:pPr>{pstyle}"
        '<w:rPr><w:del w:id="90" w:author="QA" w:date="2026-07-24T00:00:00Z"/></w:rPr>'
        "</w:pPr>"
        '<w:del w:id="91" w:author="QA" w:date="2026-07-24T00:00:00Z">'
        f"<w:r><w:delText>{text}</w:delText></w:r>"
        "</w:del>"
        "</w:p>"
    )


def test_deleted_paragraph_mark_midway_drops_block_and_separator():
    doc = Document()
    before = doc.add_paragraph("BEFORE")
    before._p.addnext(parse_xml(_deleted_mark_para_xml()))
    doc.add_paragraph("AFTER")
    data = _roundtrip(doc)
    for clean, reader in assert_twins_identical(data):
        if clean:
            # No empty container and no leftover separator for it.
            assert reader == "BEFORE\n\nAFTER", reader
        else:
            assert "{--removed entirely--}" in reader


def test_deleted_paragraph_mark_as_first_block_emits_no_leading_separator():
    doc = Document()
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)
    doc.element.body.insert(0, parse_xml(_deleted_mark_para_xml()))
    doc.add_paragraph("AFTER")
    data = _roundtrip(doc)
    for clean, reader in assert_twins_identical(data):
        if clean:
            assert reader == "AFTER", reader
            assert not reader.startswith("\n")


def test_deleted_heading_mark_drops_its_markdown_prefix_too():
    """The reader drops the whole `prefix + p_text` block, so the mapper's
    rollback must also undo the emitted '# ' heading prefix."""
    doc = Document()
    before = doc.add_paragraph("BEFORE")
    before._p.addnext(parse_xml(_deleted_mark_para_xml(style="Heading1", text="Removed Heading")))
    doc.add_paragraph("AFTER")
    data = _roundtrip(doc)
    for clean, reader in assert_twins_identical(data):
        if clean:
            assert reader == "BEFORE\n\nAFTER", reader
            assert "Removed Heading" not in reader
            assert "#" not in reader


def test_empty_header_part_contributes_nothing():
    doc = Document()
    doc.add_paragraph("Body content only.")
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False  # materializes an empty header part
    data = _roundtrip(doc)
    for _clean, reader in assert_twins_identical(data):
        assert not reader.startswith("\n")
        assert reader.startswith("Body content only.")


def test_styled_run_with_only_drawing_leaves_no_dangling_markers():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("")
    run.bold = True
    drawing = parse_xml(f"<w:drawing {nsdecls('w')}/>")
    run._r.append(drawing)
    p.add_run(" trailing text")
    data = _roundtrip(doc)
    for _clean, reader in assert_twins_identical(data):
        assert "****" not in reader
        assert "__" not in reader


def test_tracked_changes_fixture_parity():
    from tests.docx_fixtures import make_doc_with_track_changes

    doc = make_doc_with_track_changes()
    data = _roundtrip(doc)
    list(assert_twins_identical(data))


def test_comments_fixture_parity():
    from tests.docx_fixtures import make_doc_with_comments

    doc = make_doc_with_comments()
    data = _roundtrip(doc)
    list(assert_twins_identical(data))


def test_multi_section_fixture_parity():
    from tests.docx_fixtures import make_doc_with_multiple_sections

    doc = make_doc_with_multiple_sections()
    data = _roundtrip(doc)
    list(assert_twins_identical(data))


def test_tables_and_headings_parity():
    doc = Document()
    doc.add_heading("Section", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "cell one"
    table.cell(1, 1).text = "trailing space "
    doc.add_paragraph("closing")
    data = _roundtrip(doc)
    list(assert_twins_identical(data))
