import json
from pathlib import Path

from adeu.cli import handle_extract
from adeu.ingest import _extract_text_from_doc
from adeu.mcp_components._response_builders import (
    build_paginated_response,
)
from tests.fixtures_synth import build_long_docx
from tests.utils import approx_tokens


class _Args:
    """Mock argparse Namespace for CLI tests."""

    def __init__(self, **kwargs):
        self.input = kwargs.get("input", None)
        self.live = kwargs.get("live", False)
        self.output = kwargs.get("output", None)
        self.force = kwargs.get("force", False)
        self.clean_view = kwargs.get("clean_view", False)
        self.mode = kwargs.get("mode", "full")
        self.changes_author = kwargs.get("changes_author", None)
        self.changes_offset = kwargs.get("changes_offset", 0)
        self.page = kwargs.get("page", None)
        self.search_query = kwargs.get("search_query", None)
        self.search_regex = kwargs.get("search_regex", False)
        self.search_case_insensitive = kwargs.get("search_case_insensitive", False)
        self.max_matches = kwargs.get("max_matches", 20)
        self.match_offset = kwargs.get("match_offset", 0)
        self.full_paragraph = kwargs.get("full_paragraph", False)
        self.outline_max_level = kwargs.get("outline_max_level", 2)
        self.outline_verbose = kwargs.get("outline_verbose", False)
        self.json = kwargs.get("json", False)
        self.no_chrome = kwargs.get("no_chrome", False)


def test_no_chrome_drops_file_path_header_and_prose(tmp_path: Path):
    doc_path = tmp_path / "test.docx"
    build_long_docx(doc_path, pages=3)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)

    res_default = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=False)
    res_no_chrome = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=True)

    assert "**File Path:**" in str(res_default.content)
    assert "(synthetic page — a length-based chunk" in str(res_default.content)

    assert "**File Path:**" not in str(res_no_chrome.content)
    assert "(synthetic page — a length-based chunk" not in str(res_no_chrome.content)
    assert "Continues on page" not in str(res_no_chrome.content)


def test_no_chrome_page_content_is_byte_identical_apart_from_chrome(tmp_path: Path):
    doc_path = tmp_path / "test_single.docx"
    build_long_docx(doc_path, pages=1)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)

    res_default = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=False)
    res_no_chrome = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=True)

    # In single page doc with no_chrome=True, content is pure body text.
    body_text = str(res_no_chrome.content)
    # The default content includes '> **File Path:** `path`\n\n' prefix and body_text
    default_content = str(res_default.content)
    header_prefix = f"> **File Path:** `{doc_path}`\n\n"
    assert default_content.startswith(header_prefix)
    extracted_body = default_content[len(header_prefix) :]

    assert extracted_body == body_text


def test_no_chrome_keeps_bare_page_marker_on_multipage(tmp_path: Path):
    doc_path = tmp_path / "test_multi.docx"
    build_long_docx(doc_path, pages=3)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)

    res_no_chrome_p1 = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=True)
    res_no_chrome_p2 = build_paginated_response(text, page=2, file_path=str(doc_path), no_chrome=True)

    content_p1 = str(res_no_chrome_p1.content)
    content_p2 = str(res_no_chrome_p2.content)

    assert content_p1.startswith("[p1/3]\n\n")
    assert content_p2.startswith("[p2/3]\n\n")
    assert "synthetic page" not in content_p1
    assert "**File Path:**" not in content_p1


def test_no_chrome_saves_tokens(tmp_path: Path):
    doc_path = tmp_path / "test_tokens.docx"
    build_long_docx(doc_path, pages=3)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)

    res_default = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=False)
    res_no_chrome = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=True)

    default_tokens = approx_tokens(str(res_default.content))
    no_chrome_tokens = approx_tokens(str(res_no_chrome.content))

    assert default_tokens - no_chrome_tokens >= 20


def test_no_chrome_composes_with_json(tmp_path: Path, capsys):
    doc_path = tmp_path / "test_json.docx"
    build_long_docx(doc_path, pages=3)

    args = _Args(input=doc_path, page="1", json=True, no_chrome=True)
    handle_extract(args)

    captured = capsys.readouterr()
    stdout = captured.out.strip()

    data = json.loads(stdout)
    assert "markdown" in data
    md = data["markdown"]

    assert "**File Path:**" not in md
    assert "(synthetic page" not in md
    assert md.startswith("[p1/3]\n\n")
