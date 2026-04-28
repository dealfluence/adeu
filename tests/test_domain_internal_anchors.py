import io

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


def _create_bookmark_docx() -> io.BytesIO:
    doc = Document()
    p = doc.add_paragraph("Section 5. Indemnification")

    # Inject bookmarkStart
    bstart = OxmlElement("w:bookmarkStart")
    bstart.set(qn("w:name"), "_Ref12345")
    bstart.set(qn("w:id"), "0")
    p._element.append(bstart)

    # Inject bookmarkEnd
    bend = OxmlElement("w:bookmarkEnd")
    bend.set(qn("w:id"), "0")
    p._element.append(bend)

    # Inject a noise bookmark that should be ignored
    p2 = doc.add_paragraph("Some text.")
    bstart2 = OxmlElement("w:bookmarkStart")
    bstart2.set(qn("w:name"), "_GoBack")
    bstart2.set(qn("w:id"), "1")
    p2._element.append(bstart2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_internal_anchor_projection():
    """
    Bookmarks (excluding _GoBack) should project as {#_BookmarkName} at the end of the text block.
    """
    stream = _create_bookmark_docx()
    projection = extract_text_from_stream(stream)

    # The valid bookmark should be projected
    assert "{#_Ref12345}" in projection, "Bookmark projection missing."
    assert "Section 5. Indemnification{#_Ref12345}" in projection

    # The _GoBack noise bookmark should be ignored
    assert "{#_GoBack}" not in projection, "Noise bookmarks must be ignored."


def test_internal_anchor_strict_refusal():
    """
    Attempting to fabricate or modify an internal anchor syntax should raise a BatchValidationError.
    """
    stream = _create_bookmark_docx()
    engine = RedlineEngine(stream)

    # 1. Attempt to fabricate a new anchor
    edit_fabricate = ModifyText(target_text="Some text.", new_text="Some text.{#_Ref99999}")

    with pytest.raises(BatchValidationError) as excinfo:
        engine.process_batch([edit_fabricate])
    assert "Cannot modify or insert internal anchor markers" in str(excinfo.value)

    # 2. Attempt to rename an existing anchor
    edit_modify = ModifyText(
        target_text="Section 5. Indemnification{#_Ref12345}", new_text="Section 5. Indemnification{#_Ref99999}"
    )

    with pytest.raises(BatchValidationError) as excinfo:
        engine.process_batch([edit_modify])
    assert "Cannot modify or insert internal anchor markers" in str(excinfo.value)
