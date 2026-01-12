import io
import pytest
from docx import Document
from adeu.redline.engine import RedlineEngine
from adeu.models import DocumentEdit
from adeu.ingest import extract_text_from_stream

def test_delete_paragraph_with_newline():
    doc = Document()
    doc.add_paragraph("Paragraph 1.")
    doc.add_paragraph("Paragraph 2.")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    edit = DocumentEdit(
        target_text="Paragraph 1.\n\n", 
        new_text=""
    )
    
    stream.seek(0)
    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])
    
    assert applied == 1
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml
    
    assert "<w:delText>Paragraph 1.</w:delText>" in xml