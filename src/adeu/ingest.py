import io

import structlog
from docx import Document
from docx.text.run import Run

from adeu.redline.comments import CommentsManager
from adeu.utils.docx import (
    DocxEvent,
    get_paragraph_prefix,
    get_run_style_markers,
    get_run_text,
    iter_document_parts,
    iter_paragraph_content,
)

logger = structlog.get_logger(__name__)


def extract_text_from_stream(file_stream: io.BytesIO, filename: str = "document.docx") -> str:
    """
    Extracts text from a file stream using raw run concatenation.
    Includes Markdown headers (#) and CriticMarkup Comments ({==Text==}{>>Comment<<}).

    CRITICAL: This must match DocumentMapper._build_map logic exactly.
    """
    try:
        # Ensure stream is at start
        file_stream.seek(0)
        doc = Document(file_stream)

        comments_mgr = CommentsManager(doc)
        comments_map = comments_mgr.extract_comments_data()

        full_text = []

        for part in iter_document_parts(doc):
            # 1. Paragraphs
            for para in part.paragraphs:
                # Add Markdown prefix if heading
                prefix = get_paragraph_prefix(para)

                # Build content
                p_text = _build_paragraph_text(para, comments_map)

                full_text.append(prefix + p_text)

            # 2. Tables
            for table in part.tables:
                for row in table.rows:
                    row_parts = []
                    for cell in row.cells:
                        # Cell paragraphs
                        cell_text_parts = []
                        for p in cell.paragraphs:
                            prefix = get_paragraph_prefix(p)
                            p_content = _build_paragraph_text(p, comments_map)
                            cell_text_parts.append(prefix + p_content)

                        cell_text = "\n".join(cell_text_parts)
                        if cell_text:
                            row_parts.append(cell_text)

                    if row_parts:
                        full_text.append(" | ".join(row_parts))

        return "\n\n".join(full_text)

    except Exception as e:
        logger.error(f"Text extraction failed: {e}", exc_info=True)
        raise ValueError(f"Could not extract text: {str(e)}") from e


def _build_paragraph_text(paragraph, comments_map):
    """
    Flatten overlapping comments into sequential CriticMarkup blocks.
    Merges metadata for adjacent Redline blocks (Substitutions).
    """
    parts = []

    active_ins: dict[str, DocxEvent] = {}
    active_del: dict[str, DocxEvent] = {}
    active_comments: set[str] = set()

    # Buffer for deferred metadata (used for merging substitution blocks)
    # List of (active_ins_snapshot, active_del_snapshot, active_comments_snapshot)
    deferred_meta_states = []

    # Pre-calculate item list to allow lookahead
    items = list(iter_paragraph_content(paragraph))

    for i, item in enumerate(items):
        if isinstance(item, Run):
            prefix, suffix = get_run_style_markers(item)
            text = get_run_text(item)

            seg = f"{prefix}{text}{suffix}"
            if seg:
                # 1. Determine Wrappers
                start_token, end_token = _get_wrappers(active_ins, active_del, active_comments)

                # 2. Output Text Wrapped
                parts.append(f"{start_token}{seg}{end_token}")

                # 3. Handle Metadata
                current_state = (active_ins.copy(), active_del.copy(), active_comments.copy())
                deferred_meta_states.append(current_state)

                should_defer = False
                is_redline = bool(active_ins) or bool(active_del)

                if is_redline:
                    # Lookahead
                    j = i + 1
                    next_is_redline = False

                    temp_ins = bool(active_ins)
                    temp_del = bool(active_del)

                    while j < len(items):
                        next_item = items[j]
                        if isinstance(next_item, Run):
                            if temp_ins or temp_del:
                                next_is_redline = True
                            break
                        elif isinstance(next_item, DocxEvent):
                            if next_item.type == "ins_start":
                                temp_ins = True
                            elif next_item.type == "ins_end":
                                temp_ins = False
                            elif next_item.type == "del_start":
                                temp_del = True
                            elif next_item.type == "del_end":
                                temp_del = False
                        j += 1

                    if next_is_redline:
                        should_defer = True

                if not should_defer:
                    meta_block = _build_merged_meta_block(deferred_meta_states, comments_map)
                    if meta_block:
                        parts.append(f"{{>>{meta_block}<<}}")
                    deferred_meta_states = []

        elif isinstance(item, DocxEvent):
            # Update State
            if item.type == "start":
                active_comments.add(item.id)
            elif item.type == "end":
                active_comments.discard(item.id)
            elif item.type == "ins_start":
                active_ins[item.id] = item
            elif item.type == "ins_end":
                active_ins.pop(item.id, None)
            elif item.type == "del_start":
                active_del[item.id] = item
            elif item.type == "del_end":
                active_del.pop(item.id, None)

    if deferred_meta_states:
        meta_block = _build_merged_meta_block(deferred_meta_states, comments_map)
        if meta_block:
            parts.append(f"{{>>{meta_block}<<}}")

    return "".join(parts)


def _get_wrappers(active_ins, active_del, active_comments):
    if active_del:
        return "{--", "--}"
    elif active_ins:
        return "{++", "++}"
    elif active_comments:
        return "{==", "==}"
    return "", ""


def _build_merged_meta_block(states_list, comments_map) -> str:
    """
    Combines metadata from multiple states, removing duplicates.
    Canonical Order: Changes first, then Comments.
    """
    change_lines = []
    comment_lines = []
    seen_sigs = set()

    for ins_map, del_map, comments_set in states_list:
        # 1. Changes (Ins & Del)
        for map_obj in (ins_map, del_map):
            for uid, meta in map_obj.items():
                sig = f"Chg:{uid}"
                if sig not in seen_sigs:
                    auth = meta.author or "Unknown"
                    change_lines.append(f"[{sig}] {auth}")
                    seen_sigs.add(sig)

        # 2. Comments
        for c_id in sorted(comments_set):
            if c_id not in comments_map:
                continue
            sig = f"Com:{c_id}"
            if sig not in seen_sigs:
                data = comments_map[c_id]
                header = f"[{sig}] {data['author']}"
                if data["date"]:
                    header += f" @ {data['date'].split('T')[0]}"
                comment_lines.append(f"{header}: {data['text']}")
                seen_sigs.add(sig)

    # Return Changes first, then Comments
    return "\n".join(change_lines + comment_lines)
