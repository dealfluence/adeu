# FILE: repro_bug2.py
"""
Reproduction script for Bug 2: outline detector misses custom heading styles
whose name doesn't START with "Heading" but contains "Heading2" (or similar)
somewhere in the name. This pattern is common when a user does
"Save Selection as a New Quick Style" in Word — Word generates a name like
'StyleHeading2NotItalicBefore0ptAfter0ptLinespa' that is *based on* Heading 2
but has a mangled name.

This script:
  1. Builds a minimal DOCX with a custom paragraph style whose:
       - styleId = 'StyleHeading2NotItalicBefore0ptAfter0ptLinespa'
       - w:name  = same as styleId
       - basedOn = (NOT SET — to mirror the test exactly)
       - no explicit outlineLvl
  2. Builds a SECOND DOCX where the same custom style explicitly basedOn
     "Heading2" — to see what get_paragraph_prefix does in that scenario.
  3. Dumps the resolved style cache for the doc so we can see exactly what
     get_paragraph_prefix sees.
  4. Runs get_paragraph_prefix on the lone paragraph in each doc and prints
     the result.
"""

import io
import zipfile

from docx import Document

from adeu.utils.docx import _get_style_cache, get_paragraph_prefix

W_NS_DECL = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _wrap_docx(document_xml: str, styles_xml: str) -> io.BytesIO:
    rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rId1"'
        b' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"'
        b' Target="word/document.xml"/>'
        b"</Relationships>"
    )
    doc_rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rId1"'
        b' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles"'
        b' Target="styles.xml"/>'
        b"</Relationships>"
    )
    ct = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        b'<Default Extension="rels"'
        b' ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        b'<Override PartName="/word/document.xml"'
        b' ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        b'<Override PartName="/word/styles.xml"'
        b' ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        b"</Types>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("[Content_Types].xml", ct)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("word/document.xml", document_xml.encode("utf-8"))
        z.writestr("word/styles.xml", styles_xml.encode("utf-8"))
    buf.seek(0)
    return buf


def section(title):
    print()
    print("=" * 72)
    print(title)
    print("=" * 72)


def case_a_no_basedon():
    """Test scenario as written: no basedOn, no outlineLvl."""
    style_name = "StyleHeading2NotItalicBefore0ptAfter0ptLinespa"
    styles_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:styles {W_NS_DECL}>"
        f'<w:style w:type="paragraph" w:styleId="{style_name}">'
        f'<w:name w:val="{style_name}"/>'
        f"</w:style>"
        f"</w:styles>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f'<w:p><w:pPr><w:pStyle w:val="{style_name}"/></w:pPr>'
        f"<w:r><w:t>Sub Heading</w:t></w:r></w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    return (
        Document(_wrap_docx(document_xml, styles_xml)),
        style_name,
        "no basedOn, no outlineLvl",
    )


def case_b_with_basedon():
    """Same custom style but explicitly basedOn 'Heading2' (Word's typical export)."""
    style_name = "StyleHeading2NotItalicBefore0ptAfter0ptLinespa"
    styles_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:styles {W_NS_DECL}>"
        # Define Heading2 itself with outlineLvl=1
        f'<w:style w:type="paragraph" w:styleId="Heading2">'
        f'<w:name w:val="heading 2"/>'
        f'<w:pPr><w:outlineLvl w:val="1"/></w:pPr>'
        f"</w:style>"
        # Custom style basedOn Heading2
        f'<w:style w:type="paragraph" w:styleId="{style_name}">'
        f'<w:name w:val="{style_name}"/>'
        f'<w:basedOn w:val="Heading2"/>'
        f"</w:style>"
        f"</w:styles>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<w:document {W_NS_DECL}><w:body>"
        f'<w:p><w:pPr><w:pStyle w:val="{style_name}"/></w:pPr>'
        f"<w:r><w:t>Sub Heading</w:t></w:r></w:p>"
        f"<w:sectPr/></w:body></w:document>"
    )
    return (
        Document(_wrap_docx(document_xml, styles_xml)),
        style_name,
        "basedOn Heading2 (no outlineLvl on child)",
    )


def dump_case(label, doc, style_name):
    section(label)
    p = doc.paragraphs[0]
    cache, default_pstyle = _get_style_cache(p.part)
    print(f"default_pstyle = {default_pstyle!r}")
    print(f"style_cache size = {len(cache)}")
    for sid, info in cache.items():
        print(f"  styleId={sid!r}  -> {info}")

    print()
    print(f"Looking up style {style_name!r} in cache:")
    info = cache.get(style_name)
    print(f"  cache.get(...) = {info}")

    print()
    prefix = get_paragraph_prefix(p)
    print(f"get_paragraph_prefix() returned: {prefix!r}")
    expected = "## "
    marker = "PASS" if prefix == expected else "FAIL"
    print(f"  [{marker}] expected {expected!r}")


def main():
    doc_a, style_a, desc_a = case_a_no_basedon()
    dump_case(f"CASE A — {desc_a}", doc_a, style_a)

    doc_b, style_b, desc_b = case_b_with_basedon()
    dump_case(f"CASE B — {desc_b}", doc_b, style_b)


if __name__ == "__main__":
    main()
