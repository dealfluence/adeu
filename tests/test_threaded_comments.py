import io
import re

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import DocumentEdit, ReviewAction
from adeu.redline.engine import RedlineEngine


def test_threaded_comment_structure():
    """
    Verifies that replying to a comment creates a valid threaded structure
    compatible with Word (w15:p attribute and w15 namespace declaration).
    """
    doc = Document()
    doc.add_paragraph("Threaded conversation anchor.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Create Parent Comment (Force edit to ensure comment is attached)
    engine = RedlineEngine(stream, author="UserA")
    # Change "anchor" -> "Anchor" to force a tracked change with comment
    edit = DocumentEdit(target_text="anchor", new_text="Anchor", comment="Parent Topic")
    engine.apply_edits([edit])
    stream_mid = engine.save_to_stream()

    # Verify Parent Exists
    text_mid = extract_text_from_stream(stream_mid)
    match = re.search(r"\[Com:(\d+)\]", text_mid)
    assert match, f"Parent comment not created. Text: {text_mid}"
    parent_id = match.group(1)

    # 2. Create Reply (Action: REPLY)
    engine2 = RedlineEngine(stream_mid, author="UserB")
    action = ReviewAction(action="REPLY", target_id=f"Com:{parent_id}", text="Reply Content")
    applied, skipped = engine2.apply_review_actions([action])

    assert applied == 1
    assert skipped == 0

    stream_final = engine2.save_to_stream()

    # 3. XML Inspection (The Critical Check)
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc_final = Document(stream_final)
    comments_part = None
    for rel in doc_final.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break

    assert comments_part, "Comments part missing"

    xml = comments_part.blob.decode("utf-8")

    # Check 1: Namespace Declaration in Root
    assert 'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"' in xml, (
        "Missing w15 namespace declaration in comments.xml root"
    )

    # Check 2: Parent Attribute on the Reply
    # Search for w15:p="{parent_id}"
    expected_attr = f'w15:p="{parent_id}"'
    assert expected_attr in xml, f"Missing threaded attribute {expected_attr} in XML\nXML: {xml}"

    # 4. Ingestion Inspection
    text_final = extract_text_from_stream(stream_final)

    assert "Parent Topic" in text_final
    assert "Reply Content" in text_final

    sigs = re.findall(r"\[Com:(\d+)\]", text_final)
    # Depending on ingest implementation, reply might reuse parent ID in signature or have its own.
    # Current implementation gives every comment its own ID signature.
    assert len(set(sigs)) == 2, f"Should have 2 unique comments visible. Found: {sigs}"


def test_threaded_rendering_order():
    """
    Ensures that when ingesting text, replies are rendered correctly.
    """
    doc = Document()
    doc.add_paragraph("Target")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Setup Engine with forced edit
    engine = RedlineEngine(stream, author="A")
    engine.apply_edits([DocumentEdit(target_text="Target", new_text="TargetModified", comment="Root")])

    # Get ID of Root
    data = engine.comments_manager.extract_comments_data()
    assert data, "Root comment not created"
    root_id = list(data.keys())[0]

    # Reply 1
    engine.author = "B"
    engine.apply_review_actions([ReviewAction(action="REPLY", target_id=f"Com:{root_id}", text="Reply1")])

    # Get ID of Reply1
    data = engine.comments_manager.extract_comments_data()
    # Find key that isn't root_id
    others = [k for k in data.keys() if k != root_id]
    assert others, "Reply1 not created"
    reply1_id = others[0]

    # Reply 2 (Reply to Reply1)
    engine.author = "C"
    engine.apply_review_actions([ReviewAction(action="REPLY", target_id=f"Com:{reply1_id}", text="Reply2")])

    stream_final = engine.save_to_stream()
    text = extract_text_from_stream(stream_final)

    # Check that all 3 appear
    assert "Root" in text
    assert "Reply1" in text
    assert "Reply2" in text
