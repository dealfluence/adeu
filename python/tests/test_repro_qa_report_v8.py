# FILE: tests/test_repro_qa_report_v8.py
"""
Repro tests for the 2026-07-19 black-box QA and UX report on 1.26.0
(adeu 1.26.0+0741eaf, "Adeu CLI 1.26.0 Black-Box QA and UX Report").

Finding index (report finding -> test class below):
  F-01  `sanitize --baseline` with a clearly unrelated baseline REPLACES the
        working document with the baseline's content, exits 0 and prints
        "Result: CLEAN" — the 5%-shared warning does not block anything
  F-02  a failed all-or-nothing sanitize batch leaves a valid
        `.<name>.staging.tmp` DOCX (full document body) in --outdir
  F-03  a change object with no `type` field is silently applied as a
        modify by the CLI, although the CLI documents `type` as required
  F-04  replacing a full sentence that crosses bold/italic formatting runs
        leaves a partial word bold: `**The Suppli** must perform ...`
  F-05  `-o -` creates a literal file named '-' instead of streaming to
        stdout (extract, markup)
  F-06  `adeu-server --help` / `--version` start the stdio server instead
        of printing help/version and exiting
  F-07  review-action validation permits blank replies and double-counts
        duplicate accept actions ("Actions: 2 applied" for one change)
  F-08  `markup --json` prints human success/status lines on stderr
  F-09  sanitize exit codes are inconsistent (missing input exits 2) and a
        missing output parent surfaces a raw `Errno 2 ... .tmp` message
  F-10  search snippet highlighting collides with existing style markers:
        `**The **Supplier** _shall provide**_`
  F-11  the default tracked-changes author is the OS username, which is
        `root` in containers — customer-visible in outbound documents
  F-12  the major-deletion guard is disarmed below 2,000 characters: a
        short contract can be truncated to its title without any flag
  F-13  `adeu help` unsupported; `--debug` only accepted before the
        subcommand; output parent directories not created

Every test fails against the commit preceding its fix.
"""

import json
import subprocess
import sys
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from adeu.ingest import _extract_text_from_doc
from adeu.models import AcceptChange, ModifyText, RejectChange, ReplyComment
from adeu.redline.engine import BatchValidationError, RedlineEngine
from adeu.sanitize.core import SanitizeError, sanitize_docx

# ---------------------------------------------------------------------------
# Helpers / fixture builders
# ---------------------------------------------------------------------------


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from unittest.mock import patch

    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


def doc_to_stream(doc) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_docx(paragraphs, path: Path = None):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if path is not None:
        doc.save(path)
    return doc


def clean_text_of(doc) -> str:
    return _extract_text_from_doc(doc, clean_view=True, include_appendix=False)


def build_cross_run_doc():
    """The report's F-04 fixture: bold 'The Supplier ', italic 'shall
    provide ', underlined remainder — one visible sentence across three
    formatting runs."""
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("The Supplier ")
    r1.bold = True
    r2 = p.add_run("shall provide ")
    r2.italic = True
    r3 = p.add_run("the Services with reasonable skill and care.")
    r3.underline = True
    return doc


def build_tracked_change_stream() -> BytesIO:
    """A document carrying exactly one tracked modification (Chg:1 del +
    Chg:2 ins)."""
    doc = build_docx(["Payment is due in 30 days.", "Second paragraph here."])
    engine = RedlineEngine(doc_to_stream(doc), author="Editor")
    engine.apply_edits([ModifyText(target_text="30 days", new_text="60 days")])
    return engine.save_to_stream()


def build_comment_stream() -> BytesIO:
    """A document carrying exactly one comment (Com:0)."""
    doc = build_docx(["Alpha beta gamma.", "Delta epsilon."])
    engine = RedlineEngine(doc_to_stream(doc), author="Editor")
    engine.process_batch(
        [ModifyText(target_text="Alpha beta gamma.", new_text="Alpha beta gamma.", comment="Please review.")]
    )
    return engine.save_to_stream()


def first_comment_id(stream: BytesIO) -> str:
    from adeu.redline.comments import CommentsManager

    doc = Document(BytesIO(stream.getvalue()))
    data = CommentsManager(doc).extract_comments_data()
    assert data, "expected the fixture to carry a comment"
    return next(iter(data.keys()))


WORKING_PARAGRAPHS = [
    "Master Services Agreement",
    "This Agreement is entered into by Supplier and Customer for the provision of professional services.",
    "1. Services. The Supplier shall provide the Services with reasonable skill and care.",
    "2. Fees. Customer shall pay the fees set out in Schedule 1 within thirty (30) days of invoice.",
    "3. Term. This Agreement commences on the Effective Date and continues for two (2) years.",
    "4. Liability. Neither party's liability shall exceed the fees paid in the preceding twelve months.",
]

UNRELATED_BASELINE_PARAGRAPHS = [
    "Simple Contract",
    "Alpha beta gamma.",
    "Payment is due in 30 days.",
]


# ---------------------------------------------------------------------------
# F-01: a wrong sanitize baseline must block, not silently replace the doc
# ---------------------------------------------------------------------------


class TestF01WrongBaselineBlocks:
    def test_low_similarity_baseline_blocks_and_writes_nothing(self, tmp_path):
        working = tmp_path / "msa.docx"
        baseline = tmp_path / "unrelated.docx"
        build_docx(WORKING_PARAGRAPHS, working)
        build_docx(UNRELATED_BASELINE_PARAGRAPHS, baseline)
        out = tmp_path / "out.docx"

        with pytest.raises(SanitizeError) as exc_info:
            sanitize_docx(str(working), str(out), baseline_path=str(baseline))

        message = str(exc_info.value)
        assert "baseline" in message.lower()
        assert "--allow-low-similarity-baseline" in message
        assert not out.exists(), "a blocked baseline sanitize must write nothing"

    def test_cli_wrong_baseline_exits_nonzero_and_keeps_document(self, tmp_path, capsys):
        working = tmp_path / "msa.docx"
        baseline = tmp_path / "unrelated.docx"
        build_docx(WORKING_PARAGRAPHS, working)
        build_docx(UNRELATED_BASELINE_PARAGRAPHS, baseline)
        out = tmp_path / "wrong-baseline.docx"

        code, _stdout, stderr = run_cli(["sanitize", working, "-o", out, "--baseline", baseline, "--report"], capsys)

        assert code != 0, "a wrong-baseline sanitize must not exit 0"
        assert not out.exists(), "a wrong-baseline sanitize must not write output"
        assert "CLEAN" not in stderr.replace("BLOCKED", ""), "must not report a clean result"

    def test_explicit_override_flag_proceeds_with_warning(self, tmp_path, capsys):
        working = tmp_path / "msa.docx"
        baseline = tmp_path / "unrelated.docx"
        build_docx(WORKING_PARAGRAPHS, working)
        build_docx(UNRELATED_BASELINE_PARAGRAPHS, baseline)
        out = tmp_path / "forced.docx"

        code, _stdout, _stderr = run_cli(
            [
                "sanitize",
                working,
                "-o",
                out,
                "--baseline",
                baseline,
                "--allow-low-similarity-baseline",
            ],
            capsys,
        )

        assert code == 0
        assert out.exists()

    def test_related_baseline_still_works_without_flag(self, tmp_path):
        original = tmp_path / "v1.docx"
        edited = tmp_path / "v2.docx"
        build_docx(WORKING_PARAGRAPHS, original)
        modified = list(WORKING_PARAGRAPHS)
        modified[3] = "2. Fees. Customer shall pay the fees set out in Schedule 1 within sixty (60) days of invoice."
        build_docx(modified, edited)
        out = tmp_path / "redline.docx"

        result = sanitize_docx(str(edited), str(out), baseline_path=str(original))
        assert out.exists()
        assert result.status in ("clean", "clean_with_warnings")


# ---------------------------------------------------------------------------
# F-02: failed atomic sanitize batch must leave the outdir untouched
# ---------------------------------------------------------------------------


class TestF02BatchStagingCleanup:
    def test_failed_batch_leaves_no_staging_files(self, tmp_path, capsys):
        good = tmp_path / "simple.docx"
        build_docx(["Alpha beta gamma.", "Payment is due in 30 days."], good)
        disguised = tmp_path / "text-disguised.docx"
        disguised.write_text("This is not a DOCX file at all.", encoding="utf-8")
        outdir = tmp_path / "out"

        code, _stdout, _stderr = run_cli(["sanitize", good, disguised, "--outdir", outdir, "--report"], capsys)

        assert code != 0
        leftovers = sorted(p.name for p in outdir.iterdir()) if outdir.exists() else []
        assert leftovers == [], (
            f"'if any input is blocked or fails, NO outputs are written' — but the outdir contains {leftovers}"
        )

    def test_successful_batch_still_commits_outputs(self, tmp_path, capsys):
        a = tmp_path / "a.docx"
        b = tmp_path / "b.docx"
        build_docx(["Doc A content here."], a)
        build_docx(["Doc B content here."], b)
        outdir = tmp_path / "out"

        code, _stdout, _stderr = run_cli(["sanitize", a, b, "--outdir", outdir], capsys)

        assert code == 0
        assert sorted(p.name for p in outdir.iterdir()) == ["a.docx", "b.docx"]


# ---------------------------------------------------------------------------
# F-03: a change object without `type` is rejected by the CLI
# ---------------------------------------------------------------------------


class TestF03MissingTypeRejected:
    def test_cli_apply_rejects_missing_type(self, tmp_path, capsys):
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        changes = tmp_path / "changes.json"
        changes.write_text(json.dumps([{"target_text": "Alpha", "new_text": "Omega"}]), encoding="utf-8")
        out = tmp_path / "out.docx"

        code, _stdout, stderr = run_cli(["apply", doc_path, changes, "-o", out], capsys)

        assert code != 0, "the CLI documents 'type' as required on every change"
        assert "'type'" in stderr or "type" in stderr
        assert not out.exists(), "a rejected batch must write no output"

    def test_cli_markup_rejects_missing_type(self, tmp_path, capsys):
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        changes = tmp_path / "changes.json"
        changes.write_text(json.dumps([{"target_text": "Alpha", "new_text": "Omega"}]), encoding="utf-8")
        out = tmp_path / "out.md"

        code, _stdout, _stderr = run_cli(["markup", doc_path, changes, "-o", out], capsys)

        assert code != 0
        assert not out.exists()

    def test_mcp_boundary_keeps_documented_unambiguous_inference(self):
        """The MCP schema documents inference for unambiguous payloads; the
        CLI strictness must not regress that tolerance layer."""
        from pydantic import TypeAdapter

        from adeu.models import BatchChanges

        changes = TypeAdapter(BatchChanges).validate_python([{"target_text": "Alpha", "new_text": "Omega"}])
        assert len(changes) == 1
        assert changes[0].type == "modify"


# ---------------------------------------------------------------------------
# F-04: full-span replacements across formatting runs keep words whole
# ---------------------------------------------------------------------------


class TestF04CrossRunFormatting:
    TARGET = "The Supplier shall provide the Services with reasonable skill and care."
    NEW = "The Supplier must perform the Services professionally."

    def test_replacement_leaves_no_partial_word_bold(self):
        doc = build_cross_run_doc()
        engine = RedlineEngine(doc_to_stream(doc), author="QA")
        stats = engine.process_batch([ModifyText(target_text=self.TARGET, new_text=self.NEW)])
        assert stats["edits_applied"] == 1

        clean = clean_text_of(engine.doc)
        assert "**The Suppli**" not in clean
        assert "Suppli**" not in clean, f"partial-word bold artifact in {clean!r}"
        assert clean.strip() == "**The Supplier** must perform the Services professionally."

    def test_accepted_document_matches_clean_view(self):
        doc = build_cross_run_doc()
        engine = RedlineEngine(doc_to_stream(doc), author="QA")
        engine.process_batch([ModifyText(target_text=self.TARGET, new_text=self.NEW)])
        engine.accept_all_revisions(remove_comments=True)
        accepted = _extract_text_from_doc(engine.doc, clean_view=False, include_appendix=False)
        assert accepted.strip() == "**The Supplier** must perform the Services professionally."


# ---------------------------------------------------------------------------
# F-05: `-o -` streams to stdout, never creates a file named '-'
# ---------------------------------------------------------------------------


class TestF05DashOutputMeansStdout:
    def test_extract_dash_output_streams_to_stdout(self, tmp_path, capsys, monkeypatch):
        monkeypatch.chdir(tmp_path)
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)

        code, stdout, _stderr = run_cli(["extract", doc_path, "-o", "-"], capsys)

        assert code == 0
        assert "Alpha beta gamma." in stdout
        assert not (tmp_path / "-").exists(), "'-' must mean stdout, not a literal file"

    def test_markup_dash_output_streams_to_stdout(self, tmp_path, capsys, monkeypatch):
        monkeypatch.chdir(tmp_path)
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        edits = tmp_path / "edits.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "Alpha", "new_text": "Omega"}]),
            encoding="utf-8",
        )

        code, stdout, _stderr = run_cli(["markup", doc_path, edits, "-o", "-"], capsys)

        assert code == 0
        assert "{--Alpha--}" in stdout and "{++Omega++}" in stdout
        assert not (tmp_path / "-").exists()

    def test_markup_dash_output_with_json_embeds_content(self, tmp_path, capsys, monkeypatch):
        monkeypatch.chdir(tmp_path)
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        edits = tmp_path / "edits.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "Alpha", "new_text": "Omega"}]),
            encoding="utf-8",
        )

        code, stdout, _stderr = run_cli(["markup", doc_path, edits, "-o", "-", "--json"], capsys)

        assert code == 0
        payload = json.loads(stdout)
        assert payload["status"] == "ok"
        assert "{++Omega++}" in payload["content"]
        assert not (tmp_path / "-").exists()


# ---------------------------------------------------------------------------
# F-06: adeu-server --help/--version print and exit, never start the server
# ---------------------------------------------------------------------------


class TestF06ServerHelpVersion:
    @pytest.mark.parametrize("flag", ["--version", "--help"])
    def test_server_flag_prints_and_exits(self, flag):
        proc = subprocess.Popen(
            [sys.executable, "-m", "adeu.server", flag],
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        try:
            stdout, _stderr = proc.communicate(timeout=60)
        except subprocess.TimeoutExpired:
            proc.kill()
            proc.communicate()
            pytest.fail(f"adeu-server {flag} started the stdio server instead of exiting")

        assert proc.returncode == 0
        text = stdout.decode("utf-8", errors="replace")
        if flag == "--version":
            from adeu.mcp_components.shared import get_build_info

            version, _sha, _ = get_build_info()
            assert version in text
        else:
            assert "usage" in text.lower()
            assert "--scope" in text


# ---------------------------------------------------------------------------
# F-07: blank replies and duplicate review actions are rejected
# ---------------------------------------------------------------------------


class TestF07ReviewActionValidation:
    def test_blank_reply_is_rejected(self):
        stream = build_comment_stream()
        com_id = first_comment_id(stream)
        engine = RedlineEngine(stream)

        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([ReplyComment(target_id=f"Com:{com_id}", text="   ")])
        assert "empty" in "\n".join(exc_info.value.errors).lower()

    def test_duplicate_accept_is_rejected(self):
        stream = build_tracked_change_stream()
        engine = RedlineEngine(stream)

        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([AcceptChange(target_id="Chg:1"), AcceptChange(target_id="Chg:1")])
        assert "duplicate" in "\n".join(exc_info.value.errors).lower()

    def test_conflicting_accept_and_reject_are_rejected(self):
        stream = build_tracked_change_stream()
        engine = RedlineEngine(stream)

        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch([AcceptChange(target_id="Chg:1"), RejectChange(target_id="Chg:1")])
        assert "conflict" in "\n".join(exc_info.value.errors).lower()

    def test_duplicate_identical_reply_is_rejected(self):
        stream = build_comment_stream()
        com_id = first_comment_id(stream)
        engine = RedlineEngine(stream)

        with pytest.raises(BatchValidationError) as exc_info:
            engine.process_batch(
                [
                    ReplyComment(target_id=f"Com:{com_id}", text="Same reply."),
                    ReplyComment(target_id=f"Com:{com_id}", text="Same reply."),
                ]
            )
        assert "duplicate" in "\n".join(exc_info.value.errors).lower()

    def test_distinct_replies_to_same_comment_stay_allowed(self):
        stream = build_comment_stream()
        com_id = first_comment_id(stream)
        engine = RedlineEngine(stream)

        stats = engine.process_batch(
            [
                ReplyComment(target_id=f"Com:{com_id}", text="First reply."),
                ReplyComment(target_id=f"Com:{com_id}", text="Second reply."),
            ]
        )
        assert stats["actions_applied"] == 2
        assert stats["actions_skipped"] == 0

    def test_accepting_both_ids_of_one_modification_stays_allowed(self):
        """Accepting the del+ins pair of one logical modification via its two
        distinct Chg ids is the documented workflow and must keep working."""
        stream = build_tracked_change_stream()
        engine = RedlineEngine(stream)

        stats = engine.process_batch([AcceptChange(target_id="Chg:1"), AcceptChange(target_id="Chg:2")])
        assert stats["actions_skipped"] == 0
        assert "60 days" in clean_text_of(engine.doc)


# ---------------------------------------------------------------------------
# F-08: markup --json keeps stderr free of decorative success lines
# ---------------------------------------------------------------------------


class TestF08MarkupJsonMachineClean:
    def test_markup_json_success_has_clean_streams(self, tmp_path, capsys):
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        edits = tmp_path / "edits.json"
        edits.write_text(
            json.dumps([{"type": "modify", "target_text": "Alpha", "new_text": "Omega"}]),
            encoding="utf-8",
        )
        out = tmp_path / "out.md"

        code, stdout, stderr = run_cli(["markup", doc_path, edits, "-o", out, "--json"], capsys)

        assert code == 0
        payload = json.loads(stdout)  # stdout is exactly one JSON object
        assert payload["status"] == "ok"
        assert "✅" not in stderr and "Stats:" not in stderr, (
            f"--json promises machine-clean output but stderr says: {stderr!r}"
        )


# ---------------------------------------------------------------------------
# F-09: sanitize follows the CLI error contract
# ---------------------------------------------------------------------------


class TestF09SanitizeErrorContract:
    def test_missing_input_exits_1(self, tmp_path, capsys):
        code, _stdout, stderr = run_cli(["sanitize", tmp_path / "missing.docx"], capsys)
        assert code == 1, "operational failures exit 1; argument errors exit 2"
        assert "File not found" in stderr

    def test_unwritable_output_parent_reports_cleanly(self, tmp_path, capsys):
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        blocker = tmp_path / "blocker.txt"
        blocker.write_text("not a directory", encoding="utf-8")
        out = blocker / "out.docx"  # parent is a file: mkdir/write must fail

        code, _stdout, stderr = run_cli(["sanitize", doc_path, "-o", out], capsys)

        assert code == 1
        assert "Could not write output file" in stderr
        assert ".tmp" not in stderr, f"raw temp-file internals leaked into: {stderr!r}"

    def test_missing_output_parent_is_created(self, tmp_path, capsys):
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        out = tmp_path / "nested" / "dir" / "out.docx"

        code, _stdout, _stderr = run_cli(["sanitize", doc_path, "-o", out], capsys)

        assert code == 0
        assert out.exists()


# ---------------------------------------------------------------------------
# F-10: search highlighting must not collide with existing style markers
# ---------------------------------------------------------------------------


class TestF10SearchSnippetMarkers:
    def test_snippet_emphasis_strips_existing_markers(self):
        from adeu.mcp_components._response_builders import build_search_response

        doc = build_cross_run_doc()
        text = _extract_text_from_doc(doc, clean_view=False, include_appendix=False)
        res = build_search_response(text, "Supplier.*provide", True, True, None, "doc.docx", is_cli=True)
        md = res.structured_content["markdown"]

        assert "**The **" not in md, f"nested-marker artifact in snippet: {md!r}"
        assert "**_" not in md and "_**" not in md, f"colliding markers in snippet: {md!r}"
        assert "**Supplier shall provide**" in md

    def test_plain_document_snippets_unchanged(self):
        from adeu.mcp_components._response_builders import build_search_response

        doc = build_docx(["The quick brown fox jumps over the lazy dog."])
        text = _extract_text_from_doc(doc, clean_view=False, include_appendix=False)
        res = build_search_response(text, "brown fox", False, True, None, "doc.docx", is_cli=True)
        md = res.structured_content["markdown"]
        assert "**brown fox**" in md


# ---------------------------------------------------------------------------
# F-11: machine accounts never become the visible review author
# ---------------------------------------------------------------------------


class TestF11DefaultAuthor:
    def test_machine_account_falls_back_to_neutral_default(self, monkeypatch):
        import getpass

        from adeu.cli import _default_author

        monkeypatch.delenv("ADEU_AUTHOR", raising=False)
        monkeypatch.setattr(getpass, "getuser", lambda: "root")
        assert _default_author() == "Adeu AI"

    def test_env_override_wins(self, monkeypatch):
        from adeu.cli import _default_author

        monkeypatch.setenv("ADEU_AUTHOR", "Legal Review Bot")
        assert _default_author() == "Legal Review Bot"

    def test_human_username_is_kept(self, monkeypatch):
        import getpass

        from adeu.cli import _default_author

        monkeypatch.delenv("ADEU_AUTHOR", raising=False)
        monkeypatch.setattr(getpass, "getuser", lambda: "mikko")
        assert _default_author() == "mikko"

    def test_cli_apply_honors_adeu_author_env(self, tmp_path, capsys, monkeypatch):
        monkeypatch.setenv("ADEU_AUTHOR", "Env Reviewer")
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        changes = tmp_path / "changes.json"
        changes.write_text(
            json.dumps([{"type": "modify", "target_text": "Alpha", "new_text": "Omega"}]),
            encoding="utf-8",
        )
        out = tmp_path / "out.docx"

        code, _stdout, _stderr = run_cli(["apply", doc_path, changes, "-o", out], capsys)
        assert code == 0

        with zipfile.ZipFile(out) as z:
            body = z.read("word/document.xml").decode("utf-8")
        assert 'w:author="Env Reviewer"' in body


# ---------------------------------------------------------------------------
# F-12: the deletion guard also arms on small documents
# ---------------------------------------------------------------------------


class TestF12SmallDocumentDeletionGuard:
    SMALL_PARAGRAPHS = [
        "Consulting Agreement",
        "The Consultant will provide advisory services to the Client.",
        "Fees are 100 euros per hour, invoiced monthly.",
        "Either party may terminate with 30 days written notice.",
    ]

    def _extract_clean(self, tmp_path, capsys, doc_path):
        text_file = tmp_path / "doc.txt"
        code, _stdout, _stderr = run_cli(
            ["extract", doc_path, "--clean-view", "--page", "all", "-o", text_file], capsys
        )
        assert code == 0
        return text_file

    def test_truncation_to_title_requires_flag(self, tmp_path, capsys):
        doc_path = tmp_path / "small.docx"
        build_docx(self.SMALL_PARAGRAPHS, doc_path)
        text_file = self._extract_clean(tmp_path, capsys, doc_path)
        text_file.write_text("Consulting Agreement", encoding="utf-8")
        out = tmp_path / "out.docx"

        code, _stdout, stderr = run_cli(["apply", doc_path, text_file, "-o", out], capsys)

        assert code != 0, "a small document must not be silently truncated to its title"
        assert "--allow-major-deletions" in stderr
        assert not out.exists()

    def test_flag_still_allows_the_truncation(self, tmp_path, capsys):
        doc_path = tmp_path / "small.docx"
        build_docx(self.SMALL_PARAGRAPHS, doc_path)
        text_file = self._extract_clean(tmp_path, capsys, doc_path)
        text_file.write_text("Consulting Agreement", encoding="utf-8")
        out = tmp_path / "out.docx"

        code, _stdout, _stderr = run_cli(["apply", doc_path, text_file, "-o", out, "--allow-major-deletions"], capsys)
        assert code == 0
        assert out.exists()

    def test_moderate_small_doc_edit_needs_no_flag(self, tmp_path, capsys):
        doc_path = tmp_path / "small.docx"
        build_docx(self.SMALL_PARAGRAPHS, doc_path)
        text_file = self._extract_clean(tmp_path, capsys, doc_path)
        # Delete one of four paragraphs (~25% of the text): plausible edit,
        # must not demand the flag.
        content = text_file.read_text(encoding="utf-8")
        content = content.replace("Fees are 100 euros per hour, invoiced monthly.\n\n", "")
        text_file.write_text(content, encoding="utf-8")
        out = tmp_path / "out.docx"

        code, _stdout, _stderr = run_cli(["apply", doc_path, text_file, "-o", out], capsys)
        assert code == 0
        assert out.exists()


# ---------------------------------------------------------------------------
# Hunt-profile counterexample found while fixing F-12's diff-separator issue:
# a paragraph-splitting insertion at paragraph START must relocate the host
# paragraph's content into the LAST new paragraph (deterministic pin per the
# property-test convention; pre-existing on 1.26.0, both engines).
# ---------------------------------------------------------------------------


class TestParagraphStartSplitInsertion:
    def test_pinned_counterexample_00_dot(self):
        from adeu.diff import generate_edits_via_paragraph_alignment
        from adeu.ingest import extract_text_from_stream

        doc = build_docx(["00."])
        stream = doc_to_stream(doc)
        text_orig = extract_text_from_stream(BytesIO(stream.getvalue()), clean_view=True)
        text_mod = "0.\n\n0 00."

        edits = generate_edits_via_paragraph_alignment(text_orig, text_mod)
        engine = RedlineEngine(BytesIO(stream.getvalue()), author="Fuzz")
        stats = engine.process_batch(list(edits))
        assert stats["edits_skipped"] == 0

        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == text_mod

    def test_prepending_a_paragraph_via_text_apply(self, tmp_path, capsys):
        """The natural-language shape of the same bug: insert a new first
        paragraph via the extract → edit → apply workflow."""
        doc_path = tmp_path / "doc.docx"
        build_docx(["Original body paragraph."], doc_path)
        text_file = tmp_path / "doc.txt"
        code, _stdout, _stderr = run_cli(
            ["extract", doc_path, "--clean-view", "--page", "all", "-o", text_file], capsys
        )
        assert code == 0
        content = text_file.read_text(encoding="utf-8")
        # Edit below extract's "> **File Path:**" header line, like a user
        # editing the document text would.
        header, _, body = content.partition("\n\n")
        text_file.write_text(f"{header}\n\nNew opening paragraph.\n\n{body}", encoding="utf-8")
        out = tmp_path / "out.docx"

        code, _stdout, stderr = run_cli(["apply", doc_path, text_file, "-o", out, "--json"], capsys)
        assert code == 0, f"prepend apply failed: {stderr}"

        from adeu.ingest import extract_text_from_stream

        final = extract_text_from_stream(BytesIO(out.read_bytes()), clean_view=True)
        assert final.strip() == "New opening paragraph.\n\nOriginal body paragraph."


# ---------------------------------------------------------------------------
# F-13: shell-convention rough edges
# ---------------------------------------------------------------------------


class TestF13ShellConventions:
    def test_adeu_help_prints_usage(self, capsys):
        code, stdout, _stderr = run_cli(["help"], capsys)
        assert code == 0
        assert "usage" in stdout.lower()

    def test_adeu_help_subcommand(self, capsys):
        code, stdout, _stderr = run_cli(["help", "extract"], capsys)
        assert code == 0
        assert "--clean-view" in stdout

    def test_debug_flag_works_after_subcommand(self, tmp_path, capsys):
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)

        code, stdout, _stderr = run_cli(["extract", doc_path, "--debug"], capsys)
        assert code == 0
        assert "Alpha beta gamma." in stdout

    def test_output_parent_directories_are_created(self, tmp_path, capsys):
        doc_path = tmp_path / "doc.docx"
        build_docx(["Alpha beta gamma."], doc_path)
        out = tmp_path / "deep" / "nested" / "out.txt"

        code, _stdout, _stderr = run_cli(["extract", doc_path, "-o", out], capsys)
        assert code == 0
        assert out.exists()
