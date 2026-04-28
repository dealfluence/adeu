from docx import Document

from adeu.domain import build_structural_appendix


def test_defined_terms_extraction_and_diagnostics():
    doc = Document()

    # 1. Glossary Definitions
    doc.add_heading("Definitions", 1)
    doc.add_paragraph('"Agreement" means this contract.')
    doc.add_paragraph("“Party” shall mean either side.")
    doc.add_paragraph('"Agreement" means another thing.')  # Duplicate

    # 2. Inline Definitions
    doc.add_paragraph('This contract (hereinafter, the "Contract") is valid.')

    # 3. Multilingual Typographic Definitions
    doc.add_heading("Miscellaneous", 1)
    doc.add_paragraph('"Confidential Information" on salainen asia.')  # Finnish
    doc.add_paragraph('1.1 "Affiliate" tarkoittaa osakkuusyhtiötä.')  # Finnish + Numbered List
    doc.add_paragraph('We will act as the disclosing party (jäljempänä "Discloser").')  # Multilingual Inline

    # 4. Phantom avoidance (Syntax examples)
    doc.add_paragraph('This is a syntax example: ("Heading*") and ("<Term>")')

    # 5. Usages, Typos
    # "Agreement" is used 1 time (excluding declarations).
    # "Party" is used 0 times -> Should be pruned from index.
    # "Contract" is used 1 time.
    # "Agrement" is a typo of "Agreement".
    doc.add_paragraph("The Agreement is binding. The Contract is signed.")
    doc.add_paragraph("There is an Agrement here.")
    # Usages added here so the new validation rule (usages > 0) keeps them in the index!
    doc.add_paragraph("We shared Confidential Information with the Affiliate. The Discloser is happy.")

    # Build base_text mimicking text ingest
    base_text = "\n".join(p.text for p in doc.paragraphs)

    appendix = build_structural_appendix(doc, base_text)

    # Check Definitions Symbol Table
    assert '"Agreement" — used' in appendix
    assert '"Contract" — used' in appendix
    assert '"Confidential Information" — used' in appendix
    assert '"Affiliate" — used' in appendix
    assert '"Discloser" — used' in appendix

    # Ensure pruned logic works (unused terms removed)
    assert '"Party"' not in appendix

    # Ensure phantom syntax terms are strictly ignored
    assert '"Heading*"' not in appendix
    assert '"<Term>"' not in appendix

    # Check Semantic Diagnostics Linter Rules
    assert "[Error] Duplicate Definition: 'Agreement' is defined multiple times." in appendix
    assert "[Info] Possible Typos for 'Agreement': Found 'Agrement'" in appendix


def test_acronym_typo_noise_reduction():
    """
    Ensures short acronyms don't trigger false positive typos due to Levenshtein proximity.
    (e.g., 'GPUs' vs 'PSUs', 'CPU' vs 'GPU'). Validates the strict length/first-letter heuristic.
    """
    doc = Document()
    doc.add_heading("Definitions", 1)
    doc.add_paragraph('"PSUs" means power supply units.')
    doc.add_paragraph('"CPU" means central processing unit.')
    doc.add_paragraph('"Party" means the entity.')

    # Body with acronyms that are close in edit distance but differ in first letter or distance > 1
    doc.add_paragraph("We rely on ESAs, LSPs, and GPUs for the servers.")
    doc.add_paragraph("The GPU is very fast.")

    # Valid short typo (dist 1, same first letter, must be >= 4 chars to bypass the absolute min length rule)
    doc.add_paragraph("The Pary signed the contract.")

    # Terms must be used at least once to avoid being pruned
    doc.add_paragraph("We bought PSUs and a CPU.")
    doc.add_paragraph("The Party begins today.")

    base_text = "\n".join(p.text for p in doc.paragraphs)
    appendix = build_structural_appendix(doc, base_text)

    # 1. Valid short typo SHOULD be detected
    assert "[Info] Possible Typos for 'Party': Found 'Pary'" in appendix

    # 2. Acronym noise SHOULD BE suppressed
    # 'GPU' (dist 1 from 'CPU') -> Rejected because first letter differs
    assert "'GPU'" not in appendix
    # 'GPUs' (dist 2 from 'PSUs') -> Rejected because dist > 1 for short words
    assert "'GPUs'" not in appendix
    assert "'ESAs'" not in appendix
    assert "'LSPs'" not in appendix
