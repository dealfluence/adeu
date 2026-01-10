"""
Low-level utilities for manipulating DOCX XML structures.
Contains normalization logic ported from Open-Xml-PowerTools concepts.
"""
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import structlog

logger = structlog.get_logger(__name__)

def create_element(name: str):
    return OxmlElement(name)

def create_attribute(element, name: str, value: str):
    element.set(qn(name), value)

def get_visible_runs(paragraph: Paragraph):
    """
    Iterates over runs in a paragraph, including those inside <w:ins> tags.
    Effectively returns the 'Accepted Changes' view of the runs.
    """
    runs = []
    # Iterate over all children of the paragraph XML element
    for child in paragraph._element:
        tag = child.tag
        if tag == qn('w:r'):
            # Standard run
            runs.append(Run(child, paragraph))
        elif tag == qn('w:ins'):
            # Inserted runs (Track Changes)
            for subchild in child:
                if subchild.tag == qn('w:r'):
                    runs.append(Run(subchild, paragraph))
        # w:del is skipped implies we read the "Future" state (Deletions are gone)
        # w:hyperlink could be added here if needed, but skipping for MVP
        
    return runs

def _are_runs_identical(r1: Run, r2: Run) -> bool:
    """
    Compares two runs to see if they have identical formatting properties.
    """
    rPr1 = r1._r.rPr
    rPr2 = r2._r.rPr

    xml1 = rPr1.xml if rPr1 is not None else ""
    xml2 = rPr2.xml if rPr2 is not None else ""

    return xml1 == xml2

def _coalesce_runs_in_paragraph(paragraph: Paragraph):
    """
    Merges adjacent runs with identical formatting.
    This fixes issues where words are split like ["Con", "tract"] due to editing history.
    """
    i = 0
    # Safe iteration while modifying the list
    while i < len(paragraph.runs) - 1:
        current_run = paragraph.runs[i]
        next_run = paragraph.runs[i + 1]

        if _are_runs_identical(current_run, next_run):
            # Merge content
            current_run.text += next_run.text
            # Remove next_run from the XML tree
            paragraph._p.remove(next_run._r)
            # Do NOT increment i; check the *new* next_run against current_run
        else:
            i += 1

def normalize_docx(doc: DocumentObject):
    """
    Applies normalization to a DOCX document to make text mapping reliable.
    1. Removes proof errors (spellcheck squiggles).
    2. Coalesces adjacent runs.
    """
    logger.info("Normalizing DOCX structure...")
    
    # Remove proof errors (spelling/grammar tags) via XPath
    for proof_err in doc.element.xpath("//w:proofErr"):
        proof_err.getparent().remove(proof_err)

    # Coalesce body paragraphs
    for p in doc.paragraphs:
        _coalesce_runs_in_paragraph(p)

    # Coalesce table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _coalesce_runs_in_paragraph(p)