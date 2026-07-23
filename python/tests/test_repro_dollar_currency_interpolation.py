# FILE: tests/test_repro_dollar_currency_interpolation.py
"""
Repro tests for Finding 1 in eval_report.md:
"Regex Backreference Interpolation Corrupts Dollar Amounts and Currency Figures in adeu apply and adeu diff"

Finding summary:
  Editing documents containing currency amounts / dollar figures ($100, $250,000)
  must preserve the monetary values ($100, $250,000) without dollar sign stripping
  or target-not-found failures when processing currency edits.
"""

from io import BytesIO

from docx import Document

from adeu.diff import generate_structured_edits
from adeu.ingest import _extract_text_from_doc, extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def build_contract_fee_doc(text: str = "The total compensation shall be $100 per hour.") -> BytesIO:
    d = Document()
    d.add_paragraph(text)
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


class TestDollarCurrencyInterpolation:
    def test_apply_currency_edit_with_regex_flag(self):
        """
        Applying a regex edit targeting document text containing dollar signs
        (e.g., 'The total compensation shall be $100 per hour.') with regex=True
        must match the target string and substitute the new dollar figure ($250,000).
        """
        doc_stream = build_contract_fee_doc("The total compensation shall be $100 per hour.")
        engine = RedlineEngine(doc_stream)
        edit = ModifyText(
            type="modify",
            target_text="The total compensation shall be $100 per hour.",
            new_text="The total compensation shall be $250,000 per hour.",
            regex=True,
        )

        stats = engine.process_batch([edit])
        assert stats["edits_applied"] == 1, "The currency edit with regex=True must apply"

        out_stream = engine.save_to_stream()
        clean = extract_text_from_stream(out_stream, clean_view=True)
        assert "$250,000" in clean, f"Output clean text must contain '$250,000': {clean}"

    def test_apply_currency_edit_preserves_dollar_figures_in_docx(self):
        """
        Applying an edit replacing $100 with $250,000 must result in the new
        dollar figure ($250,000) being directly accessible in paragraph text
        of the output DOCX document.
        """
        doc_stream = build_contract_fee_doc("The total compensation shall be $100 per hour.")
        engine = RedlineEngine(doc_stream)
        edit = ModifyText(
            type="modify",
            target_text="The total compensation shall be $100 per hour.",
            new_text="The total compensation shall be $250,000 per hour.",
        )

        engine.process_batch([edit])
        out_stream = engine.save_to_stream()

        # Load output docx clean view text via extract_text_from_stream
        clean_text = extract_text_from_stream(out_stream, clean_view=True)
        assert "$250,000" in clean_text, f"Output docx clean text must contain '$250,000' (got: {repr(clean_text)})"

    def test_diff_currency_edit_includes_dollar_symbol(self):
        """
        adeu diff between a $100 document and $250,000 document must emit
        a target_text containing '$100' and a new_text containing '$250,000'.
        """
        doc1 = build_contract_fee_doc("The total compensation shall be $100 per hour.")
        doc2 = build_contract_fee_doc("The total compensation shall be $250,000 per hour.")

        doc_o = Document(doc1)
        doc_m = Document(doc2)
        text_o, struct_o = _extract_text_from_doc(doc_o, clean_view=True, include_appendix=False, return_structure=True)
        text_m, struct_m = _extract_text_from_doc(doc_m, clean_view=True, include_appendix=False, return_structure=True)
        edits, warnings = generate_structured_edits(text_o, struct_o, text_m, struct_m)
        assert len(edits) == 1, "Should generate exactly 1 edit hunk"
        edit = edits[0]

        assert "$100" in edit.target_text, f"target_text should include '$100': {edit.target_text}"
        assert "$250,000" in edit.new_text, f"new_text should include '$250,000': {edit.new_text}"
