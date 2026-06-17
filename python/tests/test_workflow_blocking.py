import io

import pytest
from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_repro_workflow_blocking():
    """
    Scenario:
    1. Doc has tracked changes (from Round 1).
    2. User tries to edit that tracked text (Round 2).
    3. Engine should NOT skip it, but convert it to a replacement.
    """
    doc = Document()
    doc.add_paragraph("Start ")

    stream1 = io.BytesIO()
    doc.save(stream1)
    stream1.seek(0)

    # Round 1: Insert "Round1"
    edit1 = ModifyText(target_text="Start ", new_text="Start Round1")
    engine1 = RedlineEngine(stream1, author="Party A")
    engine1.apply_edits([edit1])
    stream2 = engine1.save_to_stream()

    # 2. Verify Ingest shows markup
    text = extract_text_from_stream(stream2)
    # CriticMarkup: {++Round1++}
    assert "{++Round1++}" in text

    # 3. Round 2: Edit "Round1" to "Round2"
    # Target "Round1" inside the markup
    edit2 = ModifyText(target_text="Round1", new_text="Round2", comment="Counter proposal")

    engine2 = RedlineEngine(stream2, author="Party A")
    applied, skipped = engine2.apply_edits([edit2])

    if skipped > 0:
        pytest.fail(f"Engine skipped the edit! Blocks workflow. Applied: {applied}, Skipped: {skipped}")

    res_stream = engine2.save_to_stream()
    final_text = extract_text_from_stream(res_stream)

    # Should contain Round2
    assert "Round2" in final_text
    # Round1 should be tracked as deleted natively instead of wiped entirely
    assert "{--Round1--}" in final_text


def test_repro_workflow_blocking_target_with_markup():
    """
    Scenario: LLM includes the markup brackets in the target.
    target_text = "{++Round1++}"
    """
    doc = Document()
    doc.add_paragraph("Start ")
    stream1 = io.BytesIO()
    doc.save(stream1)
    stream1.seek(0)

    engine1 = RedlineEngine(stream1, author="A")
    engine1.apply_edits([ModifyText(target_text="Start ", new_text="Start Round1")])
    stream2 = engine1.save_to_stream()

    # Target WITH brackets
    # Note: Regex chars in brackets might cause issues if diff/find uses regex?
    # DocumentMapper.find_match_index uses string.find.
    edit2 = ModifyText(
        target_text="{++Round1++}",
        new_text="Round2",
    )

    engine2 = RedlineEngine(stream2, author="A")
    applied, skipped = engine2.apply_edits([edit2])

    # This is likely where it fails if the mapper treats brackets as virtual
    # Mapper.full_text contains virtual text?
    # Mapper._build_map adds virtual text to spans AND self.full_text.
    # So finding "{++Round1++}" should theoretically work.

    if skipped > 0:
        # If this skips, we know we need to handle markup targeting explicitly
        print("Skipped markup target")
    else:
        print("Applied markup target")

    res_stream = engine2.save_to_stream()
    final_text = extract_text_from_stream(res_stream)

    assert "Round2" in final_text


def test_repro_p1_and_p2_validation_messages():
    from adeu.redline.engine import BatchValidationError

    doc = Document()
    p = doc.add_paragraph("The quick brown fox jumps over the ")
    run = p.add_run("lazy")
    run.bold = True
    p.add_run(" dog.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Author A deletes "lazy dog" and inserts "sleepy cat"
    engine_a = RedlineEngine(stream, author="Author A")
    engine_a.apply_edits([ModifyText(target_text="lazy dog", new_text="sleepy cat")])
    stream_after_a = engine_a.save_to_stream()

    # 2. Under Author B, try to target the deleted "lazy dog"
    engine_b = RedlineEngine(stream_after_a, author="Author B")
    edit_deleted = ModifyText(target_text="lazy dog", new_text="lazy hound")
    with pytest.raises(BatchValidationError) as excinfo:
        engine_b.process_batch([edit_deleted])
    err_str = str(excinfo.value)
    assert "matches text inside a tracked deletion by Author A" in err_str
    assert "Reject/accept that change first" in err_str

    # 3. Under Author B, try to target the active insertion "sleepy cat" from Author A
    edit_inserted = ModifyText(target_text="sleepy cat", new_text="sleepy kitten")
    with pytest.raises(BatchValidationError) as excinfo_ins:
        engine_b.process_batch([edit_inserted])
    err_str_ins = str(excinfo_ins.value)
    assert "targets an active insertion from another author" in err_str_ins
    assert "Author A (e.g. Chg:2)" in err_str_ins or "Author A (e.g. Chg:1)" in err_str_ins
