import asyncio
import json
import re
import sys
from pathlib import Path

import docx
import pytest
from fastmcp.exceptions import ToolError

from adeu.cli import main
from adeu.mcp_components.tools.document import read_docx
from adeu.payloads import failure_envelope, response_budget_limit, whole_doc_guard_message


class MockContext:
    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


def _create_doc(path: Path, chars: int, add_headings: bool = True) -> Path:
    doc = docx.Document()
    if add_headings:
        doc.add_heading("First L1 Heading", level=1)
        doc.add_paragraph("Introductory text under first heading.")
        doc.add_heading("Second L1 Heading", level=1)

    chunk_size = 1000
    written = 0
    while written < chars:
        doc.add_paragraph(f"Paragraph chunk {written}: " + ("x" * chunk_size))
        written += chunk_size

    doc.save(path)
    return path


def test_guard_fires_on_oversize_all_pages(tmp_path: Path, monkeypatch, capsys):
    doc_path = _create_doc(tmp_path / "big.docx", chars=80000)
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all"])
    with pytest.raises(SystemExit) as exc:
        main()
    assert exc.value.code != 0
    captured = capsys.readouterr()
    output = captured.out + captured.err
    assert "Refused unbounded full document read" in output
    assert "Recipe to read bounded sections" in output


def test_guard_output_is_under_800_tokens():
    msg = whole_doc_guard_message(
        total_chars=120000,
        limit=76000,
        file_path="big_document.docx",
        outline="# Heading 1 (p1)\n# Heading 2 (p5)\n# Heading 3 (p10)",
        page_count=12,
    )
    assert len(msg) <= 3100
    approx_tokens = len(msg) // 4
    assert approx_tokens <= 775


def _emitted_json(msg: str) -> str:
    """The CLI --json emission of a guard message: the largest surface form."""
    return json.dumps(failure_envelope("response_budget_exceeded", [], msg), ensure_ascii=False)


def _extract_help_text(monkeypatch, capsys) -> str:
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", "--help"])
    with pytest.raises(SystemExit):
        main()
    return capsys.readouterr().out


def test_recipe_advertises_only_runnable_extract_flags(monkeypatch, capsys):
    help_text = _extract_help_text(monkeypatch, capsys)
    mode_match = re.search(r"--mode \{([^}]+)\}", help_text)
    assert mode_match, "could not read --mode choices out of `extract --help`"
    mode_choices = set(mode_match.group(1).split(","))

    msg = whole_doc_guard_message(
        total_chars=120000,
        limit=76000,
        file_path="big_document.docx",
        outline="# Heading 1 (p1)",
        page_count=12,
    )

    advertised = set(re.findall(r"--[a-z][a-z-]*", msg))
    assert advertised, "guard recipe advertises no CLI flags"
    for flag in advertised:
        # Exact option match: "--search" must not be excused by "--search-query".
        assert re.search(re.escape(flag) + r"(?![\w-])", help_text), f"guard advertises unknown flag {flag}"

    for mode in re.findall(r"--mode ([a-z]+)", msg):
        assert mode in mode_choices, f"guard advertises unknown --mode {mode}"


def test_emitted_response_stays_in_budget_on_long_path_and_outline():
    long_path = "C:\\Users\\someone\\" + ("a_very_long_directory_name\\" * 8) + "contract_final_v12.docx"
    entries = [f"# Article {i} — Obligations Of The Parties (p{i})" for i in range(1, 121)]

    msg = whole_doc_guard_message(
        total_chars=1200000,
        limit=76000,
        file_path=long_path,
        outline="\n".join(entries),
        page_count=160,
    )

    assert len(_emitted_json(msg)) // 4 <= 800
    # The recipe survives trimming, and outline entries are dropped whole —
    # never sliced mid-line.
    assert "Recipe to read bounded sections" in msg
    assert "--force" in msg
    rendered = [line for line in msg.splitlines() if line.startswith("# ")]
    assert rendered, "outline was dropped entirely before the budget was reached"
    assert all(line in entries for line in rendered)


def test_outline_section_omitted_when_no_l1_headings(tmp_path: Path, monkeypatch, capsys):
    doc_path = _create_doc(tmp_path / "no_headings.docx", chars=80000, add_headings=False)
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all"])
    with pytest.raises(SystemExit):
        main()
    output = capsys.readouterr().err
    assert "Refused unbounded full document read" in output
    assert "Outline" not in output
    assert "No headings" not in output


def test_json_stderr_suppression_is_scoped_to_the_guard(tmp_path: Path, monkeypatch, capsys):
    missing = tmp_path / "nope.docx"
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(missing), "--page", "all", "--json"])
    with pytest.raises(SystemExit):
        main()
    captured = capsys.readouterr()
    assert json.loads(captured.out)["error"] == "file_not_found"
    assert "❌" in captured.err


def test_cli_output_file_sink_exempt_from_guard(tmp_path: Path, monkeypatch):
    doc_path = _create_doc(tmp_path / "big.docx", chars=80000)
    dump_path = tmp_path / "dump.txt"
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all", "-o", str(dump_path)])
    main()
    assert dump_path.exists()
    content = dump_path.read_text(encoding="utf-8")
    assert "Paragraph chunk 0" in content


def test_cli_refusal_includes_page_count(tmp_path: Path, monkeypatch, capsys):
    doc_path = _create_doc(tmp_path / "big.docx", chars=80000)
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all"])
    with pytest.raises(SystemExit) as exc:
        main()
    assert exc.value.code != 0
    captured = capsys.readouterr()
    output = captured.out + captured.err
    assert "pages)" in output


def test_cli_json_mode_no_stderr_duplication_and_bounded(tmp_path: Path, monkeypatch, capsys):
    doc_path = _create_doc(tmp_path / "big.docx", chars=80000)
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all", "--json"])
    with pytest.raises(SystemExit) as exc:
        main()
    assert exc.value.code != 0
    captured = capsys.readouterr()
    assert captured.err == ""
    data = json.loads(captured.out)
    assert data.get("error") == "response_budget_exceeded"
    assert len(captured.out) // 4 <= 800


@pytest.mark.skipif(sys.platform != "win32", reason="Live Word only supported on Windows")
def test_live_word_guard_and_force(monkeypatch):
    from adeu.mcp_components.tools.live_word import read_active_word_document

    big_text = "x" * 80000
    doc = docx.Document()
    doc.add_paragraph(big_text)

    def mock_core(clean_view=False, file_path=None, include_appendix=False, return_paragraph_offsets=False):
        if return_paragraph_offsets:
            return big_text, "Active Document", doc, {}
        return big_text, "Active Document", doc

    monkeypatch.setattr("adeu.mcp_components.tools.live_word._read_active_word_document_core", mock_core)
    ctx = MockContext()

    with pytest.raises(ToolError) as exc_info:
        asyncio.run(read_active_word_document(ctx=ctx, page="all", force=False))
    err_msg = str(exc_info.value)
    assert "Refused unbounded full document read" in err_msg

    res = asyncio.run(read_active_word_document(ctx=ctx, page="all", force=True))
    assert res is not None
    assert big_text in str(res.content)


def test_guard_includes_outline():
    outline_text = "# First L1 Heading (p1)\n# Second L1 Heading (p3)"
    msg = whole_doc_guard_message(
        total_chars=90000,
        limit=76000,
        file_path="sample.docx",
        outline=outline_text,
        page_count=5,
    )
    assert "Outline (L1 Headings):" in msg
    assert "First L1 Heading" in msg
    assert "Second L1 Heading" in msg


def test_small_document_is_unaffected(tmp_path: Path):
    doc_path = _create_doc(tmp_path / "small.docx", chars=5000)
    ctx = MockContext()
    res = asyncio.run(read_docx(reasoning="test", file_path=str(doc_path), page="all", ctx=ctx))
    assert res is not None
    assert "Paragraph chunk 0" in str(res.content)


def test_force_overrides(tmp_path: Path, monkeypatch, capsys):
    doc_path = _create_doc(tmp_path / "big.docx", chars=80000)

    # CLI check
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all", "--force"])
    main()
    captured = capsys.readouterr()
    assert "Paragraph chunk 0" in captured.out

    # MCP check
    ctx = MockContext()
    res = asyncio.run(read_docx(reasoning="test", file_path=str(doc_path), page="all", force=True, ctx=ctx))
    assert "Paragraph chunk 0" in str(res.content)


def test_env_var_overrides_threshold(tmp_path: Path, monkeypatch, capsys):
    monkeypatch.setenv("ADEU_MAX_RESPONSE_CHARS", "2000")
    assert response_budget_limit() == 2000

    doc_path = _create_doc(tmp_path / "medium.docx", chars=3000)

    # With threshold = 2000, 3000 chars should fire the guard
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all"])
    with pytest.raises(SystemExit) as exc:
        main()
    assert exc.value.code != 0
    captured = capsys.readouterr()
    assert (
        "Refused unbounded full document read" in captured.err or "Refused unbounded full document read" in captured.out
    )

    # With threshold = 50000, 3000 chars should pass
    monkeypatch.setenv("ADEU_MAX_RESPONSE_CHARS", "50000")
    assert response_budget_limit() == 50000
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all"])
    main()


def test_guard_does_not_fire_for_search_or_ranges(tmp_path: Path, monkeypatch, capsys):
    doc_path = _create_doc(tmp_path / "big.docx", chars=80000)

    # Search with page="all"
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--search-query", "chunk", "--page", "all"])
    main()

    # Page range N-M
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "1-2"])
    main()


def test_mcp_guard_fires(tmp_path: Path):
    doc_path = _create_doc(tmp_path / "big.docx", chars=80000)
    ctx = MockContext()

    with pytest.raises(ToolError) as exc_info:
        asyncio.run(read_docx(reasoning="test", file_path=str(doc_path), page="all", ctx=ctx))

    err_msg = str(exc_info.value)
    assert "Refused unbounded full document read" in err_msg
    assert "Recipe to read bounded sections" in err_msg


def test_guard_never_fires_when_paginating_normally(tmp_path: Path, monkeypatch, capsys):
    doc_path = _create_doc(tmp_path / "big.docx", chars=80000)

    # Single page read via CLI
    monkeypatch.setattr(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "1"])
    main()

    # Single page read via MCP
    ctx = MockContext()
    res = asyncio.run(read_docx(reasoning="test", file_path=str(doc_path), page=1, ctx=ctx))
    assert res is not None
    assert "Refused unbounded" not in str(res.content)
