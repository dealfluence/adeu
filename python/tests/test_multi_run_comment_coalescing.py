# FILE: tests/test_multi_run_comment_coalescing.py
"""
Regression tests for the multi-run comment duplication bug.

Word naturally fragments runs around numbers, punctuation, and formatting
changes. A single <w:commentRangeStart>/<w:commentRangeEnd> pair can span many
<w:r> elements. The projection and mapper layers must coalesce these into one
{==...==}{>>[Com:N]...<<} block, not one per run.

These tests guard both layers (projection + mapper) — they must stay in
lockstep or the engine's find-and-replace anchoring will break.
"""

import io

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import XmlPart
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from adeu.ingest import extract_text_from_stream
from adeu.redline.mapper import DocumentMapper

# -----------------------------------------------------------------------------
# Helpers: build DOCX from hand-authored OOXML (no high-level builders that
# would coalesce runs and hide the bug)
# -----------------------------------------------------------------------------


def _run(text: str, preserve: bool = False) -> str:
    sp = ' xml:space="preserve"' if preserve else ""
    return f"<w:r><w:t{sp}>{text}</w:t></w:r>"


def _ref_run(comment_id: int) -> str:
    return f'<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:commentReference w:id="{comment_id}"/></w:r>'


def _build(paragraph_xml: str, comments: list[tuple[int, str, str]]) -> io.BytesIO:
    """
    paragraph_xml: a complete <w:p ...>...</w:p> XML string.
    comments: list of (id, author, text).
    """
    doc = Document()
    doc.add_paragraph("placeholder")

    body = doc.element.body
    for p in list(body.findall(qn("w:p"))):
        body.remove(p)

    new_p = parse_xml(paragraph_xml)
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        sectPr.addprevious(new_p)
    else:
        body.append(new_p)

    package = doc.part.package
    partname = package.next_partname("/word/comments%d.xml")
    comment_xml_parts = []
    for cid, author, text in comments:
        comment_xml_parts.append(
            f'<w:comment w:id="{cid}" w:author="{author}" '
            f'w:date="2025-01-01T00:00:00Z" w:initials="X">'
            f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
            f"</w:comment>"
        )
    comments_xml = (f"<w:comments {nsdecls('w')}>" + "".join(comment_xml_parts) + "</w:comments>").encode("utf-8")
    cpart = XmlPart(partname, CT.WML_COMMENTS, parse_xml(comments_xml), package)
    package.parts.append(cpart)
    doc.part.relate_to(cpart, RT.COMMENTS)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _project(stream: io.BytesIO) -> str:
    return extract_text_from_stream(io.BytesIO(stream.getvalue()))


def _mapper_text(stream: io.BytesIO) -> str:
    return DocumentMapper(Document(io.BytesIO(stream.getvalue()))).full_text


# -----------------------------------------------------------------------------
# Case 1: core bug — single comment spans 8 runs
# -----------------------------------------------------------------------------


def test_single_comment_across_8_runs_coalesces():
    runs = "".join(
        [
            _run("Party A shall pay ", preserve=True),
            _run("100"),
            _run("%"),
            _run(" of the total", preserve=True),
            _run(" amount", preserve=True),
            _run(" on ", preserve=True),
            _run("time"),
            _run("."),
        ]
    )
    p_xml = (
        f"<w:p {nsdecls('w')}>"
        '<w:commentRangeStart w:id="1"/>' + runs + '<w:commentRangeEnd w:id="1"/>' + _ref_run(1) + "</w:p>"
    )
    stream = _build(p_xml, [(1, "Reviewer", "Risk note.")])

    text = _project(stream)
    assert text.count("{==") == 1
    assert text.count("{>>") == 1
    assert text.count("[Com:1]") == 1

    mtext = _mapper_text(stream)
    assert mtext.count("{==") == 1
    assert mtext.count("{>>") == 1
    assert mtext.count("[Com:1]") == 1


# -----------------------------------------------------------------------------
# Case 2: independence — two adjacent but non-overlapping comments stay distinct
# -----------------------------------------------------------------------------


def test_two_adjacent_comments_stay_independent():
    # [Com:1 covers "Alpha"] [Com:2 covers "Bravo"]
    p_xml = (
        f"<w:p {nsdecls('w')}>"
        '<w:commentRangeStart w:id="1"/>'
        + _run("Alpha")
        + '<w:commentRangeEnd w:id="1"/>'
        + _ref_run(1)
        + _run(" ", preserve=True)
        + '<w:commentRangeStart w:id="2"/>'
        + _run("Bravo")
        + '<w:commentRangeEnd w:id="2"/>'
        + _ref_run(2)
        + "</w:p>"
    )
    stream = _build(p_xml, [(1, "A", "First"), (2, "B", "Second")])

    text = _project(stream)
    assert text.count("[Com:1]") == 1
    assert text.count("[Com:2]") == 1
    assert text.count("{==") == 2
    assert text.count("{>>") == 2

    mtext = _mapper_text(stream)
    assert mtext.count("[Com:1]") == 1
    assert mtext.count("[Com:2]") == 1
    assert mtext.count("{==") == 2
    assert mtext.count("{>>") == 2


# -----------------------------------------------------------------------------
# Case 3: overlap — two comments cover the exact same multi-run span; produce
# one highlight wrapper and one meta block listing both IDs
# -----------------------------------------------------------------------------


def test_two_overlapping_comments_share_one_meta_block():
    runs = _run("Hello ", preserve=True) + _run("world") + _run("!")
    p_xml = (
        f"<w:p {nsdecls('w')}>"
        '<w:commentRangeStart w:id="1"/>'
        '<w:commentRangeStart w:id="2"/>'
        + runs
        + '<w:commentRangeEnd w:id="1"/>'
        + '<w:commentRangeEnd w:id="2"/>'
        + _ref_run(1)
        + _ref_run(2)
        + "</w:p>"
    )
    stream = _build(p_xml, [(1, "A", "First note"), (2, "B", "Second note")])

    text = _project(stream)
    assert text.count("{==") == 1, text
    assert text.count("{>>") == 1, text
    assert text.count("[Com:1]") == 1, text
    assert text.count("[Com:2]") == 1, text

    mtext = _mapper_text(stream)
    assert mtext.count("{==") == 1, mtext
    assert mtext.count("{>>") == 1, mtext
    assert mtext.count("[Com:1]") == 1, mtext
    assert mtext.count("[Com:2]") == 1, mtext


# -----------------------------------------------------------------------------
# Case 4: mixed redline + comment — a <w:ins> nested inside a comment range
# produces a {++...++} wrapper, and the comment payload appears exactly once.
# -----------------------------------------------------------------------------


def test_mixed_insertion_inside_comment_range():
    p_xml = (
        f"<w:p {nsdecls('w')}>"
        '<w:commentRangeStart w:id="1"/>'
        + _run("Before ", preserve=True)
        + '<w:ins w:id="100" w:author="Editor" w:date="2025-01-01T00:00:00Z">'
        + _run("inserted")
        + "</w:ins>"
        + _run(" after.", preserve=True)
        + '<w:commentRangeEnd w:id="1"/>'
        + _ref_run(1)
        + "</w:p>"
    )
    stream = _build(p_xml, [(1, "Reviewer", "Look at this")])

    text = _project(stream)
    assert text.count("[Com:1]") == 1, text
    assert text.count("[Chg:100 insert]") == 1, text
    assert text.count("{++") == 1, text
    assert "inserted" in text

    mtext = _mapper_text(stream)
    assert mtext.count("[Com:1]") == 1, mtext
    assert mtext.count("[Chg:100 insert]") == 1, mtext
    assert mtext.count("{++") == 1, mtext


# -----------------------------------------------------------------------------
# Case 5: mapper parity — same input, mapper.full_text contains the same single
# comment block. Guards mapper independently of projection.
# -----------------------------------------------------------------------------


def test_mapper_parity_for_multi_run_comment():
    runs = _run("X") + _run("Y") + _run("Z")
    p_xml = (
        f"<w:p {nsdecls('w')}>"
        '<w:commentRangeStart w:id="42"/>' + runs + '<w:commentRangeEnd w:id="42"/>' + _ref_run(42) + "</w:p>"
    )
    stream = _build(p_xml, [(42, "Reviewer", "Triplet")])

    proj = _project(stream)
    mtext = _mapper_text(stream)

    # Both layers must agree on the single coalesced shape.
    assert proj.count("{==") == 1 == mtext.count("{==")
    assert proj.count("{>>") == 1 == mtext.count("{>>")
    assert proj.count("[Com:42]") == 1 == mtext.count("[Com:42]")
