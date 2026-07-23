# FILE: tests/test_repro_qa_customer_assessment_2026_07_23.py
"""
Failing repro tests for the customer-facing QA assessment (2026-07-23,
"Adeu Document Redlining — QA Assessment for Customers", v1.29.0) — Python side.

Findings covered here (Node-side mirrors live in
node/packages/core/src/repro_qa_customer_assessment_2026_07_23.test.ts):

- C1: "The Python server currently rejects the simplest well-formed edit
  request unless every optional setting is spelled out." Reproduces. The
  process_document_batch tool declares author_name as a REQUIRED parameter
  (mcp_components/tools/document.py, no default), while Node gives it a
  schema default ("Adeu AI (TS)", index.ts) precisely because real MCP
  clients drop primitive-typed entries from required[] — so schema-following
  models legitimately omit it (QA 2026-07-23 F3 / client-compat). Python must
  default author_name the same way ("Adeu AI", the engine's own default).

- C2: "silently wrote the literal text '$1' into a payment clause". Reproduces.
  Spec §6 (search_and_targeted_write.md) deliberately keeps each runtime's
  native backreference syntax (\\1/\\g<1> on Python, $1 on Node) and mandates
  that "the tool descriptions on both servers will instruct the LLM of the
  active environment's engine behavior". Node's new_text description documents
  "$1, $2…"; Python's documents nothing, and a $1-style replacement applies
  with status=applied and warning=None while writing literal "$1" into the
  clause. The schema must disclose the syntax, and the engine must attach a
  warning (not a hard reject — "$1,000"-style literals are everyday legal
  text) when a $N token survives a regex substitution verbatim.

- C3: "attach a comment to a sentence without supplying replacement text".
  The Node engine deletes the sentence (tracked) and hangs the comment on the
  deletion; Python rejects the same payload with "modify.new_text: Field
  required" even though the ADVERTISED flat schema (FlatSchemaDocumentChange,
  required == ["type"]) tells the model new_text is optional. By the repo's
  own client-compat principle, a payload the published schema invites must
  degrade gracefully: modify + comment + ABSENT new_text is an annotation,
  and must normalize to the pure-comment form (new_text == target_text).
  An explicit new_text="" stays a deletion (delete-with-explanation is a
  legitimate, distinct intent), and omitting new_text WITHOUT a comment stays
  an error (no meaningful interpretation exists).

- C4: "Adeu cannot see text inside floating text boxes or form-style content
  controls, and it does not warn you about this." Reproduces, twice over:
  (a) block-level w:sdt content controls are silently skipped —
  iter_block_items (utils/docx.py) matches only w:p/w:tbl direct children, so
  whole paragraphs of body text vanish from every view (Word renders them as
  ordinary flowed text; hiding them is data loss — they must be extracted);
  (b) w:txbxContent is not handled anywhere — an mc:AlternateContent-wrapped
  DrawingML text box projects NOTHING (not even an image marker), so a
  coversheet notice or boxed obligation is invisible with no disclosure. The
  read output (inline marker or structural appendix) must at least disclose
  the presence of text-box text it cannot project.

- C5: "the standalone command-line interface … its apply step currently
  fails." Does NOT reproduce — no test. Verified end-to-end on this commit
  (extract --clean-view --page all → edit → apply; JSON-batch apply; review
  actions JSON; default-output overwrite), all exit 0 with correct redlines;
  the CLI apply suite (test_cli_features.py -k apply) is green. The likely
  QA-side failure mode is an environment where the `adeu` entry point runs
  without the package importable (ModuleNotFoundError), which cli.py already
  hints about, or a tripped guardrail (markup-view text / major-deletion
  abort) read as a crash. Nothing actionable in-repo.

Every test is written test-first: it fails on current main and passes once
the finding is fixed.
"""

import asyncio
import json
import re
from io import BytesIO

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.mcp_components.tools.document import _normalize_changes
from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _batch_tool_schema() -> dict:
    from adeu.server import mcp

    tools = asyncio.run(mcp.list_tools())
    tool = next(t for t in tools if t.name == "process_document_batch")
    schema = getattr(tool, "inputSchema", None) or getattr(tool, "parameters", None)
    assert schema, "process_document_batch published no input schema"
    return schema


def _change_item_schema() -> dict:
    changes = _batch_tool_schema()["properties"]["changes"]
    # The items schema may sit directly on the property or behind anyOf.
    if "items" in changes:
        return changes["items"]
    for variant in changes.get("anyOf", []):
        if "items" in variant:
            return variant["items"]
    raise AssertionError(f"could not locate change-item schema in: {json.dumps(changes)[:400]}")


# ────────────────────────────────────────────────────────────────────────────
# C1 — the minimal well-formed batch request must be accepted.
# ────────────────────────────────────────────────────────────────────────────
class TestC1MinimalRequestParity:
    def test_author_name_is_not_required_and_carries_a_default(self):
        """Node ships author_name with a schema default because real clients
        strip primitive entries from required[]; Python must match or every
        schema-following model gets -32602 on its first edit."""
        schema = _batch_tool_schema()
        assert "author_name" not in schema.get("required", []), (
            "author_name is still a required parameter — the minimal request "
            f"Node accepts is rejected by Python. required={schema.get('required')}"
        )
        author = schema["properties"]["author_name"]
        assert author.get("default"), (
            f"author_name must carry a schema default (the engine's own 'Adeu AI'), got: {json.dumps(author)}"
        )

    def test_minimal_batch_call_without_author_name_executes(self, tmp_path):
        """The exact payload shape from the Node client-compat repro — no
        author_name — must run through the JSON-RPC boundary, not bounce off
        input validation before the handler is reached."""
        from adeu.server import mcp

        doc = Document()
        doc.add_paragraph("The lazy dog sleeps.")
        src = tmp_path / "minimal.docx"
        doc.save(str(src))

        arguments = {
            "reasoning": "minimal well-formed edit request",
            "original_docx_path": str(src),
            "dry_run": True,
            "changes": [{"type": "modify", "target_text": "lazy dog", "new_text": "sleepy cat"}],
        }
        try:
            result = asyncio.run(mcp.call_tool("process_document_batch", arguments))
        except Exception as exc:  # noqa: BLE001 — the repro IS the rejection
            pytest.fail(
                "the minimal request (author_name omitted, as the published "
                f"client-side schema allows) was rejected at the boundary: {exc}"
            )
        text = "".join(item.text for item in result.content if item.type == "text")
        assert "sleepy cat" in text or "1 applied" in text or "Dry run" in text, (
            f"minimal batch did not reach the engine: {text[:400]}"
        )


# ────────────────────────────────────────────────────────────────────────────
# C2 — $1 backreferences must not silently corrupt a clause.
# ────────────────────────────────────────────────────────────────────────────
PAYMENT_SENTENCE = "The Client shall pay a fee of 12000 euros within 30 days."


def _payment_doc() -> BytesIO:
    d = Document()
    d.add_paragraph(PAYMENT_SENTENCE)
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


class TestC2RegexBackreferenceSyntax:
    def test_schema_documents_python_backreference_syntax(self):
        """Spec §6: 'The tool descriptions on both servers will instruct the
        LLM of the active environment's engine behavior.' Node's new_text
        description says '$1, $2…'; Python's says nothing about \\1/\\g<1>."""
        items = _change_item_schema()
        props = items.get("properties", {})
        disclosure = " ".join(str(props.get(field, {}).get("description", "")) for field in ("new_text", "regex"))
        assert "\\1" in disclosure or "\\g<" in disclosure, (
            "neither new_text nor regex documents Python's backreference "
            f"syntax (\\1 / \\g<1>) as spec §6 requires. Got: {disclosure!r}"
        )

    def test_dollar_backreference_is_not_applied_silently(self):
        """A JS-style $1 replacement writes the literal text '$1' into the
        clause. That may stay platform-native behavior (spec §6), but it must
        not pass through with status=applied and warning=None — the edit
        report must flag the surviving $N token and name \\1 syntax."""
        engine = RedlineEngine(_payment_doc(), author="QA")
        try:
            stats = engine.process_batch(
                [
                    ModifyText(
                        target_text=r"fee of (\d+) euros",
                        new_text="fee of $1 euros (plus applicable VAT)",
                        regex=True,
                    )
                ]
            )
        except BatchValidationError as exc:
            # A guided rejection is an acceptable fix shape too — but it must
            # teach the caller the correct syntax.
            assert re.search(r"\\1|\\g<1>|backreference", str(exc)), (
                f"rejection must name Python's backreference syntax: {exc}"
            )
            return

        report = stats["edits"][0]
        wrote_literal = "$1" in (report.get("clean_text") or "")
        if wrote_literal:
            warning = report.get("warning") or ""
            assert warning, (
                f"literal '$1' was written into a payment clause with warning=None — silent corruption: {report}"
            )
            assert re.search(r"\\1|\\g<1>|backreference", warning), (
                f"the warning must name Python's \\1/\\g<1> syntax: {warning!r}"
            )

    def test_literal_dollar_amounts_still_apply(self):
        """'$1,000'-style literals are everyday legal text. The C2 guardrail
        must warn, never hard-block, or routine fee edits become impossible."""
        engine = RedlineEngine(_payment_doc(), author="QA")
        stats = engine.process_batch(
            [
                ModifyText(
                    target_text=r"fee of (\d+) euros",
                    new_text="fee of $1,250.00 (one thousand two hundred fifty dollars)",
                    regex=True,
                )
            ]
        )
        assert stats["edits"][0]["status"] == "applied", (
            "a literal dollar amount in new_text must remain applicable "
            f"(a warning is fine, a rejection is not): {stats['edits'][0]}"
        )


# ────────────────────────────────────────────────────────────────────────────
# C3 — comment + absent new_text is an annotation, not a rejection.
# ────────────────────────────────────────────────────────────────────────────
LIABILITY_SENTENCE = "Liability is capped at the fees paid in the preceding twelve months."


class TestC3CommentWithoutReplacement:
    @pytest.mark.parametrize("payload_extra", [{}, {"new_text": None}], ids=["omitted", "null"])
    def test_comment_with_absent_new_text_normalizes_to_pure_comment(self, payload_extra):
        """The published flat schema advertises new_text as optional
        (required == ["type"]), so a schema-following model legitimately
        omits it when it only wants to annotate. The lossless interpretation
        is the pure-comment form (new_text == target_text) — not a
        'Field required' bounce (Python today), and never a tracked deletion
        (Node today)."""
        payload = {
            "type": "modify",
            "target_text": LIABILITY_SENTENCE,
            "comment": "Please reconsider this cap.",
            **payload_extra,
        }
        changes, errors = _normalize_changes([payload])
        assert not errors, (
            f"a comment-only modify (new_text absent) was rejected instead of normalizing to a pure comment: {errors}"
        )
        assert len(changes) == 1
        change = changes[0]
        assert isinstance(change, ModifyText)
        assert change.new_text == LIABILITY_SENTENCE, (
            "absent new_text with a comment must normalize to the pure-comment "
            f"form (new_text == target_text), got new_text={change.new_text!r}"
        )

    def test_pure_comment_normalization_produces_no_redline(self):
        """End-to-end: the normalized annotation must anchor a comment on the
        intact sentence — no {--deletion--} in the raw view."""
        changes, errors = _normalize_changes(
            [
                {
                    "type": "modify",
                    "target_text": LIABILITY_SENTENCE,
                    "comment": "Please reconsider this cap.",
                }
            ]
        )
        assert not errors, f"normalization failed: {errors}"

        d = Document()
        d.add_paragraph(LIABILITY_SENTENCE)
        buf = BytesIO()
        d.save(buf)
        buf.seek(0)
        engine = RedlineEngine(buf, author="QA")
        engine.process_batch(changes)
        raw = extract_text_from_stream(engine.save_to_stream(), clean_view=False)
        assert "{--" not in raw, f"annotating a sentence must not delete it:\n{raw}"
        assert "Please reconsider this cap." in raw
        assert LIABILITY_SENTENCE in raw.replace("{==", "").replace("==}", "")

    def test_explicit_empty_new_text_stays_a_deletion(self):
        """Pin: new_text='' is the documented delete-with-explanation intent
        and must NOT be reinterpreted by the C3 fix."""
        payload = {
            "type": "modify",
            "target_text": LIABILITY_SENTENCE,
            "new_text": "",
            "comment": "Removed per our call on 2026-07-20.",
        }
        changes, errors = _normalize_changes([payload])
        assert not errors
        assert changes[0].new_text == ""

    def test_absent_new_text_without_comment_still_errors(self):
        """Pin: with no comment there is no meaningful interpretation of an
        absent new_text — the self-service error must survive the C3 fix."""
        changes, errors = _normalize_changes([{"type": "modify", "target_text": LIABILITY_SENTENCE}])
        assert errors, "an uncommented modify without new_text must stay an error"


# ────────────────────────────────────────────────────────────────────────────
# C4 — invisible containers: block-level w:sdt and w:txbxContent.
# ────────────────────────────────────────────────────────────────────────────
SDT_SENTENCE = "The Supplier shall indemnify the Client against third-party claims."
BOXED_SENTENCE = "Notice: delivery obligations are suspended during force majeure."


def _doc_with_block_sdt() -> BytesIO:
    d = Document()
    d.add_paragraph("Intro paragraph before the content control.")
    sdt_xml = (
        f'<w:sdt xmlns:w="{W_NS}">'
        '<w:sdtPr><w:alias w:val="Indemnity"/></w:sdtPr>'
        "<w:sdtContent>"
        f'<w:p><w:r><w:t xml:space="preserve">{SDT_SENTENCE}</w:t></w:r></w:p>'
        "</w:sdtContent>"
        "</w:sdt>"
    )
    body = d.element.body
    body.insert(list(body).index(body.find(qn("w:sectPr"))), parse_xml(sdt_xml))
    d.add_paragraph("Tail paragraph after the content control.")
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


_TEXTBOX_RUN_XML = (
    f'<w:r xmlns:w="{W_NS}" '
    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
    'xmlns:v="urn:schemas-microsoft-com:vml">'
    "<mc:AlternateContent>"
    '<mc:Choice Requires="wps">'
    "<w:drawing>"
    '<wp:anchor behindDoc="0" distT="0" distB="0" distL="0" distR="0" simplePos="0" '
    'relativeHeight="2" locked="0" layoutInCell="1" allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>'
    '<wp:extent cx="2540000" cy="1270000"/>'
    "<wp:wrapNone/>"
    '<wp:docPr id="7" name="Text Box 7"/>'
    '<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    "<wps:wsp><wps:txbx>"
    "<w:txbxContent>"
    f'<w:p><w:r><w:t xml:space="preserve">{BOXED_SENTENCE}</w:t></w:r></w:p>'
    "</w:txbxContent>"
    "</wps:txbx></wps:wsp>"
    "</a:graphicData></a:graphic>"
    "</wp:anchor>"
    "</w:drawing>"
    "</mc:Choice>"
    "<mc:Fallback>"
    "<w:pict>"
    '<v:shape id="Text Box 7" style="position:absolute">'
    "<v:textbox>"
    "<w:txbxContent>"
    f'<w:p><w:r><w:t xml:space="preserve">{BOXED_SENTENCE}</w:t></w:r></w:p>'
    "</w:txbxContent>"
    "</v:textbox>"
    "</v:shape>"
    "</w:pict>"
    "</mc:Fallback>"
    "</mc:AlternateContent>"
    "</w:r>"
)


def _doc_with_textbox() -> BytesIO:
    d = Document()
    d.add_paragraph("Body text before the floating shape.")
    host = d.add_paragraph("Anchor paragraph. ")
    host._p.append(parse_xml(_TEXTBOX_RUN_XML))
    d.add_paragraph("Body text after the floating shape.")
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


class TestC4InvisibleContainers:
    def test_block_level_content_control_text_is_extracted(self):
        """Word renders a block-level w:sdt as ordinary flowed body text; a
        redlining tool that silently drops it is hiding live obligations.
        The paragraph inside w:sdtContent must appear in every view."""
        stream = _doc_with_block_sdt()
        text = extract_text_from_stream(stream, clean_view=True)
        assert "Intro paragraph before the content control." in text, "setup: body walk broken"
        assert SDT_SENTENCE in text, (
            "text inside a block-level content control (w:sdt) is invisible — "
            f"the indemnity obligation vanished from the read view:\n{text}"
        )

    def test_floating_textbox_text_is_disclosed(self):
        """An mc:AlternateContent text box currently projects NOTHING — no
        text, no marker, no warning. The read output must either project the
        boxed text or disclose that a text box with text was skipped (inline
        marker or structural appendix)."""
        stream = _doc_with_textbox()
        full = extract_text_from_stream(stream, clean_view=False, include_appendix=True)
        disclosed = BOXED_SENTENCE in full or re.search(r"text\s*box", full, re.IGNORECASE)
        assert disclosed, (
            "a floating text box containing an obligation is completely "
            f"invisible in the read output, with no disclosure:\n{full}"
        )
