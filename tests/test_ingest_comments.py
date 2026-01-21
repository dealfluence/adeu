import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.redline.engine import RedlineEngine


def test_ingest_critic_markup_simple():
    """
    Test single comment: {==Target==}{>>[Author] Comment<<}
    """
    doc = Document()
    p = doc.add_paragraph("Start ")
    # IMPORTANT: Apply formatting to prevent RedlineEngine from coalescing runs
    # normalize_docx merges adjacent runs with same properties.
    p.add_run("Target").bold = True
    p.add_run(" End")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Tester")

    # Inject comment on the run "Target"
    comment_id = engine.comments_manager.add_comment("Tester", "Fix this.")

    # Use the engine's document instance
    p = engine.doc.paragraphs[0]
    target_run = p.runs[1]  # "Target" (Index 1 because 'Start ' is 0)

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), comment_id)
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), comment_id)

    ref_run = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), comment_id)
    ref_run.append(ref)

    p._element.insert(p._element.index(target_run._element), start)
    p._element.insert(p._element.index(target_run._element) + 1, end)
    p._element.insert(p._element.index(target_run._element) + 2, ref_run)

    res_stream = engine.save_to_stream()

    # Act
    text = extract_text_from_stream(res_stream)

    # Assert
    # Note: **Target** because we bolded it.
    assert "Start {==**Target**==}{>>[Tester" in text
    assert "] Fix this.<<} End" in text


def test_ingest_critic_markup_overlapping():
    """
    Test overlap flattening:
    Text: A B C
    Comment 1: A B
    Comment 2: B C
    """
    doc = Document()
    p = doc.add_paragraph()
    # Distinct formatting to prevent coalescing
    p.add_run("A ").bold = True
    p.add_run("B ").italic = True
    p.add_run("C").bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    engine = RedlineEngine(stream)

    c1 = engine.comments_manager.add_comment("U1", "C1")
    c2 = engine.comments_manager.add_comment("U2", "C2")

    p = engine.doc.paragraphs[0]
    rA = p.runs[0]
    rB = p.runs[1]
    rC = p.runs[2]
    p_elm = p._element

    # Insert markers manually
    # Start 1 before A
    s1 = OxmlElement("w:commentRangeStart")
    s1.set(qn("w:id"), c1)
    p_elm.insert(p_elm.index(rA._element), s1)

    # Start 2 before B
    s2 = OxmlElement("w:commentRangeStart")
    s2.set(qn("w:id"), c2)
    p_elm.insert(p_elm.index(rB._element), s2)

    # End 1 after B
    e1 = OxmlElement("w:commentRangeEnd")
    e1.set(qn("w:id"), c1)
    p_elm.insert(p_elm.index(rB._element) + 1, e1)

    # End 2 after C
    e2 = OxmlElement("w:commentRangeEnd")
    e2.set(qn("w:id"), c2)
    p_elm.insert(p_elm.index(rC._element) + 1, e2)

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream)

    print(f"Overlap Result: {text}")

    # Segment A
    assert "{==**A **==}{>>[U1" in text

    # Segment B (Both)
    match_b = re.search(r"\{==_B _==\}\{>>(.*?)<<\}", text, re.DOTALL)
    assert match_b
    content_b = match_b.group(1)
    assert "U1" in content_b
    assert "U2" in content_b

    # Segment C
    assert "{==**C**==}{>>[U2" in text


def test_ingest_legal_brackets_collision():
    """
    Ensure [Bracket] inside text doesn't break CriticMarkup parsing visually.
    """
    doc = Document()
    doc.add_paragraph("The [Vendor]")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    engine = RedlineEngine(stream)

    c1 = engine.comments_manager.add_comment("Lawyer", "Check this")

    p = engine.doc.paragraphs[0]

    # Wrap text manually
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), c1)
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), c1)

    p._element.insert(0, start)
    p._element.append(end)

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream)

    # Expect: {==The [Vendor]==}{>>...<<}
    assert "{==The [Vendor]==}" in text
