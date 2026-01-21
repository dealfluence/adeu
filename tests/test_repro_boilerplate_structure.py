import io

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from adeu.ingest import extract_text_from_stream
from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_visible_runs


def test_ingest_detects_structural_info():
    """
    FIXED: The extracted text should now contain Markdown headers.
    """
    doc = Document()
    doc.add_heading("1. Assignment", level=1)
    doc.add_paragraph("This is the body text.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    extracted = extract_text_from_stream(stream)

    # SUCCESS EXPECTATION: "# 1. Assignment"
    assert "# 1. Assignment" in extracted


def test_insert_boilerplate_creates_paragraphs():
    """
    FIXED: Inserting multi-paragraph text creates actual new paragraphs in DOCX.
    """
    doc = Document()
    doc.add_paragraph("Clause 1: Term.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Simulating LLM inserting a full new clause with structure
    boilerplate = "\n\nClause 2: Termination.\nEither party may terminate this agreement."

    edit = DocumentEdit(target_text="Clause 1: Term.", new_text="Clause 1: Term." + boilerplate)

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    # SUCCESS EXPECTATION:
    # Original: 1 paragraph
    # Inserted: 2 new paragraphs (Clause 2, Either party)
    # Total: 3 paragraphs
    assert len(doc_result.paragraphs) == 3

    # Check text content of new paragraphs
    # Note: docx.Paragraph.text does not see tracked changes (w:ins). We must use our helper.
    p1_text = "".join(r.text for r in get_visible_runs(doc_result.paragraphs[1]))
    p2_text = "".join(r.text for r in get_visible_runs(doc_result.paragraphs[2]))

    assert p1_text == "Clause 2: Termination."
    assert p2_text == "Either party may terminate this agreement."


def test_insert_boilerplate_with_comment_attaches_correctly():
    """
    REGRESSION: Ensure that when track_insert creates new paragraphs (Block path),
    it attaches the comment to the new content.
    Previously, track_insert returned None for blocks, bypassing _attach_comment in the caller.
    """
    doc = Document()
    doc.add_paragraph("Old Header.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # "#### New Header" triggers header detection -> Block insertion logic.
    edit = DocumentEdit(target_text="Old Header.", new_text="#### New Header", comment="Changed header style.")

    engine = RedlineEngine(stream, author="Tester")
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1

    result_stream = engine.save_to_stream()

    # Verify via Ingest that Comment Metadata exists
    text = extract_text_from_stream(result_stream)

    # Should see [Com:X] ... : Changed header style.
    assert "Changed header style." in text

    # Verify XML structure
    doc_result = Document(result_stream)

    # Check if comments part exists and has content
    comments_part = None
    for rel in doc_result.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break

    assert comments_part is not None
    comments_xml = comments_part.blob.decode("utf-8")
    assert "Changed header style." in comments_xml

    # Check if the new paragraph contains commentRangeStart/End
    # The new header is likely the second paragraph (after the deleted "Old Header.")
    # Or modification replaces content?
    # Modification: Delete runs in P1. Insert New P (Heading) AFTER P1.
    # So P1 is now empty (or contains deleted text). P2 is New Header.

    p2 = doc_result.paragraphs[1]  # "New Header"
    xml = p2._element.xml

    assert "w:commentRangeStart" in xml
    assert "w:commentRangeEnd" in xml
