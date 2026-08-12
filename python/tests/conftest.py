import io
import logging
import sys

import pytest
import structlog
from docx import Document

from adeu.utils.console import dynamic_stderr

# Unconfigured structlog prints DEBUG lines to STDOUT, so any test that calls
# engine/ingest helpers in its setup pollutes captured stdout (a `--json` CLI
# test then fails json.loads on "2026-…" log lines). The CLI and server both
# bind structlog to stderr at entry; tests only got that binding if some
# earlier test happened to trigger it — an order dependence pytest-xdist's
# fresh workers expose. Configure it here, once per worker, exactly the way
# the CLI does (WARNING level, dynamic stderr proxy so capsys replacement and
# teardown are honored — never pin the sys.stderr object itself).
structlog.configure(
    wrapper_class=structlog.make_filtering_bound_logger(logging.WARNING),
    logger_factory=structlog.PrintLoggerFactory(file=dynamic_stderr),  # type: ignore[arg-type]
)

try:
    from hypothesis import HealthCheck
    from hypothesis import settings as _hyp_settings

    # Property-test profiles (tests/test_property_invariants.py). Registered
    # here so `--hypothesis-profile=hunt` resolves at pytest configure time.
    _hyp_settings.register_profile(
        "default", deadline=None, max_examples=25, suppress_health_check=[HealthCheck.too_slow]
    )
    _hyp_settings.register_profile(
        "hunt", deadline=None, max_examples=300, suppress_health_check=[HealthCheck.too_slow]
    )
    _hyp_settings.load_profile("default")
except ImportError:
    pass


@pytest.hookimpl(tryfirst=True)
def pytest_collection_modifyitems(config, items):
    """Serialize live-Word tests under pytest-xdist.

    tryfirst matters: in each xdist worker, WorkerInteractor's own
    pytest_collection_modifyitems consumes the xdist_group marker (it rewrites
    the nodeid to "…@group" for the loadgroup scheduler). Our marker must be
    attached before that hook runs, or the grouping silently does nothing.

    Tests that drive the real Word COM instance (the `active_word_app` and
    `word_app` fixtures, plus everything in test_live_word*.py) all bind to the
    single active Word application — two xdist workers doing that concurrently
    corrupt each other's document state. `xdist_group` + `--dist loadgroup`
    (set in pyproject addopts) pins them to ONE worker, where they run
    sequentially; every other test still distributes freely.
    """
    com_fixtures = {"active_word_app", "word_app"}
    for item in items:
        is_live_word = bool(com_fixtures & set(getattr(item, "fixturenames", ())))
        if not is_live_word:
            basename = item.path.name if getattr(item, "path", None) else ""
            is_live_word = basename.startswith("test_live_word")
        if is_live_word:
            item.add_marker(pytest.mark.xdist_group("live_word"))


@pytest.fixture(scope="session", autouse=True)
def _isolate_windows_appdata(tmp_path_factory):
    """On Windows, `adeu init` resolves the Claude Desktop config via %APPDATA%.
    A test that runs init without patching _get_claude_config_path would rewrite
    the developer's real claude_desktop_config.json (this happened 2026-07-20:
    two QA-repro tests injected fake uvx entries into a live config). Pointing
    APPDATA at a throwaway directory for the whole session makes that class of
    accident impossible; tests that assert on the config still patch the path
    getter explicitly."""
    if sys.platform != "win32":
        yield
        return
    mp = pytest.MonkeyPatch()
    mp.setenv("APPDATA", str(tmp_path_factory.mktemp("appdata")))
    yield
    mp.undo()


@pytest.fixture
def simple_docx_stream():
    """Returns a BytesIO stream containing a simple DOCX."""
    doc = Document()
    doc.add_heading("Contract Agreement", 0)
    doc.add_paragraph("This is a simple contract.")
    doc.add_paragraph("The party of the first part shall be known as the Seller.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


# Only define COM fixtures on Windows
if sys.platform == "win32":
    import pythoncom
    import win32com.client

    @pytest.fixture
    def active_word_app():
        """
        Creates an ephemeral, visible MS Word instance with a fresh document.
        Ensures it is torn down properly after the test.
        """
        pythoncom.CoInitialize()

        app = None
        try:
            # Dispatch starts a new background instance if one doesn't exist.
            # GetActiveObject will then be able to hook into it in the tool.
            app = win32com.client.Dispatch("Word.Application")
            app.Visible = True  # Needs to be visible/active for GetActiveObject sometimes
            doc = app.Documents.Add()

            # Bring to front so GetActiveObject definitely binds to this instance
            app.Activate()

            # Seed initial content
            doc.Range(0, 0).Text = "Hello world! This is a live testing document.\n"

            yield app, doc

        except Exception as e:
            pytest.skip(f"Could not initialize Word COM for testing: {e}")

        finally:
            if app:
                try:
                    doc.Close(0)  # 0 = wdDoNotSaveChanges
                except Exception:
                    pass
                # We intentionally omit app.Quit() and pythoncom.CoUninitialize()
                # to avoid Windows Access Violations (0x800706be) when Pytest holds COM locals.

    @pytest.fixture(scope="session")
    def word_app():
        """
        A Word application for INSPECTING packages the engine wrote to disk
        (tests/word_com.py). Distinct from `active_word_app`, which drives an
        open document through the live-Word tools: this one opens saved files
        read-only and is the oracle for ids Word reinterprets on load —
        w14:paraId threading, w16cid:durableId anchoring
        (BUG_paraId_signed_int32_thread_collapse.md).

        Session-scoped because starting Word costs seconds and the tests only
        read. Like `active_word_app` it deliberately omits app.Quit() and
        pythoncom.CoUninitialize(): tearing COM down while pytest still holds
        proxies raises 0x800706be. Not quitting is also what makes it safe to
        attach to a developer's already-running Word.

        `Visible` is deliberately NOT touched. Dispatch attaches to the running
        Word instance, so this is the SAME application object `active_word_app`
        drives and the developer has open; forcing it hidden changes what
        `GetActiveObject` binds to in the live-Word tools. Reading document
        properties does not need visibility. `DisplayAlerts` does have to go:
        without it, opening a package Word considers corrupted — which is one
        of the outcomes under test — blocks on a modal dialog forever.
        """
        pythoncom.CoInitialize()
        try:
            app = win32com.client.Dispatch("Word.Application")
            app.DisplayAlerts = 0
        except Exception as e:
            pytest.skip(f"Could not initialize Word COM for testing: {e}")
            return
        yield app

else:  # pragma: no cover - non-Windows CI

    @pytest.fixture(scope="session")
    def word_app():
        pytest.skip("Live Word COM tests require Windows")
