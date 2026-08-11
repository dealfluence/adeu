import ast
import asyncio
from pathlib import Path

from fastmcp.tools import ToolResult

from adeu.server import mcp


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass

    async def report_progress(self, progress, total=None, message=None):
        pass


def test_reasoning_not_in_required():
    """Check tool schemas across all registered MCP tools and assert 'reasoning' is not present in required[]."""
    tools = asyncio.run(mcp.list_tools())
    assert len(tools) > 0, "No tools registered on mcp server"
    for tool in tools:
        schema = tool.parameters if getattr(tool, "parameters", None) else {}
        required = schema.get("required", []) if isinstance(schema, dict) else []
        assert "reasoning" not in required, f"'reasoning' must not be required in tool schema for {tool.name}"


def test_reasoning_still_advertised():
    """Assert 'reasoning' remains listed in properties schema for all MCP tools."""
    tools = asyncio.run(mcp.list_tools())
    assert len(tools) > 0, "No tools registered on mcp server"
    for tool in tools:
        schema = tool.parameters if getattr(tool, "parameters", None) else {}
        properties = schema.get("properties", {}) if isinstance(schema, dict) else {}
        assert "reasoning" in properties, f"'reasoning' missing from properties schema for tool {tool.name}"
        reasoning_prop = properties["reasoning"]
        assert isinstance(reasoning_prop, dict)
        assert reasoning_prop.get("description"), f"'reasoning' property missing description in tool {tool.name}"


def test_call_without_reasoning_succeeds(tmp_path):
    """Invoke MCP tools without passing reasoning and verify calls succeed."""
    from docx import Document

    from adeu.mcp_components.tools.document import (
        accept_all_changes,
        diff_docx_files,
        process_document_batch,
        read_docx,
    )
    from adeu.mcp_components.tools.sanitize import sanitize_docx

    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.save(doc_path)

    ctx = MockContext()

    # read_docx without reasoning
    res_read = asyncio.run(read_docx(file_path=str(doc_path), ctx=ctx))
    assert isinstance(res_read, ToolResult) or res_read is not None

    # diff_docx_files without reasoning
    res_diff = asyncio.run(diff_docx_files(original_path=str(doc_path), modified_path=str(doc_path), ctx=ctx))
    assert "Word Patch" in res_diff or "No text differences" in res_diff

    # accept_all_changes without reasoning
    res_accept = asyncio.run(accept_all_changes(docx_path=str(doc_path), ctx=ctx))
    assert "already clean" in res_accept or "Accepted all changes" in res_accept

    # process_document_batch without reasoning
    res_batch = asyncio.run(
        process_document_batch(
            original_docx_path=str(doc_path),
            changes=[{"type": "modify", "target_text": "Hello", "new_text": "Hi"}],
            ctx=ctx,
        )
    )
    assert "Batch complete" in res_batch

    # sanitize_docx without reasoning
    res_sanitize = asyncio.run(sanitize_docx(file_path=str(doc_path), ctx=ctx))
    assert isinstance(res_sanitize, dict)
    assert res_sanitize.get("status") in ("clean", "blocked", "sanitized") or "output_path" in res_sanitize


def test_call_with_reasoning_still_succeeds(tmp_path):
    """Invoke MCP tools with reasoning="explanation" and verify calls succeed."""
    from docx import Document

    from adeu.mcp_components.tools.document import (
        accept_all_changes,
        diff_docx_files,
        process_document_batch,
        read_docx,
    )
    from adeu.mcp_components.tools.sanitize import sanitize_docx

    doc_path = tmp_path / "sample2.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.save(doc_path)

    ctx = MockContext()

    # read_docx with reasoning
    res_read = asyncio.run(read_docx(reasoning="Checking document content", file_path=str(doc_path), ctx=ctx))
    assert isinstance(res_read, ToolResult) or res_read is not None

    # diff_docx_files with reasoning
    res_diff = asyncio.run(
        diff_docx_files(
            reasoning="Comparing file against itself",
            original_path=str(doc_path),
            modified_path=str(doc_path),
            ctx=ctx,
        )
    )
    assert "Word Patch" in res_diff or "No text differences" in res_diff

    # accept_all_changes with reasoning
    res_accept = asyncio.run(accept_all_changes(reasoning="Finalizing document", docx_path=str(doc_path), ctx=ctx))
    assert "already clean" in res_accept or "Accepted all changes" in res_accept

    # process_document_batch with reasoning
    res_batch = asyncio.run(
        process_document_batch(
            reasoning="Applying requested edits",
            original_docx_path=str(doc_path),
            changes=[{"type": "modify", "target_text": "Hello", "new_text": "Hi"}],
            ctx=ctx,
        )
    )
    assert "Batch complete" in res_batch

    # sanitize_docx with reasoning
    res_sanitize = asyncio.run(sanitize_docx(reasoning="Scrubbing document metadata", file_path=str(doc_path), ctx=ctx))
    assert isinstance(res_sanitize, dict)
    assert res_sanitize.get("status") in ("clean", "blocked", "sanitized") or "output_path" in res_sanitize


def test_no_positional_callers_remain():
    """Static test verifying keyword/default invocations across test/tool callers."""
    python_dir = Path(__file__).parent.parent
    target_tools = {
        "read_docx",
        "process_document_batch",
        "diff_docx_files",
        "accept_all_changes",
        "open_local_file",
        "sanitize_docx",
        "debug_xml_diff",
        "open_word_document",
        "save_active_word_document",
    }

    positional_reasoning_callers = []

    for search_dir in [python_dir / "src", python_dir / "tests"]:
        for py_file in search_dir.rglob("*.py"):
            try:
                content = py_file.read_text(encoding="utf-8")
                tree = ast.parse(content, filename=str(py_file))
            except Exception:
                continue

            for node in ast.walk(tree):
                if isinstance(node, ast.Call):
                    func_name = None
                    if isinstance(node.func, ast.Name):
                        func_name = node.func.id
                    elif isinstance(node.func, ast.Attribute):
                        func_name = node.func.attr

                    if func_name in target_tools:
                        # If positional arguments are supplied, check if the first
                        # argument looks like a reasoning string
                        if node.args:
                            first_arg = node.args[0]
                            # Check if reasoning keyword is also present
                            has_reasoning_kw = any(kw.arg == "reasoning" for kw in node.keywords)
                            if (
                                not has_reasoning_kw
                                and isinstance(first_arg, ast.Constant)
                                and isinstance(first_arg.value, str)
                            ):
                                # If the string argument is a reasoning message (e.g. "test", "Why...", "Checking...")
                                val = first_arg.value
                                if val in ("test", "reasoning", "explanation") or "why" in val.lower():
                                    positional_reasoning_callers.append(
                                        f"{py_file.relative_to(python_dir)}:{node.lineno} {func_name}('{val}')"
                                    )

    assert not positional_reasoning_callers, f"Found positional reasoning callers: {positional_reasoning_callers}"
