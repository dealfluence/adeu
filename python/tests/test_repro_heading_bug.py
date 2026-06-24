import io
import pytest
from docx import Document

from adeu.models import ModifyText
from adeu.redline.engine import BatchValidationError, RedlineEngine


class TestReproHeadingBug:
    def test_tc1_heading_prefix_no_duplicate(self):
        """
        TC-1: '#' prefixed target, no duplicate — Python succeeds
        Target: "# 2. Confidentiality", new_text: "## 2. Confidentiality"
        Expected: Heading demoted to level 2 (Heading 2 style), and text is 2. Confidentiality.
        """
        doc = Document()
        p = doc.add_paragraph("2. Confidentiality")
        p.style = "Heading 1"
        doc.add_paragraph("As defined in Section 1, the Recipient shall...")
        
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        engine = RedlineEngine(stream)
        edit = ModifyText(
            target_text="# 2. Confidentiality",
            new_text="## 2. Confidentiality"
        )
        
        # Python should apply this without error
        engine.process_batch([edit])
        engine.accept_all_revisions()
        
        # Save and verify
        res = Document(engine.save_to_stream())
        
        # Check that paragraph 0 is now styled as Heading 2
        assert res.paragraphs[0].style.name == "Heading 2"
        # The text of the paragraph should contain 2. Confidentiality
        assert "2. Confidentiality" in res.paragraphs[0].text

    def test_tc2_bare_target_with_duplicate(self):
        """
        TC-2: Bare target, duplicate in body — ambiguity error (both engines, correct)
        Target: "2. Confidentiality", new_text: "2. CONFIDENTIALITY"
        Expected: Ambiguity error
        """
        doc = Document()
        p1 = doc.add_paragraph("2. Confidentiality")
        p1.style = "Heading 1"
        doc.add_paragraph("As defined in Section 1, the Recipient shall...")
        
        # Add a body paragraph containing the heading text
        doc.add_paragraph("Page footer notice: subject to NDA dated 2026-01-15.")
        doc.add_paragraph("For further detail see section 2. Confidentiality above.")
        
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        engine = RedlineEngine(stream)
        edit = ModifyText(
            target_text="2. Confidentiality",
            new_text="2. CONFIDENTIALITY"
        )
        
        with pytest.raises(BatchValidationError, match="[Aa]mbiguous"):
            engine.process_batch([edit])

    def test_tc3_heading_prefix_with_duplicate(self):
        """
        TC-3: '#' prefixed target, duplicate in body — Python disambiguates
        Target: "# 2. Confidentiality", new_text: "## 2. Confidentiality"
        Expected: Applied. Paragraph style changed to Heading 2, body paragraph untouched.
        """
        doc = Document()
        p1 = doc.add_paragraph("2. Confidentiality")
        p1.style = "Heading 1"
        doc.add_paragraph("As defined in Section 1, the Recipient shall...")
        
        # Add a body paragraph containing the heading text
        doc.add_paragraph("Page footer notice: subject to NDA dated 2026-01-15.")
        doc.add_paragraph("For further detail see section 2. Confidentiality above.")
        
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        engine = RedlineEngine(stream)
        edit = ModifyText(
            target_text="# 2. Confidentiality",
            new_text="## 2. Confidentiality"
        )
        
        # Python should successfully disambiguate and apply
        engine.process_batch([edit])
        engine.accept_all_revisions()
        
        res = Document(engine.save_to_stream())
        # Paragraph 0 style is Heading 2
        assert res.paragraphs[0].style.name == "Heading 2"
        assert "2. Confidentiality" in res.paragraphs[0].text
        
        # Paragraph 3 text remains untouched (not demoted, nor modified)
        assert "section 2. Confidentiality above." in res.paragraphs[3].text
