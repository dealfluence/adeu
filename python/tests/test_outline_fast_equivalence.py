"""
Pins the cache-backed outline helpers (_safe_style_name_fast /
_determine_heading_style_fast) against the python-docx-property originals
they replace. The originals resolve the styles part through the document
part's relationship list on every access — O(rels) per call, the dominant
cost of mode='outline' on large documents — but they remain in the module
as the legacy-path implementation and as the executable specification here.
"""

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from adeu.outline import (
    _determine_heading_style,
    _determine_heading_style_fast,
    _safe_style_name,
    _safe_style_name_fast,
    extract_outline,
)
from adeu.utils.docx import _get_style_cache, iter_block_items


def _build_style_variety_doc():
    doc = Document()
    doc.add_paragraph("Title paragraph", style="Title")
    doc.add_heading("Native heading one", level=1)
    doc.add_heading("Deep heading", level=4)
    doc.add_paragraph("Plain body paragraph")
    doc.add_paragraph("List item", style="List Bullet")
    doc.add_paragraph("Quote text", style="Quote")

    # Custom style whose name embeds a heading token.
    custom = doc.styles.add_style("StyleHeading2NotItalic", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Custom quick-style heading", style=custom)

    # Direct outline level WITHOUT a heading style.
    p = doc.add_paragraph("Outline-leveled body text")
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn

    oLvl = pPr.makeelement(qn("w:outlineLvl"), {qn("w:val"): "2"})
    pPr.append(oLvl)

    # Paragraph referencing a style id that does not exist (unknown pStyle).
    p2 = doc.add_paragraph("Dangling style reference")
    pPr2 = p2._p.get_or_add_pPr()
    pStyle = pPr2.makeelement(qn("w:pStyle"), {qn("w:val"): "NoSuchStyleId"})
    pPr2.insert(0, pStyle)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


def _all_paragraphs(doc):
    for item in iter_block_items(doc):
        if type(item).__name__ == "Paragraph":
            yield item


def test_safe_style_name_fast_matches_original():
    doc = _build_style_variety_doc()
    style_cache, default_pstyle = _get_style_cache(doc.part)
    for para in _all_paragraphs(doc):
        assert _safe_style_name_fast(para, style_cache, default_pstyle) == _safe_style_name(para), para.text


def test_determine_heading_style_fast_matches_original():
    doc = _build_style_variety_doc()
    style_cache, default_pstyle = _get_style_cache(doc.part)
    for para in _all_paragraphs(doc):
        assert _determine_heading_style_fast(para, style_cache, default_pstyle) == _determine_heading_style(para), (
            para.text
        )


def test_fast_and_legacy_outline_paths_agree():
    """
    extract_outline with paragraph_offsets (fast path, what the server uses)
    must produce the same nodes as the legacy walk (paragraph_offsets=None)
    on a document with headings, tables, and footnote-free bodies.
    """
    from adeu.ingest import _extract_text_from_doc
    from adeu.pagination import paginate, split_structural_appendix

    doc = _build_style_variety_doc()
    doc.add_heading("Section with table", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "cell"
    doc.add_paragraph("After table")

    text, offsets = _extract_text_from_doc(doc, clean_view=False, include_appendix=False, return_paragraph_offsets=True)
    body, _ = split_structural_appendix(text)
    pag = paginate(body, "")

    fast_nodes = extract_outline(doc, body, pag.body_pages, pag.body_page_offsets, paragraph_offsets=offsets)
    legacy_nodes = extract_outline(doc, body, pag.body_pages, pag.body_page_offsets, paragraph_offsets=None)

    assert [(n.level, n.text, n.page, n.style, n.has_table, n.footnote_ids, n.end_page) for n in fast_nodes] == [
        (n.level, n.text, n.page, n.style, n.has_table, n.footnote_ids, n.end_page) for n in legacy_nodes
    ]
