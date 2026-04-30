import docx


def test_diff_engine_splits_markdown_tokens(tmp_path):
    """
    Reproduces Bug 8: Diff engine breaks mid-markdown token.
    """
    d1_path = tmp_path / "diff_markdown_1.docx"
    d2_path = tmp_path / "diff_markdown_2.docx"

    d1 = docx.Document()
    d1.add_paragraph("ons)**  **(In millions)**          Dear")
    d1.save(str(d1_path))

    d2 = docx.Document()
    d2.add_paragraph("ons)**  **(In millions)**\n\n\n\n\n\n\n\n\nDear shareholders:")
    d2.save(str(d2_path))

    import subprocess

    diff_output = subprocess.check_output(["uv", "run", "adeu", "diff", str(d1_path), str(d2_path)]).decode("utf-8")

    # The bug produces `- *(In millions)**` which is invalid markdown.
    assert "- *(In millions)**" not in diff_output, "BUG: Diff engine split a markdown token!"
