import io
import json
import re

import pytest

from adeu.payloads import shrink_batch_stats
from tests.utils import approx_tokens, extract_content, get_mock_ctx, run_async

_BUBBLE_RE = re.compile(r"\{--.*?--\}|\{\+\+.*?\+\+\}|\{==.*?==\}|\{>>.*?<<\}", re.DOTALL)
_DELIMITERS = ("{--", "--}", "{++", "++}", "{==", "==}", "{>>", "<<}")


def _assert_balanced_critic_markup(markup):
    """
    A shrunken preview must never leave a bare CriticMarkup delimiter behind:
    an orphaned "{++" corrupts the markup for every consumer (AI_CONTEXT.md
    display invariant).
    """
    outside_bubbles = _BUBBLE_RE.sub("", markup)
    for delim in _DELIMITERS:
        assert delim not in outside_bubbles, f"orphaned {delim!r} in preview: {markup!r}"


def _legal_batch(tmp_path):
    """A realistic legal edit: deep heading path, long comment, long spans."""
    from docx import Document

    from adeu.models import DocumentChange, ModifyText
    from adeu.redline.engine import RedlineEngine

    doc_path = tmp_path / "legal.docx"
    doc = Document()
    doc.add_heading("Article IV Representations and Warranties", level=1)
    doc.add_heading("Section 4.2 Limitation of Liability", level=2)
    doc.add_paragraph(
        "The Seller shall be liable for any and all indirect or consequential damages "
        "arising out of or relating to this Agreement, including lost profits."
    )
    doc.add_paragraph("The quick brown fox jumps over the lazy dog in a very detailed manner.")
    doc.save(doc_path)

    engine = RedlineEngine(io.BytesIO(doc_path.read_bytes()), author="Tester")
    changes: list[DocumentChange] = [
        ModifyText(
            type="modify",
            target_text="shall be liable for any and all indirect or consequential damages",
            new_text="shall not be liable for any indirect or consequential damages",
            comment="Per client instruction 2026-08-03, cap indirect liability exposure.",
        ),
        ModifyText(
            type="modify",
            target_text="Section 4.2 Limitation of Liability",
            new_text="Section 4.2 Limitation of Seller Liability",
            comment=None,
        ),
        ModifyText(type="modify", target_text="lazy dog", new_text="", comment=None),
    ]
    return engine.process_batch(changes)


def test_minimal_report_drops_echoes():
    stats = {
        "engine": "python",
        "version": "2.0.0",
        "edits": [
            {
                "status": "applied",
                "type": "modify",
                "target_text": "The quick brown fox jumps over the lazy dog.",
                "new_text": "The fast brown fox jumps over the lazy dog.",
                "clean_text": "The fast brown fox jumps over the lazy dog.",
                "critic_markup": "{--quick--}{++fast++}",
                "pages": [1],
                "heading_path": "Introduction",
                "occurrences_modified": 1,
                "match_mode": "strict",
            }
        ],
    }
    shrunk = shrink_batch_stats(stats)
    edit = shrunk["edits"][0]
    assert "target_text" not in edit
    assert "new_text" not in edit
    assert "clean_text" not in edit


def test_minimal_report_keeps_verification_fields():
    """An edit that already fits the budget is passed through untouched."""
    stats = {
        "engine": "python",
        "version": "2.0.0",
        "edits": [
            {
                "status": "applied",
                "type": "modify",
                "target_text": "old",
                "new_text": "new",
                "clean_text": "new",
                "critic_markup": "{--old--}{++new++}",
                "pages": [1, 2],
                "heading_path": "Section 1 > Part A",
                "occurrences_modified": 1,
                "match_mode": "strict",
            }
        ],
    }
    shrunk = shrink_batch_stats(stats)
    edit = shrunk["edits"][0]
    assert edit["status"] == "applied"
    assert edit["type"] == "modify"
    assert edit["pages"] == [1, 2]
    assert edit["heading_path"] == "Section 1 > Part A"
    assert edit["occurrences_modified"] == 1
    assert edit["critic_markup"] == "{--old--}{++new++}"
    assert approx_tokens(json.dumps(edit)) <= 40


def test_minimal_report_spends_locator_before_evidence():
    """
    When the budget bites, the heading path gives way first — `pages` already
    says where the edit landed — and the CriticMarkup evidence survives whole.
    """
    deep_path = "Article IV Representations and Warranties > Section 4.2 Liability"

    def report(markup):
        stats = {
            "edits": [
                {
                    "status": "applied",
                    "type": "modify",
                    "critic_markup": markup,
                    "pages": [4],
                    "heading_path": deep_path,
                    "occurrences_modified": 1,
                }
            ]
        }
        return shrink_batch_stats(stats)["edits"][0]

    # Ancestors are the first thing surrendered: the deepest heading is the
    # only specific part of the path.
    edit = report("{--old--}{++new++}")
    assert edit["critic_markup"] == "{--old--}{++new++}"
    assert edit["heading_path"] == "Section 4.2 Liability"
    assert approx_tokens(json.dumps(edit)) <= 40

    # A preview long enough to need the whole budget takes it: the locator
    # goes entirely (`pages` still says where), the evidence stays whole.
    edit = report("{--old wording--}{++new wording++}")
    assert edit["critic_markup"] == "{--old wording--}{++new wording++}"
    assert "heading_path" not in edit
    assert edit["pages"] == [4]
    assert approx_tokens(json.dumps(edit)) <= 40


def test_minimal_report_keeps_match_mode_only_when_non_strict():
    stats = {
        "edits": [
            {"status": "applied", "match_mode": "strict"},
            {"status": "applied", "match_mode": "first"},
            {"status": "applied", "match_mode": "all"},
        ]
    }
    shrunk = shrink_batch_stats(stats)
    assert "match_mode" not in shrunk["edits"][0]
    assert shrunk["edits"][1]["match_mode"] == "first"
    assert shrunk["edits"][2]["match_mode"] == "all"


def test_failed_edit_keeps_full_error_and_stub_target():
    long_target = "A" * 120
    full_error = "- Edit 1 Failed: Target text 'AAAA...' was not found anywhere in the active document projection."
    stats = {
        "edits": [
            {
                "status": "failed",
                "type": "modify",
                "target_text": long_target,
                "new_text": "replacement",
                "clean_text": "clean",
                "error": full_error,
                "match_mode": "strict",
            }
        ]
    }
    shrunk = shrink_batch_stats(stats)
    edit = shrunk["edits"][0]
    assert edit["status"] == "failed"
    assert edit["error"] == full_error
    assert "target_text" in edit
    assert len(edit["target_text"]) <= 80
    assert "new_text" not in edit
    assert "clean_text" not in edit


def test_minimal_report_token_budget(tmp_path):
    """Real engine edits — long comment, deep heading path, legal-length spans."""
    shrunk = shrink_batch_stats(_legal_batch(tmp_path))

    assert shrunk["edits"], "no edit reports to budget"
    for edit in shrunk["edits"]:
        dumped = json.dumps(edit)
        assert approx_tokens(dumped) <= 40, (
            f"Edit JSON exceeded 40 approx-tokens budget ({len(dumped)} chars): {dumped}"
        )


def _fanout_batch(tmp_path, filler, page_breaks=False, occurrences=6):
    """
    A real match_mode="all" fan-out over `occurrences` separated paragraphs.
    `filler` controls the gap between hits, and so whether the engine renders
    one merged preview window or one window per occurrence joined by
    separators; `page_breaks` spreads the hits over one page each, which is
    what makes the report's `pages` list long. Every shape must be budgeted.
    """
    from docx import Document
    from docx.enum.text import WD_BREAK

    from adeu.models import DocumentChange, ModifyText
    from adeu.redline.engine import RedlineEngine

    doc_path = tmp_path / "fanout.docx"
    doc = Document()
    doc.add_heading("Article IX Miscellaneous Provisions", level=1)
    for i in range(occurrences):
        para = doc.add_paragraph(f"Clause {i + 1}: The Consultant shall deliver the work product.{filler}")
        if page_breaks:
            para.add_run().add_break(WD_BREAK.PAGE)
    doc.save(doc_path)

    engine = RedlineEngine(io.BytesIO(doc_path.read_bytes()), author="Tester")
    changes: list[DocumentChange] = [
        ModifyText(
            type="modify",
            target_text="Consultant",
            new_text="Contractor",
            match_mode="all",
            comment=None,
        )
    ]
    return engine.process_batch(changes)


_SPACED_CLAUSE = (
    " Such delivery shall be subject to the terms of this Agreement, to any"
    " schedule agreed between the parties and to applicable law."
)


@pytest.mark.parametrize(
    ("filler", "page_breaks"),
    [
        ("", False),
        (_SPACED_CLAUSE, False),
        (_SPACED_CLAUSE, True),
    ],
    ids=["merged-window", "separate-windows", "one-page-each"],
)
def test_minimal_report_token_budget_multi_occurrence(tmp_path, filler, page_breaks):
    """
    A fan-out preview is many context windows, not one: the per-edit budget
    still holds, and what survives is still valid evidence.
    """
    shrunk = shrink_batch_stats(_fanout_batch(tmp_path, filler, page_breaks))

    edit = shrunk["edits"][0]
    dumped = json.dumps(edit)
    assert approx_tokens(dumped) <= 40, f"Fan-out edit exceeded 40 approx-tokens budget ({len(dumped)} chars): {dumped}"
    markup = edit["critic_markup"]
    _assert_balanced_critic_markup(markup)
    assert _BUBBLE_RE.search(markup), f"fan-out preview kept no changed span: {markup!r}"


def test_minimal_report_keeps_valid_critic_markup(tmp_path):
    """
    The budget is paid with echoed input and context, never by cutting a
    CriticMarkup bubble open: the preview stays balanced and still shows the
    changed span.
    """
    shrunk = shrink_batch_stats(_legal_batch(tmp_path))

    previews = [e["critic_markup"] for e in shrunk["edits"] if e.get("critic_markup")]
    assert len(previews) == len(shrunk["edits"]), "an applied edit lost its critic_markup evidence"
    for markup in previews:
        _assert_balanced_critic_markup(markup)
        assert _BUBBLE_RE.search(markup), f"preview no longer shows a changed span: {markup!r}"

    liability = shrunk["edits"][0]["critic_markup"]
    assert "{--" in liability and "--}" in liability
    assert "not be liable" in liability, f"the substantive change is gone: {liability!r}"


def test_minimal_report_drops_comment_echo(tmp_path):
    """`comment` is caller input; the minimal payload does not echo it back."""
    shrunk = shrink_batch_stats(_legal_batch(tmp_path))
    assert "comment" not in shrunk["edits"][0]


def test_minimal_report_keeps_warning_within_budget():
    """
    An engine advisory is kept — the agent must know the edit needs a second
    look — but it is budgeted like every other field of an applied edit: only
    a failed edit's `error` rides free (criterion 4). What survives still names
    the offending token.
    """
    warning = (
        "new_text contains '$1', which Python's re engine does not expand — the literal text "
        "'$1' was written into the document. For a capture-group backreference use \\1 or \\g<1>."
    )
    stats = {
        "edits": [
            {
                "status": "applied",
                "type": "modify",
                "target_text": "old",
                "new_text": "new",
                "critic_markup": "{--old--}{++new++}",
                "warning": warning,
                "occurrences_modified": 1,
            }
        ]
    }
    edit = shrink_batch_stats(stats)["edits"][0]
    dumped = json.dumps(edit)
    assert approx_tokens(dumped) <= 40, f"Warned edit exceeded 40 approx-tokens budget ({len(dumped)} chars): {dumped}"
    assert warning.startswith(edit["warning"].rstrip(".")), "the warning was rewritten, not clamped"
    assert "$1" in edit["warning"], f"the warning no longer names the offending token: {edit['warning']!r}"


@pytest.mark.parametrize(
    ("match_mode", "occurrences"),
    [("strict", 1), ("all", 6)],
    ids=["single-occurrence", "fan-out"],
)
def test_minimal_report_budgets_regex_backreference_warning(tmp_path, match_mode, occurrences):
    """
    A real surviving-$N advisory (the engine's JavaScript-backreference
    guardrail) is the fattest field an applied edit can carry: ~260 chars of
    prose, over six times the whole per-edit budget. It is clamped to fit — in
    both a single-hit edit and a fan-out, whose counters and span-counted
    preview leave the advisory least room — and what survives still names '$1'.
    """
    from docx import Document

    from adeu.models import DocumentChange, ModifyText
    from adeu.redline.engine import RedlineEngine

    doc_path = tmp_path / "backref.docx"
    doc = Document()
    for i in range(occurrences):
        doc.add_paragraph(f"Clause {i + 1}: The Consultant shall deliver the work product to the Company.")
    doc.save(str(doc_path))

    engine = RedlineEngine(io.BytesIO(doc_path.read_bytes()), author="Tester")
    changes: list[DocumentChange] = [
        ModifyText(
            type="modify",
            target_text=r"(Consultant) shall deliver",
            new_text=r"$1 must deliver",
            regex=True,
            match_mode=match_mode,
            comment=None,
        )
    ]
    stats = engine.process_batch(changes)
    assert stats["edits"][0].get("warning"), "the engine did not flag the surviving $1 backreference"

    edit = shrink_batch_stats(stats)["edits"][0]
    dumped = json.dumps(edit)
    assert approx_tokens(dumped) <= 40, f"Warned edit exceeded 40 approx-tokens budget ({len(dumped)} chars): {dumped}"
    assert "$1" in edit["warning"], f"the warning no longer names the offending token: {edit['warning']!r}"
    # Whatever preview survives the budget is still valid CriticMarkup; under
    # maximum pressure the preview goes whole rather than shipping a fragment.
    if "critic_markup" in edit:
        _assert_balanced_critic_markup(edit["critic_markup"])
        assert _BUBBLE_RE.search(edit["critic_markup"])


def test_standard_report_is_unchanged(tmp_path, capsys):
    import sys
    from unittest.mock import patch

    from docx import Document

    from adeu.cli import main
    from adeu.models import DocumentChange, ModifyText
    from adeu.redline.engine import RedlineEngine

    doc_path = tmp_path / "test_std.docx"
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.save(doc_path)

    engine = RedlineEngine(io.BytesIO(doc_path.read_bytes()), author="Tester")
    changes: list[DocumentChange] = [
        ModifyText(type="modify", target_text="quick", new_text="fast", comment=None),
        ModifyText(type="modify", target_text="fox", new_text="", comment=None),
    ]
    stats = engine.process_batch(changes)

    assert stats["engine"] == "python"
    edit1 = stats["edits"][0]
    assert edit1["target_text"] == "quick"
    assert edit1["new_text"] == "fast"
    assert "fast" in edit1["clean_text"]

    edit2 = stats["edits"][1]
    assert edit2["target_text"] == "fox"
    assert edit2["new_text"] == ""
    assert isinstance(edit2["clean_text"], str)

    changes_file = tmp_path / "changes.json"
    out_path = tmp_path / "out.docx"
    with open(changes_file, "w") as f:
        json.dump(
            [
                {"type": "modify", "target_text": "quick", "new_text": "fast"},
                {"type": "modify", "target_text": "fox", "new_text": ""},
            ],
            f,
        )

    test_args = [
        "adeu",
        "apply",
        str(doc_path),
        str(changes_file),
        "-o",
        str(out_path),
        "--report",
        "standard",
    ]
    with patch.object(sys, "argv", test_args):
        try:
            main()
        except SystemExit as e:
            assert e.code == 0 or e.code is None

    captured = capsys.readouterr()
    err_output = captured.err
    assert "Target: 'fox'" in err_output
    assert "New text: ''" in err_output

    # The standard machine payload keeps every echo and the engine field.
    with patch.object(sys, "argv", test_args + ["--json"]):
        try:
            main()
        except SystemExit as e:
            assert e.code == 0 or e.code is None
    payload = json.loads(capsys.readouterr().out)
    assert payload["engine"] == "python"
    reported = payload["edits"][1]
    assert reported["target_text"] == "fox"
    assert reported["new_text"] == ""
    assert "clean_text" in reported
    assert "comment" in reported


def test_batch_level_keeps_version_drops_engine():
    stats = {
        "engine": "python",
        "version": "2.0.0",
        "actions_applied": 1,
        "edits_applied": 1,
    }
    shrunk = shrink_batch_stats(stats)
    assert "version" in shrunk
    assert shrunk["version"] == "2.0.0"
    assert "engine" not in shrunk


def test_skipped_details_deduped_against_edit_errors():
    err_msg = "- Edit 1 Failed: target text not found"
    stats = {
        "skipped_details": [err_msg, "Other skipped detail"],
        "edits": [
            {
                "status": "failed",
                "error": err_msg,
            }
        ],
    }
    shrunk = shrink_batch_stats(stats)
    assert shrunk["skipped_details"] == ["Other skipped detail"]


def test_skipped_details_key_not_injected_when_absent():
    """A batch that never reported skipped details must not grow an empty list."""
    shrunk = shrink_batch_stats({"version": "2.0.0", "edits": [{"status": "applied"}]})
    assert "skipped_details" not in shrunk


def test_mcp_default_is_minimal(tmp_path):
    from docx import Document

    from adeu.mcp_components.tools.document import process_document_batch
    from adeu.models import ModifyText

    doc_path = tmp_path / "test.docx"
    d = Document()
    d.add_paragraph("Hello world")
    d.save(doc_path)

    ctx = get_mock_ctx()
    changes = [ModifyText(type="modify", target_text="world", new_text="earth", comment=None)]
    res = run_async(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(doc_path),
            author_name="Tester",
            changes=changes,
            ctx=ctx,
        )
    )
    res_text = extract_content(res)
    assert "*Preview (Clean):*" not in res_text
