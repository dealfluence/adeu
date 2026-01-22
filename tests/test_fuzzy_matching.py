import io

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine


def test_placeholder_fuzzy_match():
    """
    Scenario: Document has [_______], User inputs [___].
    Should match and replace correctly.
    """
    doc = Document()
    doc.add_paragraph("Sign here: [__________]")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Target has fewer underscores than doc
    edit = DocumentEdit(target_text="Sign here: [___]", new_text="Sign here: John Doe")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1
    assert skipped == 0

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream)

    # Expect deletion of full placeholder
    assert "{--[__________]--}" in text
    assert "{++John Doe++}" in text


def test_whitespace_fuzzy_match():
    """
    Scenario: Document has 'Word   Word', User inputs 'Word Word'.
    """
    doc = Document()
    doc.add_paragraph("Start   End")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="Start End", new_text="Start Middle End")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream)

    # Context trimming should handle the mismatch.
    # Logic:
    # Actual: "Start   End"
    # New:    "Start Middle End"
    # Prefix "Start" matches (fuzzy whitespace logic in trim? No, trim is strict).
    # Wait, the engine passes 'Actual Text' to trim.
    # Actual: "Start   End"
    # New:    "Start Middle End"
    # Prefix: "Start" (len 5). Matches.
    # Suffix: "End" (len 3). Matches.
    # Result: Replace "   " with " Middle ".

    # Corrected expectation based on actual trimming logic (Prefix 'Start ' and Suffix ' End' consumed)
    # Remaining target is one space. Remaining new is 'Middle'.
    # Resulting text should read "Start Middle End" when accepted.
    # Markup: Start {-- --}{++Middle++} End
    assert "{-- --}{++Middle++}" in text


def test_smart_quote_match():
    """
    Scenario: Doc has smart quotes, user uses straight quotes.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("“Hello”")  # Smart quotes

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text='"Hello"', new_text='"Hi"')

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1

    res_stream = engine.save_to_stream()
    text = extract_text_from_stream(res_stream)

    # Should replace the smart quoted version
    assert "{--“Hello”--}" in text
    assert '{++"Hi"++}' in text
