import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_visible_runs


def _add_list_paragraph(doc, text: str, num_id: str, ilvl: str):
    """Helper to manually inject list numbering properties into a paragraph."""
    p = doc.add_paragraph(text)
    pPr = p._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")

    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    numPr.append(ilvl_el)

    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), num_id)
    numPr.append(numId_el)

    pPr.append(numPr)
    return p


def _create_multilevel_list_docx() -> io.BytesIO:
    doc = Document()
    doc.add_paragraph("List Introduction:")

    # Level 1 (ilvl 0)
    _add_list_paragraph(doc, "Top level item", num_id="1", ilvl="0")

    # Level 2 (ilvl 1)
    _add_list_paragraph(doc, "Nested item", num_id="1", ilvl="1")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_val_obs_10_multilevel_list_projection():
    """
    Test the Read Path:
    OOXML `<w:ilvl>` should project as 4 spaces per level in Markdown.
    """
    stream = _create_multilevel_list_docx()
    projection = extract_text_from_stream(stream)

    # Asserting exact Markdown projection syntax
    # Level 1 should be just "* Top level item" (0 spaces)
    assert "\n* Top level item" in projection, "Failed to project base list item."

    # Level 2 should be "    * Nested item" (4 spaces)
    assert "\n    * Nested item" in projection, (
        f"Failed to project nested list item indentation.\nProjection was:\n{projection}"
    )


def test_val_obs_10_multilevel_list_write_path():
    """
    Test the Write Path:
    Leading spaces in Markdown `new_text` should translate to `<w:ilvl>` in OOXML,
    and the prefix should be stripped from the literal text.
    """
    stream = _create_multilevel_list_docx()
    engine = RedlineEngine(stream)

    # Target the nested item and add a level 3 item (8 spaces)
    # The current engine will likely just copy ilvl=1 from the anchor, but we want it to parse the 8 spaces.
    edit = ModifyText(target_text="Nested item", new_text="Nested item\n        * Deeply nested item")

    applied, _ = engine.apply_edits([edit])
    assert applied == 1, "Edit failed to apply."

    stream_edited = engine.save_to_stream()
    doc_edited = Document(stream_edited)

    # Find the inserted paragraph (should be the 4th paragraph: Intro -> Top -> Nested -> Deeply nested)
    p_inserted = doc_edited.paragraphs[3]

    # Verify the text is clean (no markdown prefixes)
    visible_text = "".join(r.text for r in get_visible_runs(p_inserted))
    assert visible_text == "Deeply nested item", f"Literal text contains markdown prefix: '{visible_text}'"

    # Verify OOXML structure
    pPr = p_inserted._element.find(qn("w:pPr"))
    assert pPr is not None, "Inserted paragraph missing properties"

    numPr = pPr.find(qn("w:numPr"))
    assert numPr is not None, "Inserted paragraph missing <w:numPr> (list tracking)"

    ilvl_el = numPr.find(qn("w:ilvl"))
    assert ilvl_el is not None, "Inserted paragraph missing <w:ilvl>"

    ilvl_val = ilvl_el.get(qn("w:val"))
    assert ilvl_val == "2", f"Expected ilvl='2' for 8 spaces, got '{ilvl_val}'"
