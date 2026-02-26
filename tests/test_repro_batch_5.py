import io
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine, _trim_common_context


def _is_bold(run_element) -> bool:
    """Helper to check if a run element has bold property active."""
    b_tags = run_element.xpath("./w:rPr/w:b")
    if not b_tags:
        return False
    val = b_tags[0].get(qn("w:val"))
    if val is None: return True  # <w:b/> is true
    return val.lower() not in ("0", "false", "off")


# --- 5.1 Tab Deletions ---

def test_repro_5_1_tab_consistency():
    """
    BUG: get_run_text() converts <w:tab/> to " " but literal \t in <w:t> is unchanged.
    This causes diffing errors when the AI provides "Word Word" vs "Word\tWord".
    """
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run()
    
    # Run contains: "A" + <w:tab/> + "B" + "\t" + "C"
    t1 = OxmlElement("w:t")
    t1.text = "A"
    r._element.append(t1)
    
    r._element.append(OxmlElement("w:tab"))
    
    t2 = OxmlElement("w:t")
    t2.text = "B\tC"
    r._element.append(t2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)
    
    # If the bug exists, text will be "A B\tC"
    # We want "A B C" for consistency with AI text processing
    assert "\t" not in text, f"Literal tab found in extracted text: {repr(text)}"
    assert text == "A B C"


# --- 5.3 Bold Formatting Inheritance ---

def test_repro_5_3_bold_inheritance_bleed():
    """
    BUG: Modifications inherit style from the last deleted run. 
    If we replace **Bold** with Plain, the result is often silently Bold.
    """
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("BOLD")
    r.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Edit: Change "BOLD" (bold) to "plain" (no markers)
    edit = DocumentEdit(target_text="BOLD", new_text="plain")
    
    engine = RedlineEngine(stream)
    engine.apply_edits([edit])
    
    res_stream = engine.save_to_stream()
    doc_res = Document(res_stream)
    
    # Find the inserted run
    # In Word XML: <w:ins><w:r><w:t>plain</w:t></w:r></w:ins>
    ins_runs = doc_res.element.xpath('//w:ins//w:r[w:t[text()="plain"]]')
    assert ins_runs, "Insertion 'plain' not found"
    
    # If bug exists, the run inherits the bold property from the deleted "BOLD" run
    assert not _is_bold(ins_runs[0]), "Inserted text 'plain' incorrectly inherited Bold style"


# --- 5.5 Markdown Markers Leaking ---

def test_repro_5_5_trim_context_marker_leak():
    """
    BUG: _trim_common_context may leave unbalanced markdown fragments.
    Example: Target "**Agreement**", New "**Agreements**".
    If it trims "**Agreement", it leaves "**" in the deletion.
    """
    target = "**Agreement**"
    new_val = "**Agreements**"
    
    prefix_len, suffix_len = _trim_common_context(target, new_val)
    
    # Prefix " **Agreement" (len 11) or Suffix "**" (len 2)
    # If prefix_len is 11, remaining target is "}" (if virtual) or empty.
    # If suffix_len is 2, the tail is safe.
    
    # The real danger is if it trims into the middle of a token:
    target_2 = "**Bold**"
    new_2 = "**Bolder**"
    # If it trims "**Bold", prefix_len is 6. 
    # Remaining target: "**". Remaining new: "er**".
    # This results in: **Bold{--**--}{++er**++}
    
    p, s = _trim_common_context(target_2, new_2)
    
    # We expect the trimmer to NOT consume the start of a markdown block 
    # if it doesn't consume the end, OR to back off to a word boundary.
    remaining_target = target_2[p : len(target_2)-s]
    remaining_new = new_2[p : len(new_2)-s]
    
    assert "**" not in remaining_target, "Unbalanced markdown marker leaked into deletion"
    assert "**" not in remaining_new, "Unbalanced markdown marker leaked into insertion"


# --- 5.6 Duplicate Overlapping Insertions ---

def test_repro_5_6_overlapping_edits_collision():
    """
    BUG: If two edits overlap, the second one targets shifted indices and causes corruption.
    """
    doc = Document()
    doc.add_paragraph("The quick brown fox")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    edits = [
        DocumentEdit(target_text="quick brown", new_text="slow red"),
        DocumentEdit(target_text="brown", new_text="tan") # Overlaps with previous
    ]
    
    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits(edits)
    
    # Engine should handle overlaps by skipping or merging. 
    # Currently, it might try to apply both and fail or garble text.
    assert applied == 1, f"Should only apply 1 of the overlapping edits, applied {applied}"
    assert skipped == 1
    
    res_text = extract_text_from_stream(engine.save_to_stream())
    
    # Check that "tan" (the second, overlapping edit) didn't create a double-deletion mess
    # e.g. "{--{--quick brown--}{++slow red++}--}{++tan++}"
    assert "{--{--" not in res_text, "Nested markup detected indicating collision"