import io

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_r1_r4_table_cell_quadruplication_and_leak():
    """
    R1: Single-edit content quadruplication when targeting Word table cells.
    R4: Bold markdown markers leak across table cell boundaries.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "CellA"
    table.cell(0, 1).text = "CellB"
    table.cell(0, 2).text = "CellC"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Target the entire row, modify only the first cell to be bold.
    edit = ModifyText(target_text="CellA | CellB | CellC", new_text="**CellA Bold** | CellB | CellC")

    engine = RedlineEngine(stream)
    stats = engine.process_batch([edit])

    # It must not skip the edit
    assert stats["edits_applied"] == 1, "Edit was skipped."

    # Verify accepted output for content duplication (R1)
    engine.accept_all_revisions()
    clean_stream = engine.save_to_stream()
    clean_text = extract_text_from_stream(clean_stream, clean_view=True)

    # Assert exact counts. R1 caused these to multiply by the number of cells.
    count_a = clean_text.count("CellA Bold")
    count_b = clean_text.count("CellB")
    count_c = clean_text.count("CellC")

    assert count_a == 1, f"R1 Regression: 'CellA Bold' appears {count_a} times. Text:\n{clean_text}"
    assert count_b == 1, f"R1 Regression: 'CellB' appears {count_b} times."
    assert count_c == 1, f"R1 Regression: 'CellC' appears {count_c} times."

    # Reload the raw redlined output to verify Markdown leakage (R4)
    # If bold leaks, the raw view might have `**CellA Bold | CellB...**`
    # We test this by looking at how the virtual format parses it back.
    engine_raw = RedlineEngine(stream)
    engine_raw.process_batch([edit])
    raw_text = extract_text_from_stream(engine_raw.save_to_stream(), clean_view=False)

    # 'CellB' should absolutely not be immediately preceded by a ` | **` because that means the bold leaked into CellB
    assert " | **CellB" not in raw_text, (
        f"R4 Regression: Bold markers leaked across table cells.\nRaw Text:\n{raw_text}"
    )
    assert "** | " not in raw_text, f"R4 Regression: Bold boundary crossed cell boundaries.\nRaw Text:\n{raw_text}"


def test_r2_table_comment_duplication():
    """
    R2: Comment duplication on no-op modify edits targeting table-cell content.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "X"
    table.cell(0, 1).text = "Y"
    table.cell(0, 2).text = "Z"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # A comment-only no-op on a table row
    edit = ModifyText(target_text="X | Y | Z", new_text="X | Y | Z", comment="Only one comment")

    engine = RedlineEngine(stream)
    engine.process_batch([edit])
    raw_text = extract_text_from_stream(engine.save_to_stream(), clean_view=False)

    # Assert exact comment count
    comment_count = raw_text.count("[Com:")
    assert comment_count == 1, f"R2 Regression: Found {comment_count} comments. Expected exactly 1.\nText:\n{raw_text}"

    # Assert no redlines were emitted (it should hit the pure COMMENT_ONLY path)
    assert "{--" not in raw_text, "R2 Regression: No-op table comment emitted unnecessary deletions."
    assert "{++" not in raw_text, "R2 Regression: No-op table comment emitted unnecessary insertions."


def test_r3_table_multi_paragraph_triplication():
    """
    R3: Multi-paragraph insertion adjacent to a Word table row triplicates the entire payload.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Val1"
    table.cell(0, 1).text = "Val2"
    table.cell(0, 2).text = "Val3"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(target_text="Val1 | Val2 | Val3", new_text="Val1 | Val2 | Val3\n\n_New Italic Note_")

    engine = RedlineEngine(stream)
    engine.process_batch([edit])
    engine.accept_all_revisions()
    clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)

    # Assert exact count of the inserted note
    note_count = clean_text.count("New Italic Note")
    val_count = clean_text.count("Val1")

    assert note_count == 1, f"R3 Regression: Paragraph was duplicated {note_count} times.\nText:\n{clean_text}"
    assert val_count == 1, f"R3 Regression: Table row content was duplicated {val_count} times."


def test_n1_valid_table_edit_not_rejected():
    """
    N1: Valid within-cell modifications containing '|' should not be incorrectly rejected
    as structural changes.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Valid edit: modifies content, but retains the same number of cell boundaries (|)
    edit = ModifyText(target_text="A | B", new_text="A (Updated) | B")
    engine = RedlineEngine(stream)
    stats = engine.process_batch([edit])

    assert stats["edits_applied"] == 1, (
        f"N1 Regression: Valid table edit was incorrectly rejected. Skipped details: {stats.get('skipped_details')}"
    )
    assert stats["edits_skipped"] == 0

    # Verify text actually changed
    engine.accept_all_revisions()
    clean_text = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
    assert "A (Updated)" in clean_text


def test_n2_consistent_skipped_details_for_tables():
    """
    N2: Structural table rejections must have consistent, non-empty error messages
    with the explanatory note on *every* failure, not just the first.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "X"
    table.cell(0, 1).text = "Y"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Two invalid edits (structural changes: removing and adding columns)
    edit1 = ModifyText(target_text="X | Y", new_text="XY")
    edit2 = ModifyText(target_text="X | Y", new_text="X | Y | Z")

    engine = RedlineEngine(stream)
    stats = engine.process_batch([edit1, edit2])

    assert stats["edits_applied"] == 0
    assert stats["edits_skipped"] == 2

    details = stats["skipped_details"]
    assert len(details) == 2

    for i, detail in enumerate(details):
        assert "(Note: Structural table changes" in detail, f"N2 Regression: Note missing from detail {i + 1}: {detail}"
        assert "Failed to apply edit targeting:" in detail
        assert "\\n\\n" not in detail, (
            "N2 Regression: Error message truncated to raw internal newlines instead of human-readable text"
        )
        assert "X | Y" in detail, f"N2 Regression: Original target text snippet missing: {detail}"
