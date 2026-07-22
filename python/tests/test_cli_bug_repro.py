import sys
from pathlib import Path

import docx

from adeu.ingest import _extract_text_from_doc
from adeu.outline import extract_outline
from adeu.pagination import paginate


def create_manual_break_doc(path: Path):
    doc = docx.Document()
    doc.add_heading("Pagination Test Document", level=1)
    for page_num in range(1, 6):
        doc.add_heading(f"Heading on Page {page_num}", level=2)
        doc.add_paragraph(f"This is paragraph content belonging strictly to page {page_num}. " * 5)
        if page_num < 5:
            doc.add_page_break()
    doc.save(str(path))


def test_manual_page_breaks_pagination(tmp_path):
    doc_path = tmp_path / "manual_breaks.docx"
    create_manual_break_doc(doc_path)

    doc = docx.Document(str(doc_path))
    projected_body = _extract_text_from_doc(doc, include_appendix=False)

    # Run paginate
    pag_res = paginate(projected_body)

    # The document must have 5 pages
    assert pag_res.total_pages == 5, f"Expected 5 pages, got {pag_res.total_pages}"

    # Assert each page has the correct heading
    for page_num in range(1, 6):
        page_content = pag_res.pages[page_num - 1].page_content
        assert f"Heading on Page {page_num}" in page_content
        assert f"This is paragraph content belonging strictly to page {page_num}." in page_content
        # Ensure other page content is NOT leaked to this page
        for other_num in range(1, 6):
            if other_num != page_num:
                assert f"Heading on Page {other_num}" not in page_content


def test_manual_page_breaks_outline(tmp_path):
    doc_path = tmp_path / "manual_breaks.docx"
    create_manual_break_doc(doc_path)

    doc = docx.Document(str(doc_path))
    projected_body = _extract_text_from_doc(doc, include_appendix=False)

    pag_res = paginate(projected_body)

    # Get outline
    # extract_outline expects (doc, projected_body, body_pages, body_page_offsets, paragraph_offsets)
    body_pages = [p.page_content for p in pag_res.pages]
    body_page_offsets = pag_res.body_page_offsets

    nodes = extract_outline(doc, projected_body, body_pages, body_page_offsets)

    # We should have headings on pages 1 to 5
    headings = [node for node in nodes if node.level == 2]
    assert len(headings) == 5
    for i, node in enumerate(headings):
        expected_page = i + 1
        assert node.text == f"Heading on Page {expected_page}"
        assert node.page == expected_page, f"Expected {node.text} to be on page {expected_page}, got page {node.page}"


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from unittest.mock import patch

    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


def test_cli_apply_large_document_major_deletions(tmp_path, capsys):
    # 1. Create the original document with 100 sections
    doc_path = tmp_path / "large_bug.docx"
    doc = docx.Document()
    for i in range(1, 101):
        doc.add_heading(f"Section {i}", level=2)
        doc.add_paragraph(f"This is the body paragraph for section {i}. " * 15)
        if i % 10 == 0:
            doc.add_paragraph("Some special marker for search in section " + str(i))
    doc.save(str(doc_path))

    # 2. Create the truncated modified text file
    txt_path = tmp_path / "large_truncated_bug.txt"
    content = f"> **File Path:** {doc_path.name}\n\n# Large Document Test\n\nOnly Section 1 is here."
    txt_path.write_text(content, encoding="utf-8")

    # 3. Execute apply command
    out_path = tmp_path / "large_applied_bug.docx"
    code, stdout, stderr = run_cli(
        ["apply", str(doc_path), str(txt_path), "-o", str(out_path), "--allow-major-deletions"], capsys
    )

    # The regression test asserts that the bug is fixed and it completes successfully
    assert code == 0, f"apply failed with code {code}\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}"
    assert out_path.exists(), "Output file was not generated."


def test_accept_all_json_response_enrichment(tmp_path, capsys):
    import json

    import docx

    # 1. Create a simple base document
    doc_path = tmp_path / "base.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a test document.")
    doc.save(str(doc_path))

    # 2. Define a modify edit with a comment
    changes_file = tmp_path / "changes.json"
    changes_file.write_text(
        json.dumps(
            [
                {
                    "type": "modify",
                    "target_text": "document",
                    "new_text": "dossier",
                    "comment": "Review note to be stripped",
                }
            ]
        ),
        encoding="utf-8",
    )

    redlined_path = tmp_path / "redlined.docx"

    # 3. Apply the edit to create tracked changes + comment
    code, stdout, stderr = run_cli(
        ["apply", str(doc_path), str(changes_file), "-o", str(redlined_path), "--json"],
        capsys,
    )
    assert code == 0, f"apply failed with code {code}\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}"
    assert redlined_path.exists()

    # 4. Run accept-all in JSON mode
    accepted_path = tmp_path / "accepted.docx"
    code, stdout, stderr = run_cli(
        ["accept-all", str(redlined_path), "-o", str(accepted_path), "--json"],
        capsys,
    )
    assert code == 0, f"accept-all failed with code {code}\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}"

    # 5. Parse JSON output and assert that keys are present and counts are correct
    result = json.loads(stdout.strip())
    assert result.get("status") == "ok"
    assert "accepted_insertions" in result, "accepted_insertions missing from JSON output"
    assert "accepted_deletions" in result, "accepted_deletions missing from JSON output"
    assert "removed_comments" in result, "removed_comments missing from JSON output"

    assert result["accepted_insertions"] == 1
    assert result["accepted_deletions"] == 1
    assert result["removed_comments"] == 1


def test_sanitize_baseline_printf_leak(tmp_path):
    """
    Asserts that sanitize_docx raises a SanitizeError with an error message
    that does not leak printf-style string formatting specifiers (i.e. percent signs are properly escaped as %%).
    This prevents Go, Python or any other printf-style logging/formatting host from breaking
    with "not enough arguments" or showing "!o(MISSING)" / "!d(MISSING)".
    """
    import docx
    import pytest

    from adeu.sanitize.core import SanitizeError, sanitize_docx

    # 1. Create two highly different documents to trigger similarity check failure
    working = tmp_path / "normal.docx"
    doc_norm = docx.Document()
    doc_norm.add_paragraph("Agreement between Alpha and Beta.")
    doc_norm.save(str(working))

    baseline = tmp_path / "unicode.docx"
    doc_uni = docx.Document()
    doc_uni.add_paragraph("Different text in Chinese 统一码.")
    doc_uni.save(str(baseline))

    out = tmp_path / "out.docx"

    # 2. Call sanitize_docx and verify the error message is safe from printf leaks
    with pytest.raises(SanitizeError) as exc_info:
        sanitize_docx(str(working), str(out), baseline_path=str(baseline))

    err_msg = str(exc_info.value)

    # The error message should contain the similarity warning/block reason
    assert "share only" in err_msg
    assert "differs" in err_msg

    # Verify that the percent signs are properly escaped as '%%' by applying Python '%' formatting on it.
    # If any plain '%' is followed by space and characters like 'o' or 'd', it will raise TypeError.
    # Once fixed (escaped to '%%'), formatting with % () will succeed and produce a clean message with single '%' signs.
    try:
        formatted = err_msg % ()
    except TypeError as e:
        pytest.fail(
            f"Printf-style leak detected in error message! Formatting the message failed with: {e}\n"
            f"Error message: {err_msg}"
        )

    # After safe formatting, it should have single percent signs and NO formatting leaks
    assert "share only 41% of" in formatted
    assert "differs" in formatted
    assert "%!" not in formatted  # Ensure no Go-style formatting errors


def test_multi_paragraph_newline_comment(tmp_path):
    import io

    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    from adeu.models import ModifyText
    from adeu.redline.engine import RedlineEngine

    # 1. Create a base document with the target text
    doc = Document()
    doc.add_paragraph("Hello world. 🚀🔥🌟 and some suffix.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 2. Define the edit with a comment and a newline (paragraph break) in new_text
    edit = ModifyText(
        type="modify",
        target_text="🚀🔥🌟",
        new_text="🚀🔥🌟\n\nAdded: This is an extra line.",
        comment="Diff: Text inserted",
    )

    # 3. Apply edits
    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    # 4. Save and load document
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)

    # 5. Assert the comment ranges and references exist in the main document XML
    doc_xml = doc.element.xml
    assert "w:commentRangeStart" in doc_xml
    assert "w:commentRangeEnd" in doc_xml
    assert "w:commentReference" in doc_xml

    # 6. Assert the comment actually exists in comments.xml
    comments_part = None
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break

    assert comments_part is not None, "Comments part was not created"
    comments_xml = comments_part.blob.decode("utf-8")
    assert "Diff: Text inserted" in comments_xml, "Comment text missing from comments.xml"


def test_process_document_batch_flat_schema():
    """
    Asserts that the JSON Schema for process_document_batch.changes is a single flat Change object schema,
    rather than a nested union (oneOf/anyOf) of separate schema definitions.
    This ensures compatibility with 100% of MCP host implementations.
    """
    import asyncio

    from adeu.server import mcp

    tools = asyncio.run(mcp.list_tools())
    process_tool = next(t for t in tools if t.name == "process_document_batch")

    assert process_tool.parameters is not None, "process_document_batch has no parameters"

    # Extract changes property schema
    properties = process_tool.parameters.get("properties", {})
    assert "changes" in properties, "changes parameter missing from process_document_batch tool"

    changes_schema = properties["changes"]
    assert changes_schema.get("type") == "array", "changes must be an array type"

    items_schema = changes_schema.get("items", {})

    # The schema must NOT use oneOf or anyOf for the items
    assert "oneOf" not in items_schema, (
        "changes.items schema uses oneOf, which breaks nested array parsing in some MCP hosts"
    )
    assert "anyOf" not in items_schema, (
        "changes.items schema uses anyOf, which breaks nested array parsing in some MCP hosts"
    )

    # Instead, it must be a single flat object
    assert items_schema.get("type") == "object", "changes.items must be a flat object type"

    # Check that individual properties from all change variants are defined as optional properties of this single model
    item_properties = items_schema.get("properties", {})

    # Must have the type discriminator
    assert "type" in item_properties, "type field missing in flat change schema"

    # Must have all the other specific properties as optional properties
    expected_fields = [
        "target_text",
        "new_text",
        "target_id",
        "text",
        "cells",
        "position",
        "regex",
        "comment",
        "match_mode",
    ]
    for field in expected_fields:
        assert field in item_properties, f"expected field {field} missing from the unified flat Change schema"

    # The other fields should NOT be listed as 'required' at the JSON schema level to ensure they are optional
    required_fields = items_schema.get("required", [])
    # Only type is required (acting as discriminator), or maybe even type is optional.
    # Let's make sure none of the variant-specific fields are required.
    for field in expected_fields:
        assert field not in required_fields, f"field {field} must be optional (not required) in the flat Change schema"
