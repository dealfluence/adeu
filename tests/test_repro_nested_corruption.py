import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine


def test_repro_nested_edit_corruption():
    """
    REPRO: Attempting to edit text that is ALREADY inside a w:ins tag
    should ideally work or fail gracefully, but currently causes corruption/truncation.
    """
    doc = Document()
    p = doc.add_paragraph("Start ")

    # Simulate an existing Tracked Change (Insertion)
    # <w:ins w:id="1" w:author="Other"><w:r><w:t>Existing Insert</w:t></w:r></w:ins>
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "Other")

    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Existing Insert"
    run.append(t)
    ins.append(run)
    p._element.append(ins)

    p_run_end = OxmlElement("w:r")
    t_end = OxmlElement("w:t")
    t_end.text = " End"
    p_run_end.append(t_end)
    p._element.append(p_run_end)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Verify Ingest sees it (currently sees "Existing Insert" as plain text)
    initial_text = extract_text_from_stream(stream)
    assert "Existing Insert" in initial_text

    # Action: Try to edit the "Existing Insert" text
    # This triggers the "edit inside edit" scenario
    edit = DocumentEdit(target_text="Existing Insert", new_text="Modified Insert")

    engine = RedlineEngine(stream, author="Me")
    engine.apply_edits([edit])

    res_stream = engine.save_to_stream()

    # Check 1: Does output text look right?
    final_text = extract_text_from_stream(res_stream)

    # Failure condition reported by user: Truncation or missing text
    # We expect "Start Modified Insert End"
    # If corruption happened, we might see "Start  End" or broken XML
    print(f"Final Text: {final_text}")

    assert "Modified Insert" in final_text, "Edit was lost or corrupted"

    # Check 2: XML Validity (Manual Inspection logic)
    # We don't want nested <w:ins><w:ins>...</w:ins></w:ins>
    doc_res = Document(res_stream)
    xml = doc_res.element.xml

    if xml.count("<w:ins") > 1:
        # Check nesting
        # A simple string check isn't perfect but nested tags usually look like:
        # <w:ins ...><w:ins ...>
        import re

        if re.search(r"<w:ins[^>]*>.*?<w:ins", xml, re.DOTALL):
            print("WARNING: Nested w:ins detected!")
            # This is technically what we want to fix, but for this repro
            # we just want to confirm if it breaks the doc content.
