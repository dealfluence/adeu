import io

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine
from adeu.redline.mapper import DocumentMapper


def test_full_roundtrip_workflow(simple_docx_stream):
    extracted_text = extract_text_from_stream(simple_docx_stream)
    assert "Contract Agreement" in extracted_text
    assert "Seller" in extracted_text

    edit = DocumentEdit(target_text="Seller", new_text="Vendor", comment="Standardizing terminology.")

    simple_docx_stream.seek(0)
    engine = RedlineEngine(simple_docx_stream)

    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml_content = doc.element.xml

    assert "w:del" in xml_content
    # Word boundary logic ensures full word is deleted
    assert "<w:delText>Seller</w:delText>" in xml_content
    assert "w:ins" in xml_content
    assert "<w:t>Vendor</w:t>" in xml_content


def test_split_run_behavior():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("The quick brown fox.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="brown", new_text="")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml_content = doc.element.xml

    assert "<w:delText>brown</w:delText>" in xml_content


def test_insertion_spacing_between_complex_runs():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("ARTICLE")
    r1.bold = True
    r2 = p.add_run("3 ")
    r2.bold = False
    r3 = p.add_run("FEES")
    r3.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit1 = DocumentEdit(target_text="ARTICLE", new_text="ARTICLE ")

    edit2 = DocumentEdit(target_text="3 ", new_text="3  ")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit1, edit2])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    assert ">3 </w:t>" in xml
    assert ">FEES</w:t>" in xml

    idx_art = xml.find(">ARTICLE</w:t>")
    idx_3 = xml.find(">3 </w:t>")
    idx_fees = xml.find(">FEES</w:t>")

    assert idx_art < idx_3, "ARTICLE before 3"
    assert idx_3 < idx_fees, "3 before FEES"

    segment_1 = xml[idx_art:idx_3]
    assert "w:ins" in segment_1, "Missing insertion between ARTICLE and 3"

    segment_2 = xml[idx_3:idx_fees]
    assert "w:ins" in segment_2, "Missing insertion between 3 and FEES"


def test_insertion_splits_coalesced_run():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("ARTICLE3")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="ARTICLE", new_text="ARTICLE ")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    idx_art = xml.find(">ARTICLE</w:t>")
    idx_3 = xml.find(">3</w:t>")
    idx_ins = xml.find('<w:t xml:space="preserve"> </w:t>')
    if idx_ins == -1:
        idx_ins = xml.find("<w:t> </w:t>")

    assert idx_art < idx_ins < idx_3, f"Order wrong! Art:{idx_art}, Ins:{idx_ins}, 3:{idx_3}"


def test_insertion_at_start_of_document():
    doc = Document()
    doc.add_paragraph("Contract")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    original_text = extract_text_from_stream(stream)
    modified_text = "Big " + original_text

    from adeu.diff import generate_edits_from_text

    edits = generate_edits_from_text(original_text, modified_text)

    assert len(edits) > 0, "Should generate an edit for start-of-doc insertion"
    assert "Big" in edits[0].new_text
    assert edits[0].target_text in original_text


def test_insertion_multiple_splits_same_run():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("ARTICLE3 FEES")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    e1 = DocumentEdit(target_text="ARTICLE", new_text="ARTICLE ")
    e2 = DocumentEdit(target_text="3", new_text="3 ")

    engine = RedlineEngine(stream)
    engine.apply_edits([e1, e2])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    assert xml.count("<w:ins") == 2


def test_complex_run_sequence_repro():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("ARTICLE3 FEES")
    p.add_run("AN")
    p.add_run("D")
    p.add_run("PAYMENT")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    e1 = DocumentEdit(target_text="ARTICLE3 FEES", new_text="ARTICLE3 FEES ")
    e2 = DocumentEdit(target_text="AND", new_text="AND ")

    engine = RedlineEngine(stream)
    engine.apply_edits([e1, e2])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    idx_fees = xml.find("ARTICLE3 FEES")
    idx_ins1 = xml.find('w:id="1"')
    idx_an = xml.find(">AND</w:t>")

    assert idx_fees != -1
    assert idx_ins1 != -1
    assert idx_an != -1

    assert idx_fees < idx_ins1 < idx_an


def test_overlapping_run_boundaries():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("HELLO")
    p.add_run("WORLD")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    mapper = DocumentMapper(Document(stream))
    runs = mapper.find_target_runs("HELLO")

    assert len(runs) == 1
    assert runs[0].text == "HELLO"


def test_split_run_ordering_repro():
    doc = Document()
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    p = doc.add_paragraph()
    p.add_run("e0")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    e1 = DocumentEdit(target_text="", new_text=" END")
    e1._match_start_index = 2
    e2 = DocumentEdit(target_text="e", new_text="")
    e2._match_start_index = 0

    engine = RedlineEngine(stream)
    engine.apply_edits([e2, e1])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    idx_0 = xml.find(">0</w:t>")
    idx_ins = xml.find("> END</w:t>")

    assert idx_0 < idx_ins, f"0 ({idx_0}) should be before END ({idx_ins})"


def test_manual_context_disambiguation():
    doc = Document()
    doc.add_paragraph("Section 1: Fee")
    doc.add_paragraph("Section 2: Fee")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="Section 2: Fee", new_text="Section 2: Price", comment="Disambiguated via context")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])

    assert applied == 1
    assert skipped == 0

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    assert "Section 1: Fee" in xml
    # Context trimming leaves "Section 2: " intact, only "Fee" -> "Price"
    assert "<w:delText>Fee</w:delText>" in xml
    assert "<w:t>Price</w:t>" in xml
