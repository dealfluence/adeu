"""
BUG_adeu_accept_all_table_row_loss: accept_all_revisions() deleted table rows.

Replacing the entire text of the only paragraph in a table cell stamped a
row-deletion mark (w:trPr/w:del) at apply time, because the "is this row now
fully deleted?" inference ran BEFORE the replacement text was inserted — at
that instant every run in the row sat inside a <w:del>. accept_all_revisions()
then faithfully honoured the mark and dropped the whole <w:tr>, taking the
inserted text with it. Silent: no exception, and accepted_deletions counts
revision marks, so the accept looked successful.

Second, related defect: paragraph removal during accept/reject had no floor,
so a cell could end up with zero <w:p> children — invalid per ECMA-376
(a <w:tc> requires at least one block-level element) and reported as a corrupt
document by Word.
"""

import io

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import DeleteTableRow, ModifyText
from adeu.redline.engine import RedlineEngine

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _two_row_table_stream() -> io.BytesIO:
    """2x1 table (ALPHA / BETA) plus a tail paragraph — the report's fixture."""
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).paragraphs[0].text = "ALPHA"
    table.cell(1, 0).paragraphs[0].text = "BETA"
    doc.add_paragraph("tail paragraph")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _rows_marked_deleted(table) -> int:
    """Rows carrying a tracked row-deletion mark (w:trPr/w:del)."""
    count = 0
    for row in table.rows:
        trPr = row._tr.find(qn("w:trPr"))
        if trPr is not None and trPr.find(qn("w:del")) is not None:
            count += 1
    return count


def _cell_paragraph_counts(table) -> list[int]:
    return [len(cell._tc.findall(qn("w:p"))) for row in table.rows for cell in row.cells]


def test_accept_all_preserves_table_row_on_full_cell_replacement():
    """Replacing a cell's entire text must not delete its row on accept."""
    engine = RedlineEngine(_two_row_table_stream(), author="T")
    stats = engine.process_batch([ModifyText(target_text="ALPHA", new_text="GAMMA")])
    assert stats["edits_applied"] == 1

    tracked = engine.save_to_stream().getvalue()

    # The tracked document must not claim the row is deleted: a replacement
    # is not a row deletion, and Word would already render it struck through.
    tracked_table = Document(io.BytesIO(tracked)).tables[0]
    assert _rows_marked_deleted(tracked_table) == 0, (
        "Replacement stamped a spurious w:trPr/w:del on a row that keeps content"
    )

    accept_engine = RedlineEngine(io.BytesIO(tracked), author="T")
    accept_engine.accept_all_revisions()
    table = Document(accept_engine.save_to_stream()).tables[0]

    cells = [cell.text for row in table.rows for cell in row.cells]
    assert len(table.rows) == 2, f"Row lost on accept: table has {len(table.rows)} row(s), cells={cells}"
    assert cells == ["GAMMA", "BETA"], f"Expected ['GAMMA', 'BETA'], got {cells}"


def test_accept_all_preserves_row_on_multi_paragraph_replacement():
    """A new_text containing "\\n" must not collapse the row either."""
    engine = RedlineEngine(_two_row_table_stream(), author="T")
    stats = engine.process_batch([ModifyText(target_text="ALPHA", new_text="GAMMA\nDELTA")])
    assert stats["edits_applied"] == 1

    accept_engine = RedlineEngine(engine.save_to_stream(), author="T")
    accept_engine.accept_all_revisions()
    table = Document(accept_engine.save_to_stream()).tables[0]

    cells = [cell.text for row in table.rows for cell in row.cells]
    assert len(table.rows) == 2, f"Row lost on accept: table has {len(table.rows)} row(s), cells={cells}"
    assert cells[0] == "GAMMA\nDELTA", f"Expected the cell to hold both paragraphs, got {cells[0]!r}"
    assert cells[1] == "BETA"


def test_accept_all_preserves_last_paragraph_in_cell():
    """A <w:tc> must never end up with zero <w:p> children."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].text = "ALPHA"
    table.cell(0, 1).paragraphs[0].text = "KEEP"
    doc.add_paragraph("tail paragraph")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Deleting all of cell 0's text empties the cell but the sibling cell keeps
    # content, so the row survives — the cell must keep an empty paragraph.
    engine = RedlineEngine(stream, author="T")
    stats = engine.process_batch([ModifyText(target_text="ALPHA", new_text="")])
    assert stats["edits_applied"] == 1

    accept_engine = RedlineEngine(engine.save_to_stream(), author="T")
    accept_engine.accept_all_revisions()
    table = Document(accept_engine.save_to_stream()).tables[0]

    counts = _cell_paragraph_counts(table)
    assert all(n >= 1 for n in counts), f"Cell left with zero <w:p> (invalid OOXML): per-cell w:p counts={counts}"
    assert table.cell(0, 0).text == ""
    assert table.cell(0, 1).text == "KEEP"


def test_accept_all_preserves_signature_block_rows():
    """
    The real-world failure: replacing the text of signature-block cells in a
    4-row stipulation table dropped both counsel rows from the executed order.
    """
    doc = Document()
    table = doc.add_table(rows=4, cols=1)
    table.cell(0, 0).paragraphs[0].text = "Attorneys for Plaintiff"
    table.cell(1, 0).paragraphs[0].text = "spacer one"
    table.cell(2, 0).paragraphs[0].text = "Attorneys for Defendant"
    table.cell(3, 0).paragraphs[0].text = "spacer two"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="T")
    stats = engine.process_batch(
        [
            ModifyText(
                target_text="Attorneys for Plaintiff",
                new_text=(
                    "Cedar & Pike LLP\nSarah Chen, WSBA No. 51234\nAttorneys for Plaintiff Asteria Systems, Inc."
                ),
            ),
            ModifyText(
                target_text="Attorneys for Defendant",
                new_text="Harbor Line PLLC\nAttorneys for Defendant Northstar Analytics LLC",
            ),
        ]
    )
    assert stats["edits_applied"] == 2

    accept_engine = RedlineEngine(engine.save_to_stream(), author="T")
    accept_engine.accept_all_revisions()
    table = Document(accept_engine.save_to_stream()).tables[0]

    assert len(table.rows) == 4, f"Signature rows lost: {len(table.rows)} of 4 rows survived"
    assert "Sarah Chen, WSBA No. 51234" in table.cell(0, 0).text
    assert table.cell(1, 0).text == "spacer one"
    assert "Harbor Line PLLC" in table.cell(2, 0).text
    assert table.cell(3, 0).text == "spacer two"


def test_accept_all_preserves_only_body_paragraph_on_full_replacement():
    """The body counterpart: the paragraph container must survive a replacement."""
    doc = Document()
    doc.add_paragraph("SOLO")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="T")
    assert engine.process_batch([ModifyText(target_text="SOLO", new_text="REPLACED")])["edits_applied"] == 1

    accept_engine = RedlineEngine(engine.save_to_stream(), author="T")
    accept_engine.accept_all_revisions()
    result = Document(accept_engine.save_to_stream())

    texts = [p.text for p in result.paragraphs if p.text]
    assert texts == ["REPLACED"], f"Expected ['REPLACED'], got {texts}"


def test_accept_all_preserves_row_when_cell_paragraph_sits_in_a_content_control():
    """
    The court-caption shape: the cell's only paragraph lives inside a
    block-level <w:sdt>, so the paragraph is a grandchild of the <w:tc>.
    """
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.cell(1, 0).paragraphs[0].text = "BETA"
    doc.add_paragraph("tail paragraph")

    tc = table.cell(0, 0)._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    tc.append(
        parse_xml(
            f'<w:sdt xmlns:w="{W_NS}">'
            '<w:sdtPr><w:alias w:val="Caption"/></w:sdtPr>'
            "<w:sdtContent><w:p><w:r><w:t>ALPHA</w:t></w:r></w:p></w:sdtContent>"
            "</w:sdt>"
        )
    )

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="T")
    assert engine.process_batch([ModifyText(target_text="ALPHA", new_text="GAMMA")])["edits_applied"] == 1

    accept_engine = RedlineEngine(engine.save_to_stream(), author="T")
    accept_engine.accept_all_revisions()
    accepted = accept_engine.save_to_stream().getvalue()

    table = Document(io.BytesIO(accepted)).tables[0]
    assert len(table.rows) == 2, f"Content-control row lost on accept: {len(table.rows)} row(s) survived"
    assert len(table.cell(0, 0)._tc.findall(qn("w:sdt"))) == 1, "The <w:sdt> content control was dropped"

    # python-docx's cell.text only reads direct <w:p> children, so assert on
    # the engine's own projection (AI_CONTEXT §12, "Testing Redlines").
    clean_text = extract_text_from_stream(io.BytesIO(accepted), clean_view=True)
    assert "GAMMA" in clean_text, f"Replacement text lost with the row. Text:\n{clean_text}"
    assert "BETA" in clean_text
    assert "ALPHA" not in clean_text


def test_explicit_delete_table_row_still_removes_the_row():
    """
    Guard against over-correcting: an explicit DeleteTableRow must still stamp
    w:trPr/w:del and drop the row on accept.
    """
    doc = Document()
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).paragraphs[0].text = "ALPHA"
    table.cell(1, 0).paragraphs[0].text = "BETA"
    table.cell(2, 0).paragraphs[0].text = "GAMMA"
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="T")
    assert engine.process_batch([DeleteTableRow(target_text="BETA")])["edits_applied"] == 1

    tracked = engine.save_to_stream().getvalue()
    assert _rows_marked_deleted(Document(io.BytesIO(tracked)).tables[0]) == 1

    accept_engine = RedlineEngine(io.BytesIO(tracked), author="T")
    accept_engine.accept_all_revisions()
    clean_text = extract_text_from_stream(accept_engine.save_to_stream(), clean_view=True)

    assert "ALPHA" in clean_text
    assert "BETA" not in clean_text
    assert "GAMMA" in clean_text
