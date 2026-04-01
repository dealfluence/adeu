# FILE: tests/test_repro_formatting_bugs.py
import io

from adeu.diff import trim_common_context
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from docx import Document
from docx.shared import Pt


def test_repro_token_slicing_mid_sentence():
    """
    Guards against Errors 2 & 3.
    Ensures that trim_common_context does NOT slice markdown tokens in half
    when the target change is inside a formatted block.
    """
    target = "Standard payment terms are **Net 90 Days** from the date"
    new_val = "Standard payment terms are **Net 60 Days** from the date"

    p, s = trim_common_context(target, new_val)

    rem_target = target[p : len(target) - s]
    rem_new = new_val[p : len(new_val) - s]

    # Prior to the fix, the trimmer would slice the first '**' into '*'
    # causing unbalanced asterisks to leak into the extracted diff.
    assert "**" not in rem_target, f"Leaked markdown marker in target: {rem_target}"
    assert "**" not in rem_new, f"Leaked markdown marker in new_val: {rem_new}"

    # The absorption loop should cleanly peel off the markers
    assert rem_target == "Net 90 Days"
    assert rem_new == "Net 60 Days"


def test_repro_token_slicing_boundary_append():
    """
    Guards against Error 4.
    Ensures that appending text immediately after a formatting block
    doesn't slice the preceding markers.
    """
    target = "**$2,000,000 per occurrence**."
    new_val = "**$2,000,000 per occurrence** and **$4,000,000 in the aggregate**."

    p, s = trim_common_context(target, new_val)

    rem_target = target[p : len(target) - s] if s else target[p:]
    rem_new = new_val[p : len(new_val) - s] if s else new_val[p:]

    # Because the period after "occurrence**" is removed in the new_val,
    # the trimmer correctly refuses to trim the prefix, backing out entirely
    # to prevent slicing the markdown tokens. This ensures a safe, clean replacement.
    assert rem_target == "**$2,000,000 per occurrence**."
    assert rem_new == "**$2,000,000 per occurrence** and **$4,000,000 in the aggregate**."


def test_repro_schema_ordering_lost_italic():
    """
    Guards against Error 6.
    Ensures that when italic/bold is applied, the <w:i> tag is inserted
    in the correct strict OpenXML sequence (e.g., before <w:sz>), otherwise Word ignores it.
    """
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Test string")

    # Force a size element to exist in the run properties
    r.font.size = Pt(12)

    # Save to a valid DOCX stream so the engine can initialize
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Apply italic via the engine
    engine = RedlineEngine(stream)
    engine._apply_run_props(r._element, {"italic": True})

    rpr_xml = r._element.rPr.xml
    i_idx = rpr_xml.find("w:i")
    sz_idx = rpr_xml.find("w:sz")

    assert i_idx != -1, "Italic tag <w:i> was not added."
    assert sz_idx != -1, "Size tag <w:sz> was lost."

    # Crucial assertion: In strict OpenXML, <w:i> MUST precede <w:sz>.
    assert i_idx < sz_idx, f"Invalid XML Schema order! <w:i> appeared after <w:sz>:\n{rpr_xml}"


def test_repro_lost_bold_suppressed_formatting():
    """
    Guards against Error 5.
    Ensures that an edit replacing "**4.0%**" with "**3.0%**" correctly
    inherits bold formatting, instead of stripping it because the trimmed string
    "3.0%" lacked markdown markers.
    """
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("4.0%")
    r.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Setup edit: target text will map to "4.0%" run, new text has bold markers
    edit = ModifyText(target_text="**4.0%**", new_text="**3.0%**")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    res_stream = engine.save_to_stream()
    doc_res = Document(res_stream)

    # Find the inserted run
    ins_runs = doc_res.element.xpath('//w:ins//w:r[w:t[text()="3.0%"]]')
    assert ins_runs, "Insertion '3.0%' not found"

    # Verify the inserted run retained bold property
    b_tags = ins_runs[0].xpath("./w:rPr/w:b")
    is_bold = False
    if b_tags:
        val = b_tags[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
        is_bold = val is None or val.lower() not in ("0", "false", "off")

    assert is_bold, "The inserted run lost its bold formatting (suppress_inherited fired incorrectly)."


def test_repro_insertion_trailing_space_omission():
    """
    Guards against the bug where an LLM appends text to a sentence but omits
    the trailing space that exists in the original document.
    Prior to the fix, this caused the engine to delete and rewrite the entire
    sentence instead of just performing a clean insertion of the new text.
    """
    doc = Document()
    # Notice the trailing space in the document text
    doc.add_paragraph("Retailer shall be named as an additional insured. ")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Target text matched in document has the space, but new_text drops the space
    # in favor of a newline (\n) before the new paragraph.
    edit = ModifyText(
        target_text="Retailer shall be named as an additional insured. ",
        new_text="Retailer shall be named as an additional insured.\n**7.3 Limitation of Liability.**",
    )

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    res_stream = engine.save_to_stream()
    doc_res = Document(res_stream)

    # 1. Verify that the engine recognized this as a pure insertion and
    # did NOT delete the original sentence.
    del_tags = doc_res.element.xpath("//w:del")
    assert (
        len(del_tags) == 0
    ), "Original sentence was needlessly deleted! Heuristic failed to classify as pure insertion."

    # 2. Verify that the new text was correctly tracked as an insertion
    ins_texts = doc_res.element.xpath('//w:ins//w:t[contains(text(), "Limitation of Liability")]')
    assert len(ins_texts) > 0, "New paragraph text was not inserted."
