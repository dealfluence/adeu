import io
import re
import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from adeu.redline.engine import RedlineEngine
from adeu.models import DocumentEdit

def test_native_comment_creation_and_linking():
    doc = Document()
    doc.add_paragraph("The quick brown fox.")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(
        target_text="quick",
        new_text="slow",
        comment="Foxes are not always quick."
    )

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    
    doc_xml = doc.element.xml
    assert "w:commentRangeStart" in doc_xml
    assert "w:commentRangeEnd" in doc_xml
    assert "w:commentReference" in doc_xml

    id_match = re.search(r'<w:commentRangeStart[^>]*w:id="(\d+)"', doc_xml)
    assert id_match
    comment_id = id_match.group(1)

    comments_part = None
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break
    
    assert comments_part is not None
    comments_xml = comments_part.blob.decode("utf-8")
    assert "Foxes are not always quick." in comments_xml
    assert f'w:id="{comment_id}"' in comments_xml

def test_multiple_comments_ids():
    doc = Document()
    doc.add_paragraph("First sentence.")
    doc.add_paragraph("Second sentence.")
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit1 = DocumentEdit(
        target_text="First",
        new_text="First The ",
        comment="Comment One"
    )
    edit2 = DocumentEdit(
        target_text="Second",
        new_text="Second The ",
        comment="Comment Two"
    )

    engine = RedlineEngine(stream)
    engine.apply_edits([edit1, edit2])
    
    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    
    comments_part = None
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break
            
    comments_xml = comments_part.blob.decode("utf-8")
    
    assert "Comment One" in comments_xml
    assert "Comment Two" in comments_xml
    
    ids = re.findall(r'<w:comment\s+[^>]*w:id="(\d+)"', comments_xml)
    assert len(set(ids)) == 2