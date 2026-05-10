import io

from docx import Document
from docx.oxml import OxmlElement

from adeu.ingest import extract_text_from_stream
from adeu.utils.docx import normalize_docx


def test_markdown_wrapping_newline_bug():
    """
    Case 1: Run contains text + br + text.
    Extract should be **Text**\n**Text**, not **Text\nText**.
    """
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run()
    r.bold = True

    # Manually inject text + br + text
    t1 = OxmlElement("w:t")
    t1.text = "Line 1"
    r._element.append(t1)

    br = OxmlElement("w:br")
    r._element.append(br)

    t2 = OxmlElement("w:t")
    t2.text = "Line 2"
    r._element.append(t2)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    # Check for correct output
    assert "**Line 1**" in text
    assert "**Line 2**" in text
    assert "**Line 1\nLine 2**" not in text


def test_trailing_newline_wrapping():
    """
    Case 2: Run contains text + br.
    Extract should be **Text**\n, not **Text\n**.
    """
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run()
    r.bold = True

    t1 = OxmlElement("w:t")
    t1.text = "Line 1"
    r._element.append(t1)

    br = OxmlElement("w:br")
    r._element.append(br)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    text = extract_text_from_stream(stream)

    assert "**Line 1**" in text
    assert "**Line 1\n**" not in text


def test_normalize_preserves_breaks():
    """
    Regression check: normalize_docx should not destroy w:br when coalescing.
    If it does, fixing the markdown wrapping is moot because the newline disappears.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run 1: Text
    r1 = p.add_run("A")
    r1.bold = True

    # Run 2: Break (in separate run, same formatting)
    r2 = p.add_run()
    r2.bold = True
    r2._element.append(OxmlElement("w:br"))

    # Run 3: Text
    r3 = p.add_run("B")
    r3.bold = True

    normalize_docx(doc)

    # Check if BR is preserved
    xml = p._element.xml
    assert "w:br" in xml, "Normalization destroyed w:br tag!"
