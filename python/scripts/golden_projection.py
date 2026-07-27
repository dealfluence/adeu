"""Golden capture/compare for the Python projection pipeline.

Protocol from docs/PERFORMANCE.md §2/§3.6: before touching projection code,
capture every view this engine can emit; after the change, byte-compare. The
projection is a contract with downstream agents (anchors, ids, offsets), so
performance work is only safe when the bytes are provably unchanged.

Views captured per document:
  reader_raw       _extract_text_from_doc(clean_view=False, include_appendix=False)
  reader_clean     _extract_text_from_doc(clean_view=True,  include_appendix=False)
  reader_appendix  _extract_text_from_doc(clean_view=False, include_appendix=True)
  mapper_raw       DocumentMapper(clean_view=False).full_text
  mapper_clean     DocumentMapper(clean_view=True).full_text
  pagination       page count + per-page lengths + page boundary offsets
  outline          one line per OutlineNode, all fields

The twin contract (reader_raw == mapper_raw, reader_clean == mapper_clean)
is asserted on every capture, so a run that breaks it fails loudly even
before the byte-compare.

Usage:
  python scripts/golden_projection.py verify [manifest]   # vs COMMITTED hashes
  python scripts/golden_projection.py capture <outdir>
  python scripts/golden_projection.py compare <baseline_dir> <new_dir>

`verify` is the durable gate: it compares against the committed
tests/golden_manifest.txt (hashes only — no multi-MB golden text in git) and
is what tests/test_projection_goldens.py runs. Use `capture` + `compare` when
you need to SEE a diff, since those keep the full text side by side.

Regenerating the manifest is an explicit, reviewable act — do it only when a
projection change is intended, and say so in the commit:
  python scripts/golden_projection.py capture tmp/g && \
      cp tmp/g/MANIFEST.txt tests/golden_manifest.txt
"""

from __future__ import annotations

import hashlib
import io
import sys
import time
from pathlib import Path

from docx import Document

from adeu.ingest import _extract_text_from_doc
from adeu.outline import extract_outline
from adeu.pagination import paginate, split_structural_appendix
from adeu.redline.mapper import DocumentMapper
from adeu.utils.docx import strip_bom_from_docx_bytes

# (name, path) — VVBIG/BIGDOC live on the user's Desktop, not in the repo.
DOCS = [
    ("cells", None),  # synthetic, built below
    ("BIGDOC", Path(r"C:\Users\Uzair\Desktop\BIGDOC.docx")),
    ("VVBIG", Path(r"C:\Users\Uzair\Desktop\VVBIG.docx")),
]


def build_cells_fixture() -> bytes:
    """Anchor/table fixture exercising every cell-anchor branch (§3.6):
    a labeled cell, a text cell without id, empty cells with unlabeled
    paragraphs, an empty cell with NO paragraph, a nested table with empty
    cells, and a second table."""
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph("Heading for cells fixture", style="Heading 1")
    doc.add_paragraph("Body text before the table.")

    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "Labeled"
    t.cell(0, 1).text = "Text cell without id"
    # (1,0) empty with an unlabeled paragraph (default from add_table)
    # (1,1): strip its only paragraph -> empty cell with NO paragraph
    tc = t.cell(1, 1)._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    # nested table inside (2,0) with empty cells
    inner = t.cell(2, 0).add_table(rows=2, cols=2)
    inner.cell(0, 0).text = "inner"

    doc.add_paragraph("Between tables.")
    t2 = doc.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "second table"

    doc.add_paragraph("Trailing paragraph.")
    p = doc.add_paragraph()
    r = p.add_run("bold text")
    r.bold = True
    r2 = p.add_run(" italic")
    r2.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def load_bytes(name, path):
    if path is None:
        return build_cells_fixture()
    return path.read_bytes()


def render_outline(doc):
    """Mirrors the production path exactly (doc_cache._fill_view): the reader
    is re-run with return_paragraph_offsets=True because the outline consumes
    those offsets, so this golden covers that branch too."""
    text, offsets = _extract_text_from_doc(doc, clean_view=False, include_appendix=False, return_paragraph_offsets=True)
    body, _ = split_structural_appendix(text)
    pg = paginate(body, structural_appendix="")
    nodes = extract_outline(doc, body, pg.body_pages, pg.body_page_offsets, paragraph_offsets=offsets)
    lines = []
    for n in nodes:
        lines.append(
            f"L{n.level}\tp{n.page}\tend={n.end_page}\tstyle={n.style}\t"
            f"table={n.has_table}\tfn={','.join(n.footnote_ids)}\t{n.text}"
        )
    return "\n".join(lines), pg, text


def render_pagination(pg):
    lines = [
        f"total_pages={pg.total_pages}",
        f"body_pages={len(pg.body_pages)}",
        f"body_page_offsets={','.join(str(o) for o in pg.body_page_offsets)}",
    ]
    for i, p in enumerate(pg.pages, 1):
        h = hashlib.sha256(p.page_content.encode("utf-8")).hexdigest()[:16]
        lines.append(f"page {i}\tlen={len(p.page_content)}\ttracked={p.tracked_change_count}\tsha={h}")
    return "\n".join(lines)


def compute_views(name: str, sanitized: bytes, verbose: bool = False) -> dict[str, str]:
    """Every projection view for one document, plus the invariant assertions.

    Importable, so the pytest gate (tests/test_projection_goldens.py) runs the
    exact same computation this CLI does — one implementation, no drift.
    """

    def _t(label, fn):
        t = time.perf_counter()
        out = fn()
        if verbose:
            print(f"  {label:16s} {time.perf_counter() - t:6.2f}s")
        return out

    views: dict[str, str] = {}
    views["reader_raw"] = _t(
        "reader_raw",
        lambda: _extract_text_from_doc(Document(io.BytesIO(sanitized)), clean_view=False, include_appendix=False),
    )
    views["reader_clean"] = _t(
        "reader_clean",
        lambda: _extract_text_from_doc(Document(io.BytesIO(sanitized)), clean_view=True, include_appendix=False),
    )
    views["reader_appendix"] = _t(
        "reader_appendix",
        lambda: _extract_text_from_doc(Document(io.BytesIO(sanitized)), clean_view=False, include_appendix=True),
    )
    views["mapper_raw"] = _t(
        "mapper_raw",
        lambda: DocumentMapper(Document(io.BytesIO(sanitized)), clean_view=False).full_text,
    )
    views["mapper_clean"] = _t(
        "mapper_clean",
        lambda: DocumentMapper(Document(io.BytesIO(sanitized)), clean_view=True).full_text,
    )

    # Twin contract (§7.3.3) — asserted on EVERY computation, so drift fails
    # loudly even when hashes are not being compared.
    assert views["reader_raw"] == views["mapper_raw"], f"{name}: TWIN DRIFT (raw view)"
    assert views["reader_clean"] == views["mapper_clean"], f"{name}: TWIN DRIFT (clean view)"

    outline_txt, pg, offsets_text = _t("outline+paginate", lambda: render_outline(Document(io.BytesIO(sanitized))))
    views["outline"] = outline_txt
    views["pagination"] = render_pagination(pg)

    # Requesting paragraph offsets must not change the projected text.
    assert offsets_text == views["reader_raw"], f"{name}: reader text differs when paragraph offsets are requested"
    return views


def iter_documents(verbose: bool = False):
    """Yields (name, sanitized_bytes) for every document available here."""
    for name, path in DOCS:
        if path is not None and not path.exists():
            if verbose:
                print(f"SKIP {name}: {path} not found")
            continue
        yield name, strip_bom_from_docx_bytes(load_bytes(name, path))


def capture(outdir: Path):
    outdir.mkdir(parents=True, exist_ok=True)
    manifest = []
    for name, sanitized in iter_documents(verbose=True):
        print(f"\n=== {name} ({len(sanitized) / 1e6:.2f} MB) ===")
        views = compute_views(name, sanitized, verbose=True)
        print("  twin contract    OK (raw + clean byte-identical)")
        for view, text in views.items():
            (outdir / f"{name}.{view}.txt").write_text(text, encoding="utf-8", newline="")
            sha = hashlib.sha256(text.encode("utf-8")).hexdigest()
            manifest.append(f"{sha}  {len(text):>10}  {name}.{view}.txt")
            print(f"    {view:16s} {len(text):>10} chars  {sha[:16]}")

    (outdir / "MANIFEST.txt").write_text("\n".join(manifest) + "\n", encoding="utf-8", newline="")
    print(f"\ncaptured {len(manifest)} views -> {outdir}")


def load_manifest(manifest_path: Path) -> dict[str, str]:
    expected = {}
    for ln in manifest_path.read_text(encoding="utf-8").splitlines():
        if not ln.strip() or ln.lstrip().startswith("#"):
            continue
        sha, _size, fname = ln.split()
        expected[fname] = sha
    return expected


def verify(manifest_path: Path) -> int:
    """Recompute every view and compare hashes to a COMMITTED manifest.

    Unlike `compare`, this needs no stored golden text — only hashes — so the
    baseline is committable (tests/golden_manifest.txt) and the
    byte-identical claim stays checkable long after capture dirs are deleted.
    Documents absent from this machine are reported, not failed.
    """
    expected = load_manifest(manifest_path)
    checked = 0
    failures = []
    seen = set()

    for name, sanitized in iter_documents(verbose=True):
        seen.add(name)
        views = compute_views(name, sanitized)
        for view, text in views.items():
            key = f"{name}.{view}.txt"
            if key not in expected:
                failures.append(f"{key}: not present in manifest (new view?)")
                continue
            got = hashlib.sha256(text.encode("utf-8")).hexdigest()
            checked += 1
            if got != expected[key]:
                failures.append(f"{key}: expected {expected[key][:16]} got {got[:16]} ({len(text)} chars)")

    for name in sorted({k.split(".")[0] for k in expected} - seen):
        print(f"NOTE: {name} unavailable here; its manifest rows were not checked")
    for f in failures:
        print(f"FAIL {f}")
    print(
        f"\n{'GOLDENS VERIFIED' if not failures else 'GOLDEN MISMATCH'} "
        f"({checked} views checked, {len(failures)} failures)"
    )
    return 0 if not failures else 1


def compare(base: Path, new: Path):
    b = (base / "MANIFEST.txt").read_text(encoding="utf-8").splitlines()
    n = (new / "MANIFEST.txt").read_text(encoding="utf-8").splitlines()
    bm = {ln.split("  ")[-1]: ln.split("  ")[0] for ln in b if ln.strip()}
    nm = {ln.split("  ")[-1]: ln.split("  ")[0] for ln in n if ln.strip()}

    only_b = sorted(set(bm) - set(nm))
    only_n = sorted(set(nm) - set(bm))
    diff = sorted(k for k in set(bm) & set(nm) if bm[k] != nm[k])

    for k in only_b:
        print(f"MISSING in new: {k}")
    for k in only_n:
        print(f"EXTRA in new:   {k}")
    for k in diff:
        print(f"DIFFERS: {k}")
        bt = (base / k).read_text(encoding="utf-8")
        nt = (new / k).read_text(encoding="utf-8")
        print(f"   baseline {len(bt)} chars, new {len(nt)} chars")
        # first divergence
        lim = min(len(bt), len(nt))
        i = next((i for i in range(lim) if bt[i] != nt[i]), lim)
        print(f"   first divergence at char {i}:")
        print(f"     baseline: {bt[max(0, i - 60) : i + 60]!r}")
        print(f"     new:      {nt[max(0, i - 60) : i + 60]!r}")

    ok = not (only_b or only_n or diff)
    print(f"\n{'ALL VIEWS BYTE-IDENTICAL' if ok else 'GOLDEN MISMATCH'} ({len(set(bm) & set(nm))} compared)")
    return 0 if ok else 1


DEFAULT_MANIFEST = Path(__file__).resolve().parent.parent / "tests" / "golden_manifest.txt"


if __name__ == "__main__":
    if sys.argv[1] == "verify":
        path = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_MANIFEST
        sys.exit(verify(path))
    if sys.argv[1] == "capture":
        capture(Path(sys.argv[2]))
    else:
        sys.exit(compare(Path(sys.argv[2]), Path(sys.argv[3])))
