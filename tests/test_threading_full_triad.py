import io

from docx import Document

from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine


def test_full_modern_comments_triad_creation():
    doc = Document()
    doc.add_paragraph("Content")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    engine.apply_edits([DocumentEdit(target_text="Content", new_text="Content Changed", comment="Modern")])
    stream_out = engine.save_to_stream()

    doc_out = Document(stream_out)
    rels = [rel.reltype for rel in doc_out.part.rels.values()]

    rel_extended = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
    rel_ids = "http://schemas.microsoft.com/office/2016/relationships/commentsIds"
    rel_extensible = "http://schemas.microsoft.com/office/2018/relationships/commentsExtensible"

    assert any(r == rel_extended for r in rels), "Missing commentsExtended"
    assert any(r == rel_ids for r in rels), "Missing commentsIds"
    assert any(r == rel_extensible for r in rels), "Missing commentsExtensible"

    # Check Extensible Content
    extensible_part = None
    for rel in doc_out.part.rels.values():
        if rel.reltype == rel_extensible:
            extensible_part = rel.target_part
            break

    xml = extensible_part.blob.decode("utf-8")
    print(xml)
    assert "w16cex:durableId" in xml
    assert "w16cex:dateUtc" in xml
