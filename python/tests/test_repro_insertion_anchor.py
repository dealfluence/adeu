"""
Reproduction for the insertion-anchor bug (2026-07, golden.docx):

    ModifyText("document" -> "finalized document") on golden.docx is
    diff-minimized to a zero-width INSERTION of "finalized " immediately
    before "document". The target paragraph already carries tracked changes
    ({--initial --}{++golden ++}) and three comment ranges, so in the raw
    projection the insertion index sits right after a virtual {>>...<<} meta
    block. DocumentMapper.get_insertion_anchor only looks for a run-backed
    span ENDING exactly at that index; every span ending there is virtual, so
    it falls back to (None, paragraph) and the engine drops the new <w:ins>
    at the START of the paragraph. After accept-all the text reads
    "finalized This is the golden document" — silent corruption.

    Desired behavior: a zero-width insertion whose index is fenced off by
    virtual markup must anchor after the nearest preceding REAL run in the
    same paragraph ("golden "), yielding "This is the golden finalized
    document".

STYLE: these tests assert the DESIRED behaviour, so they are RED while the
bug is present and turn GREEN once the engine is fixed. The controls pin the
boundary of the fix: an insertion anchored on a tracked-DELETED run must not
nest <w:ins> inside <w:del>, and true paragraph-start insertions must keep
landing at the paragraph start.
"""

import io
from pathlib import Path

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def get_fixture_path(name: str) -> Path:
    p = Path(__file__).resolve()
    for parent in p.parents:
        if (parent / "shared" / "fixtures").is_dir():
            return parent / "shared" / "fixtures" / name
    raise FileNotFoundError(f"Could not find fixtures directory for {name}")


def _apply(source: bytes, edits, author: str = "T") -> bytes:
    engine = RedlineEngine(io.BytesIO(source), author=author)
    engine.process_batch(edits)
    return engine.save_to_stream().getvalue()


def test_insertion_lands_before_target_despite_preceding_redlines_and_comments():
    """golden.docx: zero-width insertion must not teleport to paragraph start."""
    data = get_fixture_path("golden.docx").read_bytes()

    out = _apply(data, [ModifyText(target_text="document", new_text="finalized document")])

    clean = extract_text_from_stream(io.BytesIO(out), clean_view=True)
    assert "golden finalized document" in clean, f"insertion misplaced; clean view: {clean!r}"
    assert not clean.startswith("finalized This is"), (
        f"insertion anchored at paragraph start instead of before 'document': {clean!r}"
    )

    # The redlined projection must show the insertion adjacent to its target,
    # not as the first thing in the paragraph.
    red = extract_text_from_stream(io.BytesIO(out))
    assert not red.startswith("{++finalized"), f"redline starts with the insertion: {red[:80]!r}"


def test_insertion_after_tracked_deletion_does_not_nest_ins_inside_del():
    """
    Control for the anchor fix: when the nearest preceding real run is inside
    a <w:del> (deletion followed by a {>>...<<} meta block, then the target),
    the new <w:ins> must land AFTER the deletion block as a sibling — never
    nested inside <w:del>, which is invalid tracked-change XML.
    """
    doc = Document()
    doc.add_paragraph("This is the initial document")
    stream = io.BytesIO()
    doc.save(stream)

    # Pass 1: track-delete "initial " so the raw projection becomes
    # "This is the {--initial --}{>>[Chg:1 delete] T<<}document".
    deleted = _apply(stream.getvalue(), [ModifyText(target_text="initial ", new_text="")])

    # Pass 2: zero-width insertion right before "document".
    engine = RedlineEngine(io.BytesIO(deleted), author="T2")
    engine.process_batch([ModifyText(target_text="document", new_text="finalized document")])
    out = engine.save_to_stream().getvalue()

    nested = engine.doc.element.xpath("//w:del//w:ins")
    assert not nested, "insertion was nested inside a <w:del> block"

    clean = extract_text_from_stream(io.BytesIO(out), clean_view=True)
    assert "This is the finalized document" in clean, f"clean view: {clean!r}"


def test_paragraph_start_insertion_still_lands_at_paragraph_start():
    """
    Guard: an insertion at index 0 (genuine paragraph start, no preceding
    real text) must keep its existing behavior and stay at the front.
    """
    doc = Document()
    doc.add_paragraph("Alpha beta gamma")
    stream = io.BytesIO()
    doc.save(stream)

    out = _apply(stream.getvalue(), [ModifyText(target_text="Alpha", new_text="Intro Alpha")])

    clean = extract_text_from_stream(io.BytesIO(out), clean_view=True)
    assert "Intro Alpha beta gamma" in clean, f"clean view: {clean!r}"
