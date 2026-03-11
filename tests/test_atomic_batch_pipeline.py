# FILE: tests/test_atomic_batch_pipeline.py
import io
import re

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import DocumentEdit, ReviewAction
from adeu.redline.engine import RedlineEngine
from adeu.server import process_document_batch


def test_atomic_batch_prevents_cascading_misanchor(tmp_path):
    """
    Validates Issue #2 Fix:
    Ensures that processing ReviewActions (which mutate the XML DOM and shift text lengths)
    does not cause subsequent DocumentEdits in the SAME batch to misanchor.
    """
    # 1. Setup initial doc
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")

    orig_path = tmp_path / "original.docx"
    doc.save(orig_path)

    # 2. Make an initial tracked change (Simulating Round 1)
    with open(orig_path, "rb") as f:
        engine = RedlineEngine(io.BytesIO(f.read()), author="Round1")

    # Edit: "First" -> "1st" (Creates a w:del and w:ins)
    engine.apply_edits([DocumentEdit(target_text="First", new_text="1st")])

    mid_path = tmp_path / "mid.docx"
    with open(mid_path, "wb") as f:
        f.write(engine.save_to_stream().getvalue())

    # Verify intermediate state (Round 1)
    with open(mid_path, "rb") as f:
        mid_text = extract_text_from_stream(io.BytesIO(f.read()))

    assert "{--First--}" in mid_text
    assert "{++1st++}" in mid_text

    # Extract dynamically generated Change IDs for the Accept action
    chg_ids = set(re.findall(r"\[Chg:(\d+)\]", mid_text))
    assert len(chg_ids) > 0, "Tracked changes were not generated."

    # 3. Execute the Atomic Batch (Simulating Round 2)
    # We ACCEPT the previous changes. This removes the w:del and w:ins wrappers,
    # shrinking the XML and shifting the text indices of everything below it.
    actions = [ReviewAction(action="ACCEPT", target_id=f"Chg:{i}") for i in chg_ids]

    # We edit text further down the document.
    # If the mapper is not rebuilt, "Third" will look for the wrong index and fail.
    edits = [DocumentEdit(target_text="Third", new_text="3rd")]

    out_path = tmp_path / "final.docx"

    # Run the new server tool
    result_msg = process_document_batch(
        original_docx_path=str(mid_path),
        author_name="Round2",
        actions=actions,
        edits=edits,
        output_path=str(out_path),
    )

    # 4. Assertions on the Tool Execution
    assert "Batch complete" in result_msg
    assert f"Actions: {len(actions)} applied, 0 skipped" in result_msg
    assert "Edits: 1 applied, 0 skipped" in result_msg, "The edit misanchored and was skipped!"

    # 5. Assertions on the Final Document State
    with open(out_path, "rb") as f:
        final_text = extract_text_from_stream(io.BytesIO(f.read()))

    # The first paragraph should be cleanly accepted
    assert "1st paragraph." in final_text
    assert "{--First--}" not in final_text

    # The third paragraph should have the new tracked change anchored perfectly
    assert "{--Third--}" in final_text
    assert "{++3rd++}" in final_text
