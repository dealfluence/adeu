import io

import docx

from adeu import ModifyText, RedlineEngine


def test_batch_engine_swallows_trailing_deletions():
    """
    Reproduces Bug 9: When replacing text that includes a deletion at the end
    of the target, AND a multi-paragraph insertion, the trailing deletion is
    silently swallowed (ignored) and left in the document.
    """
    d = docx.Document()
    d.add_paragraph("Party A; and")
    stream = io.BytesIO()
    d.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA")
    # We remove `; and` and insert a new paragraph.
    engine.apply_edits([ModifyText(target_text="Party A; and", new_text="Party A\n\nParty B")])

    out_stream = engine.save_to_stream()
    out_stream.seek(0)
    d2 = docx.Document(out_stream)

    # The first paragraph should have '; and' deleted. Let's check the XML.
    xml = d2.paragraphs[0]._element.xml
    assert "<w:del" in xml, "BUG: The trailing deletion of '; and' was swallowed entirely!"
