import asyncio
import json
from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest
from adeu.mcp_components.tools.auth import login_to_adeu_cloud, logout_of_adeu_cloud
from adeu.mcp_components.tools.document import (
    accept_all_changes,
    diff_docx_files,
    process_document_batch,
    read_docx,
)
from adeu.mcp_components.tools.validation import validate_documents
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from docx import Document
from fastmcp.exceptions import ToolError


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


@pytest.fixture
def sample_docx(tmp_path) -> str:
    """Creates a basic DOCX file for testing."""
    doc = Document()
    doc.add_paragraph("This is the original text.")
    path = tmp_path / "sample.docx"
    doc.save(path)
    return str(path)


@pytest.fixture
def modified_docx(tmp_path) -> str:
    """Creates a slightly modified DOCX file for diff testing."""
    doc = Document()
    doc.add_paragraph("This is the modified text.")
    path = tmp_path / "modified.docx"
    doc.save(path)
    return str(path)


def test_read_docx(sample_docx):
    ctx = MockContext()
    result = asyncio.run(read_docx(file_path=sample_docx, ctx=ctx, clean_view=False))
    assert "This is the original text." in result.structured_content["markdown"]


def test_read_docx_file_not_found():
    ctx = MockContext()
    with pytest.raises(ToolError) as exc_info:
        asyncio.run(read_docx(file_path="nonexistent.docx", ctx=ctx))

    error_msg = str(exc_info.value)
    assert "Error reading file" in error_msg
    assert "not found" in error_msg


def test_diff_docx_files(sample_docx, modified_docx):
    ctx = MockContext()
    result = asyncio.run(
        diff_docx_files(
            original_path=sample_docx,
            modified_path=modified_docx,
            ctx=ctx,
            compare_clean=True,
        )
    )
    assert "@@ Word Patch @@" in result
    assert "- original" in result
    assert "+ modified" in result


def test_process_document_batch(sample_docx, tmp_path):
    ctx = MockContext()
    output_path = tmp_path / "output.docx"

    edits = [ModifyText(target_text="original text", new_text="new text", comment="Test comment")]

    result = asyncio.run(
        process_document_batch(
            original_docx_path=sample_docx,
            author_name="AI Agent",
            ctx=ctx,
            changes=edits,
            output_path=str(output_path),
        )
    )

    assert "Batch complete" in result
    assert "Edits: 1 applied, 0 skipped" in result
    assert output_path.exists()

    # Verify the edit was actually applied by reading the new file
    doc = Document(str(output_path))
    xml = doc.element.xml
    assert "<w:delText>original</w:delText>" in xml
    assert "<w:t>new</w:t>" in xml


def test_process_document_batch_validation_failure(sample_docx, tmp_path):
    ctx = MockContext()
    edits = [ModifyText(target_text="nonexistent target", new_text="new text")]

    result = asyncio.run(
        process_document_batch(
            original_docx_path=sample_docx,
            author_name="AI Agent",
            ctx=ctx,
            changes=edits,
            output_path=str(tmp_path / "fail.docx"),
        )
    )

    assert "Batch rejected" in result
    assert "Failed: Target text not found" in result


def test_accept_all_changes(sample_docx, tmp_path):
    ctx = MockContext()

    # First, programmatically create a doc with tracked changes
    with open(sample_docx, "rb") as f:
        engine = RedlineEngine(BytesIO(f.read()), author="Reviewer")
        engine.apply_edits([ModifyText(target_text="original", new_text="accepted")])
        tracked_stream = engine.save_to_stream()

    tracked_path = tmp_path / "tracked.docx"
    with open(tracked_path, "wb") as f:
        f.write(tracked_stream.getvalue())

    output_path = tmp_path / "clean.docx"

    result = asyncio.run(accept_all_changes(docx_path=str(tracked_path), ctx=ctx, output_path=str(output_path)))

    assert "Accepted all changes" in result
    assert output_path.exists()

    # Verify it is clean
    doc = Document(str(output_path))
    xml = doc.element.xml
    assert "w:ins" not in xml
    assert "w:del" not in xml
    assert "accepted" in xml
    assert "original" not in xml


# --- Cloud Auth & Validation Tool Mocks ---


@patch("adeu.auth.DesktopAuthManager.ensure_authenticated")
@patch("urllib.request.urlopen")
def test_login_to_adeu_cloud_success(mock_urlopen, mock_ensure_auth):
    ctx = MockContext()
    mock_ensure_auth.return_value = "mock_api_key"

    mock_response = MagicMock()
    mock_response.read.return_value = json.dumps({"email": "test@adeu.ai"}).encode("utf-8")
    mock_response.__enter__.return_value = mock_response
    mock_urlopen.return_value = mock_response

    result = asyncio.run(login_to_adeu_cloud(ctx=ctx))
    assert "Login successful" in result
    assert "test@adeu.ai" in result


@patch("adeu.auth.DesktopAuthManager.clear_api_key")
def test_logout_of_adeu_cloud(mock_clear_key):
    ctx = MockContext()
    result = asyncio.run(logout_of_adeu_cloud(ctx=ctx))
    assert "Successfully logged out" in result
    mock_clear_key.assert_called_once()


@patch("urllib.request.urlopen")
def test_validate_documents_success(mock_urlopen, sample_docx):
    ctx = MockContext()

    mock_response_data = {
        "report_markdown": (
            "# Validation Report\n\n"
            "## Consistency Issues\n"
            "- **Date mismatch**: Conflict in dates.\n"
            "  *Evidence: Evidence 1*\n\n"
            "## Risk Assessment\n"
            "- **BUYER - Unlimited Liability**: Buyer has no cap.\n"
            "  *Evidence: Evidence 2*"
        )
    }

    mock_response = MagicMock()
    mock_response.read.return_value = json.dumps(mock_response_data).encode("utf-8")
    mock_response.__enter__.return_value = mock_response
    mock_urlopen.return_value = mock_response

    result = asyncio.run(validate_documents(file_paths=[sample_docx], ctx=ctx, api_key="fake_key"))

    text_result = str(result.content)
    assert "Validation Report" in text_result
    assert "Date mismatch" in text_result
    assert "Unlimited Liability" in text_result
    assert "Evidence 1" in text_result
