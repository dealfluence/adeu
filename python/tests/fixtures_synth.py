import io
from pathlib import Path

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText, ReplyComment
from adeu.pagination import PAGE_TARGET_CHARS, paginate, split_structural_appendix
from adeu.redline.engine import RedlineEngine


def build_multi_author_docx(path: Path) -> Path:
    """Builds a synthetic multi-author DOCX fixture with tracked changes and comments.

    Contains:
    - H1 header (# Lease)
    - >= 3 body paragraphs
    - 2x3 table
    - 3 sequential RedlineEngine batches ("Jane Doe", "Bob Ross", "Adeu AI")
    - Including: del+ins pair, pure insertion, deletion, table-cell modify,
      modify with comment, and a comment reply action.
    """
    doc = Document()
    doc.add_heading("Lease", level=1)
    doc.add_paragraph("This Lease Agreement is made on January 1, 2026 by and between Landlord and Tenant.")
    doc.add_paragraph("The Tenant shall pay monthly rent of $2,000 on the first day of each month.")
    doc.add_paragraph("The Landlord agrees to maintain the premises in good repair and working order.")

    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Property"
    table.cell(0, 1).text = "Term"
    table.cell(0, 2).text = "Deposit"
    table.cell(1, 0).text = "123 Main St"
    table.cell(1, 1).text = "12 months"
    table.cell(1, 2).text = "2000 USD"

    doc.save(str(path))

    # Batch 1: Jane Doe
    # 1. del+ins pair with comment
    edit1 = ModifyText(
        target_text="monthly rent of $2,000",
        new_text="monthly rent of $2,500",
        comment="Increasing rent to market rate.",
    )
    # 2. pure insertion
    edit2 = ModifyText(
        target_text="premises",
        new_text="premises and grounds",
    )
    engine1 = RedlineEngine(io.BytesIO(path.read_bytes()), author="Jane Doe")
    engine1.process_batch([edit1, edit2])
    path.write_bytes(engine1.save_to_stream().getvalue())

    # Batch 2: Bob Ross
    # 3. deletion
    edit3 = ModifyText(
        target_text=" on January 1, 2026",
        new_text="",
    )
    # 4. table-cell modify
    edit4 = ModifyText(
        target_text="2000 USD",
        new_text="2500 USD",
        comment="Deposit matches new rent.",
    )
    engine2 = RedlineEngine(io.BytesIO(path.read_bytes()), author="Bob Ross")
    engine2.process_batch([edit3, edit4])
    path.write_bytes(engine2.save_to_stream().getvalue())

    # Batch 3: Adeu AI
    engine3 = RedlineEngine(io.BytesIO(path.read_bytes()), author="Adeu AI")
    comments_data = engine3.comments_manager.extract_comments_data()
    assert comments_data, "Expected comments from previous batches"
    target_com_id = list(comments_data.keys())[0]

    # 5. modify with comment
    edit5 = ModifyText(
        target_text="12 months",
        new_text="24 months",
        comment="Extended lease term.",
    )
    # 6. comment reply action
    action1 = ReplyComment(
        target_id=f"Com:{target_com_id}",
        text="Agreed to the updated rate.",
    )
    engine3.process_batch([edit5, action1])
    path.write_bytes(engine3.save_to_stream().getvalue())

    return path


def build_long_docx(path: Path, pages: int = 6) -> Path:
    """Fills paragraphs so raw Virtual Text projection paginates to `pages` synthetic pages."""
    doc = Document()
    target_para_len = PAGE_TARGET_CHARS // 19
    for i in range(1, pages + 1):
        num_paras = 18 if i < pages else 1
        for j in range(num_paras):
            prefix = f"Section {i} Para {j}: "
            padding_len = target_para_len - len(prefix)
            doc.add_paragraph(prefix + "x" * max(0, padding_len))

    doc.save(str(path))

    with open(path, "rb") as f:
        text = extract_text_from_stream(io.BytesIO(f.read()))

    body, appendix = split_structural_appendix(text)
    res = paginate(body, appendix)
    assert res.total_pages == pages, f"Expected {pages} synthetic pages, got {res.total_pages}"

    return path
