# FILE: tests/test_dry_run_validation.py
import io

from docx import Document

from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine


def test_validation_ambiguous_match_with_context():
    """
    Scenario: The exact bug reported. Target text exists in multiple clauses.
    Validates that the engine catches the duplication and extracts surrounding text for the LLM.
    """
    doc = Document()
    # Padded with introductory text to ensure the clause numbers fall fully within
    # the engine's 30-character backward context window without getting sliced.
    doc.add_paragraph("As specified in Clause 2.4 (Interest), the rate shall not exceed 150% of the base.")
    doc.add_paragraph(
        "As specified in Clause 6.1 (Liability Cap), Seller liability shall not exceed 150% of the total contract."
    )

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    # The LLM's flawed, overly generic edit
    edit = DocumentEdit(target_text="shall not exceed 150%", new_text="shall not exceed 100%")

    errors = engine.validate_edits([edit])

    assert len(errors) == 1
    error_msg = errors[0]

    # Verify the error structure
    assert "Edit 1 Failed: Ambiguous match" in error_msg
    assert "appears 2 times" in error_msg

    # Verify the context extraction is working so the LLM can see the difference
    assert "Clause 2.4" in error_msg
    assert "Clause 6.1" in error_msg
    assert "Please provide more surrounding context" in error_msg


def test_validation_not_found():
    """
    Scenario: Target text does not exist in the document (hallucination or severe typo).
    Validates that the engine catches it and reports it.
    """
    doc = Document()
    doc.add_paragraph("The liability of the Seller is strictly limited.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    edit = DocumentEdit(
        target_text="The liability of the Buyer",
        new_text="The liability of the Purchaser",
    )

    errors = engine.validate_edits([edit])

    assert len(errors) == 1
    assert "Edit 1 Failed: Target text not found" in errors[0]
    assert "The liability of the Buyer" in errors[0]


def test_validation_success_and_pure_insertion():
    """
    Scenario: Perfect targets. Includes a pure insertion (empty target text)
    which should bypass validation safely.
    """
    doc = Document()
    doc.add_paragraph("This is a unique clause.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    edit1 = DocumentEdit(target_text="unique clause", new_text="special clause")

    # Pure insertion (relies on index internally, skipping target_text validation)
    edit2 = DocumentEdit(target_text="", new_text="Inserted text.")

    errors = engine.validate_edits([edit1, edit2])

    # Should pass perfectly with zero errors
    assert len(errors) == 0


def test_validation_fuzzy_match_ambiguity():
    """
    Scenario: The target text is technically different from the doc, but the fuzzy matcher
    resolves it to multiple places (e.g., varying underscores in legal placeholders).
    """
    doc = Document()
    doc.add_paragraph("Buyer Signature: [__________]")  # 10 underscores
    doc.add_paragraph("Seller Signature: [_______]")  # 7 underscores

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    # LLM targets with 3 underscores. The fuzzy matcher should find BOTH signatures.
    edit = DocumentEdit(target_text="[___]", new_text="John Doe")

    errors = engine.validate_edits([edit])

    assert len(errors) == 1
    assert "Ambiguous match" in errors[0]
    assert "appears 2 times" in errors[0]
    assert "Buyer Signature" in errors[0]
    assert "Seller Signature" in errors[0]


def test_multiple_errors_accumulated():
    """
    Scenario: A batch with multiple bad edits.
    Validates that the dry run is exhaustive and doesn't short-circuit after the first error.
    """
    doc = Document()
    doc.add_paragraph("Duplicate word. Duplicate word.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)

    edit1 = DocumentEdit(target_text="Duplicate word.", new_text="Changed.")  # Ambiguous
    edit2 = DocumentEdit(target_text="Missing word.", new_text="Found.")  # Not Found

    errors = engine.validate_edits([edit1, edit2])

    assert len(errors) == 2
    assert "Edit 1 Failed: Ambiguous match" in errors[0]
    assert "Edit 2 Failed: Target text not found" in errors[1]
