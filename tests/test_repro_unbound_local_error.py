import io

import pytest
from docx import Document

from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine


def test_repro_unbound_local_curr_ins_id_failure():
    """
    Scenario: The FIRST run in a paragraph is empty (no text).

    This caused an UnboundLocalError in mapper.py because the 'curr_ins_id'
    variable initialization was skipped inside the 'if full_seg_text:' block,
    but the variable was accessed later in the loop for lookahead logic.
    """
    doc = Document()
    p = doc.add_paragraph()

    # 1. Empty Run FIRST
    # CRITICAL: We must give it a distinct property (like font name) so it is NOT
    # coalesced with the next run during normalize_docx().
    # We avoid Bold/Italic because those generate Markdown text ("**"), masking the bug.
    r1 = p.add_run()
    r1.font.name = "Arial"

    # 2. Subsequent content
    r2 = p.add_run("Subsequent text")
    r2.font.name = "Times New Roman"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream)
    # We define an edit just to trigger the mapping engine
    edit = DocumentEdit(target_text="Subsequent", new_text="Changed")

    try:
        # This triggers engine._apply_single_edit_heuristic -> mapper._build_map
        engine.apply_edits([edit])
    except UnboundLocalError as e:
        pytest.fail(f"Regression: UnboundLocalError raised! The fix is missing. Details: {e}")
    except Exception as e:
        # If the fix works, we might get other errors (e.g. not found), but NOT UnboundLocal
        if "local variable 'curr_ins_id' referenced before assignment" in str(e):
            pytest.fail(f"Regression: UnboundLocalError raised (wrapped)! Details: {e}")
