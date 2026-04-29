import io
import logging

from docx import Document

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from adeu.redline.mapper import DocumentMapper

logging.basicConfig(level=logging.DEBUG)

doc = Document()
doc.add_paragraph("Paragraph 1 end.")
doc.add_paragraph("Paragraph 2 start.")

stream = io.BytesIO()
doc.save(stream)
stream.seek(0)

mapper = DocumentMapper(doc)
print(repr(mapper.full_text))

edit = ModifyText(target_text="1 end.\n\nParagraph 2", new_text="1 end. Paragraph 2")

engine = RedlineEngine(stream)
applied, skipped = engine.apply_edits([edit])
print(f"Applied: {applied}, Skipped: {skipped}")
