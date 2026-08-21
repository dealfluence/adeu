"""
Parity twin of node/packages/core/src/repro_projection_parity_gaps.test.ts.

Two divergences kept the engines' projections from being byte-identical on real
documents. Both were found by running the FedRAMP SSP Moderate rev4 template
(shared/corpus) through both engines and diffing — neither is reachable from
the synthetic fixtures either suite used before.

1. EMPHASIS COALESCING. Adjacent runs with identical formatting must project as
   ONE marker span. Python already did this correctly (`build_paragraph_text`
   ignores trailing whitespace before testing for the closing marker); Node
   tested the literal tail, so hoisted boundary whitespace defeated it and it
   emitted ``**A** **B**`` where Python emitted ``**A B**``.

2. HEADER/FOOTER ENUMERATION. Python walks `w:sectPr` references, honouring
   Link-to-Previous, ``w:titlePg`` and ``w:evenAndOddHeaders``, so it projects
   what Word renders. Node listed every header/footer part in the package.

Python was the correct side of both, so these tests are characterization tests:
they pin the behaviour Node was brought into line with, so a future change to
the Python side cannot silently re-open the gap.
"""

import io

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.redline.mapper import DocumentMapper

NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

BOLD = "<w:rPr><w:b/></w:rPr>"
ITALIC = "<w:rPr><w:i/></w:rPr>"


def _paragraph_of(runs: list[tuple[str, str]]) -> str:
    body = "".join(f'<w:r>{rpr}<w:t xml:space="preserve">{t}</w:t></w:r>' for t, rpr in runs)
    return f"<w:p {NS}>{body}</w:p>"


def _docx(*body_xml: str) -> bytes:
    doc = Document()
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for xml in body_xml:
        el = parse_xml(xml)
        if sect_pr is not None:
            sect_pr.addprevious(el)
        else:
            body.append(el)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _project(data: bytes) -> str:
    return extract_text_from_stream(io.BytesIO(data), clean_view=True, include_appendix=False)


# ---------------------------------------------------------------------------
# 1. Emphasis coalescing
# ---------------------------------------------------------------------------
def test_bold_runs_separated_by_a_hoisted_space_merge():
    """The corpus shape: '**Name of Organization** **CSP Name...**' was wrong."""
    text = _project(_docx(_paragraph_of([("Name of Organization", BOLD), (" CSP Name System Connects To", BOLD)])))
    assert text.strip() == "**Name of Organization CSP Name System Connects To**"


def test_three_italic_runs_merge_into_one_span():
    """python: '_Version #.#,  Date_'; node emitted '_Version_ _#.#,_  _Date_'."""
    text = _project(_docx(_paragraph_of([("Version", ITALIC), (" #.#,", ITALIC), ("  Date", ITALIC)])))
    assert text.strip() == "_Version #.#,  Date_"


def test_runs_with_no_whitespace_between_them_merge():
    # A fully-bold paragraph also trips heading detection, hence the "## ".
    text = _project(_docx(_paragraph_of([("A", BOLD), ("B", BOLD)])))
    assert "**AB**" in text
    assert "**A****B**" not in text


def test_differing_formatting_is_not_merged():
    text = _project(_docx(_paragraph_of([("bold", BOLD), (" and ", ""), ("italic", ITALIC)])))
    assert text.strip() == "**bold** and _italic_"


def test_markers_stay_balanced_when_a_whitespace_only_same_style_run_follows():
    """Eliding the closer without a matching opener lost marker balance."""
    text = _project(_docx(_paragraph_of([("March 2012", BOLD), ("  ", BOLD)])))
    assert text.count("**") % 2 == 0, f"unbalanced emphasis markers: {text!r}"


def test_ingest_and_mapper_agree_on_merged_emphasis():
    data = _docx(_paragraph_of([("Name of Organization", BOLD), (" CSP Name", BOLD)]))
    assert DocumentMapper(Document(io.BytesIO(data)), clean_view=True).full_text == _project(data)


# ---------------------------------------------------------------------------
# 2. Header / footer enumeration
# ---------------------------------------------------------------------------
def _docx_with_header_footer(header: str | None, footer: str | None) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    if header is not None:
        sec.header.paragraphs[0].text = header
    if footer is not None:
        sec.footer.paragraphs[0].text = footer
    doc.add_paragraph("Body text.")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_referenced_header_and_footer_project():
    text = _project(_docx_with_header_footer("HEAD", "FOOT"))
    assert "HEAD" in text
    assert "Body text." in text
    assert "FOOT" in text


def test_first_page_header_is_ignored_without_title_pg():
    """python-docx creates the part, but Word renders it only with w:titlePg."""
    doc = Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "DEFAULT HEAD"
    sec.different_first_page_header_footer = False
    sec.first_page_header.paragraphs[0].text = "FIRST HEAD"
    doc.add_paragraph("Body text.")
    out = io.BytesIO()
    doc.save(out)

    text = _project(out.getvalue())
    assert "DEFAULT HEAD" in text
    assert "FIRST HEAD" not in text


def test_first_page_header_projects_once_the_section_opts_in():
    doc = Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "DEFAULT HEAD"
    sec.different_first_page_header_footer = True
    sec.first_page_header.paragraphs[0].text = "FIRST HEAD"
    doc.add_paragraph("Body text.")
    out = io.BytesIO()
    doc.save(out)

    text = _project(out.getvalue())
    assert "DEFAULT HEAD" in text
    assert "FIRST HEAD" in text


@pytest.mark.parametrize("opt_in", [False, True])
def test_even_page_header_follows_the_document_toggle(opt_in):
    doc = Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "DEFAULT HEAD"
    sec.even_page_header.paragraphs[0].text = "EVEN HEAD"
    doc.settings.odd_and_even_pages_header_footer = opt_in
    doc.add_paragraph("Body text.")
    out = io.BytesIO()
    doc.save(out)

    text = _project(out.getvalue())
    assert "DEFAULT HEAD" in text
    assert ("EVEN HEAD" in text) is opt_in


# ---------------------------------------------------------------------------
# 3. Run-level elements that used to fall through silently
# ---------------------------------------------------------------------------
@pytest.mark.parametrize(
    "name,run_xml,expected",
    [
        # A real hyphen glyph: dropping it merged the words either side.
        ("noBreakHyphen", "<w:r><w:t>e</w:t><w:noBreakHyphen/><w:t>mail</w:t></w:r>", "e-mail"),
        # Absolute-position tab: separates content, like w:tab.
        (
            "ptab",
            '<w:r><w:t>A</w:t><w:ptab w:relativeTo="margin" w:alignment="left" w:leader="none"/><w:t>B</w:t></w:r>',
            "A B",
        ),
        # Optional break hint. Word shows it only when the line actually
        # breaks, so projecting nothing is CORRECT — pinned so nobody
        # "fixes" it into a visible character.
        ("softHyphen", "<w:r><w:t>co</w:t><w:softHyphen/><w:t>operate</w:t></w:r>", "cooperate"),
    ],
)
def test_run_level_elements_project_their_glyph(name, run_xml, expected):
    assert _project(_docx(f"<w:p {NS}>{run_xml}</w:p>")).strip() == expected


def test_symbol_runs_are_still_dropped_deliberately():
    """
    w:sym is NOT projected. Symbol fonts map glyphs into the Unicode
    private-use area (Wingdings F0FE is a checked box), so the code point
    alone does not identify the character and guessing corrupts text. CC-1
    owns checkbox projection and needs a font-aware decision; this pins the
    status quo so the loss is a recorded choice rather than an oversight.
    """
    text = _project(_docx(f'<w:p {NS}><w:r><w:sym w:font="Wingdings" w:char="F0FE"/></w:r></w:p>'))
    assert text.strip() == ""
