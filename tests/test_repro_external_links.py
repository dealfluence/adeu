# FILE: tests/test_repro_external_links.py

import io

import pytest
from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def test_external_relationship_does_not_crash_comments_manager():
    """
    Validates the fix for:
    'ValueError: target_part property on _Relationship is undefined when target mode is External'

    This bug occurred when `_ensure_xml_part` looped through `doc.part.rels` looking for
    internal part targets, but encountered an external link (which has no `target_part`).
    """
    # 1. Create a base document
    doc = Document()
    doc.add_paragraph("Visit our website.")
    stream1 = io.BytesIO()
    doc.save(stream1)
    stream1.seek(0)

    # 2. Add an edit with a comment to force the creation of a Comments XML part.
    # The bug only triggers when CommentsManager tries to upgrade an EXISTING part.
    engine = RedlineEngine(stream1)
    engine.apply_edits([ModifyText(target_text="website", new_text="portal", comment="Update wording")])
    stream2 = engine.save_to_stream()

    # 3. Reload the document and inject an external relationship (e.g., a Hyperlink).
    # This perfectly mimics real-world documents with web links.
    doc_with_comments = Document(stream2)
    doc_with_comments.part.relate_to("https://kempower.com", RT.HYPERLINK, is_external=True)

    stream3 = io.BytesIO()
    doc_with_comments.save(stream3)
    stream3.seek(0)

    # 4. Trigger the bug via ingest
    # extract_text_from_stream initializes CommentsManager, which finds the existing comments part,
    # calls _ensure_xml_part, which loops over rels, hitting the new external HYPERLINK.
    try:
        text = extract_text_from_stream(stream3)
        assert "portal" in text
    except ValueError as e:
        if "target mode is External" in str(e):
            pytest.fail(f"Regression: Failed to handle external relationship in _ensure_xml_part: {e}")
        raise e
