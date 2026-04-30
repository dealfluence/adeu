import zipfile
from io import BytesIO

import docx

from adeu.sanitize.transforms import remove_all_comments


def test_sanitize_purges_empty_comment_parts():
    """
    Reproduces a bug where `remove_all_comments` physically ejects empty comment parts
    (e.g., word/comments.xml) from the OPC package.
    According to AI_CONTEXT.md (Architectural Decisions #8), this is explicitly forbidden.
    """
    doc = docx.Document()
    doc.add_paragraph("Test paragraph.")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    from adeu import ModifyText, RedlineEngine

    engine = RedlineEngine(stream, author="QA Bot")
    engine.process_batch([ModifyText(target_text="Test", new_text="Test", comment="This is a comment.")])

    commented_stream = engine.save_to_stream()
    commented_stream.seek(0)

    doc_obj = docx.Document(commented_stream)
    remove_all_comments(doc_obj)

    out_stream = BytesIO()
    doc_obj.save(out_stream)
    out_stream.seek(0)

    z = zipfile.ZipFile(out_stream)

    has_comments_part = any("word/comments" in f and f.endswith(".xml") for f in z.namelist())
    assert has_comments_part, (
        "Violation of Architectural Decision #8: Empty comment parts were violently ejected instead of left intact."
    )
