import io

import docx
import pytest

from adeu import ModifyText, RedlineEngine, ReplyComment
from adeu.redline.engine import BatchValidationError


def test_batch_engine_heading_depth_enforcement():
    """
    Reproduces a bug where the Batch Engine fails to enforce the maximum
    heading depth limit of 6, silently allowing `#` * 7 to pass through
    and corrupt the Markdown to OOXML mapping.
    """
    doc = docx.Document()
    doc.add_paragraph("Target Text")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA Bot")

    with pytest.raises(BatchValidationError):
        engine.process_batch([ModifyText(target_text="Target Text", new_text="####### Heading 7")])


def test_batch_engine_reply_to_fake_comment():
    """
    Reproduces a bug where replying to a non-existent comment
    raises an unhandled AttributeError instead of a clean validation error.
    """
    doc = docx.Document()
    doc.add_paragraph("Target Text")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="QA Bot")

    with pytest.raises(BatchValidationError):
        engine.process_batch([ReplyComment(target_id="Com:999", text="Hello")])
