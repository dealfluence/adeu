import io

import pytest
from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import XmlPart
from docx.oxml import parse_xml

from adeu.redline.engine import RedlineEngine


def inject_comments_part(doc: Document, xml_bytes: bytes):
    """Helper to inject a specific XML blob as the comments part."""
    # Ensure package has no comments part
    package = doc.part.package
    # Remove existing if any (simple approach: just append new one and relate it,
    # python-docx uses the relationship to find it)

    partname = package.next_partname("/word/comments%d.xml")
    part = XmlPart(partname, CT.WML_COMMENTS, parse_xml(xml_bytes), package)
    package.parts.append(part)
    doc.part.relate_to(part, RT.COMMENTS)


def test_ensure_namespaces_handles_self_closing_tag():
    """
    Scenario: Input is a self-closing tag (empty comments part).
    <w:comments xmlns:w="..." />

    Bug: Previously replaced with <w:comments ...> leaving no closing tag.
    Fix: Should expand to <w:comments ...></w:comments>.
    """
    doc = Document()
    doc.add_paragraph("Content")

    # Minimal self-closing tag without the required modern namespaces
    xml = b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    inject_comments_part(doc, xml)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # 1. Trigger parsing (CommentsManager.__init__ calls _ensure_namespaces)
    try:
        engine = RedlineEngine(stream)
    except Exception as e:
        pytest.fail(f"Crashed on self-closing comments tag: {e}")

    # 2. Verify XML structure
    part = engine.comments_manager.comments_part
    xml_str = part.blob.decode("utf-8")

    # Assert namespaces injected
    assert 'xmlns:w15="' in xml_str
    assert 'mc:Ignorable="' in xml_str

    # Assert valid structure (implicit if parse_xml succeeded above, but we check strings)
    # Python's lxml/ET usually serializes empty tags as <tag /> or <tag></tag>
    # As long as it parses, it's valid.
    assert "<w:comments" in xml_str


def test_ensure_namespaces_handles_populated_part():
    """
    Scenario: Input is a standard populated comments part.
    <w:comments ...><w:comment ...>...</w:comment></w:comments>

    Bug check: Ensure we don't accidentally double-close or corrupt children.
    """
    doc = Document()
    doc.add_paragraph("Content")

    xml = (
        b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:comment w:id="1" w:author="Test" w:date="2026-01-01T10:00:00Z">'
        b"<w:p><w:r><w:t>Existing</w:t></w:r></w:p>"
        b"</w:comment>"
        b"</w:comments>"
    )
    inject_comments_part(doc, xml)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    try:
        engine = RedlineEngine(stream)
    except Exception as e:
        pytest.fail(f"Crashed on populated comments tag: {e}")

    part = engine.comments_manager.comments_part
    xml_str = part.blob.decode("utf-8")

    # Assert namespaces injected
    assert 'xmlns:w15="' in xml_str

    # Assert content preserved
    assert 'w:author="Test"' in xml_str
    assert "Existing" in xml_str
