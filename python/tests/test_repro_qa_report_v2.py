import sys
from io import BytesIO

import pytest
from docx import Document

from tests.utils import get_mock_ctx, run_async


@pytest.mark.skipif(sys.platform != "win32", reason="Live Word COM tests require Windows platform")
class TestQaReportV2LiveWordParity:
    """
    Tests addressing F1 & F2 from the QA Report (Search + Targeted Write Spec).
    These tests ensure that the Live Word Python engine respects the new
    search arguments and the targeted write modes (match_mode/regex),
    achieving parity with the Node engine.
    """

    def test_live_word_f2_read_docx_search_params(self, active_word_app):
        """
        F2: read_docx search not implemented on the Python server.
        Passing search_query to read_active_word_document currently raises a TypeError.
        """
        from adeu.mcp_components.tools.live_word import read_active_word_document

        app, doc = active_word_app
        ctx = get_mock_ctx()

        doc.Range(0, doc.Content.End).Text = (
            "Test paragraph one.\n" + ("Padding line.\n" * 15) + "Test paragraph two.\n"
        )

        async def run_test():
            # This should succeed and NOT raise an unexpected keyword argument TypeError
            res = await read_active_word_document(
                ctx, search_query="two", search_regex=False, search_case_sensitive=False
            )

            content = getattr(res, "content", str(res))
            if isinstance(content, list):
                content = "".join(getattr(c, "text", str(c)) for c in content)

            assert "Test paragraph **two**" in content
            assert "Test paragraph one" not in content

        run_async(run_test())

    def test_live_word_f1_match_mode_first(self, active_word_app):
        """
        F1: match_mode="first" is a no-op on the Python server.
        Ambiguous targets should resolve to the first occurrence if match_mode="first",
        rather than raising an ambiguity failure.
        """
        from adeu.mcp_components.tools.live_word import process_active_word_batch
        from adeu.models import ModifyText

        app, doc = active_word_app
        ctx = get_mock_ctx()

        doc.Range(0, doc.Content.End).Text = "Ambiguous target. Ambiguous target.\n"

        async def run_test():
            edit = ModifyText(target_text="Ambiguous target.", new_text="Resolved target.", match_mode="first")
            res = await process_active_word_batch(ctx, changes=[edit], author_name="QA")

            assert "Applied: 1" in res
            assert "Failed: 0" in res

            doc.Revisions.AcceptAll()
            text = doc.Content.Text
            assert "Resolved target. Ambiguous target." in text

        run_async(run_test())

    def test_live_word_f1_regex_substitution(self, active_word_app):
        """
        F1: regex=True is a no-op on the Python server.
        The Live Word engine should evaluate regex patterns and backreferences.
        """
        from adeu.mcp_components.tools.live_word import process_active_word_batch
        from adeu.models import ModifyText

        app, doc = active_word_app
        ctx = get_mock_ctx()

        doc.Range(0, doc.Content.End).Text = "Article 123\n"

        async def run_test():
            edit = ModifyText(target_text=r"Article (\d+)", new_text=r"Article \1 (ref)", regex=True)
            res = await process_active_word_batch(ctx, changes=[edit], author_name="QA")

            assert "Applied: 1" in res
            assert "Failed: 0" in res

            doc.Revisions.AcceptAll()
            text = doc.Content.Text
            assert "Article 123 (ref)" in text

        run_async(run_test())


class TestQaReportV2EngineLogic:
    def test_f6_breadcrumb_resolver_no_body_dump(self):
        """
        F6: `**Path:**` breadcrumb dumps paragraph bodies instead of heading path.
        The regex `.*` incorrectly captures everything in a paragraph if it starts with `#`.
        """
        from adeu.redline.engine import RedlineEngine

        doc = Document()
        doc.add_heading("2. Confidentiality", level=1)
        # Add a huge paragraph that should definitely NOT show up in the breadcrumb path
        doc.add_paragraph("This is a huge paragraph body that should not be in the breadcrumb " * 10)
        doc.add_paragraph("Target phrase is here.")

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        engine.mapper._build_map()

        text = engine.mapper.full_text
        start_idx = text.find("Target phrase")

        # Manually invoke the breadcrumb resolver
        path, page = engine._get_heading_path_and_page(start_idx, text, [0, 1000])

        assert "2. Confidentiality" in path
        assert "This is a huge paragraph body" not in path


class TestQaReportV2Formatting:
    def test_p1_python_report_field_swap(self):
        """
        P1: Python edit report swaps `pages` and `heading_path`.
        We run a real batch and inspect the resulting stats payload to ensure
        pages is a list of ints and heading_path is a string.
        """
        from adeu.models import ModifyText
        from adeu.redline.engine import RedlineEngine

        doc = Document()
        doc.add_heading("Section 1", level=1)
        doc.add_paragraph("Find me.")

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="Find me.", new_text="Found.", match_mode="first")
        stats = engine.process_batch([edit], dry_run=False)

        report = stats["edits"][0]
        assert all(isinstance(p, int) for p in report.get("pages", [])), (
            f"P1 Bug: Pages contains strings: {report.get('pages')}"
        )
        assert isinstance(report.get("heading_path"), str), (
            f"P1 Bug: Heading path is not a string: {report.get('heading_path')}"
        )
        assert "Section 1" in report.get("heading_path")

    def test_p2_python_dry_run_stats(self):
        """
        P2: Python dry-run reports 'Mode: strict (0 occurrences modified)' regardless of mode.
        """
        from adeu.models import ModifyText
        from adeu.redline.engine import RedlineEngine

        doc = Document()
        doc.add_paragraph("Target string.")

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="Target string.", new_text="Changed.", match_mode="all")
        stats = engine.process_batch([edit], dry_run=True)

        report = stats["edits"][0]
        assert report.get("match_mode") == "all", "P2 Bug: match_mode not correctly passed in dry-run"
        assert report.get("occurrences_modified", 0) > 0, "P2 Bug: occurrences_modified is 0 in dry-run"

    def test_p4_python_preview_doubled_word_bug(self):
        """
        P4: Python preview snippet shows a doubled leading article.
        """
        from adeu.models import ModifyText
        from adeu.redline.engine import RedlineEngine

        doc = Document()
        doc.add_paragraph("He grants the Board of Directors authority.")

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="the Board of Directors", new_text="the Supervisory Board", match_mode="all")
        stats = engine.process_batch([edit], dry_run=True)

        clean_text = stats["edits"][0].get("clean_text", "")
        assert "the the Supervisory" not in clean_text
        assert "the Supervisory Board" in clean_text


class TestQaReportV2EngineSafety:
    def test_532_transactional_rollback(self):
        """
        §5.3.2: Transactional rollback on foreign-author overlaps.
        """
        from docx.oxml.ns import qn

        from adeu.models import ModifyText
        from adeu.redline.engine import BatchValidationError, RedlineEngine
        from adeu.utils.docx import create_element

        doc = Document()
        p = doc.add_paragraph("Match 1. ")
        ins = create_element("w:ins")
        ins.set(qn("w:id"), "1")
        ins.set(qn("w:author"), "Other User")
        r = create_element("w:r")
        t = create_element("w:t")
        t.text = "Match 2."
        r.append(t)
        ins.append(r)
        p._element.append(ins)

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text="Match 1. Match 2.", new_text="Replaced", match_mode="all")

        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([edit], dry_run=False)

        assert "targets an active insertion from another author" in str(exc.value)

    def test_533_double_sided_paragraph_merge(self):
        """
        §5.3.3: Regex spanning paragraph boundary with text on both sides must be rejected.
        """
        from adeu.models import ModifyText
        from adeu.redline.engine import BatchValidationError, RedlineEngine

        doc = Document()
        doc.add_paragraph("Before text.")
        doc.add_paragraph("After text.")

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine = RedlineEngine(stream)
        edit = ModifyText(target_text=r"text\.\n\nAfter", new_text="merged", regex=True)

        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([edit], dry_run=False)

        assert "spans a paragraph boundary with body text on both sides" in str(exc.value)

    def test_s1_comment_only_foreign_author_overlap(self):
        """
        S1: Transactional rollback blocks all-mode edit overlapping foreign COMMENT_ONLY edit.
        """
        from adeu.models import ModifyText
        from adeu.redline.engine import BatchValidationError, RedlineEngine

        doc = Document()
        doc.add_paragraph("This is constituting the Board of Directors today.")

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        engine_alice = RedlineEngine(stream, author="Alice")
        edit_alice = ModifyText(
            target_text="constituting the Board of Directors",
            new_text="constituting the Board of Directors",
            comment="Alice touches this clause",
        )
        engine_alice.process_batch([edit_alice], dry_run=False)

        stream_bob = engine_alice.save_to_stream()
        engine_bob = RedlineEngine(stream_bob, author="Bob")
        edit_bob = ModifyText(target_text="the Board of Directors", new_text="the Supervisory Board", match_mode="all")

        with pytest.raises(BatchValidationError) as exc:
            engine_bob.process_batch([edit_bob], dry_run=False)

        assert "another author" in str(exc.value)
