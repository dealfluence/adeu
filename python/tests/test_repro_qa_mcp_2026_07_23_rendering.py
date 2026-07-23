# FILE: tests/test_repro_qa_mcp_2026_07_23_rendering.py
"""
Repro tests for ADEU-MCP-QA-REPORT.md (2026-07-23, black-box QA of the Node
MCP server, v1.29.0+4bb70f9) — Python-engine mirror.

Findings covered (report finding -> test class below):

  F4   Outline/search rendering strips underscores from anchors and mangles
       placeholder runs (TestF4UnderscoreStripping):
         (a) outline: an anchor's leading underscore is consumed by markdown
             emphasis pairing when italic text follows in the same heading —
             `{#_Ref444615940} ... _scope_` renders as `{#Ref444615940} ...
             _scope` (outline._strip_inline_formatting). A lone anchor with no
             other underscore in the heading survives: the loss needs a
             pairing partner, exactly the emphasis-pairing mechanism the QA
             report suspected.
         (b) search: _emphasized_snippet strips single word-edge underscores
             unconditionally (_STYLE_MARKER_RE), so BOTH of two adjacent
             anchors lose their leading underscore in snippets.
         (c) outline: a placeholder run `[_________]` renders as `[___]`
             (the `__..__` bold rule plus the italic rule each eat a pair).

  F13  Outline rendering defects (TestF13OutlineRenderingDefects):
         (a) a bold heading whose runs span a line break projects as
             `**LINE ONE\nLINE TWO**` (marker elision across the break) and
             renders as one outline entry spanning two lines, each with an
             unbalanced `**`;
         (b) a 320-word paragraph styled as a heading appears IN FULL in the
             outline (a 2400+ char line) — desired: a truncated entry;
         (c) headings with empty text or bare ':' text (the visible remnant
             of auto-numbered headings) render as bare `## ` / `## : ` lines.

  F21  Table row-op rendering artifacts (TestF21TableRowOpRendering):
         (a) after an insert_row, the raw markup projection glues the change
             id to the last cell: `{++ Alice | signed |Chg:1++}` — an agent
             copying cell text picks up `|Chg:1++}` garbage.
         (b) DOES NOT REPRODUCE in Python — no test. Filling an EMPTY cell
             via its {#cell:paraId} anchor resolves to a pure insertion
             (`Op=INSERTION`, no empty tracked deletion), so no `{----}`
             token appears in the projection. Verified against a run-less
             cell, a cell with an empty w:t run, and a whitespace-only run.

  F22  Grammar + breadcrumb leaks (TestF22GrammarAndBreadcrumbs):
         (a) the defined-terms appendix says 'used 1 times' for a single-use
             term (build_structural_appendix, src/adeu/domain.py:352);
         (b) search-result Path breadcrumbs leak raw CriticMarkup when a
             heading carries a pending tracked change — the breadcrumb
             cleaner in build_search_response.get_heading strips bold/italic
             markers and {#anchors} but not {--..--}/{++..++}/{>>..<<}.

Every test is written test-first: it fails on current main and passes once
the finding is fixed.
"""

import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.ingest import _extract_text_from_doc, extract_text_from_stream
from adeu.mcp_components._response_builders import build_outline_response, build_search_response
from adeu.models import InsertTableRow, ModifyText
from adeu.redline.engine import RedlineEngine

# Rendered outline lines look like "## Heading text (p1)" / "(p1-p3)".
_OUTLINE_LINE_RE = re.compile(r"^(#{1,6}) (.*?) ?\(p\d+(?:-p\d+)?\)$")


def doc_to_stream(doc) -> io.BytesIO:
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Inserts a real w:bookmarkStart/-End pair named `name` into `paragraph`
    (before the runs), so the projection emits a `{#name}` anchor."""
    p_el = paragraph._element
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    pPr = p_el.find(qn("w:pPr"))
    p_el.insert(1 if pPr is not None else 0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p_el.append(end)


def outline_markdown(doc) -> str:
    """Runs the exact read_docx mode='outline' pipeline over `doc` and returns
    the rendered markdown (LLM-facing content)."""
    doc_obj = Document(doc_to_stream(doc))
    body = _extract_text_from_doc(doc_obj, clean_view=False, include_appendix=False)
    return build_outline_response(doc_obj, body, "doc.docx").content


def search_markdown(raw_text: str, query: str) -> str:
    """Runs the exact read_docx mode='search' pipeline over projected text."""
    return build_search_response(
        text=raw_text,
        search_query=query,
        search_regex=False,
        search_case_sensitive=True,
        page=None,
        file_path="doc.docx",
    ).content


# ---------------------------------------------------------------------------
# F4: underscore stripping in outline / search rendering
# ---------------------------------------------------------------------------


class TestF4UnderscoreStripping:
    def test_outline_preserves_anchor_underscore_before_italic_text(self):
        """F4(a): `{#_Ref444615940}` followed by an italic run in the same
        heading loses its leading underscore in the outline — the italic
        stripper pairs the anchor's `_` with the italic marker and consumes
        both. An agent copying the anchor targets `{#Ref444615940}`, which
        does not exist."""
        doc = Document()
        heading = doc.add_heading("Annex A ", level=2)
        add_bookmark(heading, "_Ref444615940", 1)
        italic_run = heading.add_run("scope")
        italic_run.italic = True
        doc.add_paragraph("Body text under the annex heading.")

        # Fixture sanity: the full-view projection carries the real anchor.
        body = _extract_text_from_doc(Document(doc_to_stream(doc)), include_appendix=False)
        assert "{#_Ref444615940}" in body, "fixture must project the underscored anchor"

        md = outline_markdown(doc)
        assert "_Ref444615940" in md, (
            "outline must preserve the anchor's leading underscore; it rendered:\n"
            + "\n".join(line for line in md.splitlines() if "Ref444615940" in line)
        )

    def test_search_snippet_preserves_underscores_of_adjacent_anchors(self):
        """F4(b): search snippets strip the leading underscore of anchors —
        with two adjacent anchors, BOTH lose it (_emphasized_snippet treats
        each word-edge underscore as an emphasis marker)."""
        doc = Document()
        doc.add_heading("Clause 1", level=1)
        p = doc.add_paragraph("This clause is anchored here for later reference.")
        add_bookmark(p, "_Ref444615940", 1)
        add_bookmark(p, "_Ref264019820", 2)

        raw = extract_text_from_stream(doc_to_stream(doc))
        assert "{#_Ref444615940}" in raw and "{#_Ref264019820}" in raw, "fixture sanity"

        md = search_markdown(raw, "anchored")
        assert "{#_Ref444615940}" in md, "search snippet must keep the anchor's leading underscore:\n" + md
        assert "{#_Ref264019820}" in md, "the adjacent anchor must keep its underscore too:\n" + md

    def test_outline_preserves_placeholder_run_literally(self):
        """F4(c): a `[_________]` fill-in placeholder inside a heading is
        mangled to `[___]` by the outline's bold/italic marker stripping."""
        doc = Document()
        doc.add_heading("Signature [_________] Block", level=2)
        doc.add_paragraph("Sign above the line.")

        md = outline_markdown(doc)
        assert "[_________]" in md, "outline must render the placeholder run literally; it rendered:\n" + "\n".join(
            line for line in md.splitlines() if "Signature" in line
        )


# ---------------------------------------------------------------------------
# F13: outline rendering defects
# ---------------------------------------------------------------------------


class TestF13OutlineRenderingDefects:
    def test_outline_lines_have_balanced_bold_markers(self):
        """F13(a): a bold (heuristic) heading whose bold runs span a line
        break projects as `**CONTRACT TERMS AND\nGENERAL CONDITIONS**` and the
        outline emits it across two lines, each carrying an unbalanced `**`."""
        doc = Document()
        p = doc.add_paragraph()
        first = p.add_run("CONTRACT TERMS AND")
        first.bold = True
        first.add_break()
        second = p.add_run("GENERAL CONDITIONS")
        second.bold = True
        doc.add_paragraph("Body follows.")

        md = outline_markdown(doc)
        assert "CONTRACT TERMS AND" in md, "fixture sanity: the heading must be in the outline"

        unbalanced = [line for line in md.splitlines() if line.count("**") % 2 != 0]
        assert not unbalanced, "every outline line must have balanced ** pairs; unbalanced lines:\n" + "\n".join(
            repr(line) for line in unbalanced
        )

    def test_outline_truncates_heading_styled_body_paragraph(self):
        """F13(b): a 320-word body paragraph styled as Heading 2 appears IN
        FULL in the outline (a 2400+ character line). A navigation outline
        entry should be truncated. 250 chars is a generous ceiling — any
        reasonable truncation (80/120/200) passes."""
        doc = Document()
        long_text = " ".join(f"word{i}" for i in range(320))
        doc.add_heading(long_text, level=2)
        doc.add_paragraph("Body.")

        md = outline_markdown(doc)
        heading_lines = [line for line in md.splitlines() if _OUTLINE_LINE_RE.match(line)]
        assert heading_lines, "fixture sanity: the heading must appear in the outline"

        too_long = [line for line in heading_lines if len(line) > 250]
        assert not too_long, "outline entries must be truncated; got a line of length " + str(
            max(len(line) for line in too_long)
        )

    def test_outline_has_no_empty_or_colon_only_heading_lines(self):
        """F13(c): headings whose visible text is empty or a bare ':' (the
        remnant of an auto-numbered heading) render as `## ` / `## : ` outline
        lines that carry no navigation signal."""
        doc = Document()
        doc.add_heading("Real Heading", level=1)
        doc.add_heading("", level=2)
        doc.add_heading(":", level=2)
        doc.add_paragraph("Body.")

        md = outline_markdown(doc)
        degenerate = []
        for line in md.splitlines():
            m = _OUTLINE_LINE_RE.match(line)
            if m and m.group(2).strip() in ("", ":"):
                degenerate.append(line)
        assert not degenerate, "outline must not emit empty or ':'-only heading lines; got:\n" + "\n".join(
            repr(line) for line in degenerate
        )


# ---------------------------------------------------------------------------
# F21: table row-op rendering artifacts
# ---------------------------------------------------------------------------


class TestF21TableRowOpRendering:
    def test_insert_row_projection_does_not_glue_change_id_to_last_cell(self):
        """F21(a): after an insert_row, the raw projection renders the tracked
        row as `{++ Alice | signed |Chg:1++}` — the ` |Chg:1++}` suffix reads
        as part of the last cell's text. Desired: no `|Chg:<id>` glue."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).paragraphs[0].text = "Name"
        table.cell(0, 1).paragraphs[0].text = "Date signed"

        engine = RedlineEngine(doc_to_stream(doc), author="QA")
        engine.process_batch([InsertTableRow(target_text="Name", cells=["Alice", "signed"], position="below")])
        raw = extract_text_from_stream(engine.save_to_stream(), clean_view=False)

        assert "Alice" in raw, "fixture sanity: the inserted row must be in the projection"
        glued = re.search(r"\|Chg:\d", raw)
        assert glued is None, "the change id must not be glued to the last cell; offending row:\n" + "\n".join(
            line for line in raw.splitlines() if "Chg" in line and "|" in line
        )

    # F21(b) — filling an EMPTY cell via its {#cell:paraId} anchor — does NOT
    # reproduce in Python: the engine resolves the fill to a pure insertion
    # (no empty tracked deletion), so no '{----}' token is emitted. Verified
    # against a run-less cell, a cell with an empty w:t run, and a
    # whitespace-only run. See module docstring; no test.


# ---------------------------------------------------------------------------
# F22: appendix grammar + Path breadcrumb CriticMarkup leaks
# ---------------------------------------------------------------------------


class TestF22GrammarAndBreadcrumbs:
    def test_appendix_uses_singular_time_for_single_use_term(self):
        """F22(a): the defined-terms appendix line for a term used exactly
        once must read 'used 1 time', not 'used 1 times'
        (build_structural_appendix, src/adeu/domain.py:352)."""
        doc = Document()
        doc.add_paragraph('"Agreement" means the contract between the parties described herein.')
        doc.add_paragraph("The Agreement shall commence on the Effective Date.")

        text = extract_text_from_stream(doc_to_stream(doc), include_appendix=True)
        term_lines = [line for line in text.splitlines() if '"Agreement"' in line and "used" in line]
        assert term_lines, "fixture sanity: the appendix must list the defined term 'Agreement'"

        line = term_lines[0]
        assert "used 1 times" not in line, "singular count must not be pluralized: " + line
        assert re.search(r"used 1 time\b", line), "expected 'used 1 time' in: " + line

    def test_search_path_breadcrumb_carries_no_raw_criticmarkup(self):
        """F22(b): when a heading carries a pending tracked change, the search
        result's Path breadcrumb leaks the raw CriticMarkup of the heading
        line ('{--Definitions--}{++Defined Terms++}{>>[Chg:1 delete] ...').
        Desired: breadcrumbs contain no {++ / {-- / {== / {>> tokens."""
        doc = Document()
        doc.add_heading("Definitions", level=1)
        doc.add_paragraph("The quick brown fox jumps over the lazy dog.")

        engine = RedlineEngine(doc_to_stream(doc), author="QA")
        engine.process_batch([ModifyText(target_text="Definitions", new_text="Defined Terms")])
        raw = extract_text_from_stream(engine.save_to_stream(), clean_view=False)
        assert "{--Definitions--}" in raw, "fixture sanity: the heading must carry a pending tracked change"

        md = search_markdown(raw, "fox")
        path_lines = [line for line in md.splitlines() if line.startswith("**Path:**")]
        assert path_lines, "the match under the heading must report a Path breadcrumb:\n" + md

        leaking = [line for line in path_lines if any(token in line for token in ("{++", "{--", "{==", "{>>"))]
        assert not leaking, "Path breadcrumbs must not leak raw CriticMarkup; got:\n" + "\n".join(
            repr(line) for line in leaking
        )
