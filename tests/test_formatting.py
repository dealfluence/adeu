import io

from docx import Document

from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine


def _is_element_bold(run_element) -> bool:
    b_tags = run_element.xpath("./w:rPr/w:b")
    if not b_tags:
        return False
    val = b_tags[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    if val is None:
        return True
    if val.lower() in ("0", "false", "off"):
        return False
    return True


def test_insertion_inherits_next_run_style_heuristic():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Start ")
    r1.bold = False
    r2 = p.add_run("Important")
    r2.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(
        target_text="",
        new_text="Very ",
    )
    edit._match_start_index = 6

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)

    runs = doc.element.xpath('//w:r[w:t[contains(text(), "Very ")]]')
    assert len(runs) >= 1, "Inserted text 'Very ' not found in any run"

    target_run = runs[0]
    assert _is_element_bold(target_run), f"Inserted text should inherit Bold. XML: {target_run.xml}"


def test_insertion_defaults_to_prev_run_style_if_no_space():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Hello ")
    r1.bold = False
    r2 = p.add_run("World")
    r2.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(
        target_text="",
        new_text="Big",
    )
    edit._match_start_index = 6

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)

    runs = doc.element.xpath('//w:r[w:t[contains(text(), "Big")]]')
    assert len(runs) >= 1
    target_run = runs[0]

    assert not _is_element_bold(target_run)
