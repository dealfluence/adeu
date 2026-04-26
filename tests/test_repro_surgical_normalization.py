import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from adeu.redline.engine import RedlineEngine


def test_engine_init_does_not_strip_proof_err():
    """
    Bug #11: RedlineEngine should not perform global document normalization
    (which strips proofErr tags) during initialization or batch processing.
    It should operate in Surgical Mode.
    """
    doc = Document()
    p = doc.add_paragraph("Some text ")
    proof_err = OxmlElement("w:proofErr")
    proof_err.set(qn("w:type"), "spellStart")
    p._element.append(proof_err)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Verify it exists before engine initialization
    test_doc = Document(stream)
    assert len(test_doc.element.xpath("//w:proofErr")) == 1
    stream.seek(0)

    # Init engine (this currently calls normalize_docx and strips the tag)
    engine = RedlineEngine(stream)

    surviving = engine.doc.element.xpath("//w:proofErr")
    assert len(surviving) == 1, "proofErr was stripped! Engine init should not trigger global normalization."
