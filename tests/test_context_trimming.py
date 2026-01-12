import io

from docx import Document

from adeu.models import DocumentEdit
from adeu.redline.engine import RedlineEngine, _trim_common_context


def test_trim_logic_basic():
    """Prefix and Suffix exist."""
    t = "Context A Context"
    n = "Context B Context"
    p, s = _trim_common_context(t, n)
    assert p == 8  # "Context "
    assert s == 8  # " Context"
    # Remainder: "A", "B"


def test_trim_logic_prefix_only():
    t = "Hello World"
    n = "Hello User"
    p, s = _trim_common_context(t, n)
    assert p == 6  # "Hello "
    assert s == 0


def test_trim_logic_suffix_only():
    t = "Old Item"
    n = "New Item"
    p, s = _trim_common_context(t, n)
    assert p == 0
    assert s == 5  # " Item"


def test_trim_logic_morph_to_insert():
    t = "Prefix"
    n = "Prefix Added"
    p, s = _trim_common_context(t, n)
    assert p == 6
    assert s == 0


def test_end_to_end_context_cleanup():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Start ")
    p.add_run("Middle")
    p.add_run(" End")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="Start Middle End", new_text="Start Center End")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    assert "w:del" not in xml.split("Start ")[0][-20:]
    assert "<w:delText>Middle</w:delText>" in xml
    assert "<w:t>Center</w:t>" in xml
    assert "<w:t>Center End</w:t>" not in xml


def test_auto_strip_insertion_duplication():
    doc = Document()
    doc.add_paragraph("Liability Cap.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = DocumentEdit(target_text="Liability Cap.", new_text="Liability Cap. SLA Clause.")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    assert xml.count("Liability Cap.") == 1
    assert "SLA Clause." in xml
