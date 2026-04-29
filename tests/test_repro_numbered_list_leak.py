import io

from docx import Document
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_visible_runs


def test_markdown_numbered_list_leak():
    """
    Test Case: Injecting a new numbered list item using Markdown syntax (`\\n\\n1. Numbered Item`)
    should trigger a paragraph split and create a proper `<w:numPr>` list item,
    without leaking the literal `1. ` into the text run.
    """
    doc = Document()
    doc.add_paragraph("Reference text.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(target_text="Reference text.", new_text="Reference text.\n\n1. Numbered Item")

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])
    assert applied == 1, "Edit should be applied"

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    assert len(doc_result.paragraphs) >= 2, "Expected a new paragraph to be created."

    p_new = doc_result.paragraphs[1]
    visible_text = "".join(r.text for r in get_visible_runs(p_new))

    # 1. The literal '1. ' should NOT be in the text.
    assert not visible_text.startswith("1. "), f"Numbered list Markdown leaked into text: '{visible_text}'"
    assert "Numbered Item" in visible_text

    # 2. It should have list properties (either numPr or a List style).
    pPr = p_new._element.pPr
    has_numPr = False
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            has_numPr = True

    is_list_style = p_new.style is not None and "List" in p_new.style.name

    assert has_numPr or is_list_style, "Paragraph is not formatted as a list."
