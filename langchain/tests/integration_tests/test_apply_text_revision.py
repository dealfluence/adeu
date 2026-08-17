from __future__ import annotations

from pathlib import Path

from docx import Document

from langchain_adeu import AdeuApplyTextRevision, AdeuReadDocx


class TestApplyTextRevisionBehavior:
    def test_revision_is_applied_as_tracked_changes(self, working_docx: Path, output_path: Path) -> None:
        read = AdeuReadDocx()
        clean = read.invoke({"reasoning": "t", "file_path": str(working_docx), "clean_view": True, "page": "all"})
        body = clean.split("\n\n", 1)[1]
        revised = body.replace("document", "agreement")
        assert revised != body, "fixture no longer contains the word this test rewrites"

        msg = AdeuApplyTextRevision().invoke(
            {
                "name": "adeu_apply_text_revision",
                "args": {
                    "reasoning": "t",
                    "file_path": str(working_docx),
                    "revised_text": revised,
                    "output_path": str(output_path),
                    "author": "AI Reviewer",
                },
                "id": "test-text-revision",
                "type": "tool_call",
            }
        )
        assert msg.artifact["success"] is True
        assert msg.artifact["verified"] is True
        assert msg.artifact["output_path"] == str(output_path)
        assert output_path.exists()
        assert msg.artifact["edits_applied"] >= 1

    def test_major_deletion_is_blocked_without_the_flag(self, working_docx: Path, output_path: Path) -> None:
        msg = AdeuApplyTextRevision().invoke(
            {
                "name": "adeu_apply_text_revision",
                "args": {
                    "reasoning": "t",
                    "file_path": str(working_docx),
                    "revised_text": "x",
                    "output_path": str(output_path),
                },
                "id": "test-major-deletion",
                "type": "tool_call",
            }
        )
        assert msg.artifact["success"] is False
        assert not output_path.exists()

    def test_verification_failure_reports_the_unverified_copy(self, tmp_path: Path, output_path: Path) -> None:
        # The golden fixture has no headings, and a heading is the cheapest way to
        # trip the engine's clean-text gate: a heading paragraph cannot be removed
        # by text replacement, so the applied document's clean text diverges from
        # the supplied text (python/src/adeu/text_revision.py:248-277).
        source = tmp_path / "headed.docx"
        doc = Document()
        doc.add_heading("Section One", level=1)
        doc.add_paragraph("The first body paragraph of the agreement text goes here.")
        doc.add_heading("Section Two", level=1)
        doc.add_paragraph("The second body paragraph of the agreement text goes here.")
        doc.save(str(source))

        clean = AdeuReadDocx().invoke({"reasoning": "t", "file_path": str(source), "clean_view": True, "page": "all"})
        body = clean.split("\n\n", 1)[1]
        revised = "\n".join(line for line in body.split("\n") if "Section Two" not in line)
        assert revised != body, "fixture no longer contains the heading this test drops"

        msg = AdeuApplyTextRevision().invoke(
            {
                "name": "adeu_apply_text_revision",
                "args": {
                    "reasoning": "t",
                    "file_path": str(source),
                    "revised_text": revised,
                    "output_path": str(output_path),
                },
                "id": "test-verification-gate",
                "type": "tool_call",
            }
        )
        assert msg.artifact["success"] is False
        assert msg.artifact["verified"] is False
        assert msg.artifact["status"] == "verification_failed"
        assert msg.artifact["output_path"] is None
        assert msg.artifact["edits_applied"] == 0
        # The requested file was NOT written; only the diagnostic copy exists.
        assert not output_path.exists()
        unverified = Path(msg.artifact["unverified_output_path"])
        assert unverified.exists()
        assert unverified.name == f"{output_path.stem}.unverified.docx"
        assert "verification failed" in msg.content
