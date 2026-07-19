# FILE: tests/test_property_invariants.py
"""
Hypothesis property tests over the core Adeu invariants.

These generalize the 2026-07-18 v6 QA findings from single repro cases to
generated families:

  P1  Text round-trip (pinned path): for any document and any paragraph-level
      rewrite, apply(orig, alignment_diff) + accept-all reproduces the edited
      text exactly. The pinned path resolves positionally and must never
      reject or corrupt.
  P2  Text round-trip (JSON path): the self-contained edits emitted for
      `diff --json` either reproduce the edited text exactly or are rejected
      loudly (transactional BatchValidationError / skipped edits with no
      output). Silently wrong output is the only failure.
  P3  DOCX-to-DOCX structured diff closure over tables: replaying
      generate_structured_edits WITHOUT pins (the `apply changes.json` shape)
      reproduces the modified document, or fails loudly. When the diff
      emitted no warnings, it must not fail at all.
  P4  Sanitize sentinel scan: metadata written into core properties and
      custom document properties never survives into any member of the
      sanitized package, whatever characters the values contain.
  P5  trim_common_context structural invariants: the trimmed prefix/suffix
      are genuinely common, never overlap, and re-compose both strings.

Alphabets deliberately exclude Markdown/CriticMarkup metacharacters
(#, *, _, |, [], {}, ^) — text containing those exercises separately
validated projection surfaces, not these invariants.

Run a deeper hunt with: uv run pytest tests/test_property_invariants.py \
    --hypothesis-profile=hunt
"""

import json
import zipfile
from io import BytesIO

from docx import Document
from hypothesis import HealthCheck, assume, given, settings
from hypothesis import strategies as st
from pydantic import TypeAdapter

from adeu.diff import (
    generate_edits_via_paragraph_alignment,
    generate_structured_edits,
    make_edits_self_contained,
    trim_common_context,
)
from adeu.ingest import _extract_text_from_doc, extract_text_from_stream
from adeu.models import BatchChanges
from adeu.redline.engine import BatchValidationError, RedlineEngine
from adeu.sanitize.core import sanitize_docx

# Profiles ("default": 25 examples, "hunt": 300) are registered in
# tests/conftest.py so --hypothesis-profile resolves at configure time.


# ---------------------------------------------------------------------------
# Strategies
# ---------------------------------------------------------------------------

WORD_ALPHABET = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789äöåéλπ€&"

word_st = st.text(alphabet=WORD_ALPHABET, min_size=1, max_size=10)

paragraph_st = st.lists(word_st, min_size=1, max_size=8).map(lambda ws: " ".join(ws) + ".")

cell_st = st.lists(word_st, min_size=1, max_size=3).map(" ".join)


@st.composite
def doc_and_mod(draw):
    """A document (paragraph list) and a paragraph-level rewrite of it."""
    paras = draw(st.lists(paragraph_st, min_size=1, max_size=6, unique=True))
    mod = list(paras)
    for _ in range(draw(st.integers(min_value=1, max_value=3))):
        op = draw(st.sampled_from(["insert", "delete", "replace", "append_word", "drop_word"]))
        if op == "insert":
            mod.insert(draw(st.integers(min_value=0, max_value=len(mod))), draw(paragraph_st))
        elif op == "delete" and mod:
            mod.pop(draw(st.integers(min_value=0, max_value=len(mod) - 1)))
        elif op == "replace" and mod:
            mod[draw(st.integers(min_value=0, max_value=len(mod) - 1))] = draw(paragraph_st)
        elif op == "append_word" and mod:
            i = draw(st.integers(min_value=0, max_value=len(mod) - 1))
            mod[i] = mod[i] + " " + draw(word_st)
        elif op == "drop_word" and mod:
            i = draw(st.integers(min_value=0, max_value=len(mod) - 1))
            words = mod[i].split(" ")
            if len(words) > 1:
                words.pop(draw(st.integers(min_value=0, max_value=len(words) - 1)))
                mod[i] = " ".join(words)
    return paras, mod


@st.composite
def table_and_mod(draw):
    """A 3-column table (unique rows) and a row-level rewrite of it."""
    n_rows = draw(st.integers(min_value=1, max_value=4))
    rows = []
    seen = set()
    for _ in range(n_rows):
        row = tuple(draw(cell_st) for _ in range(3))
        if row not in seen and all(c.strip() for c in row):
            seen.add(row)
            rows.append(list(row))
    assume(rows)

    mod = [list(r) for r in rows]
    for _ in range(draw(st.integers(min_value=1, max_value=3))):
        op = draw(st.sampled_from(["insert", "delete", "modify_cell"]))
        if op == "insert":
            new_row = [draw(cell_st) for _ in range(3)]
            mod.insert(draw(st.integers(min_value=0, max_value=len(mod))), new_row)
        elif op == "delete" and mod:
            mod.pop(draw(st.integers(min_value=0, max_value=len(mod) - 1)))
        elif op == "modify_cell" and mod:
            r = draw(st.integers(min_value=0, max_value=len(mod) - 1))
            c = draw(st.integers(min_value=0, max_value=2))
            mod[r][c] = draw(cell_st)
    return rows, mod


# ---------------------------------------------------------------------------
# Builders
# ---------------------------------------------------------------------------


def build_doc_stream(paragraphs) -> BytesIO:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    s = BytesIO()
    doc.save(s)
    s.seek(0)
    return s


def build_table_doc_stream(paragraphs, rows) -> BytesIO:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if rows:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for r, row in enumerate(rows):
            for c, cell in enumerate(row):
                table.rows[r].cells[c].text = cell
    doc.add_paragraph("Trailing paragraph after the table.")
    s = BytesIO()
    doc.save(s)
    s.seek(0)
    return s


def clean_text(stream: BytesIO) -> str:
    return extract_text_from_stream(BytesIO(stream.getvalue()), clean_view=True)


def accept_and_extract(engine: RedlineEngine) -> str:
    engine.accept_all_revisions(remove_comments=True)
    return extract_text_from_stream(engine.save_to_stream(), clean_view=True)


# ---------------------------------------------------------------------------
# P1 — pinned text round trip is exact
# ---------------------------------------------------------------------------


@given(doc_and_mod())
def test_p1_pinned_text_roundtrip_is_exact(data):
    paras, mod = data
    orig_stream = build_doc_stream(paras)
    text_orig = clean_text(orig_stream)
    text_mod = "\n\n".join(mod)

    edits = generate_edits_via_paragraph_alignment(text_orig, text_mod)
    engine = RedlineEngine(BytesIO(orig_stream.getvalue()), author="Fuzz")
    stats = engine.process_batch(list(edits))
    assert stats["edits_skipped"] == 0, f"pinned edits skipped: {stats['skipped_details']}"

    final = accept_and_extract(engine)
    assert final == text_mod


# ---------------------------------------------------------------------------
# P2 — JSON text round trip is exact or loud
# ---------------------------------------------------------------------------


@given(doc_and_mod())
def test_p2_json_text_roundtrip_is_exact_or_loud(data):
    paras, mod = data
    orig_stream = build_doc_stream(paras)
    text_orig = clean_text(orig_stream)
    text_mod = "\n\n".join(mod)

    edits = generate_edits_via_paragraph_alignment(text_orig, text_mod)
    edits = make_edits_self_contained(edits, text_orig)
    dumped = json.loads(json.dumps([e.model_dump() for e in edits]))
    changes = TypeAdapter(BatchChanges).validate_python(dumped)

    engine = RedlineEngine(BytesIO(orig_stream.getvalue()), author="Fuzz")
    try:
        stats = engine.process_batch(list(changes))
    except BatchValidationError:
        return  # loud, transactional rejection — fail-closed is compliant
    if stats["edits_skipped"] > 0:
        return  # loud failure: the CLI writes no output for this outcome

    final = accept_and_extract(engine)
    assert final == text_mod, "JSON replay applied cleanly but produced the wrong document"


# ---------------------------------------------------------------------------
# P3 — structured docx-to-docx diff closure (tables), pins dropped
# ---------------------------------------------------------------------------


@given(st.lists(paragraph_st, min_size=1, max_size=2, unique=True), table_and_mod())
def test_p3_structured_table_diff_replays_or_fails_loud(paras, tables):
    rows, mod_rows = tables
    orig_stream = build_table_doc_stream(paras, rows)
    mod_stream = build_table_doc_stream(paras, mod_rows)

    doc_o = Document(BytesIO(orig_stream.getvalue()))
    doc_m = Document(BytesIO(mod_stream.getvalue()))
    text_o, struct_o = _extract_text_from_doc(doc_o, clean_view=True, include_appendix=False, return_structure=True)
    text_m, struct_m = _extract_text_from_doc(doc_m, clean_view=True, include_appendix=False, return_structure=True)

    edits, warnings = generate_structured_edits(text_o, struct_o, text_m, struct_m)
    dumped = json.loads(json.dumps([e.model_dump() for e in edits]))
    changes = TypeAdapter(BatchChanges).validate_python(dumped)

    engine = RedlineEngine(BytesIO(orig_stream.getvalue()), author="Fuzz")
    try:
        stats = engine.process_batch(list(changes))
    except BatchValidationError:
        assert warnings or _has_texty_hazard(text_o, edits), (
            "diff emitted no warning yet its output was rejected outright"
        )
        return
    if stats["edits_skipped"] > 0:
        assert warnings, f"un-warned diff output failed to apply: {stats['skipped_details']}"
        return

    final = accept_and_extract(engine)
    want = clean_text(mod_stream)
    assert final == want, "structured diff replay applied cleanly but produced the wrong document"


def _has_texty_hazard(text_orig: str, edits) -> bool:
    """
    True when any generated text anchor occurs more than once in the original
    projection — the fail-closed ambiguity path documented on the diff.
    """
    for e in edits:
        target = getattr(e, "target_text", "") or ""
        if target and text_orig.count(target) > 1:
            return True
    return False


# ---------------------------------------------------------------------------
# P4 — no generated metadata value survives sanitize
# ---------------------------------------------------------------------------

xml_text_st = st.text(
    alphabet=st.characters(min_codepoint=0x20, max_codepoint=0x2FFF, blacklist_categories=("Cs",)),
    min_size=1,
    max_size=24,
)

prop_name_st = st.text(
    alphabet="abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789_",
    min_size=1,
    max_size=16,
)


@given(
    author=xml_text_st,
    identifier=xml_text_st,
    description=xml_text_st,
    prop_name=prop_name_st,
    prop_value=xml_text_st,
)
@settings(max_examples=25, deadline=None, suppress_health_check=[HealthCheck.too_slow])
def test_p4_sanitize_leaves_no_sentinel_in_any_member(
    tmp_path_factory, author, identifier, description, prop_name, prop_value
):
    from xml.sax.saxutils import escape, quoteattr

    tmp_path = tmp_path_factory.mktemp("fuzz_sanitize")
    marker = "ZQXSENTINEL"
    doc = Document()
    doc.add_paragraph("Ordinary body content.")
    core = doc.core_properties
    core.author = marker + author
    core.identifier = marker + identifier
    core.comments = marker + description
    src = tmp_path / "meta.docx"
    doc.save(src)

    # Inject a custom-properties part carrying the generated value.
    raw = src.read_bytes()
    out = BytesIO()
    with zipfile.ZipFile(BytesIO(raw)) as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = (
                    data.decode("utf-8").replace(
                        "</Types>",
                        '<Override PartName="/docProps/custom.xml" '
                        'ContentType="application/vnd.openxmlformats-officedocument.custom-properties+xml"/></Types>',
                    )
                ).encode("utf-8")
            elif item.filename == "_rels/.rels":
                data = (
                    data.decode("utf-8").replace(
                        "</Relationships>",
                        '<Relationship Id="rIdCP" Type="http://schemas.openxmlformats.org/officeDocument/2006/'
                        'relationships/custom-properties" Target="docProps/custom.xml"/></Relationships>',
                    )
                ).encode("utf-8")
            zout.writestr(item, data)
        custom = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
            'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
            f'<property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="2" name={quoteattr(prop_name)}>'
            f"<vt:lpwstr>{escape(marker + prop_value)}</vt:lpwstr></property>"
            "</Properties>"
        )
        zout.writestr("docProps/custom.xml", custom)
    src.write_bytes(out.getvalue())

    dest = tmp_path / "meta_sanitized.docx"
    sanitize_docx(str(src), str(dest))

    with zipfile.ZipFile(dest) as z:
        for name in z.namelist():
            member = z.read(name).decode("utf-8", errors="ignore")
            assert marker not in member, f"sanitized package member {name} still contains generated metadata"


# ---------------------------------------------------------------------------
# P5 — trim_common_context structural invariants
# ---------------------------------------------------------------------------

trim_text_st = st.text(
    alphabet="abcdefg .,*_#\n|[]{}xyzäö0123456789",
    min_size=0,
    max_size=40,
)


@given(target=trim_text_st, new=trim_text_st)
@settings(max_examples=200, deadline=None)
def test_p5_trim_common_context_invariants(target, new):
    prefix_len, suffix_len = trim_common_context(target, new)

    assert 0 <= prefix_len <= min(len(target), len(new))
    assert 0 <= suffix_len <= min(len(target), len(new))
    assert prefix_len + suffix_len <= len(target)
    assert prefix_len + suffix_len <= len(new)
    if prefix_len:
        assert target[:prefix_len] == new[:prefix_len], "trimmed prefix is not common"
    if suffix_len:
        assert target[-suffix_len:] == new[-suffix_len:], "trimmed suffix is not common"
