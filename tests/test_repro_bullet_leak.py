import io

from docx import Document
from docx.oxml.ns import qn

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine
from adeu.utils.docx import get_visible_runs


def test_markdown_bullet_leak():
    """
    Test Case: VAL-OBS-8
    Injecting a new list item using Markdown syntax (`\\n\\n* New Bullet`)
    should trigger a paragraph split and create a proper `<w:numPr>` list item,
    without leaking the literal `*` into the text run.
    """
    doc = Document()
    doc.add_paragraph("Reference is also made to this Section 2 for further detail.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Simulating the change that failed in Phase 2
    edit = ModifyText(
        target_text="Reference is also made to this Section 2 for further detail.",
        new_text="Reference is also made to this Section 2 for further detail.\n\n* New Regression Test Bullet",
    )

    engine = RedlineEngine(stream)
    applied, skipped = engine.apply_edits([edit])
    assert applied == 1, "Edit should be applied"

    result_stream = engine.save_to_stream()
    doc_result = Document(result_stream)

    # We expect 2 paragraphs.
    assert len(doc_result.paragraphs) >= 2, "Expected a new paragraph to be created."

    # Check the newly inserted paragraph
    p_new = doc_result.paragraphs[1]
    visible_text = "".join(r.text for r in get_visible_runs(p_new))

    print(f"DEBUG: Inserted paragraph visible text: '{visible_text}'")

    # 1. The literal '* ' should NOT be in the text.
    assert not visible_text.startswith("* "), f"Bullet Markdown leaked into text: '{visible_text}'"
    assert "New Regression Test Bullet" in visible_text

    # 2. It should have list properties (either numPr or a List style).
    pPr = p_new._element.pPr
    has_numPr = False
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            has_numPr = True

    is_list_style = p_new.style is not None and "List" in p_new.style.name

    assert has_numPr or is_list_style, "Paragraph is not formatted as a list."
