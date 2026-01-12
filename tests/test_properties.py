import io

from docx import Document
from docx.oxml.ns import qn
from hypothesis import given, settings
from hypothesis import strategies as st

from adeu.diff import generate_edits_from_text
from adeu.ingest import extract_text_from_stream
from adeu.redline.engine import RedlineEngine


def extract_accepted_text_from_xml(doc_stream: io.BytesIO) -> str:
    doc_stream.seek(0)
    doc = Document(doc_stream)
    full_text = []

    for p in doc.paragraphs:
        p_text = ""
        for child in p._element:
            tag = child.tag
            if tag == qn("w:r"):
                for t in child.findall(qn("w:t")):
                    if t.text:
                        p_text += t.text
            elif tag == qn("w:ins"):
                for r in child.findall(qn("w:r")):
                    for t in r.findall(qn("w:t")):
                        if t.text:
                            p_text += t.text
        full_text.append(p_text)
    return "\n\n".join(full_text)


text_strategy = st.text(alphabet=st.characters(blacklist_categories=("Cc", "Cs")), min_size=1, max_size=50)


@settings(max_examples=50, deadline=None)
@given(paragraphs=st.lists(text_strategy, min_size=1, max_size=5))
def test_fuzz_roundtrip_correctness(paragraphs):
    doc = Document()
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        p_element = doc.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    for p_text in paragraphs:
        doc.add_paragraph(p_text)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    original_text = extract_text_from_stream(stream)
    modified_text = original_text.replace("a", "XYZ").replace("e", "") + " END"

    edits = generate_edits_from_text(original_text, modified_text)

    stream.seek(0)
    engine = RedlineEngine(stream)
    try:
        engine.apply_edits(edits)
    except Exception as e:
        raise RuntimeError(f"Engine crashed on input: {paragraphs}") from e


@settings(max_examples=20)
@given(text=text_strategy)
def test_fuzz_split_run_mechanics(text):
    if len(text) < 3:
        return

    doc = Document()
    p = doc.add_paragraph()
    p.add_run(text)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    from adeu.models import DocumentEdit

    edit = DocumentEdit(target_text=text, new_text="")

    stream.seek(0)
    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    final_text = extract_accepted_text_from_xml(result_stream)

    assert final_text.strip() == ""
